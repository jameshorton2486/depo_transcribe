"""Opt-in per-stage transcript capture for the active production path.

Writes plain-text snapshots of pipeline output to
``<case_dir>/_walkthrough/<stage_name>.txt`` when the
``WALKTHROUGH_CAPTURE`` environment variable is set to a truthy value.

When the variable is unset (the production default), ``capture_stage``
is a no-op and adds no measurable cost beyond an env-var lookup.

Intended for debugging and pipeline-inspection sessions. Not a
production logging facility.
"""
from __future__ import annotations

import os
from pathlib import Path
from typing import Any

WALKTHROUGH_ENV = "WALKTHROUGH_CAPTURE"
WALKTHROUGH_DIR = "_walkthrough"

README_TEXT = """# Pipeline Walkthrough Snapshots

Plain-text snapshots of the **active Start-Transcription path** -- the
pipeline that produces the deposition DOCX. Captured at stage
boundaries when ``WALKTHROUGH_CAPTURE=1`` is set in the environment.

## Files

| File | Stage | Source |
|---|---|---|
| `01_deepgram_raw.txt` | Raw transcript text after Deepgram + chunk assembly | `core/job_runner.py` (after `raw_deepgram.txt` is written) |
| `02_after_ai_cleanup.txt` | Text returned by the Anthropic cleanup pass | `clean_format/formatter.format_transcript` return point |
| `03_docx_text.txt` | Paragraph text extracted from the produced DOCX | After `clean_format/docx_writer.write_deposition_docx` returns |

The file numbers are sequential -- they capture every instrumented
stage with no gaps. The active path has exactly these three stage
boundaries between Deepgram input and the final DOCX.

## What this does NOT capture (important)

This walkthrough is for the **Start Transcription** path only. The
repo also contains an OFFLINE correction subsystem at `spec_engine/`
which runs only when the user presses the separate **Run Corrections**
button. The spec_engine modules -- ``block_builder``, ``classifier``,
``corrections``, ``qa_fixer``, ``emitter``, ``speaker_mapper`` -- are
**not** part of the active path that produces the deposition DOCX.
They operate on saved utterance JSON and produce a ``*_corrected.txt``
sidecar file, which is a different artifact for a different purpose.

This is a key thing to know during walkthrough inspection: if you are
looking for "where did the Q/A classification come from?" or "where
did this speaker label get normalized?", those decisions were made by
the Anthropic cleanup model inside stage 02, not by the spec_engine
modules. See ``docs/audits/ACTIVE_PATH_AUDIT.md`` for the full
active-path wiring map and the offline-path documentation.

## Usage

```powershell
$env:WALKTHROUGH_CAPTURE = "1"
.\\.venv\\Scripts\\python.exe app.py
# Run a transcription through the UI; this directory will populate
# with three snapshot files plus this README.
```

Snapshots are overwritten on each run -- they represent the most
recent Start-Transcription pass for the case.
"""


def _enabled() -> bool:
    """Return True if WALKTHROUGH_CAPTURE is set to a non-empty value."""
    return bool(os.environ.get(WALKTHROUGH_ENV, "").strip())


def _coerce_text(content: Any) -> str:
    """Render content as plain text for snapshot output."""
    if isinstance(content, Path):
        path = content
        if path.suffix.lower() == ".docx":
            return _docx_paragraph_text(path)
        return path.read_text(encoding="utf-8")
    if isinstance(content, str):
        # Could be a path to a .docx file, or already plain text.
        if content.lower().endswith(".docx") and Path(content).is_file():
            return _docx_paragraph_text(Path(content))
        return content
    if hasattr(content, "paragraphs"):
        # python-docx Document instance.
        return "\n".join(p.text for p in content.paragraphs)
    return str(content)


def _docx_paragraph_text(path: Path) -> str:
    """Extract paragraph text from a DOCX file via python-docx."""
    from docx import Document  # local import — only loaded if a DOCX is captured

    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs)


def capture_stage(case_dir: Any, stage_name: str, content: Any) -> None:
    """Write a plain-text snapshot to ``<case_dir>/_walkthrough/<stage_name>.txt``.

    No-op when ``WALKTHROUGH_CAPTURE`` is unset.

    Failures (missing case_dir, write errors, malformed DOCX) are
    swallowed silently — this is a debugging tool and must never
    affect the production pipeline.

    Parameters
    ----------
    case_dir : str | Path
        Path to the case folder. The ``_walkthrough/`` subdirectory is
        created if missing.
    stage_name : str
        Snapshot filename without extension, e.g. ``"01_deepgram_raw"``.
    content : Any
        Content to snapshot. Strings are written verbatim. ``Path``
        objects pointing at ``.docx`` files (and string paths ending in
        ``.docx``) are parsed via python-docx. ``python-docx``
        ``Document`` instances are walked paragraph-by-paragraph.
        Everything else is coerced via ``str()``.
    """
    if not _enabled():
        return

    try:
        target_dir = Path(case_dir) / WALKTHROUGH_DIR
        target_dir.mkdir(parents=True, exist_ok=True)

        readme_path = target_dir / "README.md"
        if not readme_path.exists():
            readme_path.write_text(README_TEXT, encoding="utf-8")

        text = _coerce_text(content)
        (target_dir / f"{stage_name}.txt").write_text(text, encoding="utf-8")
    except Exception:
        # Never let a debugging-tool failure break the production path.
        pass
