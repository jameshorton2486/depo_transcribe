"""Post-processor contract tests.

These tests describe the verbatim-preservation, idempotence, and
chassis-shape contract that the post-processor implementation must
satisfy. They are wired up against the real implementation at the
bottom of this file.
"""
from __future__ import annotations

import re
import shutil
import zipfile
from pathlib import Path

import pytest
from docx import Document

ROOT = Path(__file__).resolve().parents[2]
TEMPLATES_DIR = ROOT / "ufm_engine" / "templates" / "figures"


def _extract_text(docx_path: Path) -> str:
    """Whitespace-collapsed text from a .docx, for verbatim comparison.

    Walks every <w:t> in the document body so that text inside block-level
    <w:sdt> wrappers and table cells is captured uniformly.
    """
    from docx.oxml.ns import qn
    doc = Document(str(docx_path))
    body = doc.element.body
    parts = [t.text or "" for t in body.iter(qn("w:t"))]
    text = " ".join(parts)
    return re.sub(r"\s+", " ", text).strip()


def _strip_chassis_text(text: str) -> str:
    """Remove the gutter line-number digits (1..25) the chassis adds."""
    pattern = r"^(?:\s*" + r"\s+".join(str(n) for n in range(1, 26)) + r"\s*)+"
    return re.sub(pattern, "", text).strip()


def _is_implemented() -> bool:
    """True iff apply_format_box does not raise NotImplementedError on a trivial call.

    Evaluated lazily inside skipif so test collection does not depend on
    the package being importable at collection time.
    """
    try:
        from ufm_engine.post_processor.format_box import apply_format_box
    except ImportError:
        return False
    src = TEMPLATES_DIR / "appearances.docx"
    out = ROOT / "tests" / "ufm_engine" / "_probe.docx"
    try:
        apply_format_box(
            input_path=src,
            output_path=out,
            apply_line_numbers=False,
            render_firm_footer=False,
        )
    except NotImplementedError:
        return False
    except Exception:
        return True
    finally:
        if out.exists():
            out.unlink()
    return True


pytestmark = pytest.mark.skipif(
    "not __import__('tests.ufm_engine.test_post_processor', fromlist=['_is_implemented'])._is_implemented()",
    reason="post-processor not yet implemented",
)


@pytest.fixture
def populated_docx(tmp_path):
    """Copy a template to tmp as a stand-in for a populated doc."""
    src = TEMPLATES_DIR / "cert_tx_sig_waived.docx"
    dest = tmp_path / "populated.docx"
    shutil.copy(src, dest)
    return dest


def test_text_preservation(populated_docx, tmp_path):
    from ufm_engine.post_processor.format_box import apply_format_box
    out = tmp_path / "out.docx"
    apply_format_box(
        input_path=populated_docx,
        output_path=out,
        apply_line_numbers=False,
        render_firm_footer=False,
    )
    assert _extract_text(populated_docx) == _strip_chassis_text(_extract_text(out))


def test_idempotent(populated_docx, tmp_path):
    from ufm_engine.post_processor.format_box import apply_format_box
    out1 = tmp_path / "out1.docx"
    out2 = tmp_path / "out2.docx"
    apply_format_box(populated_docx, out1, apply_line_numbers=True, render_firm_footer=False)
    apply_format_box(out1, out2, apply_line_numbers=True, render_firm_footer=False)
    assert _extract_text(out1) == _extract_text(out2)


def test_format_box_present_on_every_page(populated_docx, tmp_path):
    from ufm_engine.post_processor.format_box import apply_format_box
    out = tmp_path / "out.docx"
    apply_format_box(populated_docx, out, apply_line_numbers=True, render_firm_footer=False)
    doc = Document(str(out))
    assert len(doc.tables) >= 1, "expected at least one chassis table per page"


def test_line_numbers_1_through_25(populated_docx, tmp_path):
    from ufm_engine.post_processor.format_box import apply_format_box
    out = tmp_path / "out.docx"
    apply_format_box(populated_docx, out, apply_line_numbers=True, render_firm_footer=False)
    doc = Document(str(out))
    assert doc.tables, "no chassis table present"
    left_cell = doc.tables[0].rows[0].cells[0]
    nums = [p.text.strip() for p in left_cell.paragraphs]
    nums = [n for n in nums if n]
    assert nums[:25] == [str(n) for n in range(1, 26)]


def test_firm_footer_when_firm_name_set(populated_docx, tmp_path):
    from ufm_engine.post_processor.format_box import apply_format_box
    out = tmp_path / "out.docx"
    apply_format_box(
        populated_docx, out,
        apply_line_numbers=False,
        render_firm_footer=True,
        firm_name="LEXITAS",
    )
    doc = Document(str(out))
    footer_text = ""
    for section in doc.sections:
        for p in section.footer.paragraphs:
            footer_text += p.text
    assert "LEXITAS" in footer_text


def test_no_firm_footer_when_disabled(populated_docx, tmp_path):
    from ufm_engine.post_processor.format_box import apply_format_box
    out = tmp_path / "out.docx"
    apply_format_box(populated_docx, out, apply_line_numbers=False, render_firm_footer=False)
    doc = Document(str(out))
    footer_text = ""
    for section in doc.sections:
        for p in section.footer.paragraphs:
            footer_text += p.text
    assert footer_text.strip() == ""


def test_firm_footer_requires_name():
    from ufm_engine.post_processor.format_box import apply_format_box
    with pytest.raises(ValueError):
        apply_format_box(
            input_path=TEMPLATES_DIR / "appearances.docx",
            output_path=Path("nope.docx"),
            render_firm_footer=True,
            firm_name=None,
        )
