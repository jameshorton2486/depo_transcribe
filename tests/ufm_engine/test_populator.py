"""Populator tests: content controls, block toggles, end-to-end round trip."""
from __future__ import annotations

import re
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[2]
TEMPLATES_DIR = ROOT / "ufm_engine" / "templates" / "figures"


def _all_text(docx_path: Path) -> str:
    """Extract text by paragraph; runs within a paragraph are concatenated
    without a separator so adjacent runs ("$" + "275.00") stay adjacent."""
    doc = Document(str(docx_path))
    body = doc.element.body
    para_texts = []
    for p in body.iter(qn("w:p")):
        runs = [t.text or "" for t in p.iter(qn("w:t"))]
        para_texts.append("".join(runs))
    return re.sub(r"\s+", " ", "\n".join(para_texts)).strip()


def _block_sdt_tags(docx_path: Path) -> list[str]:
    doc = Document(str(docx_path))
    body = doc.element.body
    out = []
    for sdt in [c for c in body if c.tag == qn("w:sdt")]:
        tag_el = sdt.find(qn("w:sdtPr") + "/" + qn("w:tag"))
        if tag_el is not None:
            v = tag_el.get(qn("w:val"))
            if v:
                out.append(v)
    return out


def _content_control_tags(docx_path: Path) -> list[str]:
    doc = Document(str(docx_path))
    body = doc.element.body
    return [
        sdt.find(qn("w:sdtPr") + "/" + qn("w:tag")).get(qn("w:val"))
        for sdt in body.iter(qn("w:sdt"))
        if sdt.find(qn("w:sdtPr") + "/" + qn("w:tag")) is not None
    ]


def test_content_controls_get_filled(tmp_path):
    from ufm_engine.populator.populate import populate
    out = tmp_path / "out.docx"
    populate(
        TEMPLATES_DIR / "witness_setup_standard.docx",
        out,
        fields={
            "witness_name": "JOHN DOE",
            "examining_lawyer_label": "MR. SMITH",
        },
    )
    text = _all_text(out)
    assert "JOHN DOE" in text
    assert "MR. SMITH" in text
    assert "[Witness Name (caps)]" not in text
    assert "[Examining Lawyer]" not in text


def test_block_toggle_false_removes_entire_block(tmp_path):
    from ufm_engine.populator.populate import populate
    out = tmp_path / "out.docx"
    populate(
        TEMPLATES_DIR / "cert_tx_sig_waived.docx",
        out,
        fields={},
        block_toggles={
            "block_custodial_attorney": False,
            "block_cost_paragraph": False,
            "block_credentials_suffix": False,
            "block_firm_signature_block": False,
        },
    )
    text = _all_text(out)
    assert "delivered to" not in text
    assert "deposition officer's charges" not in text
    assert "Firm Registration No." not in text
    # No remaining block sdts at body level
    assert _block_sdt_tags(out) == []


def test_block_toggle_true_keeps_and_unwraps(tmp_path):
    from ufm_engine.populator.populate import populate
    out = tmp_path / "out.docx"
    populate(
        TEMPLATES_DIR / "cert_tx_sig_waived.docx",
        out,
        fields={
            "custodial_attorney_name": "ATTORNEY OF RECORD",
            "cost_amount": "275.00",
            "cost_payor_party": "PLAINTIFF",
            "credentials": "RPR",
            "firm_name": "ACME LEGAL",
            "firm_registration_number": "12345",
        },
        block_toggles={
            "block_custodial_attorney": True,
            "block_cost_paragraph": True,
            "block_credentials_suffix": True,
            "block_firm_signature_block": True,
        },
    )
    text = _all_text(out)
    assert "delivered to ATTORNEY OF RECORD" in text
    assert "$275.00" in text
    assert "ACME LEGAL" in text
    assert "Firm Registration No. 12345" in text
    # Block wrappers are gone (unwrapped), content remains
    assert _block_sdt_tags(out) == []


def test_default_on_when_toggles_omitted(tmp_path):
    """Per recipe §6: missing toggle defaults to True (block kept)."""
    from ufm_engine.populator.populate import populate
    out = tmp_path / "out.docx"
    populate(
        TEMPLATES_DIR / "cert_tx_sig_waived.docx",
        out,
        fields={"custodial_attorney_name": "ATTY"},
        block_toggles=None,
    )
    text = _all_text(out)
    assert "delivered to ATTY" in text


def test_empty_dict_toggles_still_default_on(tmp_path):
    """Empty dict means 'no explicit toggles set' → defaults still apply."""
    from ufm_engine.populator.populate import populate
    out = tmp_path / "out.docx"
    populate(
        TEMPLATES_DIR / "cert_tx_sig_waived.docx",
        out,
        fields={"custodial_attorney_name": "ATTY"},
        block_toggles={},
    )
    text = _all_text(out)
    assert "delivered to ATTY" in text


def test_missing_field_leaves_placeholder(tmp_path, caplog):
    from ufm_engine.populator.populate import populate
    import logging
    caplog.set_level(logging.WARNING)
    out = tmp_path / "out.docx"
    populate(
        TEMPLATES_DIR / "witness_setup_standard.docx",
        out,
        fields={"witness_name": "JANE DOE"},  # examining_lawyer_label missing
    )
    text = _all_text(out)
    assert "JANE DOE" in text
    assert "[Examining Lawyer]" in text
    assert any("examining_lawyer_label" in rec.getMessage() for rec in caplog.records)


def test_field_with_null_value_leaves_placeholder(tmp_path):
    from ufm_engine.populator.populate import populate
    out = tmp_path / "out.docx"
    populate(
        TEMPLATES_DIR / "witness_setup_standard.docx",
        out,
        fields={"witness_name": "JANE DOE", "examining_lawyer_label": None},
    )
    text = _all_text(out)
    assert "[Examining Lawyer]" in text


def test_populator_then_post_processor(tmp_path):
    """End-to-end: populate, then run the post-processor; round-trips cleanly."""
    from ufm_engine.populator.populate import populate
    from ufm_engine.post_processor.format_box import apply_format_box
    populated = tmp_path / "populated.docx"
    finished = tmp_path / "finished.docx"
    populate(
        TEMPLATES_DIR / "witness_setup_standard.docx",
        populated,
        fields={"witness_name": "JANE DOE", "examining_lawyer_label": "MR. ROE"},
    )
    apply_format_box(
        input_path=populated,
        output_path=finished,
        apply_line_numbers=True,
        render_firm_footer=True,
        firm_name="ACME LEGAL",
    )
    assert finished.exists()
    text = _all_text(finished)
    assert "JANE DOE" in text
    assert "MR. ROE" in text


def test_input_template_unchanged(tmp_path):
    """Populating must not mutate the source template."""
    from ufm_engine.populator.populate import populate
    src = TEMPLATES_DIR / "witness_setup_standard.docx"
    before = src.read_bytes()
    populate(src, tmp_path / "out.docx", fields={"witness_name": "JANE"})
    after = src.read_bytes()
    assert before == after


def test_inline_block_videotaped_kept(tmp_path):
    """Title-page videotaped phrase is now inline-conditional. toggle=True keeps it."""
    from ufm_engine.populator.populate import populate
    out = tmp_path / "out.docx"
    populate(
        TEMPLATES_DIR / "title_page_tx_state.docx",
        out,
        fields={},
        block_toggles={
            "block_videotaped": True,
            "block_subpoena_duces_tecum": False,
            "block_volume": False,
        },
    )
    text = _all_text(out)
    assert "AND VIDEOTAPED" in text
    assert "WITH SUBPOENA DUCES TECUM" not in text


def test_inline_block_videotaped_dropped(tmp_path):
    """toggle=False for an inline block removes the phrase entirely."""
    from ufm_engine.populator.populate import populate
    out = tmp_path / "out.docx"
    populate(
        TEMPLATES_DIR / "title_page_tx_state.docx",
        out,
        fields={},
        block_toggles={
            "block_videotaped": False,
            "block_subpoena_duces_tecum": False,
            "block_volume": False,
        },
    )
    text = _all_text(out)
    assert "AND VIDEOTAPED" not in text
    assert "WITH SUBPOENA DUCES TECUM" not in text
    assert "ORAL DEPOSITION" in text


def test_volume_line_uses_total_volumes_field(tmp_path):
    """The volume line accepts both volume_number and total_volumes."""
    from ufm_engine.populator.populate import populate
    out = tmp_path / "out.docx"
    populate(
        TEMPLATES_DIR / "title_page_tx_state.docx",
        out,
        fields={"volume_number": "1", "total_volumes": "3"},
        block_toggles={"block_volume": True},
    )
    text = _all_text(out)
    assert "VOLUME 1 OF 3" in text


def test_tx_state_body_block_remote_kept(tmp_path):
    """", via [Remote Platform]" in the body is now conditional on block_remote."""
    from ufm_engine.populator.populate import populate
    out = tmp_path / "out.docx"
    populate(
        TEMPLATES_DIR / "title_page_tx_state.docx",
        out,
        fields={"remote_platform": "Zoom"},
        block_toggles={"block_remote": True, "block_videotaped": False,
                        "block_subpoena_duces_tecum": False, "block_volume": False},
    )
    text = _all_text(out)
    assert ", via Zoom" in text


def test_tx_state_body_block_remote_dropped(tmp_path):
    """toggle=False removes the entire ", via …" segment so no awkward ", via , "."""
    from ufm_engine.populator.populate import populate
    out = tmp_path / "out.docx"
    populate(
        TEMPLATES_DIR / "title_page_tx_state.docx",
        out,
        fields={},
        block_toggles={"block_remote": False, "block_videotaped": False,
                        "block_subpoena_duces_tecum": False, "block_volume": False},
    )
    text = _all_text(out)
    assert ", via" not in text
    assert "[Remote Platform]" not in text


def test_federal_title_page_videotaped_kept(tmp_path):
    """Federal title page block_videotaped True keeps the phrase in both
    the title block and the body paragraph."""
    from ufm_engine.populator.populate import populate
    out = tmp_path / "out.docx"
    populate(
        TEMPLATES_DIR / "title_page_federal.docx",
        out,
        fields={},
        block_toggles={"block_videotaped": True, "block_interpreted": False,
                        "block_remote": False},
    )
    text = _all_text(out)
    # Two occurrences expected: title block + body opening
    assert text.count("AND VIDEOTAPED") == 2


def test_federal_title_page_videotaped_dropped(tmp_path):
    """toggle=False removes both the title and body uses."""
    from ufm_engine.populator.populate import populate
    out = tmp_path / "out.docx"
    populate(
        TEMPLATES_DIR / "title_page_federal.docx",
        out,
        fields={},
        block_toggles={"block_videotaped": False, "block_interpreted": False,
                        "block_remote": False},
    )
    text = _all_text(out)
    assert "AND VIDEOTAPED" not in text
    assert "ORAL DEPOSITION" in text


def test_federal_title_page_interpreted_block(tmp_path):
    """The interpreted notation is a full-paragraph block; toggle on/off
    keeps or removes the whole line."""
    from ufm_engine.populator.populate import populate

    on_path = tmp_path / "on.docx"
    populate(
        TEMPLATES_DIR / "title_page_federal.docx",
        on_path,
        fields={"interpreter_language": "Spanish"},
        block_toggles={"block_videotaped": False, "block_interpreted": True,
                        "block_remote": False},
    )
    on_text = _all_text(on_path)
    assert "INTERPRETED FROM Spanish TO ENGLISH" in on_text

    off_path = tmp_path / "off.docx"
    populate(
        TEMPLATES_DIR / "title_page_federal.docx",
        off_path,
        fields={},
        block_toggles={"block_videotaped": False, "block_interpreted": False,
                        "block_remote": False},
    )
    off_text = _all_text(off_path)
    assert "INTERPRETED FROM" not in off_text
