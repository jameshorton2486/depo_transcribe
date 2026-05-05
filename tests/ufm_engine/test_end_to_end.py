"""End-to-end test: synthesize a realistic case folder and drive the same
populate → apply_format_box pipeline that the Templates UI tab uses.

This is the programmatic equivalent of clicking through the Templates tab.
Mirrors ui/tab_templates.py so behavior changes there are caught here.
"""
from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[2]
TEMPLATES_DIR = ROOT / "ufm_engine" / "templates" / "figures"
PROFILES_DIR = ROOT / "data" / "reporter_profiles"


# Realistic-shape ufm_fields, modeled on what the NOD parser writes for a
# Texas state-court depo. Field names match the manifest's expected tags.
SAMPLE_UFM_FIELDS = {
    "cause_number": "DC-25-13430",
    "plaintiff_name": "JOHN SMITH",
    "plaintiff_party_label": "Plaintiff",
    "court_designation": "DISTRICT COURT, 134th JUDICIAL DISTRICT",
    "county": "DALLAS",
    "state": "TEXAS",
    "judicial_district_phrase": "134th JUDICIAL DISTRICT",
    "defendant_names_block": "ACME WIDGET CORP., et al.",
    "defendant_party_label": "Defendants",
    "witness_name": "JACK LEIFER",
    "depo_date": "April 16, 2026",
    "depo_day_ordinal": "16th day of April, 2026",
    "depo_time_start": "10:00 a.m.",
    "depo_time_end": "1:30 p.m.",
    "depo_location_full": "1234 Main Street, Suite 500, Dallas, Texas 75201",
    "instance_party": "Plaintiff",
}


def _all_text(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    body = doc.element.body
    parts = []
    for p in body.iter(qn("w:p")):
        runs = [t.text or "" for t in p.iter(qn("w:t"))]
        parts.append("".join(runs))
    return re.sub(r"\s+", " ", "\n".join(parts)).strip()


def _make_case_folder(tmp_path: Path) -> Path:
    """Build a case folder with the same shape job_runner creates on disk."""
    case_root = tmp_path / "2026" / "Apr" / "DC2513430" / "leifer_jack"
    (case_root / "source_docs").mkdir(parents=True)
    (case_root / "Deepgram").mkdir()
    job_config = {
        "version": 1,
        "ufm_fields": SAMPLE_UFM_FIELDS,
        "confirmed_spellings": {},
    }
    (case_root / "source_docs" / "job_config.json").write_text(
        json.dumps(job_config, indent=2), encoding="utf-8"
    )
    return case_root


def _resolve_fields(profile: dict, ufm_fields: dict, manual: dict) -> dict:
    """Same merge as ui/tab_templates.py::_resolved_fields."""
    merged = {}
    merged.update(ufm_fields)
    for k, v in profile.items():
        if k in {"id", "display_name", "chassis_default"}:
            continue
        if v is not None:
            merged[k] = v
    merged.update({k: v for k, v in manual.items() if v})
    return merged


def test_full_pipeline_against_synthetic_case(tmp_path):
    """Drive populate → apply_format_box exactly as the Templates tab does.

    Verifies the Jack-Leifer-style scenario the user reported: a TX state
    deposition with Miah Bardot reporting, signature waived. Selects the
    standard set, runs the full pipeline, inspects the resulting .docx.
    """
    from ufm_engine.populator.populate import populate
    from ufm_engine.post_processor.format_box import apply_format_box

    case_root = _make_case_folder(tmp_path)

    profile = json.loads(
        (PROFILES_DIR / "miah_bardot_sa_legal.json").read_text(encoding="utf-8")
    )
    ufm_fields = json.loads(
        (case_root / "source_docs" / "job_config.json").read_text(encoding="utf-8")
    )["ufm_fields"]

    manual = {
        "custodial_attorney_name": "ATTORNEY OF RECORD, ESQ.",
        "cost_amount": "275.00",
        "cost_payor_party": "PLAINTIFF",
        "served_on_date": "April 30, 2026",
        "certification_date": "30th day of April, 2026",
    }
    fields = _resolve_fields(profile, ufm_fields, manual)

    selected = [
        "title_page_tx_state",
        "appearances",
        "witness_setup_standard",
        "cert_tx_sig_waived",
    ]
    block_toggles = {
        "block_videotaped": True,            # videotaped depo
        "block_subpoena_duces_tecum": False,  # no subpoena
        "block_volume": False,                # single volume
        "block_remote": False,
        "block_interpreted": False,
        "block_also_present": False,
        "block_custodial_attorney": True,
        "block_cost_paragraph": True,
        "block_credentials_suffix": False,    # Miah has no credentials suffix
        "block_firm_signature_block": True,
    }

    draft_dir = case_root / "output" / "draft"
    draft_dir.mkdir(parents=True)
    final_dir = case_root / "output" / "final"
    final_dir.mkdir(parents=True)

    for tid in selected:
        populate(
            TEMPLATES_DIR / f"{tid}.docx",
            draft_dir / f"{tid}.docx",
            fields=fields,
            block_toggles=block_toggles,
        )

    for src in sorted(draft_dir.glob("*.docx")):
        apply_format_box(
            input_path=src,
            output_path=final_dir / src.name,
            apply_line_numbers=True,
            render_firm_footer=True,
            firm_name=profile["firm_name"],
        )

    # Every selected template made it through both stages
    assert {p.name for p in draft_dir.glob("*.docx")} == {f"{t}.docx" for t in selected}
    assert {p.name for p in final_dir.glob("*.docx")} == {f"{t}.docx" for t in selected}

    # Title page: NOD-derived caption + reporter, conditional inline blocks
    title_text = _all_text(final_dir / "title_page_tx_state.docx")
    assert "DC-25-13430" in title_text
    assert "JOHN SMITH" in title_text
    assert "JACK LEIFER" in title_text
    assert "DALLAS" in title_text
    assert "Miah Bardot" in title_text
    assert "AND VIDEOTAPED" in title_text          # block_videotaped True
    assert "WITH SUBPOENA DUCES TECUM" not in title_text  # toggle False
    assert "VOLUME" not in title_text              # block_volume False
    assert ", via " not in title_text              # block_remote False — via clause must be gone
    assert "[Remote Platform" not in title_text    # block_remote False — placeholder must be gone too
    assert "[Witness Name]" not in title_text      # placeholders gone

    # Cert: signature-waived flow with cost + custodial attorney
    cert_text = _all_text(final_dir / "cert_tx_sig_waived.docx")
    assert "JACK LEIFER" in cert_text
    assert "Miah Bardot" in cert_text
    assert "delivered to ATTORNEY OF RECORD, ESQ." in cert_text
    assert "$275.00" in cert_text
    assert "PLAINTIFF" in cert_text
    assert "Texas CSR 12129" in cert_text
    assert "SA Legal" in cert_text                # firm name from profile
    # block_credentials_suffix is False, so the "[Reporter Name], [Credentials]"
    # paragraph is dropped entirely — but reporter_name still appears in
    # the signature block via the Texas-CSR-N-line ("Miah Bardot, Texas
    # CSR 12129"... well, our template puts reporter_name above the CSR
    # line; with that block off it's removed). Confirm the Credentials
    # placeholder is NOT visible since the block was dropped.
    assert "[Credentials]" not in cert_text

    # Format box: every page has a chassis table and the firm-name footer
    for src in final_dir.glob("*.docx"):
        doc = Document(str(src))
        assert doc.tables, f"{src.name}: missing chassis table"
        # Line numbers in the gutter: 1..25 in the left cell of the first table
        left_cell = doc.tables[0].rows[0].cells[0]
        nums = [p.text.strip() for p in left_cell.paragraphs if p.text.strip()]
        assert nums[:25] == [str(n) for n in range(1, 26)], src.name
        # Firm footer
        footer_text = ""
        for section in doc.sections:
            for p in section.footer.paragraphs:
                footer_text += p.text
        assert "SA Legal" in footer_text, f"{src.name}: missing firm footer"
