"""
ASCII export utility for Depo-Pro.
"""

from __future__ import annotations

from pathlib import Path

_LINE_NUM_PREFIXES = tuple(f"{n:2d} " for n in range(1, 26))


def _strip_line_number(line: str) -> str:
    """Remove a leading transcript line number prefix if present."""
    for prefix in _LINE_NUM_PREFIXES:
        if line.startswith(prefix):
            return line[len(prefix):]
    return line


def strip_to_ascii(text: str) -> str:
    """Strip line numbers and normalize transcript text for ASCII export."""
    lines = text.splitlines()
    cleaned = []
    for line in lines:
        line = _strip_line_number(line)
        line = line.replace("\t", "    ")
        line = line.rstrip()
        cleaned.append(line)

    result = []
    blank_run = 0
    for line in cleaned:
        if line == "":
            blank_run += 1
            if blank_run <= 2:
                result.append("")
        else:
            blank_run = 0
            result.append(line)

    return "\n".join(result) + "\n"


def extract_text_from_docx(docx_path: str) -> str:
    """Extract plain text from body paragraphs and tables in a DOCX."""
    from docx import Document

    doc = Document(docx_path)
    lines = []

    for para in doc.paragraphs:
        lines.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) == 2:
                lines.append(cells[1].text)
            else:
                for cell in cells:
                    lines.append(cell.text)

    return "\n".join(lines)


def export_ascii(source: str, output_path: str, is_docx: bool = False) -> str:
    """Write a cleaned UTF-8 transcript text file."""
    raw_text = extract_text_from_docx(source) if is_docx else source
    clean = strip_to_ascii(raw_text)

    dest = Path(output_path)
    if dest.suffix.lower() != ".txt":
        dest = dest.with_suffix(".txt")

    dest.parent.mkdir(parents=True, exist_ok=True)
    dest.write_text(clean, encoding="utf-8")
    return str(dest)
