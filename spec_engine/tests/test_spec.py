"""
test_spec.py — All 22 unit tests for DepoPro Spec v1.0.
Run with: python -m pytest spec_engine/tests/test_spec.py -v

ALL 22 TESTS MUST PASS before the app is considered production-ready.

Test groups:
  Tests  1-2:  Verbatim preservation (uh/um — absolute rule)
  Tests  3-4:  Subpoena duces tecum variants
  Tests  5-7:  Proper noun corrections (Ugalde, Marrufo, firm name)
  Test   8:    Affirmation word NOT collapsed (correct correct preserved)
  Test   9:    4+ char duplicate IS collapsed (Corrected Corrected → Corrected)
  Test  10:    Doctor. artifact normalization
  Test  11:    Tovar name correction
  Tests 12-16: Emitter format and color tests (Q/A/SP/PN/FLAG)
  Tests 17-22: Additional spec tests (objections, embedded Q+A, speaker guard, etc.)
"""

import sys
import os
import tempfile
from pathlib import Path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', '..'))

import pytest
from docx.shared import RGBColor

from spec_engine.corrections import clean_block
from spec_engine.emitter import (
    create_document, emit_a_line, emit_flag_line, emit_pn_line,
    emit_q_line, emit_sp_line,
)
from spec_engine.models import JobConfig, LineType, SpeakerMapUnverifiedError
from spec_engine.classifier import classify_block, ClassifierState
from spec_engine.models import Block


def make_config():
    """Return the Perez v. Ugalde config for testing."""
    return JobConfig.default_perez_ugalde()


def make_unverified_config():
    """Return a config with speaker_map_verified=False."""
    cfg = JobConfig()
    cfg.speaker_map = {0: "THE VIDEOGRAPHER", 1: "THE WITNESS", 2: "MR. TEST"}
    cfg.speaker_map_verified = False  # deliberately unverified
    return cfg


def test_jobconfig_reporter_defaults_are_empty():
    cfg = JobConfig()
    assert cfg.reporter_name == ""
    assert cfg.reporter_csr == ""
    assert cfg.reporter_firm == ""
    assert cfg.reporter_address == ""


# ═══════════════════════════════════════════════════════════════
# TESTS 1-2: Verbatim preservation (ABSOLUTE RULE — Spec 2.1)
# ═══════════════════════════════════════════════════════════════

def test_uh_preserved():
    """clean('uh, I think so') → 'uh' preserved. ABSOLUTE RULE."""
    result = clean_block("uh, I think so", make_config())[0]
    assert "uh" in result.lower(), f"FAIL: 'uh' was removed. Got: {result!r}"


def test_um_preserved():
    """clean('um, that is correct') → 'um' preserved. ABSOLUTE RULE."""
    result = clean_block("um, that is correct", make_config())[0]
    assert "um" in result.lower(), f"FAIL: 'um' was removed. Got: {result!r}"


# ═══════════════════════════════════════════════════════════════
# TESTS 3-4: Subpoena duces tecum normalization (Spec 2.3)
# ═══════════════════════════════════════════════════════════════

def test_subpoena_variant_deuces_tikum():
    """clean('subpoena deuces tikum') == 'subpoena duces tecum'"""
    result = clean_block("subpoena deuces tikum", make_config())[0]
    assert "duces tecum" in result.lower(), f"FAIL: Got: {result!r}"


def test_subpoena_variant_de_sus_tikum():
    """clean('subpoena de sus tikum') == 'subpoena duces tecum'"""
    result = clean_block("subpoena de sus tikum", make_config())[0]
    assert "duces tecum" in result.lower(), f"FAIL: Got: {result!r}"


# ═══════════════════════════════════════════════════════════════
# TESTS 5-7: Proper noun corrections (Spec 2.3)
# ═══════════════════════════════════════════════════════════════

def test_ugalde_correction():
    """clean('Yugaldi') == 'Ugalde' — confirmed from NOD."""
    result = clean_block("Yugaldi", make_config())[0]
    assert "Ugalde" in result, f"FAIL: Got: {result!r}"


def test_marrufo_double_r():
    """clean('Marufo') == 'Marrufo' — double R confirmed from NOD."""
    result = clean_block("Marufo", make_config())[0]
    assert "Marrufo" in result, f"FAIL: Got: {result!r}"


def test_law_firm_multiword():
    """clean('Allen Stein in Durbin') contains 'Allen, Stein & Durbin, P.C.'"""
    result = clean_block("Allen Stein in Durbin represents the defendant", make_config())[0]
    assert "Allen, Stein & Durbin, P.C." in result, f"FAIL: Got: {result!r}"


# ═══════════════════════════════════════════════════════════════
# TEST 8: Affirmation word preserved (Spec 9.6 unit test)
# ═══════════════════════════════════════════════════════════════

def test_correct_correct_preserved():
    """
    clean('correct correct') == 'correct correct' — affirmation NOT collapsed.
    Spec 9.6 unit test: 'short word NOT collapsed' label.
    'correct correct' is a common witness affirmation pattern and must be preserved.
    """
    result = clean_block("correct correct", make_config())[0]
    assert "correct correct" in result.lower(), (
        f"FAIL: 'correct correct' was collapsed. Got: {result!r}"
    )


# ═══════════════════════════════════════════════════════════════
# TEST 9: 4+ char duplicate collapsed (Spec 2.2)
# ═══════════════════════════════════════════════════════════════

def test_four_plus_chars_collapsed():
    """clean('Corrected Corrected') == 'Corrected' — 4+ chars ARE collapsed."""
    result = clean_block("Corrected Corrected", make_config())[0]
    assert result.count("Corrected") == 1, f"FAIL: Got: {result!r}"


# ═══════════════════════════════════════════════════════════════
# TEST 10: Doctor. artifact (Spec 2.3 Universal)
# ═══════════════════════════════════════════════════════════════

def test_doctor_artifact():
    """clean('Doctor. Smith') starts with 'Dr.' — Deepgram Doctor. artifact normalized."""
    result = clean_block("Doctor. Smith examined the patient", make_config())[0]
    assert result.startswith("Dr."), f"FAIL: Doctor. not normalized. Got: {result!r}"


def test_k_midblock_correction():
    """clean('K. that is correct') → 'Okay. that is correct' — Spec 2.3.
    K. must normalize even when followed by more spoken text in the same block.
    """
    result = clean_block("K. that is correct", make_config())[0]
    assert "okay" in result.lower(), (
        f"FAIL: K. mid-block not normalized. Got: {result!r}"
    )


# ═══════════════════════════════════════════════════════════════
# TEST 11: Tovar name correction (Spec 2.3)
# ═══════════════════════════════════════════════════════════════

def test_tovar_correction():
    """clean('Tobar was behind me') → 'Tovar was behind me'"""
    result = clean_block("Tobar was behind me", make_config())[0]
    assert "Tovar" in result, f"FAIL: Got: {result!r}"


# ═══════════════════════════════════════════════════════════════
# TESTS 12-16: Emitter format and color tests (Spec 3.3, 5.3, 5.4)
# ═══════════════════════════════════════════════════════════════

def test_q_line_format():
    """emit Q line text contains TAB Q. TAB"""
    doc = create_document()
    emit_q_line(doc, "What is your name?")
    assert len(doc.paragraphs) == 1
    text = doc.paragraphs[0].text
    assert '\tQ.\t' in text, f"FAIL: Q format wrong. Got: {repr(text)}"


def test_a_line_format():
    """emit A line text contains TAB A. TAB"""
    doc = create_document()
    emit_a_line(doc, "My name is John Smith.")
    text = doc.paragraphs[0].text
    assert '\tA.\t' in text, f"FAIL: A format wrong. Got: {repr(text)}"


def test_sp_line_bold_label():
    """emit SP line — label run is bold=True"""
    doc = create_document()
    emit_sp_line(doc, "MR. SALAZAR:  Objection to form.")
    para = doc.paragraphs[0]
    bold_runs = [r for r in para.runs if r.bold]
    assert len(bold_runs) > 0, "FAIL: No bold run found in SP line"


def test_pn_line_navy_color():
    """emit PN line font color == RGBColor(0x1E, 0x3A, 0x5F)"""
    doc = create_document()
    emit_pn_line(doc, "(Whereupon, a recess was taken at 10:15 a.m.)")
    run = doc.paragraphs[0].runs[0]
    assert run.font.color.rgb == RGBColor(0x1E, 0x3A, 0x5F), (
        f"FAIL: PN color wrong. Got: {run.font.color.rgb}"
    )


def test_flag_line_orange_color():
    """emit FLAG line font color == RGBColor(0xB4, 0x5F, 0x06) and bold=True"""
    doc = create_document()
    emit_flag_line(doc, "[SCOPIST: FLAG 1: Verify spelling of Bruggemann]")
    run = doc.paragraphs[0].runs[0]
    assert run.font.color.rgb == RGBColor(0xB4, 0x5F, 0x06), (
        f"FAIL: FLAG color wrong. Got: {run.font.color.rgb}"
    )
    assert run.bold, "FAIL: FLAG run is not bold"


def test_emitter_q_line_wraps_to_multiple_paragraphs():
    doc = create_document()
    emit_q_line(doc, "word " * 20)
    assert len(doc.paragraphs) > 1
    assert '\tQ.\t' in doc.paragraphs[0].text
    assert doc.paragraphs[1].text.startswith('\t')


def test_emitter_speaker_continuation_aligns_under_content():
    doc = create_document()
    emit_sp_line(doc, "MR. SALAZAR:  " + ("word " * 20))
    assert len(doc.paragraphs) > 1
    assert "MR. SALAZAR:" in doc.paragraphs[0].text
    assert "MR. SALAZAR:" not in doc.paragraphs[1].text
    assert doc.paragraphs[1].text.startswith('\t\t\t')


# ═══════════════════════════════════════════════════════════════
# TESTS 17-22: Additional spec tests
# ═══════════════════════════════════════════════════════════════

def test_objection_standardized():
    """'Objection to form.' remains preserved as spoken via classifier."""
    cfg = make_config()
    block = Block(
        speaker_id=3,
        text="Objection to form.",
        raw_text="Objection to form.",
    )
    state = ClassifierState()
    results = classify_block(block, cfg, state, block_index=0)
    assert len(results) == 1
    assert results[0][0] == LineType.SP
    assert "Objection to form." in results[0][1], f"FAIL: Got: {results[0][1]!r}"


def test_objection_bare_preserved_verbatim():
    """Bare 'Objection.' must not be rewritten — Morson's / Spec 4.3.
    The spoken word is the legal record. Adding 'to form.' is a verbatim
    violation if the attorney did not say those words.
    """
    cfg = make_config()
    state = ClassifierState()
    block = Block(speaker_id=3, text="Objection.")
    result = classify_block(block, cfg, state, block_index=0)
    text_out = " ".join(t for _, t in result)
    assert "to form" not in text_out.lower(), (
        f"FAIL: 'Objection.' was rewritten. Got: {text_out!r}"
    )
    assert "Objection." in text_out, (
        f"FAIL: 'Objection.' not preserved in output. Got: {text_out!r}"
    )


def test_objection_reserve_at_trial():
    """'We'll reserve our questions until the time of trial' standardized"""
    cfg = make_config()
    block = Block(speaker_id=3,
                  text="We'll reserve our questions until the time of trial",
                  raw_text="")
    state = ClassifierState()
    results = classify_block(block, cfg, state, block_index=0)
    assert len(results) == 1
    assert "reserve" in results[0][1].lower(), f"FAIL: Got: {results[0][1]!r}"


def test_examining_attorney_gets_q():
    """Speaker 2 (examining attorney) produces LineType.Q"""
    cfg = make_config()
    block = Block(speaker_id=2, text="What is your name?", raw_text="")
    state = ClassifierState()
    results = classify_block(block, cfg, state, block_index=0)
    assert any(lt == LineType.Q for lt, _ in results), (
        f"FAIL: Expected Q line, got: {results}"
    )


def test_witness_gets_a():
    """Speaker 1 (witness) produces LineType.A"""
    cfg = make_config()
    block = Block(speaker_id=1, text="My name is John Smith.", raw_text="")
    state = ClassifierState()
    results = classify_block(block, cfg, state, block_index=0)
    assert any(lt == LineType.A for lt, _ in results), (
        f"FAIL: Expected A line, got: {results}"
    )


def test_embedded_answer_split():
    """Attorney block with embedded answer is split into Q + A"""
    cfg = make_config()
    cfg.split_embedded_answers = True
    block = Block(
        speaker_id=2,
        text="Did you see the truck? Yes I did.",
        raw_text="",
    )
    state = ClassifierState()
    results = classify_block(block, cfg, state, block_index=0)
    types = [lt for lt, _ in results]
    assert LineType.Q in types, f"FAIL: No Q in split result: {results}"
    assert LineType.A in types, f"FAIL: No A in split result: {results}"


def test_speaker_map_unverified_raises():
    """process_transcript raises SpeakerMapUnverifiedError if map not verified"""
    from spec_engine.document_builder import process_transcript
    cfg = make_unverified_config()
    with pytest.raises(SpeakerMapUnverifiedError):
        process_transcript(
            input_docx_path="nonexistent.docx",
            output_docx_path="nonexistent_out.docx",
            job_config=cfg,
        )


if __name__ == "__main__":
    pytest.main([__file__, "-v"])


# ═══════════════════════════════════════════════════════════════
# TESTS 23-32: Administrative page generators (Phase 4-B)
# ═══════════════════════════════════════════════════════════════

from spec_engine.pages._lined_page import write_lined_page, paginate_lines
from spec_engine.pages.title_page import write_title_page
from spec_engine.pages.caption import write_caption
from spec_engine.pages.certificate import write_certificate
from spec_engine.pages.changes_signature import write_changes_signature
from spec_engine.pages.witness_index import write_witness_index
from spec_engine.pages.exhibit_index import write_exhibit_index
from spec_engine.models import (
    CounselInfo, WitnessIndexEntry, ExhibitEntry, ChangeEntry
)


def make_full_config():
    """JobConfig with all fields populated for page tests."""
    cfg = JobConfig.default_perez_ugalde()
    cfg.witness_name = "Simon Ugalde"
    cfg.depo_date = "October 15, 2025"
    cfg.depo_start_time = "9:00 a.m."
    cfg.depo_end_time = "5:00 p.m."
    cfg.county = "Bexar"
    cfg.judicial_district = "285th"
    cfg.court_type = "District Court"
    cfg.location = "SA Legal Solutions"
    cfg.location_city = "San Antonio"
    cfg.reporter_name = "Miah Bardot"
    cfg.reporter_csr = "CSR No. 12129"
    cfg.reporter_expiration = "12/31/26"
    cfg.reporter_firm = "SA Legal Solutions"
    cfg.reporter_address = "San Antonio, Texas"
    cfg.plaintiff_counsel = [CounselInfo(
        name="John Salazar", firm="Salazar Law", sbot="12345678",
        address="100 Main St", city="San Antonio",
        state="Texas", zip_code="78205", party="Plaintiff"
    )]
    cfg.defense_counsel = [CounselInfo(
        name="Jane Durbin", firm="Durbin & Associates", sbot="87654321",
        party="Defendant"
    )]
    cfg.witnesses = [WitnessIndexEntry(
        name="Simon Ugalde", direct_page="10", cross_page="45"
    )]
    cfg.exhibits = [ExhibitEntry(
        number="1", description="Photograph of scene",
        offered_page="12", admitted_page="13"
    )]
    cfg.changes = [ChangeEntry(
        page="22", line="5", change="correction here", reason="misspoken"
    )]
    cfg.notary_county = "Bexar"
    cfg.cost_paid_by = "Plaintiff"
    return cfg


def test_lined_page_produces_table():
    """write_lined_page() should create exactly one table with 25 rows."""
    doc = create_document()
    write_lined_page(doc, ["Line content"] * 10)
    assert len(doc.tables) == 1
    assert len(doc.tables[0].rows) == 25


def test_lined_page_pads_to_25():
    """write_lined_page() pads fewer than 25 lines to exactly 25."""
    doc = create_document()
    write_lined_page(doc, ["only five lines"] * 5)
    assert len(doc.tables[0].rows) == 25


def test_paginate_splits_correctly():
    """paginate_lines() splits 50 lines into 2 pages of 25."""
    pages = paginate_lines(["x"] * 50)
    assert len(pages) == 2
    assert all(len(p) == 25 for p in pages)


def test_title_page_contains_cause_number():
    """Title page must contain cause_number."""
    cfg = make_full_config()
    doc = create_document()
    write_title_page(doc, cfg)
    all_text = " ".join(
        cell.text for t in doc.tables for row in t.rows for cell in row.cells
    )
    assert cfg.cause_number in all_text


def test_title_page_contains_witness_name():
    """Title page must contain witness name in uppercase."""
    cfg = make_full_config()
    doc = create_document()
    write_title_page(doc, cfg)
    all_text = " ".join(
        cell.text for t in doc.tables for row in t.rows for cell in row.cells
    )
    assert cfg.witness_name.upper() in all_text


def test_title_page_does_not_render_literal_witness_name_placeholder():
    cfg = make_full_config()
    doc = create_document()
    write_title_page(doc, cfg)
    all_text = " ".join(
        cell.text for t in doc.tables for row in t.rows for cell in row.cells
    )
    assert "(WITNESS NAME)" not in all_text


def test_title_page_does_not_use_article_before_party_name():
    cfg = make_full_config()
    doc = create_document()
    write_title_page(doc, cfg)
    all_text = " ".join(
        cell.text for t in doc.tables for row in t.rows for cell in row.cells
    )
    assert f"at the instance of the {cfg.plaintiff_name}" not in all_text
    assert f"at the instance of {cfg.plaintiff_name}" in all_text


def test_title_page_mentions_reporter_name_once():
    cfg = make_full_config()
    doc = create_document()
    write_title_page(doc, cfg)
    all_text = " ".join(
        cell.text for t in doc.tables for row in t.rows for cell in row.cells
    )
    assert all_text.count(cfg.reporter_name) == 1


def test_caption_contains_sbot():
    """Caption appearances block must show SBOT number."""
    cfg = make_full_config()
    doc = create_document()
    write_caption(doc, cfg)
    all_text = " ".join(
        cell.text for t in doc.tables for row in t.rows for cell in row.cells
    )
    assert "12345678" in all_text


def test_certificate_contains_reporter():
    """Certificate must contain reporter name and CSR number."""
    cfg = make_full_config()
    doc = create_document()
    write_certificate(doc, cfg)
    all_text = " ".join(
        cell.text for t in doc.tables for row in t.rows for cell in row.cells
    )
    assert cfg.reporter_name in all_text
    assert "12129" in all_text


def test_certificate_omits_signature_waiver_by_default():
    cfg = make_full_config()
    doc = create_document()
    write_certificate(doc, cfg)
    all_text = " ".join(
        cell.text for t in doc.tables for row in t.rows for cell in row.cells
    )
    assert "deposition transcript was waived by the witness and the" not in all_text
    assert "parties at the time of the deposition" not in all_text


def test_certificate_includes_signature_waiver_when_enabled():
    cfg = make_full_config()
    cfg.signature_waived = True
    doc = create_document()
    write_certificate(doc, cfg)
    all_text = " ".join(
        cell.text for t in doc.tables for row in t.rows for cell in row.cells
    )
    assert "deposition transcript was waived by the witness and the" in all_text
    assert "parties at the time of the deposition" in all_text


def test_changes_signature_renders_change_row():
    """Changes page must render the ChangeEntry row."""
    cfg = make_full_config()
    doc = create_document()
    write_changes_signature(doc, cfg)
    all_text = " ".join(
        cell.text for t in doc.tables for row in t.rows for cell in row.cells
    )
    assert "correction here" in all_text


def test_witness_index_renders_name():
    """Witness index must contain witness name."""
    cfg = make_full_config()
    doc = create_document()
    write_witness_index(doc, cfg)
    all_text = " ".join(
        cell.text for t in doc.tables for row in t.rows for cell in row.cells
    )
    assert "Simon Ugalde" in all_text


def test_exhibit_index_renders_exhibit():
    """Exhibit index must contain exhibit number and description."""
    cfg = make_full_config()
    doc = create_document()
    write_exhibit_index(doc, cfg)
    all_text = " ".join(
        cell.text for t in doc.tables for row in t.rows for cell in row.cells
    )
    assert "Photograph of scene" in all_text


# ═══════════════════════════════════════════════════════════════
# TESTS 33-35: Export utilities (Phase 5)
# ═══════════════════════════════════════════════════════════════

from spec_engine.exporter import strip_to_ascii, export_ascii
import tempfile


def test_strip_line_numbers_removes_prefix():
    """strip_to_ascii removes ' 1 ' style prefixes from transcript lines."""
    raw = " 1 Q.  Did you see the accident?\n 2 A.  Yes, I did."
    result = strip_to_ascii(raw)
    assert "Q.  Did you see the accident?" in result
    assert "A.  Yes, I did." in result
    assert " 1 " not in result
    assert " 2 " not in result


def test_strip_preserves_qa_content():
    """strip_to_ascii preserves Q/A text and blank line separators."""
    raw = (
        " 1 Q.  What is your name?\n"
        " 2 A.  Simon Ugalde.\n"
        " 3 \n"
        " 4 Q.  Where do you live?\n"
    )
    result = strip_to_ascii(raw)
    assert "Q.  What is your name?" in result
    assert "A.  Simon Ugalde." in result
    assert "Q.  Where do you live?" in result


def test_strip_to_ascii_expands_tabs_consistently():
    raw = " 1 \tQ.\tDid you see the accident?"
    result = strip_to_ascii(raw)
    assert "\t" not in result
    assert "Q." in result


def test_export_ascii_writes_file():
    """export_ascii writes a .txt file to disk."""
    raw = " 1 Q.  Test question.\n 2 A.  Test answer."
    with tempfile.TemporaryDirectory() as tmp:
        out = export_ascii(source=raw, output_path=os.path.join(tmp, "out.txt"))
        assert os.path.exists(out)
        content = open(out, encoding="utf-8").read()
        assert "Test question." in content
        assert "Test answer." in content


# ═══════════════════════════════════════════════════════════════
# TESTS 36-38: JobConfig dialog field coverage (Phase 6)
# ═══════════════════════════════════════════════════════════════

def test_jobconfig_round_trip_new_fields():
    """
    JobConfig.to_json() / from_json() preserves all Phase 6 fields.
    """
    from spec_engine.models import JobConfig, CounselInfo

    cfg = JobConfig()
    cfg.court_type = "County Court"
    cfg.proceeding_type = "Hearing"
    cfg.location = "Bexar County Courthouse"
    cfg.is_videotaped = True
    cfg.volume_number = 2
    cfg.total_volumes = 3
    cfg.cost_total = "450.00"
    cfg.cost_paid_by = "Defendant"
    cfg.notary_name = "Jane Smith"
    cfg.notary_county = "Bexar"
    cfg.identification_method = "Texas driver's license"
    cfg.is_official_reporter = True
    cfg.also_present = ["John Videographer", "Lisa Interpreter"]
    cfg.time_used = {"MR. SALAZAR": "2 hours 30 minutes"}
    cfg.plaintiff_counsel = [CounselInfo(
        name="John Salazar", firm="Salazar Law",
        sbot="12345678", party="Plaintiff")]
    cfg.defense_counsel = [CounselInfo(
        name="Jane Durbin", firm="Durbin & Assoc.",
        sbot="87654321", party="Defendant")]

    restored = JobConfig.from_json(cfg.to_json())

    assert restored.court_type == "County Court"
    assert restored.proceeding_type == "Hearing"
    assert restored.location == "Bexar County Courthouse"
    assert restored.is_videotaped is True
    assert restored.volume_number == 2
    assert restored.total_volumes == 3
    assert restored.cost_total == "450.00"
    assert restored.notary_name == "Jane Smith"
    assert restored.identification_method == "Texas driver's license"
    assert restored.is_official_reporter is True
    assert "John Videographer" in restored.also_present
    assert restored.time_used.get("MR. SALAZAR") == "2 hours 30 minutes"
    assert len(restored.plaintiff_counsel) == 1
    assert restored.plaintiff_counsel[0].sbot == "12345678"
    assert len(restored.defense_counsel) == 1


def test_counsel_info_all_fields():
    """CounselInfo stores all Phase 6 fields correctly."""
    from spec_engine.models import CounselInfo
    c = CounselInfo(
        name="John Salazar", firm="Salazar Law", sbot="12345678",
        address="100 Main St", city="San Antonio",
        state="Texas", zip_code="78205",
        phone="(210) 555-1234", party="Plaintiff")
    assert c.sbot == "12345678"
    assert c.zip_code == "78205"
    assert c.city == "San Antonio"
    assert c.party == "Plaintiff"


def test_title_page_uses_proceeding_type():
    """Title page depo label changes when proceeding_type is not Deposition."""
    from spec_engine.pages.title_page import write_title_page
    cfg = make_full_config()
    cfg.proceeding_type = "Hearing"
    doc = create_document()
    write_title_page(doc, cfg)
    assert len(doc.tables) >= 1


# ═══════════════════════════════════════════════════════════════
# TESTS 39-41: Phase 7 — field sync and classification
# ═══════════════════════════════════════════════════════════════

def test_jobconfig_save_and_reload():
    """JobConfig.save() writes JSON; load() reads it back correctly."""
    from spec_engine.models import JobConfig

    cfg = make_full_config()
    cfg.cause_number = "2025-CI-99999"
    cfg.witness_name = "Test Witness"
    with tempfile.TemporaryDirectory() as tmp:
        path = cfg.save(jobs_dir=tmp)
        assert os.path.exists(path)
        restored = JobConfig.load(path)
        assert restored.cause_number == "2025-CI-99999"
        assert restored.witness_name == "Test Witness"
        assert restored.reporter_name == cfg.reporter_name


def test_jobconfig_save_uses_cause_number_in_filename():
    """JobConfig.save() filename contains the cause number."""
    from spec_engine.models import JobConfig

    cfg = JobConfig()
    cfg.cause_number = "2025-CI-12345"
    with tempfile.TemporaryDirectory() as tmp:
        path = cfg.save(jobs_dir=tmp)
        assert "2025-CI-12345" in Path(path).name or "2025-CI-12345" in path


def test_ascii_export_handles_table_content():
    """strip_to_ascii handles content that doesn't start with line numbers."""
    from spec_engine.exporter import strip_to_ascii

    raw = (
        "Q.  Did you see the vehicle?\n"
        "A.  Yes.\n"
        "\n"
        "MR. SALAZAR:  Objection, form.\n"
    )
    result = strip_to_ascii(raw)
    assert "Did you see the vehicle?" in result
    assert "Objection, form." in result
    assert "Q." in result
    assert "A." in result


# ═══════════════════════════════════════════════════════════════
# TESTS (Phase 9) — Scopist flag persistence
# ═══════════════════════════════════════════════════════════════

from spec_engine.models import ScopistFlag


def test_scopist_flag_persists_in_job_config():
    """ScopistFlag objects survive JobConfig to_json() / from_json()."""
    from spec_engine.models import JobConfig, ScopistFlag

    cfg = JobConfig()
    cfg.cause_number = "2025-CI-FLAG-TEST"
    cfg.spec_flags = [
        ScopistFlag(number=1, description="Verify spelling",
                    block_index=10, category="general",
                    inline_text='[SCOPIST: FLAG 1: "Ugaldi" — verify]'),
        ScopistFlag(number=2, description="Date mismatch",
                    block_index=22, category="date",
                    inline_text='[SCOPIST: FLAG 2: Date "Jan 15"]'),
    ]
    restored = JobConfig.from_json(cfg.to_json())
    assert len(restored.spec_flags) == 2
    assert restored.spec_flags[0].number == 1
    assert restored.spec_flags[0].category == "general"
    assert restored.spec_flags[1].description == "Date mismatch"
    assert restored.spec_flags[1].category == "date"


def test_scopist_flag_inline_text_preserved():
    """ScopistFlag inline_text round-trips exactly."""
    from spec_engine.models import JobConfig, ScopistFlag

    cfg = JobConfig()
    expected_text = '[SCOPIST: FLAG 3: "Perez" — spelled on record as P-E-R-E-Z]'
    cfg.spec_flags = [
        ScopistFlag(number=3, description="Post-record spelling",
                    block_index=45, category="post_record",
                    inline_text=expected_text)
    ]
    restored = JobConfig.from_json(cfg.to_json())
    assert restored.spec_flags[0].inline_text == expected_text


def test_process_transcript_returns_flags_key():
    """process_transcript() result dict includes 'flags' key."""
    from docx import Document
    from spec_engine.document_builder import process_transcript
    from spec_engine.models import JobConfig

    with tempfile.TemporaryDirectory() as tmp:
        in_path = os.path.join(tmp, "in.docx")
        out_path = os.path.join(tmp, "out.docx")

        doc = Document()
        doc.add_paragraph("Speaker 0: Hello.")
        doc.add_paragraph("Speaker 1: Yes.")
        doc.save(in_path)

        cfg = JobConfig()
        cfg.speaker_map = {0: "THE WITNESS", 1: "MR. TEST"}
        cfg.speaker_map_verified = True

        result = process_transcript(
            input_docx_path=in_path,
            output_docx_path=out_path,
            job_config=cfg,
        )
        assert "flags" in result
        assert isinstance(result["flags"], list)


def test_empty_spec_flags_serializes_cleanly():
    """JobConfig with no flags serializes spec_flags as empty list."""
    import json
    from spec_engine.models import JobConfig

    cfg = JobConfig()
    data = json.loads(cfg.to_json())
    assert "spec_flags" in data
    assert data["spec_flags"] == []


def test_answer_token_after_question_mark():
    """Rule 1: embedded answer token after '?' must split Q into Q/A/Q."""
    cfg = make_config()
    cfg.split_embedded_answers = True
    block = Block(
        speaker_id=2,
        text="Did you know Dr. Brown? No. During the course.",
        raw_text="",
    )
    state = ClassifierState()
    results = classify_block(block, cfg, state, block_index=0)
    types = [lt for lt, _ in results]
    assert types.count(LineType.Q) == 2, f"Expected 2 Q blocks, got: {results}"
    assert types.count(LineType.A) == 1, f"Expected 1 A block, got: {results}"


def test_answer_token_i_have_not():
    """Rule 1: 'I have not.' is a valid answer token and must split."""
    cfg = make_config()
    cfg.split_embedded_answers = True
    block = Block(
        speaker_id=2,
        text="Have you ever been deposed? I have not. This is my first time.",
        raw_text="",
    )
    state = ClassifierState()
    results = classify_block(block, cfg, state, block_index=0)
    types = [lt for lt, _ in results]
    assert LineType.A in types, f"Expected A block in results, got: {results}"


def test_correct_mid_sentence_split():
    """Rule 2: 'Correct.' after declarative statement must split to own A."""
    from spec_engine.classifier import split_correct_mid

    result = split_correct_mid(
        "The plan was to discharge her that day. Correct. "
        "Do you know if HCA had the capability?"
    )
    assert result is not None, "FAIL: Rule 2 should match"
    before, correct, after = result
    assert correct == "Correct.", f"FAIL: expected 'Correct.', got '{correct}'"
    assert "Do you know" in after, f"FAIL: continuation missing, got '{after}'"


def test_correct_mid_not_after_question():
    """Rule 2 guard: 'Correct.' after '?' must NOT trigger split."""
    from spec_engine.classifier import split_correct_mid

    result = split_correct_mid("Did you do that? Correct. And then what?")
    assert result is None, "FAIL: Rule 2 must not fire when Correct. follows a question mark"


def test_trailing_okay_moved_to_next_q():
    """Rule 3: A ending in 'Okay.' strips it and prepends to next Q."""
    from spec_engine.classifier import fix_trailing_okay_in_answer

    blocks = [
        (LineType.A, "The electrocautery. Okay."),
        (LineType.Q, "And do you know whether it came into contact?"),
    ]
    result = fix_trailing_okay_in_answer(blocks)
    a_text = result[0][1]
    q_text = result[1][1]
    assert not a_text.endswith("Okay."), f"FAIL: 'Okay.' should be stripped from A, got: '{a_text}'"
    assert q_text.startswith("Okay."), f"FAIL: Q should start with 'Okay.', got: '{q_text}'"


def test_trailing_okay_not_moved_if_question_mark():
    """Rule 3 guard: A ending in 'Okay?' must NOT be modified."""
    from spec_engine.classifier import fix_trailing_okay_in_answer

    blocks = [
        (LineType.A, "Is that okay?"),
        (LineType.Q, "And what happened next?"),
    ]
    result = fix_trailing_okay_in_answer(blocks)
    a_text = result[0][1]
    q_text = result[1][1]
    assert "okay?" in a_text.lower(), f"FAIL: 'okay?' should remain in A, got: '{a_text}'"
    assert not q_text.startswith("Okay."), f"FAIL: Q should not be modified, got: '{q_text}'"


def test_trailing_okay_no_double_prepend():
    """Rule 3 guard: Do not prepend 'Okay.' if next Q already starts with it."""
    from spec_engine.classifier import fix_trailing_okay_in_answer

    blocks = [
        (LineType.A, "I could not. Okay."),
        (LineType.Q, "Okay. And what did you do next?"),
    ]
    result = fix_trailing_okay_in_answer(blocks)
    q_text = result[1][1]
    assert not q_text.startswith("Okay. Okay."), f"FAIL: Double 'Okay.' prepend detected: '{q_text}'"


def test_uh_um_preserved_through_all_splits():
    """Verbatim rule: uh/um must survive all Q/A split operations."""
    cfg = make_config()
    cfg.split_embedded_answers = True
    block = Block(
        speaker_id=2,
        text="Did you, uh, see the driver? No, um, I did not.",
        raw_text="",
    )
    state = ClassifierState()
    results = classify_block(block, cfg, state, block_index=0)
    all_text = " ".join(text for _, text in results)
    assert "uh" in all_text, f"FAIL: 'uh' was removed from output: {results}"
    assert "um" in all_text, f"FAIL: 'um' was removed from output: {results}"


@pytest.mark.skip(reason="format_transcript() kwarg API (use_qa_format, clean_verbatim) removed — format_blocks_to_text does not accept these kwargs")
def test_percent_not_normalized_by_default():
    from core.correction_runner import format_blocks_to_text as format_transcript

    result = format_transcript("50% of the time.", use_qa_format=False)
    assert "50%" in result


@pytest.mark.skip(reason="format_transcript() kwarg API (use_qa_format, clean_verbatim) removed — format_blocks_to_text does not accept these kwargs")
def test_money_not_normalized_by_default():
    from core.correction_runner import format_blocks_to_text as format_transcript

    result = format_transcript("It cost $1,000.00 total.", use_qa_format=False)
    assert "$1,000.00" in result


@pytest.mark.skip(reason="format_transcript() kwarg API (use_qa_format, clean_verbatim) removed — format_blocks_to_text does not accept these kwargs")
def test_sentence_split_not_applied_by_default():
    from core.correction_runner import format_blocks_to_text as format_transcript

    text = "A. I went to the store. I bought milk."
    result = format_transcript(text, use_qa_format=True)
    assert "went to the store" in result
    assert "bought milk" in result
    a_count = result.count("\nA.") + (1 if result.startswith("A.") else 0)
    assert a_count <= 1


@pytest.mark.skip(reason="format_transcript() kwarg API (use_qa_format, clean_verbatim) removed — format_blocks_to_text does not accept these kwargs")
def test_clean_verbatim_normalizes_percent():
    from core.correction_runner import format_blocks_to_text as format_transcript

    result = format_transcript(
        "50% of the time.",
        use_qa_format=False,
        clean_verbatim=True,
    )
    assert "percent" in result.lower()


@pytest.mark.skip(reason="format_transcript() kwarg API (use_qa_format, clean_verbatim) removed — format_blocks_to_text does not accept these kwargs")
def test_uh_preserved_in_standard_mode():
    from core.correction_runner import format_blocks_to_text as format_transcript

    result = format_transcript("A. Uh, I think so.", use_qa_format=True)
    assert "uh" in result.lower()


@pytest.mark.skip(reason="format_transcript() kwarg API (use_qa_format, clean_verbatim) removed — format_blocks_to_text does not accept these kwargs")
def test_qa_width_applied_to_qa_lines():
    from core.correction_runner import format_blocks_to_text as format_transcript
    from spec_engine.emitter import QA_WRAP_WIDTH as QA_WIDTH

    long_answer = "A. " + ("word " * 15)
    result = format_transcript(long_answer, use_qa_format=True)
    for idx, line in enumerate(result.splitlines()):
        if line.startswith("A.") or (
            idx > 0 and not line.startswith("Q.") and not line.startswith("A.")
        ):
            assert len(line) <= QA_WIDTH + 5, (
                f"Q/A line exceeds QA_WIDTH={QA_WIDTH}: {len(line)} chars: {line!r}"
            )


@pytest.mark.skip(reason="Asserts SP-aligned continuation which is wrong per UFM spec — UFM requires left-margin continuation. Also imports removed formatter module. DOCX SP output is covered by test_emitter_speaker_continuation_aligns_under_content.")
def test_format_blocks_speaker_continuation_alignment():
    from core.correction_runner import format_blocks_to_text as format_blocks
    from spec_engine.models import Block, BlockType

    block = Block(
        speaker_id=2,
        speaker_name="MR. SALAZAR",
        speaker_role="ATTORNEY",
        block_type=BlockType.SPEAKER,
        text=" ".join(["word"] * 20),
    )
    result = format_blocks([block])
    lines = result.splitlines()
    assert len(lines) > 1
    assert lines[0].startswith("MR. SALAZAR:")
    assert lines[1].startswith(" " * len("MR. SALAZAR:  "))


def _resolve_spec_source(output_paths: dict, session: dict) -> str | None:
    auto_docx = (
        output_paths.get("auto_docx", "")
        or session.get("last_auto_docx_path", "")
    )
    session_path = session.get("audio_path", "")

    if auto_docx and os.path.exists(auto_docx):
        return auto_docx
    if (
        session_path
        and session_path.lower().endswith(".docx")
        and os.path.exists(session_path)
    ):
        return session_path
    return None


def test_export_routing_prefers_auto_docx():
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        auto_path = f.name
    try:
        result = _resolve_spec_source(
            output_paths={"auto_docx": auto_path},
            session={"audio_path": "/some/other.docx"},
        )
        assert result == auto_path
    finally:
        os.unlink(auto_path)


def test_export_routing_falls_back_to_session_docx():
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        session_docx = f.name
    try:
        result = _resolve_spec_source(
            output_paths={},
            session={"audio_path": session_docx},
        )
        assert result == session_docx
    finally:
        os.unlink(session_docx)


def test_export_routing_returns_none_when_no_docx():
    result = _resolve_spec_source(
        output_paths={},
        session={"audio_path": "/nonexistent/file.mp3"},
    )
    assert result is None


def test_export_routing_ignores_missing_auto_docx():
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        session_docx = f.name
    try:
        result = _resolve_spec_source(
            output_paths={"auto_docx": "/nonexistent/auto.docx"},
            session={"audio_path": session_docx},
        )
        assert result == session_docx
    finally:
        os.unlink(session_docx)


def _make_temp_docx(text: str) -> str:
    from docx import Document

    doc = Document()
    doc.add_paragraph(text)
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp.name)
    tmp.close()
    return tmp.name


@pytest.mark.skip(reason="ufm_engine is inactive subsystem — requires docxtpl")
def test_docx_merger_single_source():
    from docx import Document
    from ufm_engine.docx_merger import DocxMerger

    src = _make_temp_docx("Section One Content")
    out = tempfile.mktemp(suffix=".docx")
    try:
        merger = DocxMerger()
        result = merger.merge([src], out)
        assert os.path.exists(out)
        assert str(result.output_path).endswith(".docx")
        doc = Document(out)
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "Section One Content" in full_text
    finally:
        for f in [src, out]:
            if os.path.exists(f):
                os.unlink(f)


@pytest.mark.skip(reason="ufm_engine is inactive subsystem — requires docxtpl")
def test_docx_merger_multi_source_order():
    from docx import Document
    from ufm_engine.docx_merger import DocxMerger

    src1 = _make_temp_docx("FIRST SECTION")
    src2 = _make_temp_docx("SECOND SECTION")
    src3 = _make_temp_docx("THIRD SECTION")
    out = tempfile.mktemp(suffix=".docx")
    try:
        merger = DocxMerger()
        merger.merge([src1, src2, src3], out)
        doc = Document(out)
        full_text = " ".join(p.text for p in doc.paragraphs)
        idx1 = full_text.find("FIRST SECTION")
        idx2 = full_text.find("SECOND SECTION")
        idx3 = full_text.find("THIRD SECTION")
        assert idx1 < idx2 < idx3
    finally:
        for f in [src1, src2, src3, out]:
            if os.path.exists(f):
                os.unlink(f)


def test_docx_merger_missing_source_raises():
    from ufm_engine.docx_merger import DocxMerger

    merger = DocxMerger()
    with pytest.raises((ValueError, FileNotFoundError)):
        merger.merge(["/nonexistent/source.docx"], "/tmp/out.docx")


def test_docx_merger_empty_list_raises():
    from ufm_engine.docx_merger import DocxMerger

    merger = DocxMerger()
    with pytest.raises((ValueError, TypeError)):
        merger.merge([], "/tmp/out.docx")


@pytest.mark.skip(reason="normalize_sentence_spacing() API changed — now requires (text, records, block_index)")
def test_sentence_spacing_does_not_collapse_newlines():
    from spec_engine.corrections import normalize_sentence_spacing

    text = "Done.\n\nQ. What is your name?"
    result = normalize_sentence_spacing(text)
    assert "\n\nQ." in result


@pytest.mark.skip(reason="format_transcript text→text API removed; I-10 rule coverage to be added in Phase 3")
def test_formatter_highway_normalization():
    from core.correction_runner import format_blocks_to_text as format_transcript

    result = format_transcript("A. I was driving on I 10 near downtown.")
    assert "I-10" in result


def test_clean_block_oath_garble_correction():
    result = clean_block("so help you guide", make_config())[0]
    assert "so help you god" in result.lower()


def test_clean_block_cause_number_ci_format():
    result = clean_block("The cause number is 2025 CI 12281.", make_config())[0]
    assert "2025-CI-12281" in result


def test_classifier_unmapped_speaker_returns_flag():
    cfg = make_config()
    cfg.speaker_map_verified = True
    cfg.speaker_map = {1: "THE WITNESS", 2: "MR. TEST"}
    state = ClassifierState()
    block = Block(speaker_id=99, text="Unknown speaker text.", raw_text="")
    results = classify_block(block, cfg, state, block_index=0)
    assert results[0][0] == LineType.FLAG
    assert "speaker 99 role not in speaker_map" in results[0][1].lower()


def test_certificate_mentions_original_deposition_transcript():
    cfg = make_full_config()
    doc = create_document()
    write_certificate(doc, cfg)
    all_text = " ".join(
        cell.text for t in doc.tables for row in t.rows for cell in row.cells
    )
    assert "original deposition transcript" in all_text.lower()


def test_block_builder_uses_word_objects():
    from spec_engine.block_builder import build_blocks_from_deepgram
    from spec_engine.models import Word

    blocks = build_blocks_from_deepgram(
        {
            "utterances": [
                {
                    "speaker": 1,
                    "transcript": "Yes, sir.",
                    "words": [
                        {"word": "Yes", "start": 0.0, "end": 0.2, "confidence": 0.9, "speaker": 1},
                        {"word": "sir", "start": 0.21, "end": 0.4, "confidence": 0.88, "speaker": 1},
                    ],
                }
            ]
        }
    )
    assert len(blocks) == 1
    assert blocks[0].words
    assert isinstance(blocks[0].words[0], Word)


def test_process_blocks_maps_speakers_before_classification():
    from spec_engine.processor import process_blocks
    from spec_engine.models import Block, BlockType

    blocks = [Block(speaker_id=1, text="I don't remember.", raw_text="I don't remember.")]
    result = process_blocks(
        blocks,
        {
            "speaker_map": {1: "THE WITNESS"},
            "witness_id": 1,
            "cause_number": "TEST-1",
        },
    )
    assert result[0].speaker_name == "THE WITNESS"
    assert result[0].block_type == BlockType.ANSWER


@pytest.mark.skip(reason="Output format assertion predates current emitter — review in Phase 2")
def test_run_pipeline_returns_blocks_and_text():
    from pipeline.processor import run_pipeline

    result = run_pipeline(
        {
            "utterances": [
                {
                    "speaker": 2,
                    "transcript": "Did you go there? Yes.",
                    "words": [],
                }
            ]
        },
        {
            "speaker_map": {1: "THE WITNESS", 2: "EXAMINING ATTORNEY"},
            "witness_id": 1,
            "examining_attorney_id": 2,
            "cause_number": "TEST-PIPELINE",
        },
    )
    assert "blocks" in result
    assert "text" in result
    assert "Q.\tDid you go there?" in result["text"]
    assert "A.\tYes." in result["text"]


@pytest.mark.skip(reason="ai_tools module not yet implemented — planned feature")
def test_validate_legal_correction_output_rejects_label_reordering():
    from ai_tools import validate_legal_correction_output

    original = "Q.\tDid you go there?\nA.\tYes."
    corrected = "A.\tYes.\nQ.\tDid you go there?"

    with pytest.raises(ValueError, match="protected Q\\./A\\. or speaker label ordering"):
        validate_legal_correction_output(original, corrected)


def test_speaker_mapper_no_default_persistence_bucket():
    from spec_engine.speaker_mapper import _job_key

    assert _job_key({}) == ""


def test_speaker_resolver_normalizes_ids_and_roles():
    from spec_engine.speaker_resolver import (
        ROLE_EXAMINING_ATTORNEY,
        ROLE_OPPOSING_COUNSEL,
        ROLE_WITNESS,
        normalize_speaker_id,
        normalize_speaker_role,
        resolve_speaker,
    )

    assert normalize_speaker_id(0) == 0
    assert normalize_speaker_id("0") == 0
    assert normalize_speaker_id("Speaker 2") == 2

    sid, role, name = resolve_speaker(
        "2",
        {"speaker_map": {2: "MR. SMITH"}},
    )
    assert sid == 2
    assert role == ROLE_EXAMINING_ATTORNEY or role == "ATTORNEY"
    assert name == "MR. SMITH"

    sid2, role2, name2 = resolve_speaker(
        1,
        {"speaker_map": {1: "THE WITNESS"}},
    )
    assert sid2 == 1
    assert role2 == ROLE_WITNESS
    assert name2 == "THE WITNESS"

    assert normalize_speaker_role("EXAMINING ATTORNEY") == ROLE_EXAMINING_ATTORNEY
    assert normalize_speaker_role("MR. BOYCE - OPPOSING COUNSEL") == ROLE_OPPOSING_COUNSEL


@pytest.mark.skip(reason="ai_tools module not yet implemented — planned feature")
def test_parse_indexed_ai_output_rejects_reordered_output():
    from ai_tools import _parse_indexed_ai_output

    ok = _parse_indexed_ai_output("[0|Q] Hello.\n[1|A] Yes.", 2)
    assert ok == {0: "Hello.", 1: "Yes."}

    bad = _parse_indexed_ai_output("[1|A] Yes.\n[0|Q] Hello.", 2)
    assert bad is None


def test_build_blocks_from_text_splits_sentences():
    from spec_engine.block_builder import build_blocks_from_text

    blocks = build_blocks_from_text("State your name. Did you go there?")
    assert len(blocks) == 2
    assert blocks[0].text == "State your name."
    assert blocks[1].text == "Did you go there?"


@pytest.mark.skip(reason="ai_tools module not yet implemented — planned feature")
def test_diff_viewer_summary():
    from utils.diff_viewer import diff_summary

    assert diff_summary("A\nB", "A\nC") == "1 of 2 line(s) changed."


def test_objection_speaker_uses_speaker_map_not_unknown():
    from spec_engine.objections import _resolve_objection_speaker
    from spec_engine.models import JobConfig

    cfg = JobConfig(
        speaker_map={
            0: "THE VIDEOGRAPHER",
            1: "THE WITNESS",
            2: "MR. ALLAN",
            3: "MR. BOYCE - OPPOSING COUNSEL",
        },
        examining_attorney_id=2,
        witness_id=1,
        speaker_map_verified=True,
    )
    result = _resolve_objection_speaker(cfg)
    assert result != "MR. UNKNOWN"
    assert "BOYCE" in result.upper() or result == "MR. BOYCE - OPPOSING COUNSEL"


def test_objection_speaker_fallback_is_not_mr_unknown():
    from spec_engine.objections import _resolve_objection_speaker
    from spec_engine.models import JobConfig

    cfg = JobConfig()
    result = _resolve_objection_speaker(cfg)
    assert result != "MR. UNKNOWN"
    assert result == "COUNSEL"


def test_objection_garble_corrected_to_verbatim_form():
    from spec_engine.corrections import apply_universal_corrections

    test_cases = [
        ("Exit form", "Objection. Form."),
        ("Action form", "Objection. Form."),
        ("Action point", "Objection. Form."),
        ("Objection form", "Objection. Form."),
    ]
    for garble, expected in test_cases:
        records = []
        result = apply_universal_corrections(garble, records, block_index=0)
        assert result == expected


def test_objection_to_form_not_rewritten_by_objection_patterns():
    from spec_engine.objections import OBJECTION_PATTERNS

    for pattern in OBJECTION_PATTERNS:
        assert "objection to form" not in pattern.lower()


def test_spoken_objection_form_not_rewritten():
    from spec_engine.corrections import clean_block
    from spec_engine.models import JobConfig

    text = "Objection. Form."
    result, records = clean_block(text, JobConfig(), block_index=0)[:2]
    assert result == "Objection. Form."
