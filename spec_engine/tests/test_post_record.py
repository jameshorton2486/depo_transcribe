"""
spec_engine/tests/test_post_record.py

Tests for post-record spellings page and derive_correct_spelling helper.
All offline and deterministic.
Run: python -m pytest spec_engine/tests/test_post_record.py -v
"""
import pytest
from docx import Document
from spec_engine.pages.post_record import (
    derive_correct_spelling,
    write_post_record_section,
)
from spec_engine.models import PostRecordSpelling


class TestDeriveCorrectSpelling:

    def test_all_caps_sequence(self):
        assert derive_correct_spelling("B-A-L-D-E-R-A-S") == "Balderas"

    def test_mixed_case_sequence(self):
        assert derive_correct_spelling("B-r-e-n-n-e-n") == "Brennen"

    def test_five_letter_name(self):
        assert derive_correct_spelling("T-O-V-A-R") == "Tovar"

    def test_short_sequence_returns_empty(self):
        assert derive_correct_spelling("A-B") == ""

    def test_empty_string_returns_empty(self):
        assert derive_correct_spelling("") == ""

    def test_non_letter_returns_empty(self):
        assert derive_correct_spelling("1-2-3") == ""

    def test_multi_char_part_returns_empty(self):
        assert derive_correct_spelling("BA-LD-ER") == ""

    def test_three_letter_minimum(self):
        assert derive_correct_spelling("A-B-C") == "Abc"

    def test_capitalizes_first_only(self):
        result = derive_correct_spelling("T-O-V-A-R")
        assert result[0].isupper()
        assert result[1:].islower()


class TestWritePostRecordSection:

    def _make_spelling(self, name, correct, letters, flag=None):
        return PostRecordSpelling(
            name=name,
            correct_spelling=correct,
            letters_as_given=letters,
            block_index=0,
            flag=flag,
        )

    def test_empty_list_adds_nothing(self):
        doc = Document()
        initial_para_count = len(doc.paragraphs)
        write_post_record_section(doc, [], None)
        assert len(doc.paragraphs) == initial_para_count

    def test_single_spelling_adds_content(self):
        doc = Document()
        spelling = self._make_spelling(
            "Balderas", "Balderas", "B-A-L-D-E-R-A-S"
        )
        write_post_record_section(doc, [spelling], None)
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "POST-RECORD SPELLINGS" in all_text

    def test_name_appears_in_output(self):
        doc = Document()
        spelling = self._make_spelling(
            "Balderas", "Balderas", "B-A-L-D-E-R-A-S"
        )
        write_post_record_section(doc, [spelling], None)
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Balderas" in all_text

    def test_letters_as_given_appears(self):
        doc = Document()
        spelling = self._make_spelling(
            "Tovar", "Tovar", "T-O-V-A-R"
        )
        write_post_record_section(doc, [spelling], None)
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "T-O-V-A-R" in all_text

    def test_flag_appears_when_set(self):
        doc = Document()
        spelling = self._make_spelling(
            "Brennen", "Brennen", "B-r-e-n-n-e-n",
            flag="Verify against audio"
        )
        write_post_record_section(doc, [spelling], None)
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Verify against audio" in all_text

    def test_multiple_spellings(self):
        doc = Document()
        spellings = [
            self._make_spelling("Balderas", "Balderas", "B-A-L-D-E-R-A-S"),
            self._make_spelling("Tovar", "Tovar", "T-O-V-A-R"),
        ]
        write_post_record_section(doc, spellings, None)
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Balderas" in all_text
        assert "T-O-V-A-R" in all_text

    def test_correct_form_appears(self):
        doc = Document()
        spelling = self._make_spelling(
            "Balderas", "Balderas", "B-A-L-D-E-R-A-S"
        )
        write_post_record_section(doc, [spelling], None)
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Balderas" in all_text
