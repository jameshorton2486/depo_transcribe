import tempfile
from pathlib import Path

from docx import Document
import pytest

from pipeline.processor import run_pipeline
from spec_engine.document_builder import _build_witness_intro_lines, process_transcript
from spec_engine.models import Block, BlockType, JobConfig, LineType
from spec_engine.qa_fixer import _merge_reporter_preamble_blocks


def test_process_blocks_preserves_spec_engine_pipeline_order(monkeypatch):
    from spec_engine import processor as processor_module

    calls = []
    blocks = [Block(speaker_id=1, text="Test.", raw_text="")]
    cfg = {"speaker_map": {1: "THE WITNESS"}, "speaker_map_verified": True}

    def _mark(name):
        def _inner(current_blocks, *args, **kwargs):
            calls.append(name)
            return current_blocks
        return _inner

    monkeypatch.setattr(processor_module, "apply_corrections", _mark("apply_corrections"))
    monkeypatch.setattr(processor_module, "classify_blocks", _mark("classify_blocks"))
    monkeypatch.setattr(processor_module, "fix_qa_structure", _mark("fix_qa_structure"))
    monkeypatch.setattr(processor_module, "extract_objections", _mark("extract_objections"))

    def _map(current_blocks, *args, **kwargs):
        calls.append("map_speakers")
        for block in current_blocks:
            block.speaker_role = "WITNESS"
            block.speaker_name = "THE WITNESS"
        return current_blocks

    monkeypatch.setattr(processor_module, "map_speakers", _map)

    class _Validation:
        errors = []
        warnings = []

    def _validate(current_blocks, speaker_map_verified=False):
        calls.append("validate_blocks")
        return _Validation()

    monkeypatch.setattr(processor_module, "validate_blocks", _validate)

    processor_module.process_blocks(blocks, cfg)

    assert calls == [
        "apply_corrections",
        "map_speakers",
        "classify_blocks",
        "fix_qa_structure",
        "extract_objections",
        "classify_blocks",
        "validate_blocks",
    ]


def test_process_blocks_logs_additional_stage_snapshots(monkeypatch):
    from spec_engine import processor as processor_module

    blocks = [Block(speaker_id=1, text="Test.", raw_text="")]
    cfg = {"speaker_map": {1: "THE WITNESS"}, "speaker_map_verified": True}
    snapshot_names = []

    monkeypatch.setattr(processor_module, "apply_corrections", lambda current_blocks, *_args, **_kwargs: current_blocks)
    monkeypatch.setattr(processor_module, "fix_qa_structure", lambda current_blocks, *_args, **_kwargs: current_blocks)
    monkeypatch.setattr(processor_module, "extract_objections", lambda current_blocks, *_args, **_kwargs: current_blocks)

    def _map(current_blocks, *_args, **_kwargs):
        for block in current_blocks:
            block.speaker_role = "WITNESS"
            block.speaker_name = "THE WITNESS"
        return current_blocks

    def _classify(current_blocks, *_args, **_kwargs):
        for block in current_blocks:
            block.block_type = BlockType.ANSWER
        return current_blocks

    class _Validation:
        errors = []
        warnings = []

    class _Run:
        def snapshot(self, name, current_blocks):
            snapshot_names.append(name)

        def log_step(self, *args, **kwargs):
            pass

        def log_corrections_from_blocks(self, current_blocks):
            pass

        def write_validation(self, validation):
            pass

    monkeypatch.setattr(processor_module, "map_speakers", _map)
    monkeypatch.setattr(processor_module, "classify_blocks", _classify)
    monkeypatch.setattr(
        processor_module,
        "validate_blocks",
        lambda current_blocks, speaker_map_verified=False: _Validation(),
    )

    processor_module.process_blocks(blocks, cfg, run_logger=_Run())

    assert "02a_blocks_speaker_mapped" in snapshot_names
    assert "03a_blocks_classified" in snapshot_names
    assert "04a_blocks_qa_fixed" in snapshot_names


def test_process_blocks_reclassifies_after_objection_extraction(monkeypatch):
    from spec_engine import processor as processor_module

    classify_calls = []
    blocks = [Block(speaker_id=3, text="Objection. Form.", raw_text="")]
    cfg = {"speaker_map": {3: "MR. BOYCE"}, "speaker_map_verified": True}

    monkeypatch.setattr(processor_module, "apply_corrections", lambda current_blocks, *_args, **_kwargs: current_blocks)
    monkeypatch.setattr(processor_module, "fix_qa_structure", lambda current_blocks, *_args, **_kwargs: current_blocks)
    monkeypatch.setattr(processor_module, "extract_objections", lambda current_blocks, *_args, **_kwargs: current_blocks)

    def _map(current_blocks, *_args, **_kwargs):
        for block in current_blocks:
            block.speaker_role = "OPPOSING_COUNSEL"
            block.speaker_name = "MR. BOYCE"
        return current_blocks

    monkeypatch.setattr(processor_module, "map_speakers", _map)

    def _classify(current_blocks, *_args, **_kwargs):
        classify_calls.append(len(classify_calls))
        return current_blocks

    class _Validation:
        errors = []
        warnings = []

    monkeypatch.setattr(processor_module, "classify_blocks", _classify)
    monkeypatch.setattr(
        processor_module,
        "validate_blocks",
        lambda current_blocks, speaker_map_verified=False: _Validation(),
    )

    processor_module.process_blocks(blocks, cfg)

    assert len(classify_calls) == 2


def test_process_blocks_raises_runtime_error_for_incomplete_speaker_mapping(monkeypatch):
    from spec_engine import processor as processor_module

    blocks = [Block(speaker_id=1, text="Test.", raw_text="")]
    cfg = {"speaker_map": {1: "THE WITNESS"}, "speaker_map_verified": True}

    monkeypatch.setattr(processor_module, "apply_corrections", lambda current_blocks, *_args, **_kwargs: current_blocks)
    monkeypatch.setattr(processor_module, "map_speakers", lambda current_blocks, *_args, **_kwargs: current_blocks)

    with pytest.raises(RuntimeError, match="Speaker mapping incomplete"):
        processor_module.process_blocks(blocks, cfg)


def test_process_blocks_raises_runtime_error_for_missing_block_type(monkeypatch):
    from spec_engine import processor as processor_module

    blocks = [Block(speaker_id=1, text="Test.", raw_text="")]
    cfg = {"speaker_map": {1: "THE WITNESS"}, "speaker_map_verified": True}

    monkeypatch.setattr(processor_module, "apply_corrections", lambda current_blocks, *_args, **_kwargs: current_blocks)

    def _map(current_blocks, *_args, **_kwargs):
        for block in current_blocks:
            block.speaker_role = "WITNESS"
            block.speaker_name = "THE WITNESS"
        return current_blocks

    monkeypatch.setattr(processor_module, "map_speakers", _map)
    def _classify(current_blocks, *_args, **_kwargs):
        for block in current_blocks:
            block.block_type = None
        return current_blocks

    monkeypatch.setattr(processor_module, "classify_blocks", _classify)

    with pytest.raises(RuntimeError, match="Classification failed"):
        processor_module.process_blocks(blocks, cfg)


def test_witness_intro_lines_use_metadata_template():
    cfg = JobConfig(
        witness_name="Matthew Allan Coger",
        speaker_map={2: "MR. ALLAN"},
        examining_attorney_id=2,
    )

    lines = _build_witness_intro_lines(cfg)

    assert lines == [
        (LineType.HEADER, "MATTHEW ALLAN COGER,"),
        (LineType.PLAIN, "having been first duly sworn, testified as follows:"),
        (LineType.HEADER, "EXAMINATION"),
        (LineType.BY, "BY MR. ALLAN:"),
    ]


def test_witness_intro_not_added_when_exam_header_already_present():
    cfg = JobConfig(
        witness_name="Matthew Allan Coger",
        speaker_map={2: "MR. ALLAN"},
        examining_attorney_id=2,
    )

    lines = _build_witness_intro_lines(cfg, [(LineType.HEADER, "EXAMINATION")])

    assert lines == []


def test_merge_reporter_preamble_blocks_joins_consecutive_reporter_paragraphs():
    blocks = [
        Block(
            speaker_id=4,
            speaker_name="THE REPORTER",
            speaker_role="REPORTER",
            block_type=BlockType.SPEAKER,
            text="This is Cause Number 2025CI19595:",
            raw_text="This is Cause Number 2025CI19595:",
        ),
        Block(
            speaker_id=4,
            speaker_name="THE REPORTER",
            speaker_role="REPORTER",
            block_type=BlockType.SPEAKER,
            text="This deposition is being taken in accordance with the Texas Rules of Civil Procedure.",
            raw_text="This deposition is being taken in accordance with the Texas Rules of Civil Procedure.",
        ),
        Block(
            speaker_id=4,
            speaker_name="THE REPORTER",
            speaker_role="REPORTER",
            block_type=BlockType.SPEAKER,
            text="Counsel, will you please state your agreement for this deposition?",
            raw_text="Counsel, will you please state your agreement for this deposition?",
        ),
    ]

    merged = _merge_reporter_preamble_blocks(blocks)

    assert len(merged) == 1
    assert "This deposition is being taken" in merged[0].text
    assert "Counsel, will you please state your agreement" in merged[0].text
    assert merged[0].meta["merged_reporter_preamble"] is True


def test_merge_reporter_preamble_blocks_stops_when_speaker_changes():
    reporter = Block(
        speaker_id=4,
        speaker_name="THE REPORTER",
        speaker_role="REPORTER",
        block_type=BlockType.SPEAKER,
        text="This is Cause Number 2025CI19595:",
        raw_text="This is Cause Number 2025CI19595:",
    )
    witness = Block(
        speaker_id=1,
        speaker_name="THE WITNESS",
        speaker_role="WITNESS",
        block_type=BlockType.ANSWER,
        text="Yes, ma'am.",
        raw_text="Yes, ma'am.",
    )

    merged = _merge_reporter_preamble_blocks([reporter, witness])

    assert len(merged) == 2
    assert merged[0].text == reporter.text
    assert merged[1].text == witness.text


def test_process_transcript_inserts_witness_intro_block_into_docx():
    cfg = JobConfig(
        witness_name="Matthew Allan Coger",
        speaker_map={1: "THE WITNESS", 2: "MR. ALLAN"},
        witness_id=1,
        examining_attorney_id=2,
        speaker_map_verified=True,
    )

    with tempfile.TemporaryDirectory() as tmp:
        input_path = Path(tmp) / "input.docx"
        output_path = Path(tmp) / "output.docx"

        doc = Document()
        doc.add_paragraph("Speaker 2: State your full name for the record.")
        doc.add_paragraph("Speaker 1: Matthew Allan Coger.")
        doc.save(input_path)

        process_transcript(str(input_path), str(output_path), job_config=cfg)

        out_doc = Document(output_path)
        full_text = "\n".join(p.text for p in out_doc.paragraphs)

        assert "MATTHEW ALLAN COGER," in full_text
        assert "having been first duly sworn, testified as follows:" in full_text
        assert "EXAMINATION" in full_text
        assert "BY MR. ALLAN:" in full_text


def test_run_pipeline_merges_reporter_preamble_before_final_output():
    result = run_pipeline(
        {
            "utterances": [
                {"speaker": 4, "transcript": "This is Cause Number 2025CI19595:", "words": []},
                {
                    "speaker": 4,
                    "transcript": "This deposition is being taken in accordance with the Texas Rules of Civil Procedure.",
                    "words": [],
                },
                {
                    "speaker": 4,
                    "transcript": "Counsel, will you please state your agreement for this deposition?",
                    "words": [],
                },
                {"speaker": 2, "transcript": "State your name for the record.", "words": []},
                {"speaker": 1, "transcript": "Matthew Allan Coger.", "words": []},
            ]
        },
        {
            "speaker_map": {
                1: "THE WITNESS",
                2: "MR. ALLAN",
                4: "THE REPORTER",
            },
            "witness_id": 1,
            "examining_attorney_id": 2,
            "speaker_map_verified": True,
            "cause_number": "TEST-PREAMBLE",
        },
    )

    reporter_blocks = [
        block for block in result["blocks"] if (block.speaker_name or "").upper() == "THE REPORTER"
    ]

    assert len(reporter_blocks) == 1
    assert reporter_blocks[0].meta.get("merged_reporter_preamble") is True
    assert result["text"].count("THE REPORTER:") == 1
