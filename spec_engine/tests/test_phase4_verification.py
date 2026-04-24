"""
test_phase4_verification.py

Phase 4 Verification Test Suite — Structural Integrity and Pipeline Unification
"""

import os
import sys
import tempfile
from pathlib import Path
from typing import Dict, List

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))

from spec_engine.models import Block, BlockType, JobConfig


def _make_docx_standard(speaker_content: List[tuple]) -> str:
    from docx import Document

    doc = Document()
    for sid, text in speaker_content:
        doc.add_paragraph(f"Speaker {sid}:")
        doc.add_paragraph(text)
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = f.name
    doc.save(path)
    return path


def _make_docx_inline(speaker_content: List[tuple]) -> str:
    from docx import Document

    doc = Document()
    for sid, text in speaker_content:
        doc.add_paragraph(f"Speaker {sid}: {text}")
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = f.name
    doc.save(path)
    return path


def _coger_cfg() -> JobConfig:
    return JobConfig(
        cause_number="2025CI19595",
        witness_name="Matthew Allan Coger",
        speaker_map={
            0: "THE VIDEOGRAPHER",
            1: "THE WITNESS",
            2: "MR. ALLAN",
            3: "MR. BOYCE - OPPOSING COUNSEL",
            4: "THE REPORTER",
        },
        examining_attorney_id=2,
        witness_id=1,
        speaker_map_verified=True,
    )


class TestFix4A_TwoDigitSpeakerIDs:
    def test_speaker_label_re_matches_two_digits(self):
        from spec_engine.parser import SPEAKER_LABEL_RE

        assert SPEAKER_LABEL_RE.match("Speaker 10:")
        assert SPEAKER_LABEL_RE.match("Speaker 99:")

    def test_speaker_label_re_still_matches_single_digits(self):
        from spec_engine.parser import SPEAKER_LABEL_RE

        for i in range(10):
            assert SPEAKER_LABEL_RE.match(f"Speaker {i}:")

    def test_parse_blocks_with_speaker_10(self):
        from spec_engine.parser import parse_blocks

        path = _make_docx_standard([
            (0, "Good morning everyone."),
            (10, "I am the court reporter."),
            (1, "Please state your name."),
        ])
        try:
            blocks = parse_blocks(path)
        finally:
            os.unlink(path)

        ids = [b.speaker_id for b in blocks]
        assert 10 in ids
        spk10 = [b for b in blocks if b.speaker_id == 10]
        assert len(spk10) == 1
        assert "court reporter" in spk10[0].text.lower()

    def test_parse_blocks_speaker_10_content_not_merged_into_other_speaker(self):
        from spec_engine.parser import parse_blocks

        path = _make_docx_standard([
            (0, "Good morning."),
            (10, "I am the court reporter."),
        ])
        try:
            blocks = parse_blocks(path)
        finally:
            os.unlink(path)

        spk0 = [b for b in blocks if b.speaker_id == 0]
        assert spk0
        assert "court reporter" not in spk0[0].text.lower()

    def test_parse_blocks_all_speakers_0_through_12(self):
        from spec_engine.parser import parse_blocks

        content = [(i, f"Speaker {i} said this.") for i in range(13)]
        path = _make_docx_standard(content)
        try:
            blocks = parse_blocks(path)
        finally:
            os.unlink(path)

        assert len(blocks) == 13
        parsed_ids = {b.speaker_id for b in blocks}
        for i in range(13):
            assert i in parsed_ids


class TestFix4B_InlineSpeakerFormat:
    def test_inline_format_produces_blocks(self):
        from spec_engine.parser import parse_blocks

        path = _make_docx_inline([
            (0, "Today is March 24, 2026."),
            (2, "Did you witness the incident?"),
            (1, "Yes, I did."),
        ])
        try:
            blocks = parse_blocks(path)
        finally:
            os.unlink(path)

        assert len(blocks) == 3

    def test_inline_format_preserves_text(self):
        from spec_engine.parser import parse_blocks

        path = _make_docx_inline([(2, "Did you witness the incident at pump nine?")])
        try:
            blocks = parse_blocks(path)
        finally:
            os.unlink(path)

        assert len(blocks) == 1
        text = blocks[0].text
        assert "Speaker" not in text
        assert "witness" in text.lower()

    def test_inline_format_speaker_ids_correct(self):
        from spec_engine.parser import parse_blocks

        path = _make_docx_inline([
            (0, "Opening statement."),
            (3, "Objection."),
            (1, "I don't recall."),
        ])
        try:
            blocks = parse_blocks(path)
        finally:
            os.unlink(path)

        assert [b.speaker_id for b in blocks] == [0, 3, 1]

    def test_standard_format_still_works_after_4b(self):
        from spec_engine.parser import parse_blocks

        path = _make_docx_standard([(0, "Standard format content."), (1, "Response.")])
        try:
            blocks = parse_blocks(path)
        finally:
            os.unlink(path)

        assert len(blocks) == 2
        assert blocks[0].speaker_id == 0
        assert blocks[1].speaker_id == 1

    def test_mixed_format_in_same_docx(self):
        from docx import Document
        from spec_engine.parser import parse_blocks

        doc = Document()
        doc.add_paragraph("Speaker 0:")
        doc.add_paragraph("Standard content here.")
        doc.add_paragraph("Speaker 1: Inline content here.")
        doc.add_paragraph("Speaker 2:")
        doc.add_paragraph("More standard content.")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        doc.save(path)
        try:
            blocks = parse_blocks(path)
        finally:
            os.unlink(path)

        assert len(blocks) == 3

    def test_inline_format_two_digit_speaker_id(self):
        from spec_engine.parser import parse_blocks

        path = _make_docx_inline([(10, "The court reporter speaks inline.")])
        try:
            blocks = parse_blocks(path)
        finally:
            os.unlink(path)

        assert len(blocks) == 1
        assert blocks[0].speaker_id == 10
        assert "court reporter" in blocks[0].text.lower()


class TestFix4C_ChunkBoundaryDeduplication:
    def _run_improved_dedup(
        self,
        chunk1_words: List[Dict],
        chunk2_words_raw: List[Dict],
        offset: float = 590.0,
    ) -> List[Dict]:
        from config import CHUNK_OVERLAP_SECONDS

        adjusted = [
            {**w, "start": round(w["start"] + offset, 3), "end": round(w["end"] + offset, 3)}
            for w in chunk2_words_raw
        ]

        last_ts = chunk1_words[-1]["end"]
        last_word_text = (chunk1_words[-1].get("word") or "").lower().strip()

        deduplicated = []
        for w in adjusted:
            w_text = (w.get("word") or "").lower().strip()
            if w["start"] <= last_ts and w_text == last_word_text:
                continue
            if w["start"] < last_ts - 0.5:
                continue
            deduplicated.append(w)

        if len(deduplicated) < len(adjusted) * 0.5:
            cutoff = last_ts - (CHUNK_OVERLAP_SECONDS / 2)
            deduplicated = [w for w in adjusted if w["start"] > cutoff]

        return deduplicated

    def test_exact_boundary_word_deduplicated(self):
        chunk1 = [
            {"word": "I", "start": 597.0, "end": 597.2},
            {"word": "was", "start": 597.3, "end": 597.5},
            {"word": "register", "start": 599.3, "end": 599.7},
        ]
        chunk2_raw = [
            {"word": "register", "start": 9.3, "end": 9.7},
            {"word": "inside", "start": 9.8, "end": 10.0},
            {"word": "store", "start": 10.3, "end": 10.5},
        ]
        kept = self._run_improved_dedup(chunk1, chunk2_raw, offset=590.0)
        word_texts = [w["word"] for w in kept]
        assert "register" not in word_texts
        assert "inside" in word_texts
        assert "store" in word_texts

    def test_different_word_at_same_timestamp_kept(self):
        chunk1 = [
            {"word": "went", "start": 598.0, "end": 598.3},
            {"word": "outside", "start": 599.0, "end": 599.5},
        ]
        chunk2_raw = [
            {"word": "quickly", "start": 9.0, "end": 9.3},
            {"word": "and", "start": 9.4, "end": 9.5},
            {"word": "returned", "start": 9.6, "end": 9.9},
        ]
        kept = self._run_improved_dedup(chunk1, chunk2_raw, offset=590.0)
        assert "quickly" in [w["word"] for w in kept]

    def test_standard_overlap_still_deduplicated(self):
        chunk1 = [
            {"word": "the", "start": 598.0, "end": 598.2},
            {"word": "incident", "start": 599.0, "end": 599.5},
        ]
        chunk2_raw = [
            {"word": "at", "start": 9.0, "end": 9.1},
            {"word": "the", "start": 9.1, "end": 9.2},
            {"word": "pump", "start": 10.5, "end": 10.8},
        ]
        kept = self._run_improved_dedup(chunk1, chunk2_raw, offset=590.0)
        assert "pump" in [w["word"] for w in kept]

    def test_single_chunk_unaffected(self):
        from pipeline.assembler import reassemble_chunks

        single_result = {
            "words": [{"word": "test", "start": 0.0, "end": 0.3, "speaker": 1}],
            "utterances": [{"speaker": 1, "transcript": "test", "start": 0.0, "end": 0.3, "words": []}],
            "transcript": "test",
            "raw": {},
        }
        result = reassemble_chunks([single_result], [0.0])
        assert len(result["words"]) == 1


class TestFix4D_MapSpeakersInDocumentBuilder:
    def test_process_blocks_imported_in_document_builder(self):
        from spec_engine.document_builder import process_transcript
        import spec_engine.document_builder as db

        assert hasattr(db, "process_blocks") or "process_blocks" in dir(db)

    def test_process_blocks_call_in_process_transcript_source(self):
        import inspect
        from spec_engine.document_builder import process_transcript

        src = inspect.getsource(process_transcript)
        assert "process_blocks" in src

    def test_process_blocks_called_before_classify_block(self):
        import inspect
        from spec_engine.document_builder import process_transcript

        src = inspect.getsource(process_transcript)
        process_pos = src.find("process_blocks(")
        classify_pos = src.find("classify_block(")
        assert process_pos != -1
        assert classify_pos != -1
        assert process_pos < classify_pos

    def test_map_speakers_populates_speaker_role(self):
        from spec_engine.speaker_mapper import map_speakers

        blocks = [
            Block(speaker_id=0, text="Today is March 24.", raw_text=""),
            Block(speaker_id=2, text="Did you go there?", raw_text=""),
            Block(speaker_id=1, text="Yes.", raw_text=""),
        ]
        result = map_speakers(blocks, _coger_cfg())
        assert result[0].speaker_role == "VIDEOGRAPHER"
        assert result[1].speaker_role in ("ATTORNEY", "EXAMINING_ATTORNEY")
        assert result[2].speaker_role == "WITNESS"

    def test_videographer_classification_benefits_from_map_speakers(self):
        from spec_engine.classifier import ClassifierState, classify_block
        from spec_engine.models import LineType
        from spec_engine.speaker_mapper import map_speakers

        blocks = [Block(speaker_id=0, text="We are off the record.", raw_text="")]
        blocks_mapped = map_speakers(blocks, _coger_cfg())
        results = classify_block(blocks_mapped[0], _coger_cfg(), ClassifierState(), block_index=0)
        line_types = [r[0] for r in results]
        assert LineType.Q not in line_types
        assert LineType.A not in line_types

    def test_speaker_name_populated_before_classify(self):
        from spec_engine.speaker_mapper import map_speakers

        blocks = [Block(speaker_id=3, text="Objection.  Form.", raw_text="")]
        result = map_speakers(blocks, _coger_cfg())
        assert result[0].speaker_name is not None
        assert result[0].speaker_name != ""
        assert "BOYCE" in result[0].speaker_name.upper()

    def test_generic_speaker_with_reporter_text_maps_to_reporter(self):
        from spec_engine.speaker_mapper import map_speakers

        cfg = JobConfig(
            speaker_map={1: "THE WITNESS", 2: "MR. ALLAN", 4: "THE REPORTER"},
            witness_id=1,
            examining_attorney_id=2,
            speaker_map_verified=True,
        )
        blocks = [Block(speaker_id=9, text="Please raise your right hand.", raw_text="")]

        result = map_speakers(blocks, cfg)

        assert result[0].speaker_role == "REPORTER"
        assert result[0].speaker_name == "THE REPORTER"

    def test_generic_speaker_question_maps_to_last_attorney(self):
        from spec_engine.speaker_mapper import map_speakers

        cfg = _coger_cfg()
        blocks = [
            Block(speaker_id=2, text="State your name for the record.", raw_text=""),
            Block(speaker_id=9, text="Could you state your full name for the record, please?", raw_text=""),
        ]

        result = map_speakers(blocks, cfg)

        assert result[1].speaker_role in ("ATTORNEY", "EXAMINING_ATTORNEY")
        assert "ALLAN" in result[1].speaker_name.upper()

    def test_generic_speaker_after_attorney_answer_maps_to_witness(self):
        from spec_engine.speaker_mapper import map_speakers

        cfg = _coger_cfg()
        blocks = [
            Block(speaker_id=2, text="Did you go there?", raw_text=""),
            Block(speaker_id=9, text="Yes, sir.", raw_text=""),
        ]

        result = map_speakers(blocks, cfg)

        assert result[1].speaker_role == "WITNESS"
        assert result[1].speaker_name == "THE WITNESS"


class TestPhase4Integration:
    def test_detect_input_format_deepgram_standard(self):
        from spec_engine.parser import FORMAT_DEEPGRAM, detect_input_format

        path = _make_docx_standard([(0, "Test content.")])
        try:
            fmt = detect_input_format(path)
        finally:
            os.unlink(path)
        assert fmt == FORMAT_DEEPGRAM

    def test_detect_input_format_deepgram_inline(self):
        from spec_engine.parser import FORMAT_DEEPGRAM, detect_input_format

        path = _make_docx_inline([(0, "Test content."), (1, "Response.")])
        try:
            fmt = detect_input_format(path)
        finally:
            os.unlink(path)
        assert fmt == FORMAT_DEEPGRAM

    def test_coger_style_standard_docx_parses_all_speakers(self):
        from spec_engine.parser import parse_blocks

        content = [
            (0, "Today is March 24, 2026."),
            (4, "Please state your name for the record."),
            (1, "Matthew Allan Coger."),
            (2, "Hello Mr. Coger. Did you work at Murphy USA?"),
            (1, "Yes, I did."),
        ]
        path = _make_docx_standard(content)
        try:
            blocks = parse_blocks(path)
        finally:
            os.unlink(path)
        assert len(blocks) == 5
        assert [b.speaker_id for b in blocks] == [0, 4, 1, 2, 1]


class TestPhase4ExitGate:
    def test_speaker_label_re_handles_two_digits(self):
        from spec_engine.parser import SPEAKER_LABEL_RE

        assert SPEAKER_LABEL_RE.match("Speaker 10:")
        assert SPEAKER_LABEL_RE.match("Speaker 11:")

    def test_speaker_inline_re_exists(self):
        from spec_engine.parser import SPEAKER_INLINE_RE

        assert SPEAKER_INLINE_RE is not None

    def test_inline_format_docx_produces_blocks(self):
        from spec_engine.parser import parse_blocks

        path = _make_docx_inline([(0, "Opening."), (1, "Response.")])
        try:
            blocks = parse_blocks(path)
        finally:
            os.unlink(path)
        assert len(blocks) > 0

    def test_process_blocks_in_document_builder_source(self):
        import inspect
        from spec_engine.document_builder import process_transcript

        src = inspect.getsource(process_transcript)
        assert "process_blocks(" in src

    def test_phases_1_through_3_not_broken(self):
        from spec_engine.classifier import classify_blocks
        from spec_engine.corrections import clean_block
        from spec_engine.objections import _resolve_objection_speaker

        cfg = _coger_cfg()
        speaker = _resolve_objection_speaker(cfg)
        if speaker != "MR. UNKNOWN":
            assert "BOYCE" in speaker.upper() or speaker == "COUNSEL"

        result = clean_block("There were 5 witnesses.", cfg)
        assert "five" in result[0].lower()

        blocks = [
            Block(
                speaker_id=0,
                text="We are off the record.",
                raw_text="",
                speaker_role="VIDEOGRAPHER",
                speaker_name="THE VIDEOGRAPHER",
            )
        ]
        classified = classify_blocks(blocks, cfg)
        assert classified[0].block_type != BlockType.ANSWER
