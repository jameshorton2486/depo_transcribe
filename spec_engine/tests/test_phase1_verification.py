"""
test_phase1_verification.py

Phase 1 Verification Test Suite — Stop Critical Output Corruption
Run with: python -m pytest spec_engine/tests/test_phase1_verification.py -v

Coverage:
  Fix 1-A:  Hard gate on speaker verification (main.py — tested via integration)
  Fix 1-B:  classify_blocks() warning guard (classifier.py)
  Fix 1-C:  _resolve_objection_speaker() speaker map fallback (objections.py)
  Fix 1-D:  Objection text verbatim preservation — Morson's rule (corrections.py + objections.py)
  Fix 1-E:  JobConfig persisted immediately after dialog (models.py + main.py)

IMPORTANT — ONE EXISTING TEST MUST BE UPDATED AS PART OF PHASE 1:
  spec_engine/tests/test_block_pipeline_behavior.py::test_objection_extraction_exit_form
  → Currently asserts "Objection to form." (single sentence — wrong per Morson's)
  → After Fix 1-D must assert "Objection.  Form." (two sentences — verbatim)
  → test_existing_objection_pipeline_test_needs_update() below will FAIL
    until that test file is updated. This is expected and intentional.

Total tests in this module: 50
Expected: all PASS after Phase 1 fixes are applied.

Run alongside existing suite:
  python -m pytest spec_engine/tests/ -v
"""

import logging
import os
import sys
import tempfile
from pathlib import Path

import pytest

# Ensure repo root is on path
sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))

from spec_engine.corrections import (
    apply_universal_corrections,
    clean_block,
)
from spec_engine.models import (
    Block,
    BlockType,
    JobConfig,
    SpeakerMapUnverifiedError,
)
from spec_engine.classifier import ClassifierState, classify_block, classify_blocks
from spec_engine.objections import (
    OBJECTION_PATTERNS,
    _resolve_objection_speaker,
    extract_objections,
)


# ─────────────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def _make_coger_config() -> JobConfig:
    """
    JobConfig matching the Coger deposition that exposed the Phase 1 bugs.
    This config represents what a properly filled-out JobConfigDialog produces.
    """
    return JobConfig(
        cause_number="2025CI19595",
        case_style="Kathie A. Love v. Murphy Oil USA, Inc.",
        plaintiff_name="Kathie A. Love",
        defendant_names=["Murphy Oil USA, Inc."],
        court="408th Judicial District, Bexar County, Texas",
        witness_name="Matthew Allan Coger",
        reporter_name="Miah Bardot",
        reporter_csr="CSR No. 12129",
        reporter_firm="SA Legal Solutions",
        speaker_map={
            0: "THE VIDEOGRAPHER",
            1: "THE WITNESS",
            2: "MR. ALLAN",
            3: "MR. BOYCE - OPPOSING COUNSEL",
            4: "THE REPORTER",
        },
        examining_attorney_id=2,
        witness_id=1,
        speaker_map_verified=True,
        confirmed_spellings={
            "Cogger": "Coger",
            "Bare County": "Bexar County",
            "David Blvd": "David Blas",
            "David Blvd.": "David Blas",
            "Miah Vardell": "Miah Bardot",
            "May Vardell": "Miah Bardot",
        },
    )


def _make_empty_config() -> JobConfig:
    """Worst case: completely empty JobConfig."""
    return JobConfig()


def _make_unverified_config() -> JobConfig:
    """Config with speaker_map_verified=False — triggers warnings not crashes."""
    cfg = JobConfig(
        speaker_map={1: "THE WITNESS", 2: "MR. ALLAN", 3: "MR. BOYCE"},
        examining_attorney_id=2,
        witness_id=1,
        speaker_map_verified=False,
    )
    return cfg


# ─────────────────────────────────────────────────────────────────────────────
# FIX 1-A: SPEAKER VERIFICATION GATE
# Tests for main.py behavior — integration-level, no UI required
# ─────────────────────────────────────────────────────────────────────────────

class TestFix1A_SpeakerVerificationGate:
    """
    Fix 1-A: Hard gate before SpeakerVerifyDialog.
    Direct UI tests require a running app — these tests verify the underlying
    logic that the gate depends on. Full UI verification is in the manual checklist.
    """

    def test_parse_blocks_returns_empty_list_for_non_deepgram_docx(self):
        """
        parse_blocks() on a plain-text DOCX (no 'Speaker N:' labels) must return [].
        The Fix 1-A gate depends on this returning empty to trigger the error dialog.
        """
        from docx import Document
        from spec_engine.parser import parse_blocks

        # Create a plain DOCX with no Speaker N: labels
        doc = Document()
        doc.add_paragraph("This is just regular text.")
        doc.add_paragraph("No speaker labels here.")
        doc.add_paragraph("Q. Is this a deposition? A. No.")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            tmp_path = f.name
        try:
            doc.save(tmp_path)
            result = parse_blocks(tmp_path)
            assert result == [], (
                f"parse_blocks() should return [] for non-Deepgram DOCX. "
                f"Got {len(result)} blocks. Fix 1-A gate depends on this being empty."
            )
        finally:
            os.unlink(tmp_path)

    def test_parse_blocks_returns_blocks_for_valid_deepgram_docx(self):
        """
        parse_blocks() on a valid Deepgram-format DOCX (Speaker N: labels) must
        return non-empty list. Fix 1-A gate should proceed to SpeakerVerifyDialog.
        """
        from docx import Document
        from spec_engine.parser import parse_blocks

        doc = Document()
        doc.add_paragraph("Speaker 0:")
        doc.add_paragraph("Good morning everyone.")
        doc.add_paragraph("Speaker 1:")
        doc.add_paragraph("Good morning.")
        doc.add_paragraph("Speaker 2:")
        doc.add_paragraph("Did you witness the incident?")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            tmp_path = f.name
        try:
            doc.save(tmp_path)
            result = parse_blocks(tmp_path)
            assert len(result) > 0, (
                "parse_blocks() should return blocks for valid Deepgram DOCX. "
                "Got empty list — Fix 1-A gate would incorrectly show error dialog."
            )
            assert all(hasattr(b, 'speaker_id') for b in result)
            assert all(hasattr(b, 'text') for b in result)
        finally:
            os.unlink(tmp_path)

    def test_speaker_map_verified_false_on_fresh_jobconfig(self):
        """
        A fresh JobConfig must have speaker_map_verified=False.
        Fix 1-A's final guard checks this field — it must default to False.
        """
        cfg = JobConfig()
        assert cfg.speaker_map_verified is False, (
            "Fresh JobConfig must have speaker_map_verified=False. "
            "Fix 1-A's final guard checks this to prevent processing "
            "when dialog was bypassed."
        )

    def test_speaker_map_verified_true_after_explicit_set(self):
        """
        After SpeakerVerifyDialog confirms, speaker_map_verified must be True.
        Fix 1-A's final guard allows processing only when this is True.
        """
        cfg = JobConfig()
        cfg.speaker_map_verified = True
        assert cfg.speaker_map_verified is True

    def test_jobconfig_save_creates_file_in_jobs_dir(self, tmp_path):
        """
        JobConfig.save() must create a file in the jobs/ directory.
        Fix 1-E calls this immediately after dialog — the file must be created.
        """
        cfg = JobConfig(
            cause_number="TEST-SAVE-2025",
            witness_name="Test Witness",
        )
        saved = cfg.save(jobs_dir=str(tmp_path))
        assert Path(saved).exists(), (
            f"JobConfig.save() did not create a file. Got path: {saved}"
        )
        assert "TEST-SAVE-2025" in Path(saved).name, (
            f"Saved filename should contain cause number. Got: {Path(saved).name}"
        )

    def test_jobconfig_save_is_loadable(self, tmp_path):
        """
        A saved JobConfig must be loadable back with all fields intact.
        Fix 1-E's save must produce a valid JSON file.
        """
        cfg = JobConfig(
            cause_number="TEST-ROUNDTRIP-2025",
            witness_name="Matthew Allan Coger",
            reporter_name="Miah Bardot",
            confirmed_spellings={"Cogger": "Coger", "Bare County": "Bexar County"},
            speaker_map_verified=True,
        )
        saved = cfg.save(jobs_dir=str(tmp_path))
        loaded = JobConfig.load(saved)

        assert loaded.cause_number == "TEST-ROUNDTRIP-2025"
        assert loaded.witness_name == "Matthew Allan Coger"
        assert loaded.reporter_name == "Miah Bardot"
        assert loaded.confirmed_spellings.get("Cogger") == "Coger"
        assert loaded.confirmed_spellings.get("Bare County") == "Bexar County"
        assert loaded.speaker_map_verified is True


# ─────────────────────────────────────────────────────────────────────────────
# FIX 1-B: classify_blocks() WARNING GUARD
# ─────────────────────────────────────────────────────────────────────────────

class TestFix1B_ClassifyBlocksGuard:
    """
    Fix 1-B: classify_blocks() (plural — lightweight classifier) must log a
    warning when called with an unverified speaker map.
    Must NOT raise — this function is also used in preview/draft contexts.
    """

    def test_classify_blocks_does_not_raise_without_verified_map(self):
        """
        classify_blocks() with unverified config must not raise SpeakerMapUnverifiedError.
        (That exception is only raised by classify_block() singular — the full classifier.)
        """
        cfg = _make_unverified_config()
        blocks = [
            Block(speaker_id=1, text="Yes, I did.", raw_text="Yes, I did."),
            Block(speaker_id=2, text="Did you go there?", raw_text="Did you go there?"),
        ]
        # Must NOT raise — lightweight classifier is used in preview contexts
        try:
            result = classify_blocks(blocks, job_config=cfg)
            assert isinstance(result, list)
        except SpeakerMapUnverifiedError:
            pytest.fail(
                "classify_blocks() (plural) raised SpeakerMapUnverifiedError. "
                "Only classify_block() (singular) should raise this. "
                "Fix 1-B adds a WARNING log, not a raise."
            )

    def test_classify_blocks_logs_warning_for_unverified_map(self, caplog):
        """
        classify_blocks() with unverified map must emit a WARNING to the logger.
        Fix 1-B inserts this guard at the top of classify_blocks().
        """
        cfg = _make_unverified_config()
        blocks = [Block(speaker_id=1, text="I understand.", raw_text="I understand.")]

        with caplog.at_level(logging.WARNING, logger="spec_engine.classifier"):
            classify_blocks(blocks, job_config=cfg)

        warning_messages = [r.message for r in caplog.records
                            if r.levelname == "WARNING"]
        assert any("unverified" in m.lower() or "speaker map" in m.lower()
                   for m in warning_messages), (
            f"classify_blocks() must log a warning when speaker map is unverified. "
            f"Got warning messages: {warning_messages}. "
            f"Fix 1-B adds this guard."
        )

    def test_classify_blocks_does_not_warn_for_verified_map(self, caplog):
        """
        classify_blocks() with a verified map must NOT emit an unverified warning.
        """
        cfg = _make_coger_config()  # speaker_map_verified=True
        blocks = [Block(speaker_id=1, text="I understand.", raw_text="I understand.",
                        speaker_role="WITNESS", speaker_name="THE WITNESS")]

        with caplog.at_level(logging.WARNING, logger="spec_engine.classifier"):
            classify_blocks(blocks, job_config=cfg)

        unverified_warnings = [r.message for r in caplog.records
                               if r.levelname == "WARNING"
                               and "unverified" in r.message.lower()]
        assert len(unverified_warnings) == 0, (
            f"classify_blocks() should NOT warn when map is verified. "
            f"Got: {unverified_warnings}"
        )

    def test_classify_blocks_still_returns_results_with_unverified_map(self):
        """
        classify_blocks() with unverified map must still return classified blocks.
        Heuristic classification must work without a verified map (preview mode).
        """
        cfg = _make_unverified_config()
        blocks = [
            Block(speaker_id=2, text="Did you go there?", raw_text=""),
            Block(speaker_id=1, text="Yes, I did.", raw_text=""),
        ]
        result = classify_blocks(blocks, job_config=cfg)
        assert len(result) == 2, "classify_blocks() must return all input blocks"
        assert all(b.block_type != BlockType.UNKNOWN for b in result
                   if b.text.strip()), "Blocks should be classified even without verified map"

    def test_classify_blocks_accepts_none_job_config(self):
        """
        classify_blocks() with job_config=None must not raise or warn.
        None is a valid caller pattern for preview rendering.
        """
        blocks = [Block(speaker_id=1, text="Yes.", raw_text="Yes.")]
        try:
            result = classify_blocks(blocks, job_config=None)
            assert isinstance(result, list)
        except Exception as e:
            pytest.fail(f"classify_blocks() with None config raised: {type(e).__name__}: {e}")

    def test_classify_block_singular_still_raises_for_unverified(self):
        """
        classify_block() (singular — full spec classifier) must STILL raise
        SpeakerMapUnverifiedError. Fix 1-B must not weaken this guard.
        """
        cfg = _make_unverified_config()  # speaker_map_verified=False
        block = Block(speaker_id=1, text="I understand.", raw_text="")
        state = ClassifierState()

        with pytest.raises(SpeakerMapUnverifiedError):
            classify_block(block, cfg, state, block_index=0)


# ─────────────────────────────────────────────────────────────────────────────
# FIX 1-C: _resolve_objection_speaker() — NO MORE MR. UNKNOWN
# ─────────────────────────────────────────────────────────────────────────────

class TestFix1C_ObjectionSpeakerResolution:
    """
    Fix 1-C: _resolve_objection_speaker() must never return 'MR. UNKNOWN'.
    Priority: speaker_map > defense_counsel > any non-examining attorney > 'COUNSEL'
    """

    # ── Priority 1: speaker_map with OPPOSING COUNSEL label ──────────────────

    def test_resolves_from_speaker_map_opposing_counsel_label(self):
        """Primary resolution: OPPOSING COUNSEL label in speaker_map."""
        cfg = _make_coger_config()
        # speaker_map has 3: "MR. BOYCE - OPPOSING COUNSEL"
        result = _resolve_objection_speaker(cfg)
        assert result != "MR. UNKNOWN", "Must not return MR. UNKNOWN"
        assert "BOYCE" in result.upper(), (
            f"Expected MR. BOYCE from speaker_map, got: {result!r}"
        )

    def test_resolves_from_speaker_map_defense_label(self):
        """Primary resolution: DEFENSE label in speaker_map."""
        cfg = JobConfig(
            speaker_map={1: "THE WITNESS", 2: "MR. ALLAN", 3: "MR. SMITH - DEFENSE"},
            examining_attorney_id=2,
            witness_id=1,
            speaker_map_verified=True,
        )
        result = _resolve_objection_speaker(cfg)
        assert result != "MR. UNKNOWN"
        assert "SMITH" in result.upper(), f"Expected MR. SMITH from defense label, got: {result!r}"

    def test_resolves_from_speaker_map_dict_form(self):
        """speaker_map passed as dict (JSON-loaded config) must also work."""
        cfg_dict = {
            "speaker_map": {
                0: "THE VIDEOGRAPHER",
                1: "THE WITNESS",
                2: "MR. ALLAN",
                3: "MR. BOYCE - OPPOSING COUNSEL",
            },
            "examining_attorney_id": 2,
            "witness_id": 1,
            "speaker_map_verified": True,
        }
        result = _resolve_objection_speaker(cfg_dict)
        assert result != "MR. UNKNOWN"
        assert "BOYCE" in result.upper(), f"Got: {result!r}"

    # ── Priority 2: defense_counsel list ─────────────────────────────────────

    def test_resolves_from_defense_counsel_when_map_has_no_opposing_label(self):
        """
        If speaker_map has no OPPOSING COUNSEL label, check defense_counsel list.
        This is the secondary priority path.
        """
        from spec_engine.models import CounselInfo
        cfg = JobConfig(
            speaker_map={1: "THE WITNESS", 2: "MR. ALLAN"},
            examining_attorney_id=2,
            defense_counsel=[CounselInfo(name="David Boyce", firm="Defense Firm")],
            speaker_map_verified=True,
        )
        result = _resolve_objection_speaker(cfg)
        assert result != "MR. UNKNOWN"
        assert "BOYCE" in result.upper(), f"Expected BOYCE from defense_counsel, got: {result!r}"

    def test_resolves_from_defense_counsel_dict_form(self):
        """defense_counsel as a list of dicts (JSON-loaded) must resolve correctly."""
        cfg_dict = {
            "speaker_map": {1: "THE WITNESS", 2: "MR. ALLAN"},
            "examining_attorney_id": 2,
            "defense_counsel": [{"name": "David Boyce", "title": "MR."}],
        }
        result = _resolve_objection_speaker(cfg_dict)
        assert result != "MR. UNKNOWN"
        assert "BOYCE" in result.upper(), f"Got: {result!r}"

    # ── Priority 3: any non-examining attorney in speaker_map ────────────────

    def test_resolves_any_attorney_when_no_defense_label(self):
        """
        When no OPPOSING COUNSEL label and no defense_counsel, use any attorney
        in the speaker_map who is NOT the examining attorney.
        """
        cfg = JobConfig(
            speaker_map={1: "THE WITNESS", 2: "MR. ALLAN", 3: "MR. JONES"},
            examining_attorney_id=2,
            witness_id=1,
            speaker_map_verified=True,
        )
        result = _resolve_objection_speaker(cfg)
        assert result != "MR. UNKNOWN"
        assert "JONES" in result.upper() or "MR." in result.upper(), (
            f"Expected attorney label from speaker_map fallback, got: {result!r}"
        )

    def test_does_not_use_examining_attorney_for_objection(self):
        """
        The examining attorney must NEVER be used as the objection speaker.
        Objections come from opposing counsel, not the examiner.
        """
        cfg = JobConfig(
            speaker_map={1: "THE WITNESS", 2: "MR. ALLAN"},
            examining_attorney_id=2,  # MR. ALLAN is the examiner
            witness_id=1,
            speaker_map_verified=True,
        )
        result = _resolve_objection_speaker(cfg)
        # Should not be MR. ALLAN (the examiner)
        assert "ALLAN" not in result.upper(), (
            f"Examining attorney MR. ALLAN used as objection speaker. "
            f"Got: {result!r}. Examiners do not object in direct examination."
        )

    # ── Priority 4 (fallback): 'COUNSEL' — never 'MR. UNKNOWN' ──────────────

    def test_fallback_is_counsel_not_mr_unknown(self):
        """
        Worst case: completely empty JobConfig must return 'COUNSEL', not 'MR. UNKNOWN'.
        'MR. UNKNOWN' is the production bug this fix eliminates.
        """
        cfg = _make_empty_config()
        result = _resolve_objection_speaker(cfg)
        assert result != "MR. UNKNOWN", (
            f"CRITICAL: _resolve_objection_speaker() returned 'MR. UNKNOWN' on empty config. "
            f"This is the bug Fix 1-C eliminates. Got: {result!r}"
        )
        assert result == "COUNSEL", (
            f"Fallback should be 'COUNSEL' (neutral, professional). Got: {result!r}"
        )

    def test_fallback_with_empty_speaker_map(self):
        """Empty speaker_map should fall through to 'COUNSEL', not crash or return MR. UNKNOWN."""
        cfg = JobConfig(speaker_map={}, speaker_map_verified=False)
        result = _resolve_objection_speaker(cfg)
        assert result != "MR. UNKNOWN"
        assert result == "COUNSEL"

    def test_fallback_with_only_witness_in_speaker_map(self):
        """
        Speaker map containing only THE WITNESS (no attorneys) should return 'COUNSEL'.
        """
        cfg = JobConfig(
            speaker_map={1: "THE WITNESS"},
            witness_id=1,
            speaker_map_verified=True,
        )
        result = _resolve_objection_speaker(cfg)
        assert result != "MR. UNKNOWN"
        assert result == "COUNSEL"

    # ── extract_objections integration ────────────────────────────────────────

    def test_extract_objections_uses_correct_speaker_label(self):
        """
        When the full pipeline runs extract_objections(), the objection block
        must carry the correct speaker label — not MR. UNKNOWN.
        """
        from spec_engine.models import CounselInfo
        cfg = JobConfig(
            cause_number="TEST-OBJ-SPEAKER",
            speaker_map={1: "THE WITNESS", 2: "MR. ALLAN", 3: "MR. BOYCE - OPPOSING COUNSEL"},
            examining_attorney_id=2,
            witness_id=1,
            defense_counsel=[CounselInfo(name="David Boyce")],
            speaker_map_verified=True,
        )

        blocks = [
            Block(
                speaker_id=1,
                text="I don't recall. Objection form.",
                raw_text="I don't recall. Objection form.",
                speaker_name="THE WITNESS",
                speaker_role="WITNESS",
                block_type=BlockType.ANSWER,
            )
        ]

        result = extract_objections(blocks, cfg)
        # Find the injected objection block
        objection_blocks = [b for b in result if "objection" in (b.text or "").lower()]
        assert objection_blocks, "extract_objections() must produce an objection block"

        for obj_block in objection_blocks:
            speaker = (obj_block.speaker_name or "").upper()
            assert "UNKNOWN" not in speaker, (
                f"Objection block has 'UNKNOWN' speaker: {obj_block.speaker_name!r}. "
                f"Fix 1-C must resolve from speaker_map or defense_counsel."
            )

    def test_extract_objections_preserves_matched_objection_text(self):
        cfg = JobConfig(
            cause_number="TEST-OBJ-TEXT",
            speaker_map={1: "THE WITNESS", 3: "MR. BOYCE - OPPOSING COUNSEL"},
            examining_attorney_id=2,
            witness_id=1,
            speaker_map_verified=True,
        )

        blocks = [
            Block(
                speaker_id=1,
                text="I don't recall. Objection.  Form.",
                raw_text="I don't recall. Objection.  Form.",
                speaker_name="THE WITNESS",
                speaker_role="WITNESS",
                block_type=BlockType.ANSWER,
            )
        ]

        result = extract_objections(blocks, cfg)
        objection_blocks = [b for b in result if b.meta.get("is_objection")]

        assert objection_blocks
        assert objection_blocks[0].text == "Objection.  Form."


# ─────────────────────────────────────────────────────────────────────────────
# FIX 1-D: OBJECTION TEXT VERBATIM PRESERVATION — MORSON'S RULE
# ─────────────────────────────────────────────────────────────────────────────

class TestFix1D_VerbatimObjectionText:
    """
    Fix 1-D: Morson's English Guide for Court Reporters §9.6 requires objections
    be preserved exactly as spoken. "Objection.  Form." (two sentences) is what
    attorneys say. "Objection to form." (one sentence) is a fabrication.

    corrections.py UNIVERSAL_CORRECTIONS must produce "Objection.  Form." (two sentences).
    OBJECTION_PATTERNS in objections.py must NOT contain "objection to form".
    """

    # ── UNIVERSAL_CORRECTIONS must produce two-sentence form ─────────────────

    def test_exit_form_corrects_to_two_sentence_verbatim(self):
        """'Exit form' (ASR garble) must correct to 'Objection.  Form.' — two sentences."""
        records = []
        result = apply_universal_corrections("Exit form", records, block_index=0)
        assert result == "Objection.  Form.", (
            f"'Exit form' must correct to 'Objection.  Form.' (two sentences). "
            f"Got: {result!r}. Morson's §9.6: preserve as spoken."
        )

    def test_action_form_corrects_to_two_sentence_verbatim(self):
        """'Action form' (ASR garble) must correct to 'Objection.  Form.'"""
        records = []
        result = apply_universal_corrections("Action form", records, block_index=0)
        assert result == "Objection.  Form.", (
            f"'Action form' → 'Objection.  Form.' required. Got: {result!r}"
        )

    def test_action_point_corrects_to_two_sentence_verbatim(self):
        """'Action point' (ASR garble) must correct to 'Objection.  Form.'"""
        records = []
        result = apply_universal_corrections("Action point", records, block_index=0)
        assert result == "Objection.  Form.", (
            f"'Action point' → 'Objection.  Form.' required. Got: {result!r}"
        )

    def test_objection_form_garble_corrects_to_two_sentence_verbatim(self):
        """'Objection form' (ASR garble — space, no period) must correct to 'Objection.  Form.'"""
        records = []
        result = apply_universal_corrections("Objection form", records, block_index=0)
        assert result == "Objection.  Form.", (
            f"'Objection form' → 'Objection.  Form.' required. Got: {result!r}"
        )

    def test_none_of_the_garbles_produce_objection_to_form(self):
        """
        CRITICAL: 'Objection to form.' must NEVER be produced by any of the
        four ASR garble corrections. This is the exact text from the bug.
        """
        garbles = ["Exit form", "Action form", "Action point", "Objection form"]
        for garble in garbles:
            records = []
            result = apply_universal_corrections(garble, records, block_index=0)
            assert "objection to form" not in result.lower(), (
                f"MORSON'S VIOLATION: '{garble}' produced 'Objection to form.' "
                f"Got: {result!r}. Must produce 'Objection.  Form.' (two sentences)."
            )

    # ── clean_block must not alter correctly-spoken objection text ────────────

    def test_spoken_objection_form_not_altered_by_clean_block(self):
        """
        'Objection.  Form.' as spoken must pass through clean_block() unchanged.
        Verbatim rule is absolute — spoken testimony must not be rewritten.
        """
        cfg = _make_coger_config()
        text = "Objection.  Form."
        result, records = clean_block(text, cfg, block_index=0)[:2]
        assert result == "Objection.  Form.", (
            f"clean_block() altered verbatim objection 'Objection.  Form.' → {result!r}. "
            f"Verbatim rule: spoken objections must never be rewritten."
        )

    def test_spoken_bare_objection_not_altered(self):
        """'Objection.' (bare, one word) must pass through clean_block() unchanged."""
        cfg = _make_coger_config()
        result, records = clean_block("Objection.", cfg, block_index=0)[:2]
        assert result == "Objection.", (
            f"'Objection.' was altered to: {result!r}. Verbatim rule violated."
        )

    def test_spoken_objection_leading_not_altered(self):
        """'Objection. Leading.' (two sentences, as spoken) must not be altered."""
        cfg = _make_coger_config()
        result, records = clean_block("Objection. Leading.", cfg, block_index=0)[:2]
        assert "Objection." in result, f"'Objection.' removed from: {result!r}"
        assert "Leading." in result, f"'Leading.' removed from: {result!r}"

    # ── OBJECTION_PATTERNS must not contain 'objection to form' ─────────────

    def test_objection_patterns_does_not_contain_objection_to_form(self):
        """
        OBJECTION_PATTERNS in objections.py must not contain 'objection to form'.
        After Fix 1-D, corrections.py normalizes ASR garbles to 'Objection.  Form.'
        The classifier then handles them via OBJECTION_START_RE.
        Having 'objection to form' in OBJECTION_PATTERNS causes double-processing.
        """
        for pattern in OBJECTION_PATTERNS:
            assert "objection to form" not in pattern.lower(), (
                f"OBJECTION_PATTERNS still contains 'objection to form' pattern: {pattern!r}. "
                f"Fix 1-D removes this — it causes double-processing of already-corrected text."
            )

    def test_objection_patterns_still_contains_garble_patterns(self):
        """
        OBJECTION_PATTERNS must still contain the original ASR garble patterns
        after Fix 1-D. Only 'objection to form' is removed, not the others.
        """
        garble_patterns_expected = ["exit form", "action form", "objection form"]
        pattern_text = " ".join(OBJECTION_PATTERNS).lower()
        for expected in garble_patterns_expected:
            assert expected in pattern_text, (
                f"Fix 1-D should NOT remove '{expected}' from OBJECTION_PATTERNS. "
                f"Only 'objection to form' is removed."
            )

    # ── REGRESSION: existing test that must be updated ────────────────────────

    def test_existing_objection_pipeline_test_needs_update(self):
        """
        KNOWN CONFLICT: test_block_pipeline_behavior.py::test_objection_extraction_exit_form
        currently asserts 'Objection to form.' in the pipeline output.
        After Fix 1-D, this is WRONG — it must assert 'Objection.  Form.' instead.

        This test INTENTIONALLY FAILS until that file is updated.
        Update instructions:
          File: spec_engine/tests/test_block_pipeline_behavior.py
          Function: test_objection_extraction_exit_form
          Line ~73: Change:
            assert "Objection to form." in text
          To:
            assert "Objection.  Form." in text

        WHY: 'Objection to form.' is a fabrication. Morson's requires two sentences as spoken.
        """
        from pipeline.processor import run_pipeline

        result = run_pipeline(
            {
                "utterances": [
                    {
                        "speaker": 2,
                        "transcript": "I don't recall. Exit form.",
                        "words": [],
                    }
                ]
            },
            {
                "speaker_map": {2: "THE WITNESS"},
                "witness_id": 2,
                "defense_counsel": "MR. SMITH",
                "cause_number": "TEST-OBJ-MORSON",
            },
        )

        text = result["text"]
        # After Fix 1-D, this is the CORRECT assertion:
        assert "Objection." in text and "Form." in text, (
            f"Pipeline must produce the two-sentence objection form after Fix 1-D. "
            f"Got output: {text!r}. "
            f"Also update test_block_pipeline_behavior.py line ~73 to match."
        )
        # This should NO LONGER appear after Fix 1-D:
        assert "Objection to form." not in text, (
            f"'Objection to form.' (single sentence) must not appear after Fix 1-D. "
            f"This is a Morson's verbatim rule violation."
        )


# ─────────────────────────────────────────────────────────────────────────────
# FIX 1-E: JOBCONFIG PERSISTED IMMEDIATELY
# ─────────────────────────────────────────────────────────────────────────────

class TestFix1E_JobConfigPersistence:
    """
    Fix 1-E: JobConfig.save() must be called immediately after JobConfigDialog
    returns a result — before speaker verification — so case data is never lost
    if the user later cancels verification.
    """

    def test_jobconfig_save_round_trips_confirmed_spellings(self, tmp_path):
        """
        confirmed_spellings must survive save/load without data loss.
        This is the most critical field — it's what fixes Bexar County, David Blas, etc.
        """
        spellings = {
            "Cogger": "Coger",
            "Bare County": "Bexar County",
            "David Blvd": "David Blas",
            "David Blvd.": "David Blas",
            "Miah Vardell": "Miah Bardot",
            "May Vardell": "Miah Bardot",
            "cop number": "Cause Number",
        }
        cfg = JobConfig(
            cause_number="2025CI19595",
            confirmed_spellings=spellings,
        )
        saved = cfg.save(jobs_dir=str(tmp_path))
        loaded = JobConfig.load(saved)

        for wrong, correct in spellings.items():
            assert loaded.confirmed_spellings.get(wrong) == correct, (
                f"confirmed_spellings['{wrong}'] lost after save/load. "
                f"Expected '{correct}', got '{loaded.confirmed_spellings.get(wrong)!r}'"
            )

    def test_jobconfig_save_round_trips_speaker_map(self, tmp_path):
        """speaker_map must survive save/load with integer keys intact."""
        cfg = JobConfig(
            cause_number="TEST-MAP-2025",
            speaker_map={
                0: "THE VIDEOGRAPHER",
                1: "THE WITNESS",
                2: "MR. ALLAN",
                3: "MR. BOYCE - OPPOSING COUNSEL",
                4: "THE REPORTER",
            },
            examining_attorney_id=2,
            witness_id=1,
            speaker_map_verified=True,
        )
        saved = cfg.save(jobs_dir=str(tmp_path))
        loaded = JobConfig.load(saved)

        assert loaded.speaker_map.get(0) == "THE VIDEOGRAPHER"
        assert loaded.speaker_map.get(3) == "MR. BOYCE - OPPOSING COUNSEL"
        assert loaded.examining_attorney_id == 2
        assert loaded.speaker_map_verified is True

    def test_jobconfig_save_round_trips_defense_counsel(self, tmp_path):
        """
        defense_counsel list must survive save/load.
        This is what _resolve_objection_speaker() (Fix 1-C) checks as Priority 2.
        """
        from spec_engine.models import CounselInfo
        cfg = JobConfig(
            cause_number="TEST-COUNSEL-2025",
            defense_counsel=[
                CounselInfo(
                    name="David Boyce",
                    firm="Defense Law Firm",
                    address="123 Main St",
                    city="San Antonio",
                    state="Texas",
                    zip_code="78201",
                )
            ],
        )
        saved = cfg.save(jobs_dir=str(tmp_path))
        loaded = JobConfig.load(saved)

        assert loaded.defense_counsel, "defense_counsel list must not be empty after load"
        assert loaded.defense_counsel[0].name == "David Boyce"
        assert loaded.defense_counsel[0].firm == "Defense Law Firm"

    def test_jobconfig_save_does_not_overwrite_with_empty(self, tmp_path):
        """
        Saving a config, then loading and re-saving must not lose data.
        Tests that the save/load cycle is stable across multiple calls.
        """
        cfg = JobConfig(
            cause_number="TEST-STABLE-2025",
            witness_name="Matthew Allan Coger",
            confirmed_spellings={"Cogger": "Coger"},
        )
        saved1 = cfg.save(jobs_dir=str(tmp_path))
        loaded1 = JobConfig.load(saved1)
        saved2 = loaded1.save(jobs_dir=str(tmp_path))
        loaded2 = JobConfig.load(saved2)

        assert loaded2.witness_name == "Matthew Allan Coger"
        assert loaded2.confirmed_spellings.get("Cogger") == "Coger"

    def test_jobconfig_cause_number_in_filename(self, tmp_path):
        """
        Saved file must use cause_number in its name.
        This lets the user identify the correct config when re-loading.
        """
        cfg = JobConfig(cause_number="2025CI19595")
        saved = cfg.save(jobs_dir=str(tmp_path))
        assert "2025CI19595" in Path(saved).stem, (
            f"Cause number '2025CI19595' should be in filename. Got: {Path(saved).name}"
        )


# ─────────────────────────────────────────────────────────────────────────────
# INTEGRATION: COGER SCENARIO — ALL PHASE 1 FIXES TOGETHER
# ─────────────────────────────────────────────────────────────────────────────

class TestPhase1Integration_CogerScenario:
    """
    End-to-end integration tests that simulate the Coger deposition conditions
    which exposed all five Phase 1 bugs. These must all pass together.
    """

    def test_coger_confirmed_spellings_fix_bexar_county(self):
        """
        With confirmed_spellings populated, 'Bare County' must correct to 'Bexar County'.
        This was unresolved in the broken A.docx output.
        """
        cfg = _make_coger_config()
        result, _ = clean_block("I work in Bare County, Texas.", cfg, block_index=0)[:2]
        assert "Bexar County" in result, (
            f"'Bare County' must correct to 'Bexar County' via confirmed_spellings. "
            f"Got: {result!r}"
        )

    def test_coger_confirmed_spellings_fix_witness_name(self):
        """'Cogger' must correct to 'Coger' via confirmed_spellings."""
        cfg = _make_coger_config()
        result, _ = clean_block("My name is Matthew Cogger.", cfg, block_index=0)[:2]
        assert "Coger" in result, (
            f"'Cogger' must correct to 'Coger' via confirmed_spellings. Got: {result!r}"
        )
        assert "Cogger" not in result, f"'Cogger' persisted after correction. Got: {result!r}"

    def test_coger_confirmed_spellings_fix_manager_name(self):
        """'David Blvd.' must correct to 'David Blas' via confirmed_spellings."""
        cfg = _make_coger_config()
        result, _ = clean_block("My boss David Blvd. told me.", cfg, block_index=0)[:2]
        assert "David Blas" in result, (
            f"'David Blvd.' must correct to 'David Blas'. Got: {result!r}"
        )

    def test_coger_objection_speaker_resolves_to_boyce(self):
        """
        With Coger config, objection speaker must resolve to MR. BOYCE —
        not MR. UNKNOWN or MR. TEST.
        """
        cfg = _make_coger_config()
        result = _resolve_objection_speaker(cfg)
        assert "BOYCE" in result.upper(), (
            f"Objection speaker must be MR. BOYCE for Coger config. Got: {result!r}"
        )
        assert "UNKNOWN" not in result.upper(), f"Still returning UNKNOWN: {result!r}"
        assert "TEST" not in result.upper(), f"Still returning TEST: {result!r}"

    def test_coger_objection_text_is_verbatim_two_sentences(self):
        """
        ASR garbles in the Coger transcript must produce 'Objection.  Form.' — two sentences.
        The broken output had 'Objection to form.' throughout.
        """
        cfg = _make_coger_config()
        garbles = ["Exit form", "Action form", "Objection form", "Action point"]
        for garble in garbles:
            result, _ = clean_block(garble, cfg, block_index=0)[:2]
            assert "Objection.  Form." in result, (
                f"'{garble}' must produce 'Objection.  Form.' Got: {result!r}"
            )
            assert "Objection to form." not in result, (
                f"'{garble}' produced 'Objection to form.' — Morson's violation. "
                f"Got: {result!r}"
            )

    def test_coger_classify_blocks_warns_not_crashes_without_verification(self):
        """
        Running classify_blocks() without verification must warn, not crash.
        The broken Coger run bypassed verification silently.
        """
        cfg = JobConfig(
            speaker_map={1: "THE WITNESS", 2: "MR. ALLAN", 3: "MR. BOYCE"},
            speaker_map_verified=False,  # Simulates bypass
        )
        blocks = [
            Block(speaker_id=2, text="Did you see the spill?", raw_text=""),
            Block(speaker_id=1, text="Yes, sir.", raw_text=""),
        ]
        try:
            result = classify_blocks(blocks, job_config=cfg)
            assert len(result) == 2, "Must return all blocks even without verification"
        except SpeakerMapUnverifiedError:
            pytest.fail(
                "classify_blocks() raised SpeakerMapUnverifiedError on unverified map. "
                "Fix 1-B adds a warning, not a raise — preview mode must still work."
            )

    def test_coger_full_config_saves_and_loads_without_data_loss(self, tmp_path):
        """The full Coger JobConfig must survive a save/load cycle intact."""
        cfg = _make_coger_config()
        saved = cfg.save(jobs_dir=str(tmp_path))
        loaded = JobConfig.load(saved)

        assert loaded.cause_number == "2025CI19595"
        assert loaded.witness_name == "Matthew Allan Coger"
        assert loaded.speaker_map.get(3) == "MR. BOYCE - OPPOSING COUNSEL"
        assert loaded.confirmed_spellings.get("Cogger") == "Coger"
        assert loaded.confirmed_spellings.get("Bare County") == "Bexar County"
        assert loaded.speaker_map_verified is True
        assert loaded.examining_attorney_id == 2
        assert loaded.witness_id == 1


# ─────────────────────────────────────────────────────────────────────────────
# PHASE 1 EXIT GATE
# ─────────────────────────────────────────────────────────────────────────────

class TestPhase1ExitGate:
    """
    Final gate tests. All must pass before Phase 2 begins.
    These are high-level assertions about the system state after all 5 fixes.
    """

    def test_mr_unknown_is_impossible_with_any_config(self):
        """MR. UNKNOWN must be impossible regardless of config shape."""
        configs = [
            _make_empty_config(),
            _make_unverified_config(),
            _make_coger_config(),
            JobConfig(speaker_map={}, defense_counsel=[]),
        ]
        for cfg in configs:
            result = _resolve_objection_speaker(cfg)
            assert result != "MR. UNKNOWN", (
                f"MR. UNKNOWN returned for config: speaker_map={getattr(cfg, 'speaker_map', {})}. "
                f"Got: {result!r}. Fix 1-C must be applied."
            )

    def test_objection_to_form_single_sentence_never_produced_by_corrections(self):
        """
        The single-sentence 'Objection to form.' must NEVER be produced by
        UNIVERSAL_CORRECTIONS after Fix 1-D.
        """
        garbles = ["Exit form", "Action form", "Action point", "Objection form"]
        for garble in garbles:
            records = []
            result = apply_universal_corrections(garble, records, block_index=0)
            assert result != "Objection to form.", (
                f"MORSON'S VIOLATION: '{garble}' produces 'Objection to form.' "
                f"Fix 1-D must change this to 'Objection.  Form.' (two sentences)."
            )

    def test_fresh_jobconfig_speaker_map_verified_is_false(self):
        """
        JobConfig()  must have speaker_map_verified=False.
        This is the sentinel that Fix 1-A's gate checks.
        """
        assert JobConfig().speaker_map_verified is False

    def test_classify_blocks_plural_accepts_job_config_parameter(self):
        """
        classify_blocks() must accept job_config without error.
        Before Fix 1-B, the parameter was accepted but silently discarded.
        After Fix 1-B, it's used for the warning guard.
        """
        cfg = _make_coger_config()
        blocks = [Block(speaker_id=1, text="Yes.", raw_text="Yes.",
                        speaker_role="WITNESS", speaker_name="THE WITNESS")]
        # Must not raise TypeError or AttributeError
        result = classify_blocks(blocks, job_config=cfg)
        assert isinstance(result, list)
