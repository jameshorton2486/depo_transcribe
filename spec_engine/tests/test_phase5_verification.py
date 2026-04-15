"""
test_phase5_verification.py

Phase 5 Verification Test Suite — Corrections, Emitter, and Validation
Run with: python -m pytest spec_engine/tests/test_phase5_verification.py -v
"""

import inspect
import sys
from pathlib import Path

import pytest
from docx.shared import Twips

sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))

from spec_engine.models import Block, BlockType
from spec_engine.validator import _normalize_for_compare, validate_blocks


class TestFix5A_CurrentSentenceSpacingContract:
    def _normalize(self, text: str) -> tuple[str, list]:
        from spec_engine.corrections import normalize_sentence_spacing

        records = []
        return normalize_sentence_spacing(text, records, 0), records

    def test_ellipsis_not_double_spaced_before_capital(self):
        text = "Answer: . . . I do not recall."
        result, _records = self._normalize(text)
        assert ". . ." in result
        assert ". . .  " not in result

    def test_ellipsis_count_unchanged(self):
        cases = [
            "He said . . . and paused.",
            "The answer was . . . unclear. Continuing.",
            "First . . . then . . . finally.",
        ]
        for text in cases:
            result, _records = self._normalize(text)
            assert result.count(". . .") == text.count(". . .")

    def test_ellipsis_before_uppercase_preserved(self):
        text = "He said . . . Correct. And moved on."
        result, _records = self._normalize(text)
        assert ". . ." in result
        assert ". . .  Correct" not in result

    def test_normal_sentence_spacing_still_works(self):
        text = "He answered. The question was clear."
        result, records = self._normalize(text)
        assert "answered.  The" in result
        assert records
        assert records[0].pattern == "sentence_spacing_two_spaces"

    def test_question_mark_spacing_still_works(self):
        text = "Did you see it? Yes, I did."
        result, _records = self._normalize(text)
        assert "it?  Yes" in result

    def test_abbreviation_spacing_not_double_spaced_mid_name(self):
        text = "Dr. Smith testified. He left."
        result, _records = self._normalize(text)
        assert "Dr.  Smith" not in result
        assert "testified.  He" in result

class TestFix5B_FormatBlocksToTextCurrentContract:
    def test_signature_accepts_single_blocks_parameter(self):
        from core.correction_runner import format_blocks_to_text

        sig = inspect.signature(format_blocks_to_text)
        assert list(sig.parameters.keys()) == ["blocks"]

    def test_legacy_skip_parameter_absent(self):
        from core.correction_runner import format_blocks_to_text

        sig = inspect.signature(format_blocks_to_text)
        assert "skip_corrections_already_applied" not in sig.parameters

    def test_formats_question_answer_and_speaker_blocks(self):
        from core.correction_runner import format_blocks_to_text

        blocks = [
            Block(
                speaker_id=2,
                text="Did you witness the incident?",
                raw_text="",
                block_type=BlockType.QUESTION,
                speaker_role="EXAMINING_ATTORNEY",
                speaker_name="MR. TEST",
            ),
            Block(
                speaker_id=1,
                text="Yes, I did.",
                raw_text="",
                block_type=BlockType.ANSWER,
                speaker_role="WITNESS",
                speaker_name="THE WITNESS",
            ),
            Block(
                speaker_id=3,
                text="Objection. Form.",
                raw_text="",
                block_type=BlockType.SPEAKER,
                speaker_role="OPPOSING_COUNSEL",
                speaker_name="MR. BOYCE",
            ),
        ]

        result = format_blocks_to_text(blocks)
        assert "\tQ.  Did you witness the incident?" in result
        assert "\tA.  Yes, I did." in result
        assert "\t\t\tMR. BOYCE:  Objection. Form." in result

    def test_empty_blocks_returns_empty_string(self):
        from core.correction_runner import format_blocks_to_text

        assert format_blocks_to_text([]) == ""

    def test_plain_unknown_block_falls_back_to_text(self):
        from core.correction_runner import format_blocks_to_text

        blocks = [Block(speaker_id=99, text="Loose text", raw_text="", block_type=BlockType.UNKNOWN)]
        assert "Loose text" in format_blocks_to_text(blocks)

    def test_non_qa_paragraphs_use_three_tabs_and_single_newline(self):
        from core.correction_runner import format_blocks_to_text

        blocks = [
            Block(
                speaker_id=4,
                text="This deposition is taking place via Zoom.",
                raw_text="",
                block_type=BlockType.SPEAKER,
                speaker_name="THE REPORTER",
            ),
            Block(
                speaker_id=1,
                text="Yes, ma'am.",
                raw_text="",
                block_type=BlockType.ANSWER,
                speaker_name="THE WITNESS",
            ),
        ]

        result = format_blocks_to_text(blocks)
        assert "\t\t\tTHE REPORTER:  This deposition is taking place via Zoom." in result
        assert "\n\tA.  Yes, ma'am." in result
        assert "\n\n" not in result


class TestFix5C_EmitSpLineSingleSpace:
    def test_two_space_colon_bold_label(self):
        from spec_engine.emitter import create_document, emit_sp_line
        doc = create_document()
        emit_sp_line(doc, "MR. BOYCE:  Objection. Form.")
        bold_runs = [r.text for r in doc.paragraphs[0].runs if r.bold]
        assert bold_runs
        assert any("BOYCE" in r for r in bold_runs)

    def test_single_space_colon_bold_label(self):
        from spec_engine.emitter import create_document, emit_sp_line
        doc = create_document()
        emit_sp_line(doc, "MR. BOYCE: Objection. Form.")
        bold_runs = [r.text for r in doc.paragraphs[0].runs if r.bold]
        assert bold_runs
        assert any("BOYCE" in r for r in bold_runs)

    def test_single_space_text_uses_two_spaces_after_colon_in_output(self):
        from spec_engine.emitter import create_document, emit_sp_line
        doc = create_document()
        emit_sp_line(doc, "MR. BOYCE: Objection. Form.")
        plain_runs = [r.text for r in doc.paragraphs[0].runs if not r.bold]
        text_runs = [r for r in plain_runs if "Objection" in r]
        assert text_runs
        assert text_runs[0].startswith("  ")

    def test_no_colon_emits_without_error(self):
        from spec_engine.emitter import create_document, emit_sp_line
        doc = create_document()
        emit_sp_line(doc, "THE REPORTER continued speaking.")
        assert len(doc.paragraphs) == 1

    def test_reporter_label_bold(self):
        from spec_engine.emitter import create_document, emit_sp_line
        for sep in [":  ", ": "]:
            doc = create_document()
            emit_sp_line(doc, f"THE REPORTER{sep}You are under oath.")
            bold_runs = [r.text for r in doc.paragraphs[0].runs if r.bold]
            assert bold_runs

    def test_split_speaker_text_handles_extra_spaces_after_colon(self):
        from spec_engine.emitter import _split_speaker_text

        label, content = _split_speaker_text("MR. BOYCE:   Objection. Form.")

        assert label == "MR. BOYCE:"
        assert content == "Objection. Form."


class TestFix5D_SharedTabStopHelper:
    def test_standard_tabs_constant_exists(self):
        from spec_engine.emitter import _STANDARD_TABS
        assert _STANDARD_TABS is not None

    def test_standard_tabs_contains_correct_values(self):
        from spec_engine.emitter import (
            TAB_720,
            TAB_1440,
            TAB_2160,
            TAB_CENTER,
            _STANDARD_TABS,
        )
        assert _STANDARD_TABS == [TAB_720, TAB_1440, TAB_2160, TAB_CENTER]

    def test_apply_standard_tabs_helper_exists(self):
        from spec_engine.emitter import _apply_standard_tabs
        assert callable(_apply_standard_tabs)

    def test_emit_line_numbered_uses_shared_helper(self):
        from spec_engine.emitter import emit_line_numbered
        src = inspect.getsource(emit_line_numbered)
        assert "_apply_standard_tabs(" in src
        assert "for stop_twips in [TAB_720" not in src

    def test_numbered_emitter_still_produces_correct_output(self):
        from spec_engine.emitter import (
            LineNumberTracker,
            QAPairTracker,
            create_document,
            emit_line_numbered,
        )
        from spec_engine.models import LineType

        doc = create_document()
        tracker = LineNumberTracker()
        qa = QAPairTracker()
        emit_line_numbered(doc, LineType.Q, "Did you witness this?", tracker, qa)
        emit_line_numbered(doc, LineType.A, "Yes, I did.", tracker, qa)
        assert len(doc.paragraphs) == 2
        q_text = "".join(r.text for r in doc.paragraphs[0].runs)
        a_text = "".join(r.text for r in doc.paragraphs[1].runs)
        assert "Q." in q_text
        assert "A." in a_text

    def test_header_line_uses_center_tab_stop(self):
        from spec_engine.emitter import TAB_CENTER, create_document, emit_header_line

        doc = create_document()
        emit_header_line(doc, "EXAMINATION")
        tab_positions = [stop.position for stop in doc.paragraphs[0].paragraph_format.tab_stops]
        assert Twips(TAB_CENTER) in tab_positions

    def test_question_emitter_rejects_empty_text(self):
        from spec_engine.emitter import create_document, emit_q_line

        doc = create_document()
        with pytest.raises(ValueError, match="Q line cannot be empty"):
            emit_q_line(doc, "   ")

    def test_speaker_emitter_rejects_missing_label_before_colon(self):
        from spec_engine.emitter import create_document, emit_sp_line

        doc = create_document()
        with pytest.raises(ValueError, match="missing a label"):
            emit_sp_line(doc, ":  Objection.")


class TestFix5E_LinesPerPageDocumentation:
    def test_lines_per_page_value_unchanged(self):
        from spec_engine.emitter import LineNumberTracker
        assert LineNumberTracker.LINES_PER_PAGE == 25

    def test_lines_per_page_has_derivation_comment(self):
        from spec_engine.emitter import LineNumberTracker
        src = inspect.getsource(LineNumberTracker)
        has_derivation = any(
            keyword in src
            for keyword in ["margin", "Usable", "derivation", "11 inch", "11.0", "9.0"]
        )
        assert has_derivation

    def test_line_number_tracker_still_functions(self):
        from spec_engine.emitter import LineNumberTracker
        tracker = LineNumberTracker(start_page=3)
        for expected_line in range(1, 26):
            page, line = tracker.next()
            assert page == 3
            assert line == expected_line
        page, line = tracker.next()
        assert page == 4
        assert line == 1


class TestFix5F_NormalizeForCompare:
    def test_period_vs_question_mark_differ(self):
        assert _normalize_for_compare("Yes.") != _normalize_for_compare("Yes?")

    def test_identical_text_still_matches(self):
        assert _normalize_for_compare("Yes.") == _normalize_for_compare("Yes.")
        assert _normalize_for_compare("I do not recall.") == _normalize_for_compare(
            "I do not recall."
        )

    def test_period_preserved_in_normalized(self):
        assert _normalize_for_compare("I don't know.").endswith(".")

    def test_question_mark_preserved_in_normalized(self):
        assert _normalize_for_compare("Did you see it?").endswith("?")

    def test_other_punctuation_still_stripped(self):
        result = _normalize_for_compare("I don't, recall—it.")
        assert "'" not in result
        assert "," not in result
        assert "—" not in result
        assert "." in result

    def test_near_duplicate_detection_distinguishes_punctuation(self):
        b1 = Block(speaker_id=1, text="Yes.", raw_text="", block_type=BlockType.ANSWER)
        b2 = Block(speaker_id=1, text="Yes?", raw_text="", block_type=BlockType.ANSWER)
        result = validate_blocks([b1, b2])
        dup_warnings = [w for w in result.warnings if "duplicate" in w.lower()]
        assert not dup_warnings


class TestFix5G_ColloquySpeakerRoleValidation:
    def test_colloquy_with_witness_role_warns(self):
        b = Block(
            speaker_id=1,
            text="I do not recall seeing that.",
            raw_text="",
            block_type=BlockType.COLLOQUY,
            speaker_role="WITNESS",
        )
        result = validate_blocks([b])
        assert [w for w in result.warnings if "COLLOQUY" in w]

    def test_speaker_with_witness_role_warns(self):
        b = Block(
            speaker_id=1,
            text="THE WITNESS:  Something was said.",
            raw_text="",
            block_type=BlockType.SPEAKER,
            speaker_role="WITNESS",
        )
        result = validate_blocks([b])
        assert [w for w in result.warnings if "SPEAKER" in w]

    def test_colloquy_with_attorney_role_no_warning(self):
        b = Block(
            speaker_id=2,
            text="Let me rephrase that.",
            raw_text="",
            block_type=BlockType.COLLOQUY,
            speaker_role="EXAMINING_ATTORNEY",
        )
        result = validate_blocks([b])
        assert not [w for w in result.warnings if "COLLOQUY" in w]

    def test_answer_with_witness_role_no_warning(self):
        b = Block(
            speaker_id=1,
            text="Yes, I did.",
            raw_text="",
            block_type=BlockType.ANSWER,
            speaker_role="WITNESS",
        )
        result = validate_blocks([b])
        role_warns = [
            w
            for w in result.warnings
            if "WITNESS" in w and "COLLOQUY" not in w and "SPEAKER" not in w
        ]
        assert not role_warns

    def test_colloquy_with_empty_role_no_warning(self):
        b = Block(
            speaker_id=0,
            text="Let me clarify.",
            raw_text="",
            block_type=BlockType.COLLOQUY,
            speaker_role="",
        )
        result = validate_blocks([b])
        assert not [w for w in result.warnings if "COLLOQUY" in w]


class TestFix5H_TerminalPunctuationValidation:
    def test_question_without_mark_warns(self):
        b = Block(speaker_id=2, text="Did you see the spill", raw_text="", block_type=BlockType.QUESTION)
        result = validate_blocks([b])
        q_warns = [w for w in result.warnings if "Question" in w]
        assert q_warns
        assert "?" in q_warns[0]

    def test_question_with_mark_no_warning(self):
        b = Block(speaker_id=2, text="Did you see the spill?", raw_text="", block_type=BlockType.QUESTION)
        result = validate_blocks([b])
        q_warns = [w for w in result.warnings if "Question" in w and "?" in w]
        assert not q_warns

    def test_answer_without_terminal_punctuation_warns(self):
        b = Block(speaker_id=1, text="Yes I did", raw_text="", block_type=BlockType.ANSWER)
        result = validate_blocks([b])
        assert [w for w in result.warnings if "Answer" in w]

    def test_answer_with_period_no_warning(self):
        b = Block(speaker_id=1, text="Yes, I did.", raw_text="", block_type=BlockType.ANSWER)
        result = validate_blocks([b])
        a_warns = [w for w in result.warnings if "Answer" in w and "punctuation" in w]
        assert not a_warns

    def test_answer_with_exclamation_no_warning(self):
        b = Block(speaker_id=1, text="Absolutely not!", raw_text="", block_type=BlockType.ANSWER)
        result = validate_blocks([b])
        a_warns = [w for w in result.warnings if "Answer" in w and "punctuation" in w]
        assert not a_warns

    def test_answer_with_question_mark_no_warning(self):
        b = Block(
            speaker_id=1,
            text="Are you sure that's right?",
            raw_text="",
            block_type=BlockType.ANSWER,
        )
        result = validate_blocks([b])
        a_warns = [w for w in result.warnings if "Answer" in w and "punctuation" in w]
        assert not a_warns

    def test_empty_block_text_not_warned(self):
        b = Block(speaker_id=1, text="", raw_text="", block_type=BlockType.ANSWER)
        result = validate_blocks([b])
        a_warns = [w for w in result.warnings if "Answer" in w and "punctuation" in w]
        assert not a_warns

    def test_all_correct_blocks_produce_no_punctuation_warnings(self):
        q = Block(
            speaker_id=2,
            text="Did you witness the incident?",
            raw_text="",
            block_type=BlockType.QUESTION,
            speaker_role="EXAMINING_ATTORNEY",
        )
        a = Block(
            speaker_id=1,
            text="Yes, I was there.",
            raw_text="",
            block_type=BlockType.ANSWER,
            speaker_role="WITNESS",
        )
        result = validate_blocks([q, a])
        punc_warns = [
            w
            for w in result.warnings
            if "Question" in w and "?" in w or "Answer" in w and "punctuation" in w
        ]
        assert not punc_warns


class TestPhase5ExitGate:
    def test_ellipsis_not_double_spaced(self):
        from spec_engine.corrections import normalize_sentence_spacing
        result = normalize_sentence_spacing("Witness said . . . Yes.", [], 0)
        assert ". . .  " not in result
        assert ". . ." in result

    def test_format_blocks_to_text_current_signature(self):
        from core.correction_runner import format_blocks_to_text
        sig = inspect.signature(format_blocks_to_text)
        assert list(sig.parameters.keys()) == ["blocks"]

    def test_sp_single_space_produces_bold(self):
        from spec_engine.emitter import create_document, emit_sp_line
        doc = create_document()
        emit_sp_line(doc, "MR. BOYCE: Objection.")
        assert any(r.bold for r in doc.paragraphs[0].runs)

    def test_standard_tabs_constant_exists(self):
        from spec_engine.emitter import _STANDARD_TABS
        assert len(_STANDARD_TABS) == 4

    def test_lines_per_page_is_25(self):
        from spec_engine.emitter import LineNumberTracker
        assert LineNumberTracker.LINES_PER_PAGE == 25

    def test_yes_period_vs_question_differ(self):
        assert _normalize_for_compare("Yes.") != _normalize_for_compare("Yes?")

    def test_colloquy_witness_warns(self):
        b = Block(speaker_id=1, text="Test.", raw_text="", block_type=BlockType.COLLOQUY, speaker_role="WITNESS")
        result = validate_blocks([b])
        assert any("COLLOQUY" in w for w in result.warnings)

    def test_question_without_mark_warns(self):
        b = Block(speaker_id=2, text="Did you see it", raw_text="", block_type=BlockType.QUESTION)
        result = validate_blocks([b])
        assert any("Question" in w for w in result.warnings)
