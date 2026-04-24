"""
emitter.py

Writes formatted paragraphs to a python-docx Document.
Spec Section 5: Page Layout and Typography
Spec Section 3.3: Five Line Types

TYPOGRAPHY (must match exactly — Spec Section 5.2):
  Font:         Courier New, 12pt
  Line spacing: Double
  Tab stops:    720 / 1440 / 2160 twips plus centered header tab
  Margins:      Left 1.25" / Right 1.0" / Top 1.0" / Bottom 1.0"

COLOR USAGE (Spec Section 5.4):
  Orange  #B45F06 / RGBColor(0xB4,0x5F,0x06) — Scopist flags (bold)
  Navy    #1E3A5F / RGBColor(0x1E,0x3A,0x5F) — All parenthetical lines
  Black   Default                              — All other transcript text

LINE BREAK RULE:
  One paragraph per logical block. Word handles visual word-wrap.
  Hard line breaks only occur between blocks (new speaker or new Q/A).
  textwrap is intentionally NOT used anywhere in this file.
"""

import re
from docx import Document
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_BREAK,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
)
from docx.shared import Pt, RGBColor, Twips

from .models import LineType


# ── Color constants ────────────────────────────────────────────────────────────
COLOR_ORANGE = RGBColor(0xB4, 0x5F, 0x06)
COLOR_NAVY   = RGBColor(0x1E, 0x3A, 0x5F)
COLOR_BLACK  = RGBColor(0x00, 0x00, 0x00)

# ── Typography ─────────────────────────────────────────────────────────────────
FONT_NAME = "Courier New"
FONT_SIZE = Pt(12)

# ── Tab stop positions in twips (Spec Section 5.3) ────────────────────────────
TAB_720  = 720    # 0.5" — Q./A. line marker
TAB_1440 = 1440   # 1.0" — Q./A. text start
TAB_2160 = 2160   # 1.5" — speaker line text start
TAB_CENTER = 4680  # 3.25" — centered headers / witness-name line
_STANDARD_TABS = [TAB_720, TAB_1440, TAB_2160, TAB_CENTER]


# ── Q/A pair safety tracker (Spec Section 5.5) ────────────────────────────────
class QAPairTracker:
    """Spec 5.5: Never split Q and A pairs across pages."""
    def __init__(self):
        self.last_was_q = False

    def record_q(self):
        self.last_was_q = True

    def record_other(self):
        self.last_was_q = False

    def safe_to_break(self) -> bool:
        return not self.last_was_q


# ── Internal helpers ──────────────────────────────────────────────────────────

def _apply_standard_tabs(para) -> None:
    _set_paragraph_tabs(para, _STANDARD_TABS)


def _set_paragraph_tabs(para, tab_stops) -> None:
    for stop_twips in tab_stops or []:
        alignment = (
            WD_TAB_ALIGNMENT.CENTER if stop_twips == TAB_CENTER else WD_TAB_ALIGNMENT.LEFT
        )
        para.paragraph_format.tab_stops.add_tab_stop(
            Twips(stop_twips),
            alignment=alignment,
        )


def _set_paragraph_format(para, tab_stops=None):
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    if tab_stops:
        _set_paragraph_tabs(para, tab_stops)


def _add_run(para, text, bold=False, color=COLOR_BLACK):
    run = para.add_run(text)
    run.font.name      = FONT_NAME
    run.font.size      = FONT_SIZE
    run.font.bold      = bold
    run.font.color.rgb = color
    return run


def _clean(text: str) -> str:
    """Normalize whitespace and remove any carriage returns/newlines.

    Preserves the two-space separator after sentence-ending punctuation
    (Morson's Rule 14). A naive " ".join(text.split()) would collapse that
    deliberate double space to a single space and undo the two-space rule
    that corrections.py enforces upstream.
    """
    normalized = (text or "").replace("\r", " ").replace("\n", " ")
    normalized = normalized.replace("\t", " ")
    # Cap 3-or-more runs of spaces to 2 (keeps .?! double-space intact).
    normalized = re.sub(r" {3,}", "  ", normalized)
    # Collapse runs of exactly 2 spaces to 1 ONLY when they do not follow
    # sentence-ending punctuation.
    normalized = re.sub(r"(?<![.?!])  ", " ", normalized)
    return normalized.strip()


def _split_speaker_text(text: str) -> tuple[str, str]:
    match = re.split(r":\s+", text, maxsplit=1)
    if len(match) == 2:
        label, content = match
        return label + ":", content
    if ":" in text:
        label, content = text.split(":", 1)
        return label + ":", content.strip()
    return "", text


def _normalize_speaker_label(label: str) -> str:
    normalized = (label or "").strip().upper().rstrip(":")
    if normalized == "THE COURT REPORTER":
        return "THE REPORTER"
    return normalized


def _format_plain_speaker_line(label: str, text: str) -> str:
    normalized_label = _normalize_speaker_label(label) or "SPEAKER"
    clean_text = _clean(text)
    existing_label, content = _split_speaker_text(clean_text)

    if _normalize_speaker_label(existing_label) == normalized_label:
        if content:
            return f"\t\t\t{normalized_label}:  {content}"
        return f"\t\t\t{normalized_label}:"

    return f"\t\t\t{normalized_label}:  {clean_text}"


def _validate_emit_input(line_type: LineType, text: str) -> str:
    raw_text = text or ""
    clean_text = _clean(raw_text)

    if line_type in {LineType.Q, LineType.A, LineType.SP} and not clean_text:
        raise ValueError(f"{line_type.value} line cannot be empty")

    if raw_text.startswith("\t\t\t\t\t"):
        raise ValueError(f"{line_type.value} line contains more than four leading tabs")

    if line_type in {LineType.Q, LineType.A, LineType.SP} and raw_text.startswith("    "):
        raise ValueError(f"{line_type.value} line uses spaces instead of tabs for indentation")

    if line_type == LineType.SP and ":" in raw_text:
        label, _content = _split_speaker_text(raw_text)
        if not label.rstrip(":").strip():
            raise ValueError("Speaker line is missing a label before ':'")

    return clean_text


def _qa_visual_text(prefix: str, text: str) -> str:
    return f"\t{prefix}.  {text}"


def _speaker_visual_text(text: str) -> tuple[str, str, str]:
    label, content = _split_speaker_text(text)
    if label:
        normalized_label = _normalize_speaker_label(label) + ":"
        return normalized_label, content, f"\t\t\t{normalized_label}  {content}"
    return "", text, f"\t\t\t{text}"


# ── Plain-text block output (textbox + _corrected.txt) ───────────────────────

def emit_blocks(blocks: list) -> str:
    """
    Convert processed blocks into plain-text transcript output.
    One line per block. No hard wraps inside a block.
    New line only at block boundaries (new speaker or new Q/A paragraph).

    Speaker label rule (CLAUDE.md section 18):
    Consecutive SP/COLLOQUY/SPEAKER blocks from the same speaker emit the
    label only on the first block. Subsequent same-speaker blocks are
    emitted at the SP indent (three tabs) with no label. Any intervening
    block type (Q, A, PN, FLAG, etc.) resets the chain.
    """
    lines: list[str] = []
    last_speaker_label: str | None = None

    for block in blocks:
        bt = getattr(block, "block_type", None)
        block_value = getattr(bt, "value", str(bt)) if bt else "UNKNOWN"
        text = _clean(block.text or "")
        role = (getattr(block, "speaker_role", "") or "").strip()
        name = (getattr(block, "speaker_name", "") or "").strip()

        if not text:
            continue

        if block_value == "Q":
            lines.append(f"\tQ.  {text}")
            last_speaker_label = None
        elif block_value == "A":
            lines.append(f"\tA.  {text}")
            last_speaker_label = None
        elif block_value in ("COLLOQUY", "SPEAKER", "SP"):
            label = name or role or "SPEAKER"
            normalized = _normalize_speaker_label(label) or "SPEAKER"
            if last_speaker_label == normalized:
                # Same speaker as the previous SP block - emit continuation
                # at the SP indent with no label (per UFM speaker label rule).
                lines.append(f"\t\t\t{text}")
            else:
                lines.append(_format_plain_speaker_line(label, text))
                last_speaker_label = normalized
        elif block_value in ("PAREN", "PARENTHETICAL", "PN"):
            lines.append(text)
            last_speaker_label = None
        elif block_value == "FLAG":
            lines.append(text)
            last_speaker_label = None
        else:
            if name or role:
                label = (name or role).upper()
                lines.append(f"\t\t\t{label}:  {text}")
            else:
                lines.append(text)
            last_speaker_label = None

    return "\n".join(lines)


# ── DOCX line emitters (Spec Section 3.3) ────────────────────────────────────
# One paragraph per block. Word wraps to the page margin automatically.
# No textwrap — carriage returns only at block boundaries.

def emit_q_line(doc: Document, text: str):
    """Spec 3.3 Type 1 — Question: Tab + Q.  + text (one paragraph)."""
    clean_text = _validate_emit_input(LineType.Q, text)
    para = doc.add_paragraph()
    _set_paragraph_format(para, [TAB_720, TAB_1440])
    _add_run(para, _qa_visual_text("Q", clean_text))


def emit_a_line(doc: Document, text: str):
    """Spec 3.3 Type 2 — Answer: Tab + A.  + text (one paragraph)."""
    clean_text = _validate_emit_input(LineType.A, text)
    para = doc.add_paragraph()
    _set_paragraph_format(para, [TAB_720, TAB_1440])
    _add_run(para, _qa_visual_text("A", clean_text))


def emit_sp_line(doc: Document, text: str):
    """
    Spec 3.3 Type 3 — Speaker Label: 3 tabs + BOLD LABEL: + two spaces + text.
    One paragraph. Word wraps naturally.
    """
    clean_text = _validate_emit_input(LineType.SP, text)
    label, content, _visual_text = _speaker_visual_text(clean_text)
    para = doc.add_paragraph()
    _set_paragraph_format(para, [TAB_720, TAB_1440, TAB_2160])
    if not label:
        _add_run(para, _visual_text)
    else:
        _add_run(para, '\t\t\t', bold=False)
        _add_run(para, label, bold=True)
        _add_run(para, '  ' + content, bold=False)


def emit_pn_line(doc: Document, text: str):
    """Spec 3.3 Type 4 — Parenthetical: 4 tabs, navy color (one paragraph)."""
    clean_text = _validate_emit_input(LineType.PN, text)
    para = doc.add_paragraph()
    _set_paragraph_format(para, [TAB_2160])
    _add_run(para, f'\t\t\t\t{clean_text}', color=COLOR_NAVY)


def emit_flag_line(doc: Document, text: str):
    """Spec 3.3 Type 5 — Scopist Flag: bold orange (one paragraph)."""
    clean_text = _validate_emit_input(LineType.FLAG, text)
    para = doc.add_paragraph()
    _set_paragraph_format(para)
    _add_run(para, clean_text, bold=True, color=COLOR_ORANGE)


def emit_header_line(doc: Document, text: str):
    """Spec 4.1 — Examination header: centered bold."""
    clean_text = _validate_emit_input(LineType.HEADER, text)
    para = doc.add_paragraph()
    _set_paragraph_format(para, [TAB_CENTER])
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(para, clean_text, bold=True)


def emit_by_line(doc: Document, text: str):
    """Spec 4.1 — BY attribution line: left-aligned."""
    clean_text = _validate_emit_input(LineType.BY, text)
    para = doc.add_paragraph()
    _set_paragraph_format(para)
    _add_run(para, clean_text)


def emit_plain_line(doc: Document, text: str):
    """Plain transcript text — no label, no special styling."""
    clean_text = _validate_emit_input(LineType.PLAIN, text)
    para = doc.add_paragraph()
    _set_paragraph_format(para)
    _add_run(para, clean_text)


def emit_line(doc: Document, line_type: LineType, text: str):
    """Master dispatch — routes to correct emitter based on LineType."""
    dispatch = {
        LineType.Q:      emit_q_line,
        LineType.A:      emit_a_line,
        LineType.SP:     emit_sp_line,
        LineType.PN:     emit_pn_line,
        LineType.FLAG:   emit_flag_line,
        LineType.HEADER: emit_header_line,
        LineType.BY:     emit_by_line,
        LineType.PLAIN:  emit_plain_line,
    }
    fn = dispatch.get(line_type)
    if fn:
        fn(doc, text)


def add_page_break(doc: Document):
    """Insert a hard page break (Spec Section 5.5)."""
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(WD_BREAK.PAGE)


def create_document() -> Document:
    """Create Document with correct page setup (Spec Section 5.1)."""
    doc = Document()
    section = doc.sections[0]
    section.page_width    = Twips(12240)  # 8.5"
    section.page_height   = Twips(15840)  # 11"
    section.left_margin   = Twips(1800)   # 1.25"
    section.right_margin  = Twips(1440)   # 1.0"
    section.top_margin    = Twips(1440)   # 1.0"
    section.bottom_margin = Twips(1440)   # 1.0"
    if doc.paragraphs:
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)
    return doc


# ── Line number tracking (Spec Section 6 — UFM requirement) ──────────────────

class LineNumberTracker:
    """
    Tracks line numbers for Texas UFM transcript body pages.
    Lines restart at 1 at the start of each transcript page.
    Courier New 12pt double-spaced = approximately 25 lines per page.
    Derivation: 11.0 inch page height minus 1.0 inch top/bottom margins yields 9.0 inch usable space, so margin math gives ~25 lines per page.

    IMPORTANT: Line numbers apply to Pages 3+ (transcript body) only.
    Pages 1 (Corrections Log) and 2 (Caption) are NOT numbered.
    """
    LINES_PER_PAGE = 25

    def __init__(self, start_page: int = 3):
        self.current_line = 1
        self.current_page = start_page

    def next(self) -> tuple:
        page = self.current_page
        line = self.current_line
        self.current_line += 1
        if self.current_line > self.LINES_PER_PAGE:
            self.current_line = 1
            self.current_page += 1
        return page, line

    def reset_for_new_page(self):
        self.current_line = 1
        self.current_page += 1

    def safe_to_break_before(self) -> bool:
        return True


def emit_line_numbered(
    doc: Document,
    line_type: LineType,
    text: str,
    tracker: LineNumberTracker,
    qa_tracker: QAPairTracker,
):
    """
    Emit a line with a left-margin line number. One paragraph per block.
    Word handles visual word-wrap. No textwrap inside blocks.
    """
    clean = _validate_emit_input(line_type, text)

    if line_type == LineType.Q:
        visual_line = _qa_visual_text("Q", clean)
    elif line_type == LineType.A:
        visual_line = _qa_visual_text("A", clean)
    elif line_type == LineType.SP:
        label, content, visual_line = _speaker_visual_text(clean)
    elif line_type == LineType.PN:
        visual_line = '\t\t\t\t' + clean
    else:
        visual_line = clean

    page, line_num = tracker.next()

    if line_type == LineType.Q:
        qa_tracker.record_q()
    else:
        qa_tracker.record_other()

    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)

    line_num_str = f"{line_num:>2} "
    num_run = para.add_run(line_num_str)
    num_run.font.name      = FONT_NAME
    num_run.font.size      = Pt(10)
    num_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    if line_type == LineType.SP:
        label, content, _visual_text = _speaker_visual_text(clean)
        if label and (':' in visual_line):
            _add_run(para, '\t\t\t', bold=False)
            _add_run(para, label, bold=True)
            _add_run(para, '  ' + content, bold=False)
        else:
            content_run = para.add_run(visual_line)
            content_run.font.color.rgb = COLOR_BLACK
    elif line_type == LineType.PN:
        content_run = para.add_run(visual_line)
        content_run.font.color.rgb = COLOR_NAVY
    elif line_type == LineType.FLAG:
        content_run = para.add_run(visual_line)
        content_run.bold = True
        content_run.font.color.rgb = COLOR_ORANGE
    elif line_type == LineType.HEADER:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        content_run = para.add_run(visual_line)
        content_run.bold = True
        content_run.font.color.rgb = COLOR_BLACK
    else:
        content_run = para.add_run(visual_line)
        content_run.font.color.rgb = COLOR_BLACK

    for run in para.runs:
        if not run.font.name:
            run.font.name = FONT_NAME
        if not run.font.size:
            run.font.size = FONT_SIZE

    _apply_standard_tabs(para)
