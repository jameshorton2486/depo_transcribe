"""
corrections_log.py
Page 1 — Scopist corrections and changes log.
Spec Section 3.1: Working reference for the scopist. Not part of official transcript.

Contents:
  - Deposition metadata
  - VERBATIM NOTE (uh/um preserved)
  - CONFIRMED SPELLINGS section
  - CORRECTIONS APPLIED section
  - SCOPIST FLAGS section
"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from ..models import CorrectionRecord, JobConfig, ScopistFlag

COLOR_ORANGE = RGBColor(0xB4, 0x5F, 0x06)
FONT = "Courier New"


def _para(doc, text, bold=False, centered=False, color=None):
    p = doc.add_paragraph()
    if centered:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(12)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def _rule(doc):
    _para(doc, "_" * 65)


def write_corrections_log(
    doc: Document,
    job_config: JobConfig,
    corrections: list,
    flags: list,
):
    """
    Write Page 1 — Corrections and Changes Log.
    Spec Section 3.1.
    """
    _para(doc, "CORRECTIONS AND CHANGES LOG", bold=True, centered=True)
    _para(doc, "(Scopist Working Notes — Not Part of Official Transcript)",
          centered=True)
    _rule(doc)

    # Deposition metadata
    _para(doc, f"Witness:        {job_config.witness_name}", bold=True)
    _para(doc, f"Date:           {job_config.depo_date}")
    _para(doc, f"Time:           {job_config.depo_start_time} – {job_config.depo_end_time}")
    _para(doc, f"Cause Number:   {job_config.cause_number}")  # NOTE: "Cause Number" per TX terminology
    _para(doc, f"Court:          {job_config.court}")
    _para(doc, f"Method:         {job_config.method}")
    _para(doc, f"Court Reporter: {job_config.reporter_name}, {job_config.reporter_csr}")
    _rule(doc)

    # Verbatim note (Spec 3.1 — explicit requirement)
    _para(doc, "VERBATIM NOTE:", bold=True)
    _para(doc, '  "uh" and "um" are preserved throughout this transcript.')
    _para(doc, "  Filler words are part of the verbatim legal record and are never removed.")
    _para(doc, "  Stutter repetitions preserved unless confirmed Deepgram artifact.")
    _rule(doc)

    # Speaker map
    _para(doc, "SPEAKER MAP (Verified before processing):", bold=True)
    for sid, role in sorted(job_config.speaker_map.items()):
        _para(doc, f"  Speaker {sid}: {role}")
    _rule(doc)

    # Confirmed spellings
    _para(doc, f"CONFIRMED SPELLINGS ({len(job_config.confirmed_spellings)}):", bold=True)
    if job_config.confirmed_spellings:
        for wrong, correct in job_config.confirmed_spellings.items():
            _para(doc, f"  {wrong!r} → {correct!r}")
    else:
        _para(doc, "  (none for this job)")
    _rule(doc)

    # Corrections applied
    _para(doc, f"CORRECTIONS APPLIED: {len(corrections)}", bold=True)
    for rec in corrections[:100]:  # Show up to 100
        _para(doc, f"  [{rec.pattern}]  {rec.original!r} → {rec.corrected!r}")
    if len(corrections) > 100:
        _para(doc, f"  ... and {len(corrections) - 100} more (see full log in jobs/)")
    _rule(doc)

    # Scopist flags
    _para(doc, f"SCOPIST FLAGS: {len(flags)}", bold=True)
    for flag in flags:
        p = doc.add_paragraph()
        run = p.add_run(f"  FLAG {flag.number}: {flag.description}")
        run.bold = True
        run.font.color.rgb = COLOR_ORANGE
        run.font.name = FONT
        run.font.size = Pt(12)

    # Post-record spellings summary
    if job_config.post_record_spellings:
        _rule(doc)
        _para(doc, f"POST-RECORD SPELLINGS: {len(job_config.post_record_spellings)}", bold=True)
        _para(doc, "  (These are authoritative — all prior uses corrected retroactively)")
        for prs in job_config.post_record_spellings:
            _para(doc, f"  {prs.name} → {prs.correct_spelling} (spelled: {prs.letters_as_given})")
