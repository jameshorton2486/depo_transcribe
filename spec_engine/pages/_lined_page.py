"""
_lined_page.py

Shared helper: builds a 25-line bordered table per UFM page format.
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

FONT = "Courier New"
FONT_SIZE = Pt(12)
COLOR_GRAY = RGBColor(0x99, 0x99, 0x99)
COLOR_BLK = RGBColor(0x00, 0x00, 0x00)

LINES_PER_PAGE = 25
NUM_COL_TWIPS = 540
TOTAL_TWIPS = 9000
CONTENT_TWIPS = TOTAL_TWIPS - NUM_COL_TWIPS
ROW_HEIGHT_TWIPS = 360


def _tbl_border_xml():
    """Full table border including inside grid lines."""
    el = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        child = OxmlElement(f"w:{side}")
        child.set(qn("w:val"), "single")
        child.set(qn("w:sz"), "8")
        child.set(qn("w:space"), "0")
        child.set(qn("w:color"), "000000")
        el.append(child)
    return el


def _set_cell_width(cell, twips: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(twips))
    tc_w.set(qn("w:type"), "dxa")
    tc_pr.append(tc_w)


def _set_row_height(row, twips: int):
    tr_pr = row._tr.get_or_add_trPr()
    tr_h = OxmlElement("w:trHeight")
    tr_h.set(qn("w:val"), str(twips))
    tr_h.set(qn("w:hRule"), "exact")
    tr_pr.append(tr_h)


def _write_cell_text(cell, text: str, *, align: str = "left", color=None, bold: bool = False):
    para = cell.paragraphs[0]
    para.clear()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT if align == "right" else WD_ALIGN_PARAGRAPH.LEFT

    run = para.add_run(text)
    run.font.name = FONT
    run.font.size = FONT_SIZE
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def write_lined_page(doc: Document, content_lines: list[str]) -> None:
    """
    Write one 25-line page as a bordered two-column table.

    Content longer than 25 lines is truncated. Callers must paginate first.
    """
    lines = list(content_lines[:LINES_PER_PAGE])
    while len(lines) < LINES_PER_PAGE:
        lines.append("")

    table = doc.add_table(rows=LINES_PER_PAGE, cols=2)
    table.autofit = False

    tbl_pr = table._tbl.tblPr
    tbl_pr.append(_tbl_border_xml())
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(TOTAL_TWIPS))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)

    for idx, text in enumerate(lines):
        row = table.rows[idx]
        _set_row_height(row, ROW_HEIGHT_TWIPS)

        left = row.cells[0]
        right = row.cells[1]
        _set_cell_width(left, NUM_COL_TWIPS)
        _set_cell_width(right, CONTENT_TWIPS)
        _write_cell_text(left, str(idx + 1), align="right", color=COLOR_GRAY)
        _write_cell_text(right, text, color=COLOR_BLK)


def paginate_lines(lines: list[str]) -> list[list[str]]:
    """Split a flat list of strings into 25-line pages."""
    if not lines:
        return [[]]
    return [lines[i:i + LINES_PER_PAGE] for i in range(0, len(lines), LINES_PER_PAGE)]
