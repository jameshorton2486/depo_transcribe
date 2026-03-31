"""
post_record.py
Post-record spellings colloquy block.
Spec Section 8

CRITICAL: Post-record spellings are AUTHORITATIVE.
They override all prior uses of the name in the transcript.
All prior instances must be corrected retroactively.
"""
from typing import List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from ..models import JobConfig, PostRecordSpelling

FONT = "Courier New"


def _c(doc, text, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(12)
    run.bold = bold


def _l(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(12)
    run.bold = bold


def derive_correct_spelling(letters_as_given: str) -> str:
    """
    Derive the probable correct spelling from a hyphenated
    letter sequence produced by Morson's Rule 157.

    Examples:
      B-A-L-D-E-R-A-S  → Balderas
      B-r-e-n-n-e-n    → Brennen
      T-O-V-A-R        → Tovar

    This is a SUGGESTION only. The scopist must verify
    against the audio before accepting.

    Returns empty string if input is not a valid
    hyphenated letter sequence.
    """
    parts = letters_as_given.split('-')
    if len(parts) < 3:
        return ''
    if not all(len(p) == 1 and p.isalpha() for p in parts):
        return ''
    return ''.join(parts).capitalize()


def write_post_record_section(
    doc: Document,
    spellings: List[PostRecordSpelling],
    job_config: JobConfig,
):
    """
    Write the post-record spellings colloquy block (Spec Section 8).
    Formatted as Speaker Label lines (not Q/A format).
    """
    if not spellings:
        return

    _l(doc, "_" * 65)
    _c(doc, "POST-RECORD SPELLINGS", bold=True)
    _c(doc, "(Court Reporter — confirmed on record)")
    _l(doc, "_" * 65)
    doc.add_paragraph()

    for prs in spellings:
        _l(doc, f"  Name:           {prs.name}")
        _l(doc, f"  Correct Form:   {prs.correct_spelling}")
        _l(doc, f"  Spelled on Record: {prs.letters_as_given}")
        if prs.flag:
            _l(doc, f"  NOTE: {prs.flag}")
        doc.add_paragraph()


def apply_retroactive_corrections(
    output_docx_path: str,
    spellings: List[PostRecordSpelling],
    job_config: JobConfig,
) -> List[str]:
    """
    After the main transcript is saved, re-open the DOCX and correct all
    prior uses of any name confirmed in post-record spellings.
    Spec Section 8: Post-record spellings are authoritative — override all prior uses.

    Returns list of corrections made for the Corrections Log.
    """
    import re
    from docx import Document as DocxDoc

    if not spellings:
        return []

    doc = DocxDoc(output_docx_path)
    corrections_made = []

    for prs in spellings:
        # Use correct_spelling if set; otherwise skip (scopist must confirm)
        if not prs.correct_spelling:
            continue
        if prs.correct_spelling == prs.name:
            continue
        pattern = re.compile(r'\b' + re.escape(prs.name) + r'\b', re.IGNORECASE)
        for para in doc.paragraphs:
            for run in para.runs:
                if pattern.search(run.text):
                    new_text = pattern.sub(prs.correct_spelling, run.text)
                    if new_text != run.text:
                        corrections_made.append(
                            f"Post-record: {run.text!r} → {new_text!r}"
                        )
                        run.text = new_text

    if corrections_made:
        doc.save(output_docx_path)

    return corrections_made
