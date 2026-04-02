"""
parser.py

Parses Deepgram speech-to-text .docx output into Block objects.
Spec Section 1: Input Parsing

Input format:
  Each paragraph is either:
  - A speaker label line: "Speaker N:"  (N = 0-9)
  - A text fragment belonging to the most recent speaker label

Spec 1.2 NOTE: Do NOT discard empty paragraphs before joining.
Strip only AFTER joining the full block.
"""

import re
from collections import Counter
from typing import Dict, List, Optional

from docx import Document

from .models import Block


SPEAKER_LABEL_RE = re.compile(r'^speaker\s+(\d+)\s*:\s*$', re.IGNORECASE)
SPEAKER_INLINE_RE = re.compile(r'^speaker\s+(\d+)\s*:\s*(.+)$', re.IGNORECASE)
FORMAT_DEEPGRAM = 'deepgram'
FORMAT_PLAIN_QA = 'plain_qa'
FORMAT_UNKNOWN  = 'unknown'


def detect_input_format(docx_path: str) -> str:
    """Auto-detect whether input is Deepgram-format or plain Q/A transcript."""
    doc = Document(docx_path)
    first_20 = [p.text.strip() for p in doc.paragraphs[:20] if p.text.strip()]
    if any(SPEAKER_LABEL_RE.match(t) or SPEAKER_INLINE_RE.match(t) for t in first_20):
        return FORMAT_DEEPGRAM
    qa_pattern = re.compile(r'^[QA][.:\-]\s*\S')
    if any(qa_pattern.match(t) for t in first_20):
        return FORMAT_PLAIN_QA
    return FORMAT_UNKNOWN


def parse_blocks(docx_path: str) -> List[Block]:
    """
    Parse a Deepgram output .docx file into a list of Block objects.
    Spec 1.1: Speaker label detection — matches "Speaker N:"
    Spec 1.2: Join ALL fragments then strip. Do not discard empty paragraphs.
    """
    doc = Document(docx_path)
    blocks: List[Block] = []
    current_speaker_id: Optional[int] = None
    current_fragments: List[str] = []
    saw_speaker_label = False
    preamble_lines: List[str] = []

    def flush_block() -> None:
        nonlocal current_speaker_id, current_fragments
        if current_speaker_id is not None and current_fragments:
            raw_text = ' '.join(current_fragments)
            text = re.sub(r' {2,}', ' ', raw_text).strip()
            if text:
                blocks.append(Block(
                    speaker_id=current_speaker_id,
                    text=text,
                    raw_text=raw_text,
                ))
        current_fragments = []

    for para in doc.paragraphs:
        line = para.text

        match = SPEAKER_LABEL_RE.match(line.strip())
        if match:
            saw_speaker_label = True
            flush_block()
            current_speaker_id = int(match.group(1))
            continue

        inline_match = SPEAKER_INLINE_RE.match(line.strip())
        if inline_match:
            saw_speaker_label = True
            flush_block()
            current_speaker_id = int(inline_match.group(1))
            current_fragments.append(inline_match.group(2))
            continue

        if current_speaker_id is not None:
            current_fragments.append(line)
        elif line.strip():
            preamble_lines.append(line.strip())

    flush_block()
    if saw_speaker_label and preamble_lines:
        sample = preamble_lines[0][:50]
        raise ValueError(f"Text found before any speaker label: '{sample}'")
    return blocks


def detect_speaker_map(blocks: List[Block]) -> Dict[int, str]:
    """
    Analyze first 30 blocks to suggest speaker ID assignments.
    Spec 1.3: Map MUST be verified by user before processing.
    """
    sample = blocks[:30]
    counts = Counter(b.speaker_id for b in sample)
    suggestions: Dict[int, str] = {}

    for sid in sorted(counts.keys()):
        speaker_text = ' '.join(b.text for b in sample if b.speaker_id == sid).lower()
        if 'i represent' in speaker_text or 'my name is' in speaker_text:
            suggestions[sid] = f"Speaker {sid}: LIKELY ATTORNEY ({counts[sid]} blocks)"
        elif 'raise your right hand' in speaker_text or 'you may proceed' in speaker_text:
            suggestions[sid] = f"Speaker {sid}: LIKELY COURT REPORTER ({counts[sid]} blocks)"
        elif counts[sid] == min(counts.values()) and counts[sid] <= 5:
            suggestions[sid] = f"Speaker {sid}: LIKELY VIDEOGRAPHER ({counts[sid]} blocks — lowest count)"
        elif counts[sid] == max(counts.values()):
            suggestions[sid] = f"Speaker {sid}: LIKELY WITNESS or ATTORNEY ({counts[sid]} blocks — highest)"
        else:
            suggestions[sid] = f"Speaker {sid}: UNKNOWN ({counts[sid]} blocks)"

    return suggestions


def show_speaker_preview(blocks: List[Block], max_blocks: int = 50) -> str:
    """Return formatted string of first line per speaker for verification."""
    lines = [
        "SPEAKER VERIFICATION — Confirm speaker map before processing",
        "=" * 65,
    ]
    seen: set = set()
    for block in blocks[:max_blocks]:
        if block.speaker_id not in seen:
            preview = block.text[:80] + ("..." if len(block.text) > 80 else "")
            lines.append(f"Speaker {block.speaker_id}:  {preview}")
            seen.add(block.speaker_id)
        if len(seen) >= 8:
            break
    lines.append("=" * 65)
    lines.append("Assign each Speaker ID to the correct role before clicking Process.")
    return '\n'.join(lines)
