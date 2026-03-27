"""DOCX merge coordination for UFM.

This module will assemble multiple rendered DOCX outputs into a single final
document while preserving document order and formatting rules.
"""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from .data_models import MergeResult


class DocxMerger:
    """Coordinates future DOCX merge operations."""

    def merge(self, source_paths: list[str], output_path: str) -> MergeResult:
        """
        Merge multiple DOCX files into one output document.

        Appends body content from each source in order, inserting a page break
        between source documents and skipping section properties from appended
        documents so the base document remains authoritative.
        """
        if not source_paths:
            raise ValueError("source_paths must not be empty.")

        missing = [p for p in source_paths if not Path(p).is_file()]
        if missing:
            raise ValueError(f"Source file(s) not found: {missing}")

        dest = Path(output_path)
        dest.parent.mkdir(parents=True, exist_ok=True)

        try:
            base_doc = Document(source_paths[0])
            for src_path in source_paths[1:]:
                page_break_para = base_doc.add_paragraph()
                page_break_para.add_run().add_break()

                src_doc = Document(src_path)
                for child in src_doc.element.body:
                    if child.tag == qn("w:sectPr"):
                        continue
                    base_doc.element.body.append(deepcopy(child))

            base_doc.save(str(dest))
        except Exception as exc:
            raise RuntimeError(
                f"DocxMerger.merge() failed - output={output_path}: {exc}"
            ) from exc

        return MergeResult(
            output_path=dest,
            merged_documents=tuple(Path(p) for p in source_paths),
        )
