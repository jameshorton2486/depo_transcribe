"""
ufm_engine/post_processor/format_box.py

Post-population formatter: takes a populated .docx and produces a
UFM-compliant .docx by wrapping each page in a format-box table chassis
with line numbers in the gutter.

CONTRACT:
- Pure structural transform. NEVER modifies text content.
- Idempotent: running on an already-chassis'd doc reproduces the same shape.
- text(input, after stripping chassis line-number digits) == text(output).
- 25 lines per page (UFM Fig 1: 28pt exact spacing × 25 lines fits 8.5×11
  with 1" top/bottom margins).

The harvester treats each top-level paragraph as one "line". Block-level
<w:sdt> conditional wrappers are flattened (their inner paragraphs are
extracted; the wrapper is dropped). Non-trivial wrap-aware line counting
is out of scope: keep paragraphs short or expect minor overflow on the
last page (acceptable per UFM §2.13).
"""
from __future__ import annotations

from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

LINES_PER_PAGE = 25
LINE_NUMBER_GUTTER_WIDTH = Inches(0.4)
FORMAT_BOX_WIDTH = Inches(6.5)
CHASSIS_MARKER = "ufm_chassis"


def apply_format_box(
    input_path: Path,
    output_path: Path,
    *,
    apply_line_numbers: bool = True,
    render_firm_footer: bool = True,
    firm_name: Optional[str] = None,
) -> None:
    """
    Wrap each page of `input_path` in a UFM format-box chassis and save
    to `output_path`. Pure layout transform — no text changes.
    """
    if not input_path.exists():
        raise FileNotFoundError(input_path)
    if render_firm_footer and not firm_name:
        raise ValueError("render_firm_footer requires firm_name")

    doc = Document(str(input_path))

    body_paragraphs = _harvest_body_paragraphs(doc)
    _clear_body(doc)
    _enforce_page_geometry(doc)

    pages = _chunk(body_paragraphs, LINES_PER_PAGE)
    if not pages:
        pages = [[]]

    for i, page in enumerate(pages):
        if i > 0:
            _add_page_break(doc)
        _build_chassis_table(doc, page, apply_line_numbers=apply_line_numbers)

    _set_firm_footer(doc, firm_name if render_firm_footer else None)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# Harvest: extract body paragraphs, descending into block sdts and chassis
# tables so idempotency holds.
# ---------------------------------------------------------------------------

def _harvest_body_paragraphs(doc: Document) -> List:
    """Return all body paragraph elements in document order.

    - Top-level <w:p> are taken as-is.
    - Block-level <w:sdt> conditional blocks are flattened (their <w:p>
      children are taken; the sdt wrapper is dropped).
    - Tables marked as chassis (or any table whose first cell looks like
      a line-number gutter) have only their right-cell paragraphs taken.
    - Other tables are skipped (they're not part of the depo content
      flow that the chassis is meant to wrap).
    """
    body = doc.element.body
    paragraphs: List = []

    for child in list(body):
        tag = child.tag
        if tag == qn("w:p"):
            paragraphs.append(child)
        elif tag == qn("w:sdt"):
            for p in child.iter(qn("w:p")):
                paragraphs.append(p)
        elif tag == qn("w:tbl"):
            if _is_chassis_table(child):
                paragraphs.extend(_paragraphs_from_chassis_right_cell(child))
            # non-chassis tables are dropped by design
        elif tag == qn("w:sectPr"):
            continue

    return paragraphs


def _is_chassis_table(tbl) -> bool:
    """Detect a chassis table by its caption marker or by structure."""
    for caption in tbl.iter(qn("w:tblCaption")):
        if caption.get(qn("w:val")) == CHASSIS_MARKER:
            return True
    rows = tbl.findall(qn("w:tr"))
    if len(rows) != 1:
        return False
    cells = rows[0].findall(qn("w:tc"))
    if len(cells) != 2:
        return False
    left_paras = cells[0].findall(qn("w:p"))
    texts = ["".join(t.text or "" for t in p.iter(qn("w:t"))).strip() for p in left_paras]
    nums = [t for t in texts if t]
    return len(nums) >= 1 and all(t.isdigit() for t in nums)


def _paragraphs_from_chassis_right_cell(tbl) -> List:
    rows = tbl.findall(qn("w:tr"))
    if not rows:
        return []
    cells = rows[0].findall(qn("w:tc"))
    if len(cells) < 2:
        return []
    right = cells[-1]
    return list(right.findall(qn("w:p")))


# ---------------------------------------------------------------------------
# Body management
# ---------------------------------------------------------------------------

def _clear_body(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _enforce_page_geometry(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(0.5)


def _add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def _chunk(items: List, n: int) -> List[List]:
    return [items[i:i + n] for i in range(0, len(items), n)]


# ---------------------------------------------------------------------------
# Chassis table construction
# ---------------------------------------------------------------------------

def _build_chassis_table(doc: Document, page_paragraphs: List,
                          *, apply_line_numbers: bool) -> None:
    """Build a 1×2 table holding line-numbers (left) and body (right)."""
    tbl = OxmlElement("w:tbl")

    tbl_pr = OxmlElement("w:tblPr")
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), "0")
    tbl_w.set(qn("w:type"), "auto")
    tbl_pr.append(tbl_w)
    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_pr.append(tbl_layout)

    caption = OxmlElement("w:tblCaption")
    caption.set(qn("w:val"), CHASSIS_MARKER)
    tbl_pr.append(caption)
    tbl.append(tbl_pr)

    grid = OxmlElement("w:tblGrid")
    for w in (LINE_NUMBER_GUTTER_WIDTH, FORMAT_BOX_WIDTH):
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(w / Twips(1))))
        grid.append(gc)
    tbl.append(grid)

    tr = OxmlElement("w:tr")

    left_cell = _make_cell(LINE_NUMBER_GUTTER_WIDTH, bordered=False)
    if apply_line_numbers:
        for n in range(1, LINES_PER_PAGE + 1):
            left_cell.append(_line_number_paragraph(n))
    else:
        left_cell.append(_empty_paragraph())
    tr.append(left_cell)

    right_cell = _make_cell(FORMAT_BOX_WIDTH, bordered=True)
    if not page_paragraphs:
        right_cell.append(_empty_paragraph())
    else:
        for p in page_paragraphs:
            parent = p.getparent()
            if parent is not None:
                parent.remove(p)
            right_cell.append(p)
    tr.append(right_cell)

    tbl.append(tr)

    body = doc.element.body
    sect_pr = body.find(qn("w:sectPr"))
    if sect_pr is not None:
        body.insert(list(body).index(sect_pr), tbl)
    else:
        body.append(tbl)


def _make_cell(width, *, bordered: bool):
    tc = OxmlElement("w:tc")
    tc_pr = OxmlElement("w:tcPr")

    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(int(width / Twips(1))))
    tc_w.set(qn("w:type"), "dxa")
    tc_pr.append(tc_w)

    tc_mar = OxmlElement("w:tcMar")
    for side in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:w"), "0")
        e.set(qn("w:type"), "dxa")
        tc_mar.append(e)
    tc_pr.append(tc_mar)

    tc_borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{side}")
        if bordered:
            e.set(qn("w:val"), "single")
            e.set(qn("w:sz"), "4")
            e.set(qn("w:space"), "0")
            e.set(qn("w:color"), "000000")
        else:
            e.set(qn("w:val"), "nil")
        tc_borders.append(e)
    tc_pr.append(tc_borders)

    tc.append(tc_pr)
    return tc


def _line_number_paragraph(n: int):
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), "560")  # 28pt = 560 twentieths-of-a-point
    spacing.set(qn("w:lineRule"), "exact")
    pPr.append(spacing)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "right")
    pPr.append(jc)
    p.append(pPr)

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Courier New")
    rFonts.set(qn("w:hAnsi"), "Courier New")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")
    rPr.append(sz)
    r.append(rPr)

    t = OxmlElement("w:t")
    t.text = str(n)
    r.append(t)
    p.append(r)
    return p


def _empty_paragraph():
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), "560")
    spacing.set(qn("w:lineRule"), "exact")
    pPr.append(spacing)
    p.append(pPr)
    return p


# ---------------------------------------------------------------------------
# Footer
# ---------------------------------------------------------------------------

def _set_firm_footer(doc: Document, firm_name: Optional[str]) -> None:
    """Set or clear the centered firm-name footer on every section."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        for p in list(footer.paragraphs):
            p._element.getparent().remove(p._element)
        p = footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if firm_name:
            run = p.add_run(firm_name)
            run.font.name = "Courier New"
            run.font.size = Pt(12)
