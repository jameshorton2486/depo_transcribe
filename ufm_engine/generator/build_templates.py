"""
ufm_engine/generator/build_templates.py

Generates all UFM-compliant deposition templates as .docx files with
content controls and conditional blocks. Each template ships WITHOUT
format box and WITHOUT line numbers — those are applied by the
ufm_engine.post_processor at finish time (Pipeline B).

Run from project root:
    python -m ufm_engine.generator.build_templates

Output: ufm_engine/templates/figures/*.docx (13 files)
"""
from __future__ import annotations

from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUTPUT_DIR = Path(__file__).resolve().parents[1] / "templates" / "figures"


# ---------------------------------------------------------------------------
# Helpers — page setup, content controls, conditional blocks
# ---------------------------------------------------------------------------

def _new_document() -> Document:
    """Create a blank document with UFM page geometry."""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(0.5)

    style = doc.styles["Normal"]
    style.font.name = "Courier New"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    style.paragraph_format.line_spacing = Pt(28)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    return doc


def _add_paragraph(doc: Document, text: str = "", *,
                   align: Optional[int] = None,
                   bold: bool = False) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        run.font.name = "Courier New"
        run.font.size = Pt(12)
        if bold:
            run.bold = True


def _make_content_control_sdt(tag: str, alias: str,
                               placeholder: Optional[str] = None):
    """Build a Plain Text Content Control sdt element. Caller appends it.

    Used by _add_content_control (paragraph append) and by
    _add_inline_field_to_content (sdtContent append, for content
    controls nested inside inline conditional blocks).
    """
    if placeholder is None:
        placeholder = f"[{alias}]"

    sdt = OxmlElement("w:sdt")
    sdt_pr = OxmlElement("w:sdtPr")

    tag_el = OxmlElement("w:tag")
    tag_el.set(qn("w:val"), tag)
    sdt_pr.append(tag_el)

    alias_el = OxmlElement("w:alias")
    alias_el.set(qn("w:val"), alias)
    sdt_pr.append(alias_el)

    showing_plchdr = OxmlElement("w:showingPlcHdr")
    sdt_pr.append(showing_plchdr)

    sdt.append(sdt_pr)

    sdt_content = OxmlElement("w:sdtContent")
    sdt_content.append(_make_courier_run(placeholder))
    sdt.append(sdt_content)
    return sdt


def _make_courier_run(text: str):
    """Build a Courier New 12pt <w:r> element with the given text."""
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Courier New")
    rfonts.set(qn("w:hAnsi"), "Courier New")
    rpr.append(rfonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")  # 24 half-points = 12pt
    rpr.append(sz)
    r.append(rpr)

    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _add_content_control(paragraph, tag: str, alias: str,
                         placeholder: Optional[str] = None) -> None:
    """Append a Plain Text Content Control to the given paragraph."""
    paragraph._p.append(_make_content_control_sdt(tag, alias, placeholder))


def _add_field(doc: Document, tag: str, alias: str,
               *, prefix: str = "", suffix: str = "",
               align: Optional[int] = None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if prefix:
        run = p.add_run(prefix)
        run.font.name = "Courier New"
        run.font.size = Pt(12)
    _add_content_control(p, tag, alias)
    if suffix:
        run = p.add_run(suffix)
        run.font.name = "Courier New"
        run.font.size = Pt(12)


def _add_inline_field(paragraph, tag: str, alias: str) -> None:
    _add_content_control(paragraph, tag, alias)


def _open_inline_block(paragraph, block_tag: str):
    """Append an empty inline conditional sdt and return its sdtContent.

    The caller adds runs and/or nested content controls to the returned
    sdtContent element. Used for inline conditional phrases that may
    need to contain a content control (e.g., ", via [Remote Platform]"
    where the entire phrase disappears when the toggle is False).
    """
    sdt = OxmlElement("w:sdt")
    sdt_pr = OxmlElement("w:sdtPr")

    tag_el = OxmlElement("w:tag")
    tag_el.set(qn("w:val"), block_tag)
    sdt_pr.append(tag_el)

    alias_el = OxmlElement("w:alias")
    alias_el.set(qn("w:val"), f"Conditional Block: {block_tag}")
    sdt_pr.append(alias_el)

    sdt.append(sdt_pr)

    sdt_content = OxmlElement("w:sdtContent")
    sdt.append(sdt_content)

    paragraph._p.append(sdt)
    return sdt_content


def _add_inline_block(paragraph, block_tag: str, text: str) -> None:
    """Wrap a literal text run in an inline conditional sdt.

    Convenience wrapper around _open_inline_block for the common case
    where the conditional content is a single literal phrase.
    """
    content = _open_inline_block(paragraph, block_tag)
    content.append(_make_courier_run(text))


def _wrap_in_block_sdt(doc: Document, block_tag: str,
                       paragraph_count: int) -> None:
    """
    Wrap the last `paragraph_count` body paragraphs in a block-level sdt
    with the given tag. The sdt is inserted at the position where the
    first wrapped paragraph currently lives — preserving document order.
    """
    body = doc.element.body
    paragraphs = body.findall(qn("w:p"))
    if len(paragraphs) < paragraph_count:
        raise ValueError(
            f"Cannot wrap {paragraph_count} paragraphs in block "
            f"{block_tag}: only {len(paragraphs)} exist"
        )
    target_paragraphs = paragraphs[-paragraph_count:]

    body_children = list(body)
    insert_index = body_children.index(target_paragraphs[0])

    sdt = OxmlElement("w:sdt")
    sdt_pr = OxmlElement("w:sdtPr")

    tag_el = OxmlElement("w:tag")
    tag_el.set(qn("w:val"), block_tag)
    sdt_pr.append(tag_el)

    alias_el = OxmlElement("w:alias")
    alias_el.set(qn("w:val"), f"Conditional Block: {block_tag}")
    sdt_pr.append(alias_el)

    sdt.append(sdt_pr)

    sdt_content = OxmlElement("w:sdtContent")
    for p in target_paragraphs:
        body.remove(p)
        sdt_content.append(p)
    sdt.append(sdt_content)

    body.insert(insert_index, sdt)


def _save(doc: Document, filename: str) -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    out = OUTPUT_DIR / filename
    doc.save(out)
    print(f"Wrote {out}")


# ---------------------------------------------------------------------------
# Caption blocks
# ---------------------------------------------------------------------------

def _add_caption_tx_state(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CAUSE NO. ")
    run.font.name = "Courier New"
    run.font.size = Pt(12)
    _add_content_control(p, "cause_number", "Cause Number")

    p = doc.add_paragraph()
    _add_content_control(p, "plaintiff_name", "Plaintiff Name(s)")
    p.add_run(",\t)\tIN THE ")
    _add_content_control(p, "court_designation", "Court Designation")

    _add_paragraph(doc, "    Plaintiff,                 )")

    p = doc.add_paragraph()
    p.add_run("VS.\t)\t")
    _add_content_control(p, "county", "County")
    p.add_run(" COUNTY, ")
    _add_content_control(p, "state", "State")

    p = doc.add_paragraph()
    _add_content_control(p, "defendant_names_block", "Defendant Names")
    p.add_run("\t)\t")
    _add_content_control(p, "judicial_district_phrase", "Judicial District Phrase")

    _add_paragraph(doc, "    Defendant(s).              )")


def _add_caption_federal(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("IN THE UNITED STATES DISTRICT COURT")
    run.font.name = "Courier New"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("FOR THE ")
    _add_content_control(p, "federal_district_name", "Federal District Name")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_content_control(p, "federal_division", "Division Name")
    p.add_run(" DIVISION")

    p = doc.add_paragraph()
    _add_content_control(p, "plaintiff_name", "Plaintiff Name(s)")
    p.add_run(",\t)")

    _add_paragraph(doc, "    Plaintiff,                 )")

    p = doc.add_paragraph()
    p.add_run("VS.\t)\tCIVIL ACTION NO.")

    p = doc.add_paragraph()
    p.add_run("\t)\t")
    _add_content_control(p, "federal_civil_action_number", "Civil Action Number")

    p = doc.add_paragraph()
    _add_content_control(p, "defendant_names_block", "Defendant Names")
    p.add_run(",\t)")

    _add_paragraph(doc, "    Defendants.                )")


SEPARATOR = "*" * 56


# ===========================================================================
# Template builders
# ===========================================================================

def build_title_page_tx_state() -> None:
    doc = _new_document()
    _add_caption_tx_state(doc)
    _add_paragraph(doc, "")
    _add_paragraph(doc, SEPARATOR, align=WD_ALIGN_PARAGRAPH.CENTER)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("ORAL ")
    _add_inline_block(p, "block_videotaped", "AND VIDEOTAPED ")
    p.add_run("DEPOSITION ")
    _add_inline_block(p, "block_subpoena_duces_tecum", "WITH SUBPOENA DUCES TECUM ")
    p.add_run("OF")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_content_control(p, "witness_name", "Witness Name")

    _add_field(doc, "depo_date", "Deposition Date", align=WD_ALIGN_PARAGRAPH.CENTER)

    # Volume line: "VOLUME N OF M". Both volume_number and total_volumes
    # are content controls; the entire paragraph is wrapped in
    # block_volume so the toggle removes the line outright when there's
    # only one volume.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("VOLUME ")
    _add_inline_field(p, "volume_number", "Volume Number")
    p.add_run(" OF ")
    _add_inline_field(p, "total_volumes", "Total Volumes")
    _wrap_in_block_sdt(doc, "block_volume", 1)

    _add_paragraph(doc, SEPARATOR, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "")

    p = doc.add_paragraph()
    p.add_run("ORAL ")
    _add_inline_block(p, "block_videotaped", "AND VIDEOTAPED ")
    p.add_run("DEPOSITION OF ")
    _add_inline_field(p, "witness_name", "Witness Name")
    p.add_run(", produced as a witness at the instance of the ")
    _add_inline_field(p, "instance_party", "Party (Plaintiff/Defendant)")
    p.add_run(", and duly sworn, was taken in the above-styled and numbered cause on the ")
    _add_inline_field(p, "depo_day_ordinal", "Day Ordinal (e.g., 9th day of April, 2026)")
    p.add_run(", from ")
    _add_inline_field(p, "depo_time_start", "Start Time")
    p.add_run(" to ")
    _add_inline_field(p, "depo_time_end", "End Time")
    # Remote platform notation — entire ", via [PLATFORM]" segment is
    # conditional on block_remote so non-remote depos don't render an
    # awkward ", via ," with a missing platform.
    remote_block = _open_inline_block(p, "block_remote")
    remote_block.append(_make_courier_run(", via "))
    remote_block.append(_make_content_control_sdt("remote_platform", "Remote Platform"))
    p.add_run(", before ")
    _add_inline_field(p, "reporter_name", "Reporter Name")
    p.add_run(", CSR in and for the State of Texas, reported by ")
    _add_inline_field(p, "method", "Method")
    p.add_run(", at ")
    _add_inline_field(p, "depo_location_full", "Deposition Location (full)")
    p.add_run(", pursuant to the Texas Rules of Civil Procedure and the provisions stated on the record or attached hereto.")

    _save(doc, "title_page_tx_state.docx")


def build_title_page_federal() -> None:
    doc = _new_document()
    _add_caption_federal(doc)
    _add_paragraph(doc, "")
    _add_paragraph(doc, SEPARATOR, align=WD_ALIGN_PARAGRAPH.CENTER)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("ORAL ")
    _add_inline_block(p, "block_videotaped", "AND VIDEOTAPED ")
    p.add_run("DEPOSITION OF")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_content_control(p, "witness_name", "Witness Name")

    _add_field(doc, "depo_date", "Deposition Date", align=WD_ALIGN_PARAGRAPH.CENTER)

    # Optional interpreted notation, paragraph-level so the populator
    # can drop the whole line when the depo is not interpreted.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("(INTERPRETED FROM ")
    _add_inline_field(p, "interpreter_language", "Interpreter Language")
    p.add_run(" TO ENGLISH)")
    _wrap_in_block_sdt(doc, "block_interpreted", 1)

    _add_paragraph(doc, SEPARATOR, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "")

    p = doc.add_paragraph()
    p.add_run("ORAL ")
    _add_inline_block(p, "block_videotaped", "AND VIDEOTAPED ")
    p.add_run("DEPOSITION OF ")
    _add_inline_field(p, "witness_name", "Witness Name")
    p.add_run(", produced as a witness at the instance of the ")
    _add_inline_field(p, "instance_party", "Party")
    p.add_run(", and duly sworn, was taken in the above-styled and numbered cause on the ")
    _add_inline_field(p, "depo_day_ordinal", "Day Ordinal")
    p.add_run(", from ")
    _add_inline_field(p, "depo_time_start", "Start Time")
    p.add_run(" to ")
    _add_inline_field(p, "depo_time_end", "End Time")
    remote_block = _open_inline_block(p, "block_remote")
    remote_block.append(_make_courier_run(", via "))
    remote_block.append(_make_content_control_sdt("remote_platform", "Remote Platform"))
    p.add_run(", before ")
    _add_inline_field(p, "reporter_name", "Reporter Name")
    p.add_run(", CSR in and for the State of Texas, reported by ")
    _add_inline_field(p, "method", "Method")
    p.add_run(", at ")
    _add_inline_field(p, "depo_location_full", "Deposition Location")
    p.add_run(", pursuant to the Federal Rules of Civil Procedure and the provisions stated on the record or attached hereto.")

    _save(doc, "title_page_federal.docx")


def build_appearances() -> None:
    doc = _new_document()
    _add_paragraph(doc, "A P P E A R A N C E S",
                   align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    _add_paragraph(doc, "")

    p = doc.add_paragraph()
    p.add_run("FOR THE PLAINTIFF(S): ")
    _add_inline_field(p, "plaintiff_name", "Plaintiff Name")

    _add_field(doc, "plaintiff_counsel_1_name", "Plaintiff Counsel Name")
    _add_field(doc, "plaintiff_counsel_1_firm", "Firm Name")
    _add_field(doc, "plaintiff_counsel_1_address", "Address")
    _add_field(doc, "plaintiff_counsel_1_csz", "City, State ZIP")
    _add_field(doc, "plaintiff_counsel_1_phone", "Phone", prefix="Phone: ")
    _add_field(doc, "plaintiff_counsel_1_email", "Email", prefix="Email: ")

    _add_paragraph(doc, "")

    p = doc.add_paragraph()
    p.add_run("FOR THE DEFENDANT(S): ")
    _add_inline_field(p, "defendant_name", "Defendant Name")

    _add_field(doc, "defendant_counsel_1_name", "Defendant Counsel Name")
    _add_field(doc, "defendant_counsel_1_firm", "Firm Name")
    _add_field(doc, "defendant_counsel_1_address", "Address")
    _add_field(doc, "defendant_counsel_1_csz", "City, State ZIP")
    _add_field(doc, "defendant_counsel_1_phone", "Phone", prefix="Phone: ")
    _add_field(doc, "defendant_counsel_1_email", "Email", prefix="Email: ")

    _add_paragraph(doc, "")
    _add_paragraph(doc, "ALSO PRESENT:")
    _add_field(doc, "also_present_1", "Also Present (name and role)")
    _wrap_in_block_sdt(doc, "block_also_present", 2)

    _save(doc, "appearances.docx")


def build_index_chronological() -> None:
    doc = _new_document()
    _add_paragraph(doc, "I N D E X",
                   align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    _add_paragraph(doc, "")
    _add_paragraph(doc, "EXAMINATIONS                                        PAGE",
                   bold=True)
    _add_paragraph(doc, "")

    p = doc.add_paragraph()
    p.add_run("    A P P E A R A N C E S                       ")
    _add_inline_field(p, "appearances_page", "Appearances Page")

    _add_paragraph(doc, "")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_content_control(p, "witness_name", "Witness Name")
    p.add_run(",")

    p = doc.add_paragraph()
    p.add_run("    EXAMINATION BY ")
    _add_inline_field(p, "examiner_1", "First Examiner")
    p.add_run("                ")
    _add_inline_field(p, "examination_1_page", "Page")

    _add_paragraph(doc, "")
    _add_paragraph(doc, "E X H I B I T S",
                   align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    _add_paragraph(doc, "")
    _add_paragraph(doc, "Exhibit No.        Description                      Page",
                   bold=True)
    _add_paragraph(doc, "")

    p = doc.add_paragraph()
    _add_inline_field(p, "exhibit_1_number", "Exhibit Number")
    p.add_run("            ")
    _add_inline_field(p, "exhibit_1_desc", "Description")
    p.add_run("                       ")
    _add_inline_field(p, "exhibit_1_page", "Page")

    _save(doc, "index_chronological.docx")


def build_witness_setup_standard() -> None:
    doc = _new_document()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_content_control(p, "witness_name", "Witness Name (caps)")
    p.add_run(",")

    _add_paragraph(doc, "having been first duly sworn, testified as follows:")
    _add_paragraph(doc, "EXAMINATION", align=WD_ALIGN_PARAGRAPH.CENTER)

    p = doc.add_paragraph()
    p.add_run("BY ")
    _add_inline_field(p, "examining_lawyer_label", "Examining Lawyer")
    p.add_run(":")

    _save(doc, "witness_setup_standard.docx")


def build_witness_setup_interpreter() -> None:
    doc = _new_document()
    _add_paragraph(doc, "(INTERPRETER SWORN)", align=WD_ALIGN_PARAGRAPH.CENTER)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_content_control(p, "witness_name", "Witness Name (caps)")
    p.add_run(",")

    _add_paragraph(doc, "having been first duly sworn, testified through the duly sworn interpreter as follows:")
    _add_paragraph(doc, "EXAMINATION", align=WD_ALIGN_PARAGRAPH.CENTER)

    p = doc.add_paragraph()
    p.add_run("BY ")
    _add_inline_field(p, "examining_lawyer_label", "Examining Lawyer")
    p.add_run(":")

    _save(doc, "witness_setup_interpreter.docx")


def build_changes_signature_grid() -> None:
    doc = _new_document()
    _add_paragraph(doc, "CHANGES AND SIGNATURE",
                   align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    _add_paragraph(doc, "")

    p = doc.add_paragraph()
    p.add_run("WITNESS NAME: ")
    _add_inline_field(p, "witness_name", "Witness Name")
    p.add_run("        DATE OF DEPOSITION: ")
    _add_inline_field(p, "depo_date", "Deposition Date")

    _add_paragraph(doc, "")
    _add_paragraph(doc, "PAGE            LINE            CHANGE          REASON",
                   bold=True)

    for _ in range(25):
        _add_paragraph(doc, "_______________________________________________________")

    _save(doc, "changes_signature_grid.docx")


def build_witness_acknowledgment_notary() -> None:
    doc = _new_document()
    p = doc.add_paragraph()
    p.add_run("    I, ")
    _add_inline_field(p, "witness_name", "Witness Name")
    p.add_run(", have read the foregoing deposition and hereby affix my signature that same is true and correct, except as noted above.")

    _add_paragraph(doc, "")
    _add_paragraph(doc, "_______________________________________",
                   align=WD_ALIGN_PARAGRAPH.CENTER)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_content_control(p, "witness_name", "Witness Name")

    _add_paragraph(doc, "")
    _add_paragraph(doc, "")

    p = doc.add_paragraph()
    p.add_run("THE STATE OF __________)")
    p = doc.add_paragraph()
    p.add_run("COUNTY OF _____________)")

    _add_paragraph(doc, "")

    p = doc.add_paragraph()
    p.add_run("    Before me, ___________________________, on this day personally appeared ")
    _add_inline_field(p, "witness_name", "Witness Name")
    p.add_run(", known to me (or proved to me under oath or through ___________________________) (description of identity card or other document) to be the person whose name is subscribed to the foregoing instrument and acknowledged to me that they executed the same for the purposes and consideration therein expressed.")

    _add_paragraph(doc, "")

    p = doc.add_paragraph()
    p.add_run("    Given under my hand and seal of office this __________ day of ________________________, __________.")

    _add_paragraph(doc, "")
    _add_paragraph(doc, "")

    p = doc.add_paragraph()
    p.add_run("                NOTARY PUBLIC IN AND FOR")
    p = doc.add_paragraph()
    p.add_run("                THE STATE OF ______________________")

    _save(doc, "witness_acknowledgment_notary.docx")


# ---------------------------------------------------------------------------
# Reporter signature block — reused across cert templates
# ---------------------------------------------------------------------------

def _add_reporter_signature_block(doc: Document) -> None:
    _add_paragraph(doc, "")
    _add_paragraph(doc, "_______________________________________",
                   align=WD_ALIGN_PARAGRAPH.CENTER)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_content_control(p, "reporter_name", "Reporter Name")
    p.add_run(", ")
    _add_inline_field(p, "credentials", "Credentials")
    _wrap_in_block_sdt(doc, "block_credentials_suffix", 1)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Texas CSR ")
    _add_inline_field(p, "csr_number", "CSR Number")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Expiration Date: ")
    _add_inline_field(p, "csr_expiration", "CSR Expiration")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_content_control(p, "firm_name", "Firm Name")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Firm Registration No. ")
    _add_inline_field(p, "firm_registration_number", "Firm Reg No.")

    _wrap_in_block_sdt(doc, "block_firm_signature_block", 2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_content_control(p, "reporter_address_line1", "Address Line 1")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_content_control(p, "reporter_city", "City")
    p.add_run(", ")
    _add_inline_field(p, "reporter_state", "State")
    p.add_run(" ")
    _add_inline_field(p, "reporter_zip", "ZIP")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Phone: ")
    _add_inline_field(p, "reporter_phone", "Phone")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Email: ")
    _add_inline_field(p, "reporter_email", "Email")


# ---------------------------------------------------------------------------
# Certification templates
# ---------------------------------------------------------------------------

def build_cert_tx_sig_required() -> None:
    doc = _new_document()
    _add_caption_tx_state(doc)
    _add_paragraph(doc, "")
    _add_paragraph(doc, SEPARATOR, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "REPORTER'S CERTIFICATION",
                   align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("DEPOSITION OF ")
    _add_inline_field(p, "witness_name", "Witness Name")

    _add_field(doc, "depo_date", "Deposition Date", align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "")

    p = doc.add_paragraph()
    p.add_run("    I, ")
    _add_inline_field(p, "reporter_name", "Reporter Name")
    p.add_run(", Certified Shorthand Reporter in and for the State of Texas, hereby certify to the following:")

    p = doc.add_paragraph()
    p.add_run("    That the witness, ")
    _add_inline_field(p, "witness_name", "Witness Name")
    p.add_run(", was duly sworn by the officer and that the transcript of the oral deposition is a true record of the testimony given by the witness;")

    p = doc.add_paragraph()
    p.add_run("    That the deposition transcript was submitted on ")
    _add_inline_field(p, "transcript_submitted_date", "Submitted Date")
    p.add_run(" to the witness or to the attorney for the witness for examination, signature and return to me by ")
    _add_inline_field(p, "transcript_returned_date", "Return-By Date")
    p.add_run(";")

    _add_paragraph(doc, "    That the amount of time used by each party at the deposition is as follows:")
    _add_field(doc, "time_used_block", "Time Used Per Attorney (multi-line)")

    _add_paragraph(doc, "    That pursuant to information given to the deposition officer at the time said testimony was taken, the following includes counsel for all parties of record:")
    _add_field(doc, "attorney_party_pairs_block", "Attorney/Party Pairs (multi-line)")

    _add_paragraph(doc, "    I further certify that I am neither counsel for, related to, nor employed by any of the parties or attorneys in the action in which this proceeding was taken, and further that I am not financially or otherwise interested in the outcome of the action.")

    _add_paragraph(doc, "    Further certification requirements pursuant to Rule 203 of TRCP will be certified to after they have occurred.")

    p = doc.add_paragraph()
    p.add_run("    Certified to by me this ")
    _add_inline_field(p, "certification_date", "Certification Date")
    p.add_run(".")

    _add_reporter_signature_block(doc)
    _save(doc, "cert_tx_sig_required.docx")


def build_cert_tx_sig_waived() -> None:
    doc = _new_document()
    _add_caption_tx_state(doc)
    _add_paragraph(doc, "")
    _add_paragraph(doc, SEPARATOR, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "REPORTER'S CERTIFICATION",
                   align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("DEPOSITION OF ")
    _add_inline_field(p, "witness_name", "Witness Name")

    _add_field(doc, "depo_date", "Deposition Date", align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "")

    p = doc.add_paragraph()
    p.add_run("    I, ")
    _add_inline_field(p, "reporter_name", "Reporter Name")
    p.add_run(", Certified Shorthand Reporter in and for the State of Texas, hereby certify to the following:")

    p = doc.add_paragraph()
    p.add_run("    That the witness, ")
    _add_inline_field(p, "witness_name", "Witness Name")
    p.add_run(", was duly sworn by the officer and that the transcript of the oral deposition is a true record of the testimony given by the witness;")

    _add_paragraph(doc, "    That examination and signature of the witness to the deposition transcript was waived by the witness and agreement of the parties at the time of the deposition;")

    p = doc.add_paragraph()
    p.add_run("    That the original deposition was delivered to ")
    _add_inline_field(p, "custodial_attorney_name", "Custodial Attorney")
    p.add_run(";")
    _wrap_in_block_sdt(doc, "block_custodial_attorney", 1)

    _add_paragraph(doc, "    That the amount of time used by each party at the deposition is as follows:")
    _add_field(doc, "time_used_block", "Time Used Per Attorney")

    p = doc.add_paragraph()
    p.add_run("    That $")
    _add_inline_field(p, "cost_amount", "Cost Amount")
    p.add_run(" is the deposition officer's charges to the ")
    _add_inline_field(p, "cost_payor_party", "Cost Payor Party")
    p.add_run(" for preparing the original deposition transcript and any copies of exhibits;")
    _wrap_in_block_sdt(doc, "block_cost_paragraph", 1)

    _add_paragraph(doc, "    That pursuant to information given to the deposition officer at the time said testimony was taken, the following includes counsel for all parties of record:")
    _add_field(doc, "attorney_party_pairs_block", "Attorney/Party Pairs")

    p = doc.add_paragraph()
    p.add_run("    That a copy of this certificate was served on all parties shown herein on ")
    _add_inline_field(p, "served_on_date", "Served-On Date")
    p.add_run(" and filed with the Clerk pursuant to Rule 203.3.")

    _add_paragraph(doc, "    I further certify that I am neither counsel for, related to, nor employed by any of the parties or attorneys in the action in which this proceeding was taken, and further that I am not financially or otherwise interested in the outcome of the action.")

    p = doc.add_paragraph()
    p.add_run("    Certified to by me this ")
    _add_inline_field(p, "certification_date", "Certification Date")
    p.add_run(".")

    _add_reporter_signature_block(doc)
    _save(doc, "cert_tx_sig_waived.docx")


def build_cert_federal_frcp() -> None:
    doc = _new_document()
    _add_caption_federal(doc)
    _add_paragraph(doc, "")
    _add_paragraph(doc, "REPORTER'S CERTIFICATE",
                   align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("ORAL DEPOSITION OF ")
    _add_inline_field(p, "witness_name", "Witness Name")

    _add_field(doc, "depo_date", "Deposition Date", align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "")

    p = doc.add_paragraph()
    p.add_run("    I, ")
    _add_inline_field(p, "reporter_name", "Reporter Name")
    p.add_run(", Certified Shorthand Reporter in and for the State of Texas, hereby certify to the following:")

    p = doc.add_paragraph()
    p.add_run("    That the witness, ")
    _add_inline_field(p, "witness_name", "Witness Name")
    p.add_run(", was duly sworn by the officer and that the transcript of the deposition is a true record of the testimony given by the witness;")

    p = doc.add_paragraph()
    p.add_run("    That the original deposition transcript was delivered to ")
    _add_inline_field(p, "custodial_attorney_name", "Custodial Attorney")
    p.add_run(";")
    _wrap_in_block_sdt(doc, "block_custodial_attorney", 1)

    p = doc.add_paragraph()
    p.add_run("    That a copy of this certificate was served on all parties and/or the witness shown herein on ")
    _add_inline_field(p, "served_on_date", "Served-On Date")
    p.add_run(".")

    _add_paragraph(doc, "    I further certify that, pursuant to FRCP No. 30(f)(i), that the signature of the deponent was requested by the deponent or a party before the completion of the deposition and that the signature is to be returned within 30 days from date of receipt of the transcript. If returned, the attached Changes and Signature page contains any changes and the reasons therefor;")

    _add_paragraph(doc, "    That pursuant to information given to the deposition officer at the time said testimony was taken, the following includes counsel for all parties of record and the amount of time used by each party at the time of the deposition:")
    _add_field(doc, "attorney_party_pairs_with_time_block", "Attorney/Party Pairs with Time")

    _add_paragraph(doc, "    I further certify that I am neither attorney or counsel for, related to, nor employed by any parties to the action in which this testimony is taken and, further, that I am not a relative or employee of any counsel employed by the parties hereto or financially interested in the action.")

    p = doc.add_paragraph()
    p.add_run("    SUBSCRIBED AND SWORN TO under my hand and seal of office on this the ")
    _add_inline_field(p, "certification_date", "Certification Date")
    p.add_run(".")

    _add_reporter_signature_block(doc)
    _save(doc, "cert_federal_frcp.docx")


def build_cert_nonappearance() -> None:
    doc = _new_document()
    _add_caption_tx_state(doc)
    _add_paragraph(doc, "")
    _add_paragraph(doc, SEPARATOR, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "CERTIFICATE OF NONAPPEARANCE",
                   align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("FOR THE ")
    _add_inline_field(p, "depo_modality", "Depo Modality (e.g., VIDEOCONFERENCE)")
    p.add_run(" DEPOSITION OF")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_content_control(p, "witness_name", "Witness Name")

    _add_field(doc, "depo_date", "Deposition Date", align=WD_ALIGN_PARAGRAPH.CENTER)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("(")
    _add_inline_field(p, "remote_modality_label", "Remote Modality Label (e.g., Reported Remotely)")
    p.add_run(")")
    _wrap_in_block_sdt(doc, "block_remote", 1)

    _add_paragraph(doc, SEPARATOR, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "")

    p = doc.add_paragraph()
    p.add_run("    I, ")
    _add_inline_field(p, "reporter_name", "Reporter Name")
    p.add_run(", a Certified Shorthand Reporter in and for the State of Texas, certify:")

    p = doc.add_paragraph()
    p.add_run("    That I appeared ")
    _add_inline_field(p, "appearance_modality", "Appearance Modality")
    p.add_run(", on the ")
    _add_inline_field(p, "depo_day_ordinal", "Day Ordinal")
    p.add_run(", to report the ")
    _add_inline_field(p, "depo_modality_lc", "Depo Modality (lowercase)")
    p.add_run(" deposition of ")
    _add_inline_field(p, "witness_name", "Witness Name")
    p.add_run(" pursuant to Notice, scheduled for ")
    _add_inline_field(p, "depo_time_start", "Scheduled Time")
    p.add_run(";")

    _add_paragraph(doc, "    That present for the deposition were:")
    _add_field(doc, "attorney_party_pairs_block", "Attorney/Party Pairs")

    p = doc.add_paragraph()
    p.add_run("    That by ")
    _add_inline_field(p, "no_show_time", "No-Show Time")
    p.add_run(", ")
    _add_inline_field(p, "witness_name", "Witness Name")
    p.add_run(" had not appeared for the deposition and the following proceedings were had:")

    _add_field(doc, "proceedings_summary", "Proceedings Summary")

    _add_paragraph(doc, "    I further certify that I am neither employed by nor related to any attorney or party in this matter and have no interest, financial or otherwise, in its outcome.")

    p = doc.add_paragraph()
    p.add_run("    SUBSCRIBED AND SWORN TO UNDER MY HAND on this the ")
    _add_inline_field(p, "certification_date", "Certification Date")
    p.add_run(".")

    _add_reporter_signature_block(doc)
    _save(doc, "cert_nonappearance.docx")


def build_further_cert_trcp_203() -> None:
    doc = _new_document()
    _add_paragraph(doc, "Further Certification under TRCP Rule 203",
                   align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    _add_paragraph(doc, "")

    p = doc.add_paragraph()
    p.add_run("    The original deposition was/was not returned to the deposition officer on ")
    _add_inline_field(p, "transcript_returned_date", "Returned Date")
    p.add_run(".")

    _add_paragraph(doc, "    If returned, the attached Changes and Signature page(s) contain(s) any changes and the reasons therefor.")

    p = doc.add_paragraph()
    p.add_run("    If returned, the original deposition was delivered to ")
    _add_inline_field(p, "custodial_attorney_name", "Custodial Attorney")
    p.add_run(", Custodial Attorney;")

    p = doc.add_paragraph()
    p.add_run("    That $")
    _add_inline_field(p, "cost_amount", "Cost Amount")
    p.add_run(" is the deposition officer's charges to the ")
    _add_inline_field(p, "cost_payor_party", "Cost Payor Party")
    p.add_run(" for preparing the original deposition and any copies of exhibits;")

    p = doc.add_paragraph()
    p.add_run("    That the deposition was delivered in accordance with Rule 203.3, and that a copy of this certificate was served on all parties shown herein on ")
    _add_inline_field(p, "served_on_date", "Served-On Date")
    p.add_run(" and filed with the Clerk.")

    p = doc.add_paragraph()
    p.add_run("    Certified to by me on this ")
    _add_inline_field(p, "certification_date", "Certification Date")
    p.add_run(".")

    _add_reporter_signature_block(doc)
    _save(doc, "further_cert_trcp_203.docx")


# ===========================================================================
# Main
# ===========================================================================

def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    builders = [
        build_title_page_tx_state,
        build_title_page_federal,
        build_appearances,
        build_index_chronological,
        build_witness_setup_standard,
        build_witness_setup_interpreter,
        build_changes_signature_grid,
        build_witness_acknowledgment_notary,
        build_cert_tx_sig_required,
        build_cert_tx_sig_waived,
        build_cert_federal_frcp,
        build_cert_nonappearance,
        build_further_cert_trcp_203,
    ]

    for fn in builders:
        try:
            fn()
        except Exception as exc:
            print(f"FAILED: {fn.__name__}: {exc}")
            raise

    print(f"\nGenerated {len(builders)} templates in {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
