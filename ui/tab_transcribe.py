"""
ui/tab_transcribe.py

Main transcription UI — file picker, settings, and case info intake.
Includes IntakeReviewDialog for reviewing/editing AI-extracted case data.
"""

import json
import os
import re
import shutil
import threading
from datetime import datetime
from pathlib import Path
from tkinter import filedialog, messagebox
from typing import Any

import customtkinter as ctk

from ui._components import (
    AUDIO_VIDEO_EXTENSIONS,
    BTN_SAFE_GREEN,
    BTN_SAFE_GREEN_HOVER,
    BTN_UTILITY_BLUE,
    BTN_UTILITY_BLUE_HOVER,
    TEXT_MUTED,
    TEXT_PRIMARY,
    TEXT_SECONDARY,
    TOOLBAR_BTN_H,
    TOOLBAR_BTN_W,
    make_section_header,
)

from app_logging import get_logger

logger = get_logger(__name__)


# Default output directory
_DEFAULT_BASE_DIR = r"C:\Users\james\Depositions"

# ── Row colours for the review dialog ────────────────────────────────────────
_ROW_BG_A = "#1A1A2E"
_ROW_BG_B = "#16213E"
_LABEL_COLOR = "#8888AA"
_ENTRY_FG = "#0F3460"
_ENTRY_BORDER = "#1E3A5F"
_HEADER_BG = "#1E3A5F"
_AMBER = "#B8860B"
_AMBER_DARK = "#2A1A00"
_LABEL_TEXT_COLOR = "#8CA0B3"
_PRIMARY_ACTION_TEXT = "START TRANSCRIPTION"
_SOURCE_DOC_SUFFIXES = {".pdf", ".docx", ".txt"}
_SECTION_GAP_Y = 8
_SECTION_HEADER_PADY = (8, 4)
_SECTION_HEADER_PADX = 4
_SECTION_INNER_PADX = 16
_SECTION_INNER_PADY = (8, 12)
_TAB_BG = "#0F172A"
_CARD_BG = "#1E293B"
_CARD_BORDER = "#334155"
_INPUT_BG = "#0F172A"
_INPUT_BORDER = "#334155"
_PRIMARY_BLUE = "#2563EB"
_PRIMARY_BLUE_HOVER = "#3B82F6"
_EMERALD = "#10B981"
_REVIEW_BG = "#1D4ED8"
_REVIEW_HOVER = "#2563EB"
_INFO_CHIP_BG = "#172554"


def _normalize_ui_speaker_map(raw: dict | None) -> dict[int, str]:
    """Normalize persisted speaker_map keys to int IDs for UI use."""
    normalized: dict[int, str] = {}
    if not isinstance(raw, dict):
        return normalized

    for key, value in raw.items():
        try:
            speaker_id = int(str(key).strip())
        except (TypeError, ValueError):
            continue

        label = " ".join(str(value or "").split()).strip()
        if label:
            normalized[speaker_id] = label

    return normalized


def _normalize_ui_speaker_suggestion(raw: dict | None) -> dict[str, Any]:
    """Keep only supported speaker suggestion fields in a stable shape."""
    if not isinstance(raw, dict):
        return {}

    normalized: dict[str, Any] = {}
    for key in (
        "reporter",
        "witness",
        "deponent",
        "ordering_attorney",
        "filing_attorney",
    ):
        value = str(raw.get(key) or "").strip()
        if value:
            normalized[key] = value

    copy_attorneys = raw.get("copy_attorneys", [])
    if isinstance(copy_attorneys, list):
        cleaned = [str(item).strip() for item in copy_attorneys if str(item).strip()]
        if cleaned:
            normalized["copy_attorneys"] = cleaned

    return normalized


def _build_ui_speaker_defaults(
    speaker_ids: list[str],
    saved_map: dict[int, str] | dict[str, str] | None,
    suggestion: dict[str, Any] | None,
) -> dict[str, str]:
    """
    Build UI defaults for speaker entry rows.

    Priority:
    1. Exact saved speaker_map entries from prior confirmation
    2. Ordered NOD suggestions as editable defaults only
    """
    defaults: dict[str, str] = {}
    normalized_saved = _normalize_ui_speaker_map(saved_map)

    for sid_text in speaker_ids:
        try:
            speaker_id = int(str(sid_text).strip())
        except (TypeError, ValueError):
            continue
        if speaker_id in normalized_saved:
            defaults[f"Speaker {sid_text}"] = normalized_saved[speaker_id]

    if defaults:
        return defaults

    normalized_suggestion = _normalize_ui_speaker_suggestion(suggestion)
    ordered_defaults: list[str] = []

    if normalized_suggestion.get("reporter"):
        ordered_defaults.append("THE REPORTER")

    for key in ("ordering_attorney", "filing_attorney"):
        name = str(normalized_suggestion.get(key) or "").strip()
        if name:
            ordered_defaults.append(" ".join(name.split()).strip())

    for name in normalized_suggestion.get("copy_attorneys", []):
        value = " ".join(str(name).split()).strip()
        if value:
            ordered_defaults.append(value)

    if normalized_suggestion.get("witness") or normalized_suggestion.get("deponent"):
        ordered_defaults.append("THE WITNESS")

    deduped_defaults: list[str] = []
    seen: set[str] = set()
    for value in ordered_defaults:
        if value not in seen:
            seen.add(value)
            deduped_defaults.append(value)

    for sid_text, default_value in zip(speaker_ids, deduped_defaults):
        defaults[f"Speaker {sid_text}"] = default_value

    return defaults


def _build_ui_speaker_reference_text(suggestion: dict[str, Any] | None) -> str:
    """Build a compact read-only hint from NOD speaker suggestions."""
    normalized = _normalize_ui_speaker_suggestion(suggestion)
    parts: list[str] = []

    reporter_name = str(normalized.get("reporter") or "").strip()
    witness_name = str(normalized.get("witness") or normalized.get("deponent") or "").strip()

    if reporter_name:
        parts.append(f"Reporter: {reporter_name.upper()}")
    if witness_name:
        parts.append(f"Witness: {witness_name.upper()}")

    attorney_names: list[str] = []
    for key in ("ordering_attorney", "filing_attorney"):
        name = str(normalized.get(key) or "").strip()
        if name:
            attorney_names.append(name.upper())
    for name in normalized.get("copy_attorneys", []):
        value = str(name).strip().upper()
        if value:
            attorney_names.append(value)

    if attorney_names:
        deduped_attorneys = list(dict.fromkeys(attorney_names))
        parts.append("Attorneys: " + "; ".join(deduped_attorneys))

    return "  |  ".join(parts)


def _build_ui_quickfill_labels(suggestion: dict[str, Any] | None) -> list[str]:
    """Return canonical safe labels that can be applied with one click."""
    normalized = _normalize_ui_speaker_suggestion(suggestion)
    labels: list[str] = []

    if normalized.get("reporter"):
        labels.append("THE REPORTER")
    if normalized.get("witness") or normalized.get("deponent"):
        labels.append("THE WITNESS")

    return labels


_SPEAKER_LINE_RE = re.compile(r"(^|\n)(Speaker\s+(\d+)):\s*", re.MULTILINE)


def _apply_speaker_labels_to_text(text: str, speaker_map: dict[int, str]) -> str:
    def _replace(match: re.Match) -> str:
        speaker_id = int(match.group(3))
        replacement = speaker_map.get(speaker_id)
        if not replacement:
            return match.group(0)
        prefix = match.group(1) or ""
        return f"{prefix}{replacement}: "

    return _SPEAKER_LINE_RE.sub(_replace, text or "")


# ═══════════════════════════════════════════════════════════════════════════════
#  IntakeReviewDialog
# ═══════════════════════════════════════════════════════════════════════════════

class IntakeReviewDialog(ctk.CTkToplevel):
    """Modal dialog for reviewing and editing AI-extracted case intake data.
    Field labels map directly to UFM page generator field names."""

    # ── UFM Title Page (Fig. 03) ─────────────────────────────────
    _TITLE_PAGE_FIELDS = [
        ("cause_number", "Cause Number"),
        ("plaintiff_name", "Plaintiff"),
        ("defendant_name", "Defendant"),
        ("case_style", "Case Style"),
        ("court_type", "Court Type"),
        ("county", "County"),
        ("state", "State"),
        ("judicial_district", "Judicial District"),
    ]
    # ── Deposition Details ───────────────────────────────────────
    _DEPO_DETAIL_FIELDS = [
        ("depo_date", "Date of Deposition"),
        ("depo_time_start", "Scheduled Start Time"),
        ("depo_location", "Location"),
        ("depo_method", "Method"),
    ]
    # ── Witness / Deponent ───────────────────────────────────────
    _WITNESS_FIELDS = [
        ("witness_name", "Witness Full Name"),
    ]
    # ── Appearances Page (Fig. 04) — per-attorney ────────────────
    _COUNSEL_FIELDS = [
        ("name", "Attorney Name"),
        ("sbot", "State Bar No. (SBOT)"),
        ("firm", "Firm Name"),
        ("address", "Address"),
        ("phone", "Phone"),
        ("party", "Represents"),
    ]
    # ── Reporter's Certificate (Fig. 05) ─────────────────────────
    _REPORTER_FIELDS = [
        ("reporter_name", "Reporter Name"),
        ("csr_number", "CSR Number"),
        ("reporter_agency", "Agency"),
    ]
    # ── Copy / Billing ───────────────────────────────────────────
    _BILLING_FIELDS = [
        ("ordered_by", "Ordered By"),
        ("ordering_firm", "Ordering Firm"),
    ]

    def __init__(self, parent, case_data: dict):
        super().__init__(parent)
        self._parent = parent
        self._data = case_data
        self._entries: dict[str, ctk.CTkEntry] = {}  # flat key -> entry widget
        self._row_index = 0  # alternating row counter

        self.title("Case Intake \u2014 Review & Edit")
        self.geometry("900x700")
        self.resizable(True, True)
        self.grab_set()

        self._build_header()
        self._build_content()
        self._build_buttons()

    # ── Header ───────────────────────────────────────────────────────────────

    def _build_header(self):
        header = ctk.CTkFrame(self, height=46, fg_color=_HEADER_BG, corner_radius=0)
        header.pack(fill="x")
        header.pack_propagate(False)

        ctk.CTkLabel(
            header, text="CASE INTAKE REVIEW",
            font=ctk.CTkFont(size=16, weight="bold"), text_color="white",
        ).pack(side="left", padx=16)

        ctk.CTkLabel(
            header, text="Edit any field \u2014 changes apply immediately",
            font=ctk.CTkFont(size=13), text_color="#AABBCC",
        ).pack(side="right", padx=16)

    # ── Scrollable content ───────────────────────────────────────────────────

    def _build_content(self):
        from core.ufm_field_mapper import map_intake_to_ufm

        self._scroll = ctk.CTkScrollableFrame(self, fg_color="transparent")
        self._scroll.pack(fill="both", expand=True, padx=0, pady=0)
        self._row_index = 0

        # Map extracted data to UFM fields
        self._ufm = map_intake_to_ufm(self._data)

        # ── Title Page (UFM Fig. 03) ─────────────────────────────
        self._add_section_header("TITLE PAGE  (UFM Fig. 03)")
        for key, label in self._TITLE_PAGE_FIELDS:
            self._add_row(key, label, self._ufm.get(key, ""))

        # ── Deposition Details ───────────────────────────────────
        self._add_section_header("DEPOSITION DETAILS")
        for key, label in self._DEPO_DETAIL_FIELDS:
            self._add_row(key, label, self._ufm.get(key, ""))

        # ── Witness / Deponent ───────────────────────────────────
        self._add_section_header("WITNESS / DEPONENT")
        for key, label in self._WITNESS_FIELDS:
            self._add_row(key, label, self._ufm.get(key, ""))

        # ── Appearances Page (UFM Fig. 04) ───────────────────────
        self._add_section_header("APPEARANCES PAGE  (UFM Fig. 04)")

        p_counsel = self._ufm.get("plaintiff_counsel", [])
        for i, atty in enumerate(p_counsel):
            self._add_sub_header(f"Plaintiff Counsel {i + 1}")
            for key, label in self._COUNSEL_FIELDS:
                self._add_row(f"pc.{i}.{key}", label, atty.get(key, ""))

        d_counsel = self._ufm.get("defense_counsel", [])
        for i, atty in enumerate(d_counsel):
            self._add_sub_header(f"Defense Counsel {i + 1}")
            for key, label in self._COUNSEL_FIELDS:
                self._add_row(f"dc.{i}.{key}", label, atty.get(key, ""))

        # ── Reporter's Certificate (UFM Fig. 05) ────────────────
        self._add_section_header("REPORTER'S CERTIFICATE  (UFM Fig. 05)")
        for key, label in self._REPORTER_FIELDS:
            self._add_row(key, label, self._ufm.get(key, ""))

        # ── Copy / Billing ───────────────────────────────────────
        self._add_section_header("COPY / BILLING")
        for key, label in self._BILLING_FIELDS:
            self._add_row(key, label, self._ufm.get(key, ""))

        # ── Discrepancies ────────────────────────────────────────
        discreps = self._data.get("discrepancies", [])
        if discreps:
            self._add_section_header("DISCREPANCIES")
            for d in discreps:
                self._add_discrepancy_block(d)

    def _add_section_header(self, title: str):
        frame = ctk.CTkFrame(self._scroll, fg_color=_HEADER_BG, corner_radius=4, height=28)
        frame.pack(fill="x", pady=(4, 2))
        frame.pack_propagate(False)
        # Amber left border via a small colored frame
        ctk.CTkFrame(frame, width=4, fg_color=_AMBER, corner_radius=0).pack(
            side="left", fill="y"
        )
        ctk.CTkLabel(
            frame, text=title, font=ctk.CTkFont(size=12, weight="bold"),
            text_color="white",
        ).pack(side="left", padx=12)

    def _add_sub_header(self, title: str):
        sep = ctk.CTkFrame(self._scroll, height=1, fg_color="#333355")
        sep.pack(fill="x", pady=(8, 2), padx=8)
        ctk.CTkLabel(
            self._scroll, text=title, font=ctk.CTkFont(size=13, weight="bold"),
            text_color="#CCCCDD",
        ).pack(anchor="w", padx=16, pady=(2, 2))

    def _add_row(self, key: str, label: str, value: str):
        bg = _ROW_BG_A if self._row_index % 2 == 0 else _ROW_BG_B
        self._row_index += 1

        row = ctk.CTkFrame(self._scroll, fg_color=bg, corner_radius=2)
        row.pack(fill="x", padx=4, pady=1)

        ctk.CTkLabel(
            row, text=label, width=220, anchor="w",
            font=ctk.CTkFont(size=13), text_color=_LABEL_COLOR,
        ).pack(side="left", padx=(12, 8), pady=4)

        entry = ctk.CTkEntry(
            row, fg_color=_ENTRY_FG, border_color=_ENTRY_BORDER, border_width=1,
        )
        entry.pack(side="left", fill="x", expand=True, padx=(0, 12), pady=4)
        entry.insert(0, value or "")
        self._entries[key] = entry

    def _add_discrepancy_block(self, d: dict):
        block = ctk.CTkFrame(
            self._scroll, fg_color=_AMBER_DARK, border_color=_AMBER,
            border_width=1, corner_radius=6,
        )
        block.pack(fill="x", padx=8, pady=4)

        for label_text, key in [
            ("Field", "field"),
            ("Document 1 says", "value_1"),
            ("Document 2 says", "value_2"),
            ("Recommendation", "recommendation"),
        ]:
            row = ctk.CTkFrame(block, fg_color="transparent")
            row.pack(fill="x", padx=10, pady=1)
            ctk.CTkLabel(
                row, text=f"{label_text}:", width=160, anchor="w",
                font=ctk.CTkFont(size=13, weight="bold"), text_color=_AMBER,
            ).pack(side="left")
            ctk.CTkLabel(
                row, text=d.get(key, ""), anchor="w", wraplength=600,
                font=ctk.CTkFont(size=13), text_color="#DDCCAA",
            ).pack(side="left", fill="x", expand=True)

    # ── Bottom buttons ───────────────────────────────────────────────────────

    def _build_buttons(self):
        bar = ctk.CTkFrame(self, fg_color="transparent", height=50)
        bar.pack(fill="x", padx=12, pady=(4, 10))

        # Left — export buttons
        ctk.CTkButton(
            bar, text="Export as PDF", width=130, command=self._export_pdf,
        ).pack(side="left", padx=(0, 8))

        ctk.CTkButton(
            bar, text="Export as Word", width=130, command=self._export_word,
        ).pack(side="left")

        # Right — save / cancel
        ctk.CTkButton(
            bar, text="Save Changes", fg_color=_AMBER, hover_color="#9A7209",
            font=ctk.CTkFont(weight="bold"), width=140, command=self._save_changes,
        ).pack(side="right", padx=(8, 0))

        ctk.CTkButton(
            bar, text="Cancel", width=100,
            fg_color="transparent", border_width=1, border_color="#555555",
            command=self.destroy,
        ).pack(side="right")

    # ── Collect edits ────────────────────────────────────────────────────────

    def _collect_ufm_fields(self) -> dict:
        """Read all entry widgets back into a UFM field dict."""
        ufm = dict(self._ufm)  # start from the mapped version

        # Rebuild counsel lists
        p_counsel = list(ufm.get("plaintiff_counsel", []))
        d_counsel = list(ufm.get("defense_counsel", []))

        for flat_key, entry in self._entries.items():
            val = entry.get()
            parts = flat_key.split(".")

            if parts[0] == "pc":
                idx = int(parts[1])
                while len(p_counsel) <= idx:
                    p_counsel.append({})
                p_counsel[idx][parts[2]] = val
            elif parts[0] == "dc":
                idx = int(parts[1])
                while len(d_counsel) <= idx:
                    d_counsel.append({})
                d_counsel[idx][parts[2]] = val
            else:
                # Top-level UFM field
                ufm[flat_key] = val

        ufm["plaintiff_counsel"] = p_counsel
        ufm["defense_counsel"] = d_counsel
        ufm["discrepancies"] = self._data.get("discrepancies", [])
        return ufm

    # ── Save Changes ─────────────────────────────────────────────────────────

    def _save_changes(self):
        ufm = self._collect_ufm_fields()

        # Store on parent
        self._parent._ufm_fields = ufm
        self._parent._extracted_case_data = self._data  # keep raw data too

        # Sync filing-path vars from UFM fields
        if ufm.get("cause_number"):
            self._parent._cause_var.set(ufm["cause_number"])
        witness = ufm.get("witness_name", "")
        if witness:
            parts = witness.strip().split()
            if parts:
                self._parent._lastname_var.set(parts[-1])
        if ufm.get("depo_date"):
            self._parent._date_var.set(ufm["depo_date"])

        # Also push into _populate_case_fields for any UI refresh
        self._parent._populate_case_fields(self._data)

        # Save ufm_fields into job_config.json → source_docs/ (overwrites, no timestamp)
        out_dir = self._parent._last_output_dir
        if out_dir and os.path.isdir(out_dir):
            from core.job_config_manager import merge_and_save
            try:
                merge_and_save(out_dir, ufm_fields=ufm)
                self._parent.winfo_toplevel().transcript_tab.append_log(
                    "Saved UFM fields → source_docs/job_config.json"
                )
            except Exception as exc:
                self._parent.winfo_toplevel().transcript_tab.append_log(
                    f"Warning: could not save UFM fields: {exc}"
                )

        self.destroy()

        # Brief green status
        self._parent.winfo_toplevel().transcript_tab.set_status(
            text="UFM fields saved \u2014 ready for transcript generation",
            color="#44FF44",
        )
        self._parent.after(
            3000,
            lambda: self._parent.winfo_toplevel().transcript_tab.set_status("Ready", "gray"),
        )

    # ── Helper: gather flat rows for export ──────────────────────────────────

    def _gather_sections(self) -> list:
        """Return [(section_title, [(label, value), ...]), ...] for export."""
        ufm = self._collect_ufm_fields()
        sections = []

        # Title Page (UFM Fig. 03)
        rows = [(lbl, ufm.get(k, "")) for k, lbl in self._TITLE_PAGE_FIELDS]
        sections.append(("TITLE PAGE  (UFM Fig. 03)", rows))

        # Deposition Details
        rows = [(lbl, ufm.get(k, "")) for k, lbl in self._DEPO_DETAIL_FIELDS]
        sections.append(("Deposition Details", rows))

        # Witness
        rows = [(lbl, ufm.get(k, "")) for k, lbl in self._WITNESS_FIELDS]
        sections.append(("Witness / Deponent", rows))

        # Appearances Page (UFM Fig. 04) — Plaintiff Counsel
        for i, pc in enumerate(ufm.get("plaintiff_counsel", [])):
            rows = [(lbl, pc.get(k, "")) for k, lbl in self._COUNSEL_FIELDS]
            sections.append((f"APPEARANCES (UFM Fig. 04) \u2014 Plaintiff Counsel {i + 1}", rows))

        # Appearances Page — Defense Counsel
        for i, dc in enumerate(ufm.get("defense_counsel", [])):
            rows = [(lbl, dc.get(k, "")) for k, lbl in self._COUNSEL_FIELDS]
            sections.append((f"APPEARANCES (UFM Fig. 04) \u2014 Defense Counsel {i + 1}", rows))

        # Reporter's Certificate (UFM Fig. 05)
        rows = [(lbl, ufm.get(k, "")) for k, lbl in self._REPORTER_FIELDS]
        sections.append(("REPORTER'S CERTIFICATE  (UFM Fig. 05)", rows))

        # Copy / Billing
        rows = [(lbl, ufm.get(k, "")) for k, lbl in self._BILLING_FIELDS]
        sections.append(("Copy / Billing", rows))

        return sections

    # ── Export PDF ────────────────────────────────────────────────────────────

    def _export_pdf(self):
        data = self._collect_ufm_fields()
        cause = data.get("cause_number", "intake") or "intake"
        witness = data.get("witness_name", "").split()
        lname = witness[-1] if witness else "unknown"
        default_name = f"{cause}_{lname}_intake.pdf"

        path = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            initialfile=default_name,
            filetypes=[("PDF files", "*.pdf")],
        )
        if not path:
            return

        try:
            self._write_pdf(path, data)
            messagebox.showinfo("Exported", f"PDF saved:\n{path}")
        except Exception as exc:
            messagebox.showerror("PDF Export Failed", str(exc))

    def _write_pdf(self, path: str, data: dict):
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.units import inch
        from reportlab.platypus import (
            SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer,
        )
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

        doc = SimpleDocTemplate(path, pagesize=letter,
                                leftMargin=0.75 * inch, rightMargin=0.75 * inch,
                                topMargin=0.6 * inch, bottomMargin=0.6 * inch)
        styles = getSampleStyleSheet()
        elements = []

        # Custom styles
        title_style = ParagraphStyle(
            "IntakeTitle", parent=styles["Title"],
            fontSize=18, textColor=colors.HexColor("#1E3A5F"),
            spaceAfter=6,
        )
        section_style = ParagraphStyle(
            "SectionHead", parent=styles["Heading2"],
            fontSize=12, textColor=colors.HexColor("#1E3A5F"),
            spaceBefore=14, spaceAfter=4,
        )
        label_style = ParagraphStyle(
            "FieldLabel", parent=styles["Normal"],
            fontSize=9, textColor=colors.grey,
        )
        value_style = ParagraphStyle(
            "FieldValue", parent=styles["Normal"],
            fontSize=10, textColor=colors.black,
        )
        warn_style = ParagraphStyle(
            "WarnText", parent=styles["Normal"],
            fontSize=9, textColor=colors.HexColor("#B8860B"),
        )

        # Header
        elements.append(Paragraph("SA Legal Solutions \u2014 Court Reporting", styles["Normal"]))
        elements.append(Paragraph(
            f"Generated: {datetime.now().strftime('%m/%d/%Y %I:%M %p')}",
            styles["Normal"],
        ))
        elements.append(Spacer(1, 6))
        elements.append(Paragraph("CASE INTAKE SHEET", title_style))
        elements.append(Spacer(1, 8))

        # Sections
        for section_title, rows in self._gather_sections():
            elements.append(Paragraph(section_title, section_style))
            table_data = []
            for label, value in rows:
                table_data.append([
                    Paragraph(label, label_style),
                    Paragraph(value or "\u2014", value_style),
                ])
            if table_data:
                t = Table(table_data, colWidths=[2.2 * inch, 4.8 * inch])
                t.setStyle(TableStyle([
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ("LINEBELOW", (0, 0), (-1, -1), 0.25, colors.HexColor("#CCCCCC")),
                ]))
                elements.append(t)

        # Discrepancies
        discreps = self._data.get("discrepancies", [])
        if discreps:
            elements.append(Paragraph("Discrepancies", section_style))
            for d in discreps:
                field = d.get("field", "")
                flag = f"\u26a0 FLAG: {field}"
                elements.append(Paragraph(flag, warn_style))
                elements.append(Paragraph(
                    f"Doc 1: {d.get('value_1', '')}  |  Doc 2: {d.get('value_2', '')}",
                    styles["Normal"],
                ))
                elements.append(Paragraph(
                    f"Recommendation: {d.get('recommendation', '')}",
                    ParagraphStyle("Italic", parent=styles["Normal"], fontSize=9,
                                   textColor=colors.grey),
                ))
                elements.append(Spacer(1, 6))

        # Footer
        cause_num = data.get("cause_number", "")
        witness_name = data.get("witness_name", "")
        elements.append(Spacer(1, 20))
        elements.append(Paragraph(
            f"{cause_num}  \u00b7  {witness_name}  \u00b7  SA Legal Solutions",
            ParagraphStyle("Footer", parent=styles["Normal"],
                           fontSize=8, textColor=colors.grey, alignment=1),
        ))

        doc.build(elements)

    # ── Export Word ──────────────────────────────────────────────────────────

    def _export_word(self):
        data = self._collect_ufm_fields()
        cause = data.get("cause_number", "intake") or "intake"
        witness = data.get("witness_name", "").split()
        lname = witness[-1] if witness else "unknown"
        default_name = f"{cause}_{lname}_intake.docx"

        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile=default_name,
            filetypes=[("Word documents", "*.docx")],
        )
        if not path:
            return

        try:
            self._write_docx(path, data)
            messagebox.showinfo("Exported", f"Word document saved:\n{path}")
        except Exception as exc:
            messagebox.showerror("Word Export Failed", str(exc))

    def _write_docx(self, path: str, data: dict):
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        # Title
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run("CASE INTAKE SHEET")
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

        def _set_cell_shading(cell, hex_color: str):
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), hex_color)
            shading.set(qn("w:val"), "clear")
            cell._tc.get_or_add_tcPr().append(shading)

        def _remove_cell_borders(cell):
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right"):
                el = OxmlElement(f"w:{edge}")
                el.set(qn("w:val"), "none")
                el.set(qn("w:sz"), "0")
                borders.append(el)
            tc_pr.append(borders)

        # Sections
        for section_title, rows in self._gather_sections():
            # Section heading
            hp = doc.add_paragraph()
            hr = hp.add_run(f"  {section_title}")
            hr.bold = True
            hr.font.size = Pt(12)
            hr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            # Shading on paragraph
            p_pr = hp._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "1E3A5F")
            shd.set(qn("w:val"), "clear")
            p_pr.append(shd)

            # Table
            table = doc.add_table(rows=len(rows), cols=2)
            table.columns[0].width = Inches(2.0)
            table.columns[1].width = Inches(4.5)

            for i, (label, value) in enumerate(rows):
                cell_label = table.cell(i, 0)
                cell_value = table.cell(i, 1)

                cell_label.text = label
                cell_value.text = value or "\u2014"

                # Style label
                for p in cell_label.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(10)
                        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

                # Style value
                for p in cell_value.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(10)

                # Alternating shading + borderless
                shade = "FFFFFF" if i % 2 == 0 else "F5F5F5"
                _set_cell_shading(cell_label, shade)
                _set_cell_shading(cell_value, shade)
                _remove_cell_borders(cell_label)
                _remove_cell_borders(cell_value)

        # Discrepancies
        discreps = self._data.get("discrepancies", [])
        if discreps:
            hp = doc.add_paragraph()
            hr = hp.add_run("  DISCREPANCIES")
            hr.bold = True
            hr.font.size = Pt(12)
            hr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            p_pr = hp._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "B8860B")
            shd.set(qn("w:val"), "clear")
            p_pr.append(shd)

            d_table = doc.add_table(rows=len(discreps) + 1, cols=3)
            # Header row
            for ci, hdr in enumerate(["Field", "Conflict", "Recommendation"]):
                cell = d_table.cell(0, ci)
                cell.text = hdr
                _set_cell_shading(cell, "B8860B")
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        r.font.size = Pt(10)

            for ri, d in enumerate(discreps, start=1):
                d_table.cell(ri, 0).text = d.get("field", "")
                d_table.cell(ri, 1).text = (
                    f"Doc 1: {d.get('value_1', '')}\nDoc 2: {d.get('value_2', '')}"
                )
                d_table.cell(ri, 2).text = d.get("recommendation", "")
                for ci in range(3):
                    _set_cell_shading(d_table.cell(ri, ci), "FFF3CD")

        # Footer
        footer_p = doc.add_paragraph()
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = footer_p.add_run(
            f"{data.get('cause_number', '')}  \u00b7  "
            f"{data.get('witness_name', '')}  \u00b7  SA Legal Solutions"
        )
        fr.font.size = Pt(8)
        fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        doc.save(path)


# ═══════════════════════════════════════════════════════════════════════════════
#  TranscribeTab
# ═══════════════════════════════════════════════════════════════════════════════

class TranscribeTab(ctk.CTkFrame):
    def __init__(self, master, **kwargs):
        super().__init__(master, fg_color="transparent", **kwargs)

        self._selected_file: str = ""
        self._last_transcript_path: str | None = None
        self._current_txt_path: str | None = None
        self._transcript_text: str = ""
        self._last_output_dir: str = ""
        self._running = False
        self._speaker_entries: dict[str, ctk.CTkEntry] = {}
        self._speaker_map_suggestion: dict[str, Any] = {}
        self._saved_speaker_map: dict[int, str] = {}
        self._extracted_case_data: dict = {}
        self._ufm_fields: dict = {}
        self._source_docs_text: str = ""
        self._source_docs_keyterms: list[str] = []
        self._pdf_keyterms: list[str] = []
        self._confirmed_spellings: dict = {}
        self._last_pdf_path: str | None = None
        self._source_doc_paths: list[str] = []
        self._pdf_already_loaded = False
        self._current_case_path: str | None = None
        self._formatted_docx_path: str | None = None
        self._correction_mode: bool = False
        self._loaded_transcript_path: str | None = None
        self._loaded_case_folder: str | None = None
        self._case_files_expanded: bool = True

        self._build_ui()

    # ── UI Construction ──────────────────────────────────────────────────────

    def _build_ui(self):
        self.configure(fg_color=_TAB_BG)

        label_font = ctk.CTkFont(size=10, weight="bold")
        body_font = ctk.CTkFont(size=12)
        primary_font = ctk.CTkFont(size=16, weight="bold")
        utility_font = ctk.CTkFont(size=11, weight="bold")
        status_font = ctk.CTkFont(size=11)
        footer_font = ctk.CTkFont(size=9)

        def _make_card(parent, *, corner_radius: int = 12) -> ctk.CTkFrame:
            return ctk.CTkFrame(
                parent,
                fg_color=_CARD_BG,
                border_color=_CARD_BORDER,
                border_width=1,
                corner_radius=corner_radius,
            )

        def _style_entry(entry: ctk.CTkEntry) -> None:
            entry.configure(
                fg_color=_INPUT_BG,
                border_color=_INPUT_BORDER,
                text_color=TEXT_PRIMARY,
                corner_radius=8,
                font=body_font,
            )

        def _style_combo(combo: ctk.CTkComboBox) -> None:
            combo.configure(
                fg_color=_INPUT_BG,
                border_color=_INPUT_BORDER,
                button_color=_CARD_BORDER,
                button_hover_color=_PRIMARY_BLUE,
                dropdown_fg_color=_CARD_BG,
                dropdown_text_color=TEXT_PRIMARY,
                text_color=TEXT_PRIMARY,
                corner_radius=8,
                font=body_font,
            )

        # ── Outer skeleton: header / body / (speaker) / footer rows ────────────
        # Body row stretches; speaker card is gridded into row 2 dynamically
        # when a run finishes.
        self.grid_rowconfigure(1, weight=1)
        self.grid_columnconfigure(0, weight=1)

        # ── Header strip ───────────────────────────────────────────────────────
        header_strip = ctk.CTkFrame(self, fg_color=_CARD_BG, height=32, corner_radius=0)
        header_strip.grid(row=0, column=0, sticky="ew")
        header_strip.grid_propagate(False)
        ctk.CTkLabel(
            header_strip,
            text="DEPO-PRO TRANSCRIBE",
            font=ctk.CTkFont(size=14, weight="bold"),
            text_color=TEXT_PRIMARY,
        ).pack(side="left", padx=14, pady=4)
        ctk.CTkLabel(
            header_strip,
            text="● SYSTEM ONLINE",
            font=ctk.CTkFont(size=10, weight="bold"),
            text_color=_EMERALD,
        ).pack(side="right", padx=14, pady=4)

        # ── Body: 3-column grid ────────────────────────────────────────────────
        body = ctk.CTkFrame(self, fg_color="transparent")
        body.grid(row=1, column=0, sticky="nsew", padx=10, pady=(4, 4))
        body.grid_columnconfigure(0, weight=2, uniform="cols")
        body.grid_columnconfigure(1, weight=3, uniform="cols")
        body.grid_columnconfigure(2, weight=3, uniform="cols")
        body.grid_rowconfigure(0, weight=1)

        # ── Column 1: Source Media ─────────────────────────────────────────────
        source_card = _make_card(body)
        source_card.grid(row=0, column=0, sticky="nsew", padx=(0, 6))

        make_section_header(
            source_card,
            "► 1. Source Media",
            font_size=14,
        ).pack(anchor="w", padx=12, pady=(6, 2))

        source_body = ctk.CTkFrame(source_card, fg_color="transparent")
        source_body.pack(fill="both", expand=True, padx=14, pady=(0, 6))

        ctk.CTkLabel(
            source_body,
            text="AUDIO / VIDEO FILE",
            font=label_font,
            text_color=TEXT_MUTED,
        ).pack(anchor="w", pady=(0, 2))

        self._file_entry = ctk.CTkEntry(
            source_body,
            placeholder_text="No file selected",
            state="disabled",
            height=32,
        )
        _style_entry(self._file_entry)
        self._file_entry.pack(fill="x", pady=(0, 4))

        ctk.CTkLabel(
            source_body,
            text="MP3 · MP4 · WAV · M4A · MOV · MKV · FLAC",
            font=ctk.CTkFont(size=10),
            text_color=TEXT_MUTED,
        ).pack(anchor="w", pady=(0, 6))

        ctk.CTkButton(
            source_body,
            text="Browse",
            height=32,
            font=utility_font,
            fg_color=BTN_UTILITY_BLUE,
            hover_color=BTN_UTILITY_BLUE_HOVER,
            command=self._browse_file,
            corner_radius=8,
        ).pack(fill="x", pady=(0, 4))

        ctk.CTkButton(
            source_body,
            text="Combine Multiple Files",
            height=32,
            font=utility_font,
            fg_color=_INPUT_BG,
            hover_color=_CARD_BORDER,
            border_color=_CARD_BORDER,
            border_width=1,
            command=self._open_combine_dialog,
            corner_radius=8,
        ).pack(fill="x", pady=(0, 6))

        # Spacer pushes the save-location preview to the bottom of the column.
        ctk.CTkFrame(source_body, fg_color="transparent").pack(fill="both", expand=True)

        ctk.CTkLabel(
            source_body,
            text="SAVE LOCATION",
            font=label_font,
            text_color=TEXT_MUTED,
        ).pack(anchor="w", pady=(0, 2))
        preview_shell = ctk.CTkFrame(
            source_body,
            fg_color=_INPUT_BG,
            border_color=_INPUT_BORDER,
            border_width=1,
            corner_radius=8,
        )
        preview_shell.pack(fill="x", pady=(0, 4))
        self._path_preview_label = ctk.CTkLabel(
            preview_shell,
            text="",
            font=ctk.CTkFont(size=11),
            text_color=TEXT_PRIMARY,
            anchor="w",
            justify="left",
            wraplength=240,
        )
        self._path_preview_label.pack(fill="x", padx=10, pady=4)

        # _base_dir_entry is kept (some callers reach through it) but isn't
        # placed in the new layout — the "Change Base Folder" button below
        # opens a directory picker that writes through _base_dir_var.
        self._base_dir_var = ctk.StringVar(value=_DEFAULT_BASE_DIR)
        self._base_dir_entry = ctk.CTkEntry(
            source_body, textvariable=self._base_dir_var, height=1
        )
        _style_entry(self._base_dir_entry)
        self._base_dir_var.trace_add("write", lambda *_: self._update_path_preview())

        ctk.CTkButton(
            source_body,
            text="Change Base Folder",
            height=30,
            font=ctk.CTkFont(size=11, weight="bold"),
            fg_color="transparent",
            hover_color=_CARD_BORDER,
            border_color=_CARD_BORDER,
            border_width=1,
            text_color=TEXT_SECONDARY,
            command=self._browse_base_dir,
            corner_radius=8,
        ).pack(fill="x")

        # ── Column 2: Configuration (engine + deposition details) ──────────────
        columns_row = ctk.CTkFrame(body, fg_color="transparent")
        columns_row.grid(row=0, column=1, sticky="nsew", padx=6)
        columns_row.grid_columnconfigure(0, weight=1)
        columns_row.grid_rowconfigure(1, weight=1)

        settings_card = _make_card(columns_row)
        settings_card.grid(row=0, column=0, sticky="ew", pady=(0, 6))

        make_section_header(
            settings_card,
            "⚙ 2. Engine Config",
            font_size=14,
        ).pack(anchor="w", padx=12, pady=(6, 2))

        settings_body = ctk.CTkFrame(settings_card, fg_color="transparent")
        settings_body.pack(fill="x", padx=14, pady=(0, 6))
        settings_body.grid_columnconfigure(0, weight=1)
        settings_body.grid_columnconfigure(1, weight=1)

        ctk.CTkLabel(
            settings_body,
            text="MODEL",
            font=label_font,
            text_color=TEXT_MUTED,
        ).grid(row=0, column=0, sticky="w", padx=(0, 6), pady=(0, 2))
        ctk.CTkLabel(
            settings_body,
            text="PROCESSING MODE",
            font=label_font,
            text_color=TEXT_MUTED,
        ).grid(row=0, column=1, sticky="w", pady=(0, 2))

        self._model_var = ctk.StringVar(value="nova-3")
        self._model_combo = ctk.CTkComboBox(
            settings_body,
            values=["nova-3", "nova-3-medical"],
            variable=self._model_var,
            state="readonly",
            height=30,
        )
        _style_combo(self._model_combo)
        self._model_combo.grid(row=1, column=0, sticky="ew", padx=(0, 6), pady=(0, 4))

        self._quality_var = ctk.StringVar(value="ENHANCED (fair audio)")
        self._quality_combo = ctk.CTkComboBox(
            settings_body,
            values=[
                "Default",
                "ENHANCED (fair audio)",
                "RESCUE (noisy/poor audio)",
            ],
            variable=self._quality_var,
            state="readonly",
            height=30,
        )
        _style_combo(self._quality_combo)
        self._quality_combo.grid(row=1, column=1, sticky="ew", pady=(0, 4))

        info_chip = ctk.CTkFrame(settings_body, fg_color=_INFO_CHIP_BG, corner_radius=6)
        info_chip.grid(row=2, column=0, columnspan=2, sticky="ew", pady=(2, 0))
        ctk.CTkLabel(
            info_chip,
            text="ⓘ Estimated processing time: ~4m 30s",
            font=ctk.CTkFont(size=11),
            text_color="#BFDBFE",
        ).pack(anchor="w", padx=10, pady=3)

        self._audio_tier_label = ctk.CTkLabel(
            settings_body,
            text="",
            font=ctk.CTkFont(size=11),
            text_color=TEXT_SECONDARY,
            anchor="w",
        )
        self._audio_tier_label.grid(
            row=3, column=0, columnspan=2, sticky="ew", pady=(2, 0)
        )

        details_card = _make_card(columns_row)
        details_card.grid(row=1, column=0, sticky="nsew")

        details_header = ctk.CTkFrame(details_card, fg_color="transparent")
        details_header.pack(fill="x", padx=12, pady=(6, 2))
        make_section_header(
            details_header,
            "▣ 3. Deposition Details",
            font_size=14,
        ).pack(side="left", anchor="w")

        self._rescan_btn = ctk.CTkButton(
            details_header,
            text="↻ Auto-fill",
            fg_color="transparent",
            hover_color="#1D4ED8",
            text_color=_PRIMARY_BLUE_HOVER,
            font=ctk.CTkFont(size=11, weight="bold"),
            border_width=0,
            width=90,
            height=24,
            command=lambda: (
                self._apply_filename_extraction(self._selected_file)
                if self._selected_file
                else self._set_transcript_status("Select an audio file first", "#CCAA44")
            ),
        )
        self._rescan_btn.pack(side="right")

        details_body = ctk.CTkFrame(details_card, fg_color="transparent")
        details_body.pack(fill="x", padx=14, pady=(0, 6))
        details_body.grid_columnconfigure(0, weight=1)
        details_body.grid_columnconfigure(1, weight=1)

        cause_frame = ctk.CTkFrame(details_body, fg_color="transparent")
        cause_frame.grid(row=0, column=0, sticky="ew", padx=(0, 6), pady=(0, 4))
        cause_lbl_row = ctk.CTkFrame(cause_frame, fg_color="transparent")
        cause_lbl_row.pack(fill="x")
        ctk.CTkLabel(cause_lbl_row, text="CAUSE NUMBER", font=label_font, text_color=TEXT_MUTED).pack(side="left", pady=(0, 2))
        self._cause_badge = ctk.CTkLabel(cause_lbl_row, text="", font=ctk.CTkFont(size=11), width=0, text_color=TEXT_SECONDARY)
        self._cause_badge.pack(side="left", padx=(6, 0))
        self._cause_var = ctk.StringVar()
        cause_entry = ctk.CTkEntry(cause_frame, textvariable=self._cause_var, placeholder_text="e.g. 2025CI19595", height=30)
        _style_entry(cause_entry)
        cause_entry.pack(fill="x")
        self._cause_var.trace_add("write", self._on_cause_changed)

        date_frame = ctk.CTkFrame(details_body, fg_color="transparent")
        date_frame.grid(row=0, column=1, sticky="ew", pady=(0, 4))
        date_lbl_row = ctk.CTkFrame(date_frame, fg_color="transparent")
        date_lbl_row.pack(fill="x")
        ctk.CTkLabel(date_lbl_row, text="DEPOSITION DATE", font=label_font, text_color=TEXT_MUTED).pack(side="left", pady=(0, 2))
        self._date_badge = ctk.CTkLabel(date_lbl_row, text="", font=ctk.CTkFont(size=11), width=0, text_color=TEXT_SECONDARY)
        self._date_badge.pack(side="left", padx=(6, 0))
        self._date_var = ctk.StringVar()
        date_entry = ctk.CTkEntry(date_frame, textvariable=self._date_var, placeholder_text="From NOD PDF", height=30)
        _style_entry(date_entry)
        date_entry.pack(fill="x")
        self._date_var.trace_add("write", lambda *_: self._update_path_preview())

        first_frame = ctk.CTkFrame(details_body, fg_color="transparent")
        first_frame.grid(row=1, column=0, sticky="ew", padx=(0, 6))
        ctk.CTkLabel(first_frame, text="WITNESS FIRST NAME", font=label_font, text_color=TEXT_MUTED).pack(anchor="w", pady=(0, 2))
        self._firstname_var = ctk.StringVar()
        first_entry = ctk.CTkEntry(first_frame, textvariable=self._firstname_var, placeholder_text="e.g. Matthew", height=30)
        _style_entry(first_entry)
        first_entry.pack(fill="x")
        self._firstname_var.trace_add("write", lambda *_: self._update_path_preview())

        name_frame = ctk.CTkFrame(details_body, fg_color="transparent")
        name_frame.grid(row=1, column=1, sticky="ew")
        name_lbl_row = ctk.CTkFrame(name_frame, fg_color="transparent")
        name_lbl_row.pack(fill="x")
        ctk.CTkLabel(name_lbl_row, text="WITNESS LAST NAME", font=label_font, text_color=TEXT_MUTED).pack(side="left", pady=(0, 2))
        self._witness_badge = ctk.CTkLabel(name_lbl_row, text="", font=ctk.CTkFont(size=11), width=0, text_color=TEXT_SECONDARY)
        self._witness_badge.pack(side="left", padx=(6, 0))
        self._lastname_var = ctk.StringVar()
        last_entry = ctk.CTkEntry(name_frame, textvariable=self._lastname_var, placeholder_text="e.g. Coger", height=30)
        _style_entry(last_entry)
        last_entry.pack(fill="x")
        self._lastname_var.trace_add("write", self._on_lastname_changed)

        self._update_path_preview()

        # ── Column 3: Run (start / progress / log / utilities) ─────────────────
        run_card = _make_card(body)
        run_card.grid(row=0, column=2, sticky="nsew", padx=(6, 0))

        make_section_header(
            run_card,
            "▶ 4. Run",
            font_size=14,
        ).pack(anchor="w", padx=12, pady=(6, 2))

        run_body = ctk.CTkFrame(run_card, fg_color="transparent")
        run_body.pack(fill="both", expand=True, padx=14, pady=(0, 6))

        self._create_btn = ctk.CTkButton(
            run_body,
            text="▶ Start Transcription",
            height=42,
            font=primary_font,
            fg_color=_PRIMARY_BLUE,
            hover_color=_PRIMARY_BLUE_HOVER,
            command=self.start_transcription,
            corner_radius=10,
        )
        self._create_btn.pack(fill="x", pady=(0, 6))

        status_row = ctk.CTkFrame(run_body, fg_color="transparent")
        status_row.pack(fill="x", pady=(0, 4))
        self._status_progress = ctk.CTkProgressBar(
            status_row, height=6, progress_color=_PRIMARY_BLUE
        )
        self._status_progress.pack(fill="x", pady=(0, 2))
        self._status_progress.set(0)
        self._extract_status_label = ctk.CTkLabel(
            status_row,
            text="Ready",
            font=status_font,
            text_color=TEXT_SECONDARY,
            anchor="w",
            justify="left",
        )
        self._extract_status_label.pack(fill="x")

        self._run_log_card = ctk.CTkFrame(
            run_body,
            fg_color=_INPUT_BG,
            border_color=_INPUT_BORDER,
            border_width=1,
            corner_radius=8,
        )
        self._run_log_card.pack(fill="both", expand=True, pady=(0, 6))
        self._run_log = ctk.CTkTextbox(
            self._run_log_card,
            wrap="char",
            fg_color=_INPUT_BG,
            border_width=0,
            text_color=TEXT_PRIMARY,
            font=ctk.CTkFont(family="Consolas", size=11),
        )
        self._run_log.pack(fill="both", expand=True, padx=6, pady=6)
        self._run_log.insert("1.0", "> System ready.\n> Waiting for input...\n")
        self._run_log.configure(state="disabled")

        utility_grid = ctk.CTkFrame(run_body, fg_color="transparent")
        utility_grid.pack(fill="x")
        utility_grid.grid_columnconfigure(0, weight=1, uniform="util")
        utility_grid.grid_columnconfigure(1, weight=1, uniform="util")

        # Hidden trigger preserved for callers that invoke it programmatically.
        self._upload_pdf_btn = ctk.CTkButton(
            run_body,
            text="Upload NOD / PDF",
            width=1,
            height=1,
            fg_color=_CARD_BG,
            hover_color=_CARD_BG,
            text_color=_CARD_BG,
            command=self._handle_pdf_upload,
        )

        self._upload_reporter_notes_btn = ctk.CTkButton(
            utility_grid,
            text="NOD and Notes",
            height=30,
            font=utility_font,
            fg_color=BTN_UTILITY_BLUE,
            hover_color=BTN_UTILITY_BLUE_HOVER,
            command=self._upload_nod_and_notes,
            corner_radius=8,
        )
        self._upload_reporter_notes_btn.grid(
            row=0, column=0, sticky="ew", padx=(0, 4), pady=(0, 4)
        )

        self._review_btn = ctk.CTkButton(
            utility_grid,
            text="Review & Edit",
            height=30,
            font=utility_font,
            fg_color=_REVIEW_BG,
            hover_color=_REVIEW_HOVER,
            state="disabled",
            command=self._open_review_dialog,
            corner_radius=8,
        )
        self._review_btn.grid(row=0, column=1, sticky="ew", padx=(4, 0), pady=(0, 4))

        self._open_folder_btn = ctk.CTkButton(
            utility_grid,
            text="Output Folder",
            height=30,
            font=utility_font,
            fg_color=_INPUT_BG,
            hover_color=_CARD_BORDER,
            border_color=_CARD_BORDER,
            border_width=1,
            command=self._open_output_folder,
            corner_radius=8,
        )
        self._open_folder_btn.grid(row=1, column=0, sticky="ew", padx=(0, 4))

        self._open_transcript_btn = ctk.CTkButton(
            utility_grid,
            text="View Document",
            height=30,
            font=utility_font,
            fg_color=_INPUT_BG,
            hover_color=_CARD_BORDER,
            border_color=_CARD_BORDER,
            border_width=1,
            command=self._open_transcript,
            corner_radius=8,
        )
        self._open_transcript_btn.grid(row=1, column=1, sticky="ew", padx=(4, 0))

        # ── Speaker labels card (gridded into row 2 only after a run) ──────────
        # Created here but not gridded; _populate_speaker_rows places it.
        self._speaker_card = _make_card(self)
        make_section_header(
            self._speaker_card,
            "☷ Speaker Labels — Rename before saving",
            font_size=14,
        ).pack(fill="x", padx=12, pady=(6, 2))
        ctk.CTkLabel(
            self._speaker_card,
            text="Replace generic labels with correct names throughout the transcript",
            font=body_font,
            text_color=TEXT_SECONDARY,
        ).pack(anchor="w", padx=16, pady=(0, 2))
        self._speaker_hint_label = ctk.CTkLabel(
            self._speaker_card,
            text="",
            font=ctk.CTkFont(size=11),
            text_color="#7DAACC",
            justify="left",
            wraplength=760,
        )
        self._speaker_hint_label.pack(anchor="w", padx=16, pady=(0, 2))
        self._speaker_rows_frame = ctk.CTkFrame(self._speaker_card, fg_color="transparent")
        self._speaker_rows_frame.pack(fill="x", padx=16)
        self._apply_save_btn = ctk.CTkButton(
            self._speaker_card,
            text="\u2713  Apply Speaker Labels",
            height=36,
            font=ctk.CTkFont(size=13, weight="bold"),
            fg_color=BTN_SAFE_GREEN,
            hover_color=BTN_SAFE_GREEN_HOVER,
            text_color="white",
            command=self._apply_and_save_labels,
            corner_radius=8,
        )
        self._apply_save_btn.pack(anchor="e", padx=16, pady=(4, 6))

        # ── Footer strip ───────────────────────────────────────────────────────
        self._footer_strip = ctk.CTkFrame(
            self,
            fg_color=_CARD_BG,
            border_color=_CARD_BORDER,
            border_width=1,
            corner_radius=0,
            height=22,
        )
        self._footer_strip.grid(row=3, column=0, sticky="ew")
        self._footer_strip.grid_propagate(False)
        ctk.CTkLabel(
            self._footer_strip,
            text="SERVER: US-CENTRAL-1   ● SYSTEM ONLINE",
            font=footer_font,
            text_color=_EMERALD,
        ).pack(side="left", padx=12, pady=2)
        ctk.CTkLabel(
            self._footer_strip,
            text="License: Professional Enterprise • Exp: 12/2026",
            font=footer_font,
            text_color=TEXT_MUTED,
        ).pack(side="right", padx=12, pady=2)

    # ── Helpers ──────────────────────────────────────────────────────────────

    @staticmethod
    def _set_entry_text(entry: ctk.CTkEntry, text: str):
        """Set text on a disabled CTkEntry."""
        entry.configure(state="normal")
        entry.delete(0, "end")
        entry.insert(0, text)
        entry.configure(state="disabled")

    def _append_transcript_log(self, msg: str):
        self._run_log.configure(state="normal")
        self._run_log.insert("end", f"{msg}\n")
        self._run_log.see("end")
        self._run_log.configure(state="disabled")

    def _set_transcript_status(self, text: str, color: str = "gray"):
        self._extract_status_label.configure(text=text, text_color=color)

    def append_log(self, msg: str):
        self._append_transcript_log(msg)

    def set_status(self, text: str, color: str = "gray"):
        self._set_transcript_status(text, color)

    def set_transcription_running(self):
        self._status_progress.set(0.05)
        self._set_transcript_status("Transcribing audio...", "white")

    def set_transcription_complete(self, transcript_path: str, folder_path: str):
        self._last_transcript_path = transcript_path
        self._current_txt_path = transcript_path
        self._current_case_path = folder_path
        self._status_progress.set(0.85)
        self._set_transcript_status("Deepgram transcript complete", "#44FF44")

    def set_transcription_failed(self, error_msg: str):
        self._status_progress.set(0)
        self._set_transcript_status(f"Transcription failed: {error_msg}", "#FF4444")

    def load_transcript(self, filepath: str):
        self._last_transcript_path = filepath
        self._current_txt_path = filepath
        self._append_transcript_log(f"Transcript available: {filepath}")

    def _get_preferred_document_path(self) -> str | None:
        """Return the formatted deposition document when one is available."""
        candidate_paths: list[Path] = []

        if self._formatted_docx_path:
            candidate_paths.append(Path(self._formatted_docx_path))

        if self._last_transcript_path and str(self._last_transcript_path).lower().endswith(".docx"):
            candidate_paths.append(Path(self._last_transcript_path))

        if self._current_case_path and os.path.isdir(self._current_case_path):
            case_dir = Path(self._current_case_path)
            candidate_paths.extend(sorted(case_dir.glob("*_Deposition_*.docx")))

        existing_paths = [path for path in candidate_paths if path.is_file()]
        if not existing_paths:
            return None

        return str(max(existing_paths, key=lambda path: path.stat().st_mtime))

    def set_audio_file(self, filepath: str):
        self._selected_file = filepath

    def _set_create_buttons(self, state: str, text: str):
        self._create_btn.configure(state=state, text=text)

    def _update_path_preview(self):
        """Update the live path preview label based on case info fields."""
        from core.file_manager import build_case_path

        cause = self._cause_var.get().strip()
        last = self._lastname_var.get().strip()
        first = self._firstname_var.get().strip()
        date = self._date_var.get().strip()
        base = self._base_dir_var.get().strip()

        if cause and last:
            path = build_case_path(base, cause, last, first, date)
            self._path_preview_label.configure(text=f"Will save to: {path}")
            self._current_case_path = path
        else:
            self._path_preview_label.configure(
                text="Will save to: (fill in Cause Number and Witness Name)"
            )
            self._current_case_path = None

    def _set_case_files_panel_expanded(self, expanded: bool) -> None:
        # The accordion was retired in favor of a single inline toolbar.
        # The toolbar is always visible, so there is nothing to expand or
        # collapse. We keep this method as a no-op so existing callers (the
        # upload handlers) do not break.
        self._case_files_expanded = bool(expanded)
        toggle_btn = getattr(self, "_case_files_toggle_btn", None)
        if toggle_btn is not None:
            toggle_btn.configure(
                text="▼ Case Files" if self._case_files_expanded else "▶ Case Files"
            )

    def _toggle_case_files_panel(self) -> None:
        self._set_case_files_panel_expanded(not self._case_files_expanded)

    def _get_current_save_path(self):
        return self._current_case_path or ""

    def _on_cause_changed(self, *_):
        self._update_path_preview()
        if not self._cause_var.get().strip():
            self._reset_case_state()

    def _on_lastname_changed(self, *_):
        self._update_path_preview()
        if not self._lastname_var.get().strip():
            self._reset_case_state()

    def _reset_case_state(self):
        """Reset auto-detected case data when switching to a new case."""
        self._pdf_already_loaded = False
        self._source_docs_text = ""
        self._source_docs_keyterms = []
        self._pdf_keyterms = []
        self._confirmed_spellings = {}
        self._speaker_map_suggestion = {}
        self._saved_speaker_map = {}
        self._last_pdf_path = None
        self._source_doc_paths = []
        self._extracted_case_data = {}
        self._current_case_path = None
        self._last_transcript_path = None
        self._formatted_docx_path = None
        self._model_var.set("nova-3")
        self._quality_var.set("ENHANCED (fair audio)")
        self._audio_tier_label.configure(text="", text_color="gray")
        self._review_btn.configure(state="disabled")
        self._upload_pdf_btn.configure(
            text="\U0001f4c4  Upload NOD / PDF",
            fg_color=BTN_UTILITY_BLUE,
            state="normal",
        )
        # Open Output Folder + Open Transcript are pickers now — always enabled.
        self._set_create_buttons(state="normal", text=_PRIMARY_ACTION_TEXT)
        self._upload_reporter_notes_btn.configure(
            text="\U0001f4dd  NOD and Notes",
            fg_color=BTN_UTILITY_BLUE,
        )
        self._set_case_files_panel_expanded(True)
        for badge in (self._cause_badge, self._witness_badge, self._date_badge):
            badge.configure(text="")

    def _apply_saved_transcription_settings(self, config_data: dict | None):
        """Restore persisted transcription settings from job_config.json."""
        if not isinstance(config_data, dict):
            return

        self._speaker_map_suggestion = _normalize_ui_speaker_suggestion(
            config_data.get("speaker_map_suggestion", {})
        )
        ufm = config_data.get("ufm_fields", {})
        self._saved_speaker_map = _normalize_ui_speaker_map(
            ufm.get("speaker_map", {}) if isinstance(ufm, dict) else {}
        )

        model = config_data.get("model")
        if model in {"nova-3", "nova-3-medical"}:
            self._model_var.set(model)

        audio_quality = config_data.get("audio_quality")
        if audio_quality in {
            "Auto-detect (recommended)",
            "Clean (good/excellent audio)",
            "Default (fair audio)",
            "Aggressive (noisy/poor audio)",
        }:
            self._quality_var.set(audio_quality)
        if self._correction_mode:
            self._set_create_buttons(state="disabled", text="CORRECTION MODE — Transcription Disabled")

    def _force_rescan(self):
        """Force re-detection of source documents from source_docs."""
        self._pdf_already_loaded = False
        self._source_docs_text = ""
        self._source_docs_keyterms = []
        self._source_doc_paths = []
        self._pdf_keyterms = []
        self._extracted_case_data = {}
        self._review_btn.configure(state="disabled")
        self._upload_pdf_btn.configure(
            text="\U0001f4c4  Upload NOD / PDF",
            fg_color=BTN_UTILITY_BLUE,
        )
        self._upload_reporter_notes_btn.configure(
            text="\U0001f4dd  NOD and Notes",
            fg_color=BTN_UTILITY_BLUE,
        )
        self._extract_status_label.configure(
            text="Rescanning for source documents\u2026",
            text_color="gray",
        )
        self._auto_detect_source_docs()

    def _build_keyterms_from_intake(self, intake_result) -> list[str]:
        """
        Build persistent Deepgram/UFM keyterms from intake parse data.

        Sources:
        - intake_result.all_proper_nouns
        - intake_result.vocabulary_terms[].term
        - confirmed_spellings values (correct forms)

        Returns a deterministic, deduplicated, capped list.
        """
        terms: set[str] = set()

        for t in getattr(intake_result, "all_proper_nouns", []) or []:
            if isinstance(t, str) and len(t.strip()) >= 3:
                terms.add(" ".join(t.split()).strip())

        for item in getattr(intake_result, "vocabulary_terms", []) or []:
            term = getattr(item, "term", None)
            if isinstance(item, dict):
                term = item.get("term")
            if isinstance(term, str) and len(term.strip()) >= 3:
                terms.add(" ".join(term.split()).strip())

        for correct in (getattr(intake_result, "confirmed_spellings", {}) or {}).values():
            if isinstance(correct, str) and len(correct.strip()) >= 3:
                terms.add(" ".join(correct.split()).strip())

        return sorted(terms)[:100]

    def _auto_detect_source_docs(self):
        """
        Detect source documents from the resolved case folder.
        Only runs when the case folder exists on disk.
        """
        from core.pdf_extractor import find_case_pdf
        from core.job_config_manager import load_job_config

        base_path = self._get_current_save_path()
        if not base_path or not os.path.isdir(base_path):
            return

        existing_config = load_job_config(base_path)
        if existing_config and not self._pdf_already_loaded:
            self._apply_saved_transcription_settings(existing_config)
            self._confirmed_spellings = existing_config.get("confirmed_spellings", {})
            self._pdf_keyterms = list(existing_config.get("deepgram_keyterms", []) or [])
            ufm = existing_config.get("ufm_fields", {})
            if ufm.get("cause_number") and not self._cause_var.get().strip():
                self._cause_var.set(ufm["cause_number"])
            if not self._date_var.get().strip() and ufm.get("depo_date"):
                self._date_var.set(ufm["depo_date"])

            self._extract_status_label.configure(
                text="Reloaded case settings from job_config.json",
                text_color="#44AA66",
            )
            self._pdf_already_loaded = True

        if not self._pdf_already_loaded:
            pdf_path = find_case_pdf(str(base_path))
            if pdf_path:
                logger.info("[Auto-Detect] PDF found: %s", pdf_path)
                self._upload_pdf_btn.configure(
                    text="PDF Auto-Detected",
                    fg_color="#2A6F3A",
                )
                self._pdf_already_loaded = True
                self._handle_pdf_upload(pdf_path=pdf_path, auto_detected=True)
            else:
                logger.info("[Auto-Detect] No PDF found in source_docs.")

        if not self._source_docs_text:
            source_doc_paths = self._list_supported_source_docs(str(base_path))
            non_pdf_paths = [path for path in source_doc_paths if Path(path).suffix.lower() != ".pdf"]
            if non_pdf_paths:
                try:
                    self._load_source_documents(source_doc_paths, auto_detected=True)
                    logger.info("[Auto-Detect] Source documents found: %s", source_doc_paths)
                except Exception as exc:
                    logger.warning("[Auto-Detect] Could not load source documents: %s", exc)
            else:
                logger.info("[Auto-Detect] No non-PDF source documents found in source_docs.")

    def _build_case_data_from_ufm_fields(self, ufm_fields: dict, witness_fallback: str = "") -> dict:
        """Convert flat saved UFM fields back into the intake-shaped data the review dialog expects."""
        case_data = {
            "deposition_details": {
                "cause_number": ufm_fields.get("cause_number", ""),
                "witness": ufm_fields.get("witness_name", witness_fallback),
                "date": ufm_fields.get("depo_date", ""),
                "court": ufm_fields.get("court_type", ""),
                "court_caption": ufm_fields.get("court_caption", ""),
                "case_style": ufm_fields.get("case_style", ""),
                "method": ufm_fields.get("depo_method", ""),
                "county": ufm_fields.get("county", ""),
                "state": ufm_fields.get("state", ""),
                "ordered_by": ufm_fields.get("ordered_by", ""),
                "amendment": ufm_fields.get("amendment", ""),
                "location": ufm_fields.get("depo_location", ""),
                "scheduled_time": ufm_fields.get("depo_time_start", ""),
            },
            "ordering_attorney": {
                "name": ufm_fields.get("ordering_attorney_name", ""),
                "firm": ufm_fields.get("ordering_firm", ""),
            },
            "filing_attorney": {
                "name": ufm_fields.get("filing_attorney_name", ""),
                "firm": ufm_fields.get("filing_attorney_firm", ""),
            },
            "copy_attorneys": list(ufm_fields.get("copy_attorneys", [])),
            "all_attorneys": [],
            "court_reporter": {
                "name": ufm_fields.get("reporter_name", ""),
                "csr_number": ufm_fields.get("csr_number", ""),
                "agency": ufm_fields.get("reporter_agency", ""),
            },
            "discrepancies": list(ufm_fields.get("discrepancies", [])),
        }

        for role, key in (("plaintiff", "plaintiff_counsel"), ("defense", "defense_counsel")):
            for atty in ufm_fields.get(key, []):
                case_data["all_attorneys"].append({
                    "name": atty.get("name", ""),
                    "firm": atty.get("firm", ""),
                    "bar_no": atty.get("sbot", ""),
                    "address": atty.get("address", ""),
                    "phone": atty.get("phone", ""),
                    "email": atty.get("email", ""),
                    "party_represented": atty.get("party", ""),
                    "role": role,
                })

        return case_data

    def _populate_case_fields(self, data: dict):
        """Refresh main UI fields from extracted case data."""
        if "deposition_details" not in data and any(
            key in data for key in ("cause_number", "witness_name", "depo_date")
        ):
            data = self._build_case_data_from_ufm_fields(data)

        depo = data.get("deposition_details", {})
        if depo.get("cause_number"):
            self._cause_var.set(depo["cause_number"])
        witness = depo.get("witness", "")
        if witness:
            parts = witness.strip().split()
            if parts:
                self._firstname_var.set(parts[0])
                self._lastname_var.set(parts[-1])
        if depo.get("date"):
            self._date_var.set(depo["date"])

    # ── Actions ──────────────────────────────────────────────────────────────

    def _browse_file(self):
        path = filedialog.askopenfilename(filetypes=AUDIO_VIDEO_EXTENSIONS)
        if path:
            self._ingest_selected_audio(path)

    def _open_combine_dialog(self):
        """Open the multi-file Combine dialog and ingest the result, if any.

        The dialog returns dialog.result_path on success (combined audio
        written under {case_root}/source_docs/_combined/) or None on
        cancel. Cancel leaves the existing file selection unchanged.
        """
        from ui.dialog_combine_audio import CombineAudioDialog

        # The dialog appends "_combined/" to whatever directory we pass,
        # so source_docs/ here yields source_docs/_combined/<file>.
        case_audio_dir = self._resolve_combine_output_dir()

        dialog = CombineAudioDialog(parent=self, case_audio_dir=case_audio_dir)
        self.wait_window(dialog)

        if dialog.result_path:
            self._ingest_selected_audio(str(dialog.result_path))

    def _ingest_selected_audio(self, path: str):
        """Single ingestion path for both Browse… and Multiple Files…
        — keeps the two entry points aligned on case-state reset and
        downstream metadata extraction."""
        if self._correction_mode:
            self._clear_correction_mode()
        self._reset_case_state()
        self._selected_file = path
        self._set_entry_text(self._file_entry, path)
        self._apply_filename_extraction(path)
        self.after(300, self._auto_detect_source_docs)

    def _resolve_combine_output_dir(self) -> Path:
        """Decide where the combined audio should land.

        Preferred: {case_root}/source_docs/. The dialog appends a
        "_combined/" subfolder. We deliberately reuse source_docs/ rather
        than introducing a new top-level audio/ folder so the case
        structure stays predictable across cases (REQUIRED_SUBFOLDERS in
        core/file_manager.py only knows about source_docs/ and Deepgram/).

        Fallback: when case info isn't fully filled in yet, drop the
        combined output under TEMP_DIR with a unique subfolder name. The
        user can move it manually later if needed; the underlying
        combine_audio_files call still succeeds.
        """
        case_root = (self._get_current_save_path() or "").strip()
        if case_root and Path(case_root).is_dir():
            # Make sure source_docs/ exists for cases that haven't had
            # a NOD upload yet — the combiner will create _combined/
            # under it.
            self._create_case_folders_now()
            return Path(case_root) / "source_docs"

        # Fallback path. Use a per-session unique subfolder so two
        # multi-file combines in the same session don't clobber each
        # other.
        from config import TEMP_DIR
        import time
        return Path(TEMP_DIR) / f"combined_{int(time.time())}"

    def _apply_filename_extraction(self, filepath: str):
        """Parse the audio filename for witness name."""
        from core.pdf_extractor import extract_from_filename

        results = extract_from_filename(filepath)

        _BADGE = {
            "filename": ("\U0001f7e9 Filename", "#44AA66"),
            "failed":   ("\u26a0\ufe0f Manual", "#888888"),
        }

        # Witness Last Name
        witness_val, witness_src = results.get("witness_last", (None, "failed"))
        if witness_val:
            self._lastname_var.set(witness_val)
            txt, col = _BADGE.get(witness_src, _BADGE["failed"])
            self._witness_badge.configure(text=txt, text_color=col)

        witness_first, _ = results.get("witness_first", (None, "failed"))
        if witness_first:
            self._firstname_var.set(witness_first)

        # Cause number is never in filename — prompt PDF upload if empty
        if not self._cause_var.get().strip():
            self._extract_status_label.configure(
                text="Cause number and deposition date not in filename \u2014 upload the NOD PDF to extract them.",
                text_color="#CCAA44",
            )
            self._cause_badge.configure(text="\u26a0\ufe0f Manual", text_color="#888888")
        else:
            self._extract_status_label.configure(text="")

    def _browse_base_dir(self):
        path = filedialog.askdirectory(initialdir=self._base_dir_var.get())
        if path:
            self._base_dir_var.set(path)

    def _create_case_folders_now(self):
        """Ensure the current case folder structure exists before writing outputs."""
        from core.file_manager import resolve_or_create_case

        if not self._current_case_path:
            return

        _, status = resolve_or_create_case(
            self._base_dir_var.get().strip(),
            self._cause_var.get().strip(),
            self._lastname_var.get().strip(),
            self._firstname_var.get().strip(),
            self._date_var.get().strip(),
        )
        if status["errors"]:
            logger.error("Folder creation errors: %s", status["errors"])
        if status["created"]:
            logger.info("Folders created: %s", status["created"])

    def _persist_source_doc(self, source_path: str, preferred_name: str | None = None) -> str:
        """
        Copy an uploaded source document into {case_root}/source_docs/.
        Returns the destination path. If case context is incomplete, returns source_path.
        """
        if not source_path or not os.path.isfile(source_path):
            return source_path
        self._create_case_folders_now()
        case_root = self._get_current_save_path()
        if not case_root or not os.path.isdir(case_root):
            return source_path

        source_docs = os.path.join(case_root, "source_docs")
        os.makedirs(source_docs, exist_ok=True)
        filename = preferred_name or os.path.basename(source_path)
        dest = os.path.join(source_docs, filename)
        if os.path.abspath(source_path) != os.path.abspath(dest):
            shutil.copy2(source_path, dest)
        return dest

    @staticmethod
    def _is_supported_source_doc(path: str | Path) -> bool:
        return Path(path).suffix.lower() in _SOURCE_DOC_SUFFIXES

    def _list_supported_source_docs(self, case_root: str) -> list[str]:
        source_docs = Path(case_root) / "source_docs"
        if not source_docs.is_dir():
            return []

        excluded_txt = {
            "transcript.txt",
            "transcript_corrected.txt",
            "deepgram_raw_transcript.txt",
        }

        paths: list[Path] = []
        for path in source_docs.iterdir():
            if not path.is_file():
                continue
            if path.suffix.lower() not in _SOURCE_DOC_SUFFIXES:
                continue
            if path.suffix.lower() == ".txt" and path.name.lower() in excluded_txt:
                continue
            paths.append(path)

        return [str(path) for path in sorted(paths, key=lambda item: item.stat().st_mtime, reverse=True)]

    @staticmethod
    def _summarize_loaded_filenames(file_paths: list[str], prefix: str = "Loaded") -> str:
        names = [Path(path).name for path in file_paths]
        text = f"{prefix} {len(names)} file(s): {', '.join(names)}"
        max_len = 110
        if len(text) <= max_len:
            return text

        truncated = text[: max_len - 3].rstrip(", ")
        return f"{truncated}..."

    def _sync_source_docs_to_case_folder(self) -> None:
        if not self._source_doc_paths:
            return

        synced_paths: list[str] = []
        for path in self._source_doc_paths:
            try:
                synced_paths.append(self._persist_source_doc(path))
            except Exception as exc:
                logger.warning("[SourceDocs] Could not persist source document %s: %s", path, exc)
                synced_paths.append(path)
        self._source_doc_paths = synced_paths

    def _open_output_folder(self):
        """
        Folder picker for the BASE save folder. Per-case folders are
        derived from this root + cause/witness/date, so picking here
        determines where every output goes.
        """
        path = filedialog.askdirectory(
            title="Choose folder to save transcripts to",
            initialdir=self._base_dir_var.get(),
            mustexist=True,
        )
        logger.info("[UI] Output folder picked: %s", path or "(canceled)")
        if not path:
            return
        self._base_dir_var.set(path)

    def _open_transcript(self):
        """
        Open the formatted deposition document when one exists.
        Otherwise fall back to a picker for an existing transcript/document.
        """
        preferred_path = self._get_preferred_document_path()
        if preferred_path:
            logger.info("[UI] Open preferred document: %s", preferred_path)
            self._last_transcript_path = preferred_path
            os.startfile(preferred_path)
            return

        initial_dir = ""
        base_dir = self._base_dir_var.get().strip()
        if base_dir and os.path.isdir(base_dir):
            initial_dir = base_dir
        elif self._last_transcript_path and os.path.isfile(self._last_transcript_path):
            initial_dir = os.path.dirname(self._last_transcript_path)
        elif self._current_case_path and os.path.isdir(self._current_case_path):
            initial_dir = self._current_case_path

        path = filedialog.askopenfilename(
            title="Open Transcript or Document",
            initialdir=initial_dir,
            filetypes=[
                ("Word Documents", "*.docx"),
                ("Text Files", "*.txt"),
                ("All Files", "*.*"),
            ],
        )
        logger.info("[UI] Open transcript: %s", path or "(canceled)")
        if not path:
            return
        self._last_transcript_path = path
        self.load_transcript(path)
        os.startfile(path)

    # ── Load Existing Transcript (driven by Tab 2's "Load Case" button) ─────

    def _load_project_folder(self, folder: str):
        """
        Given a project folder (e.g. ...\\coger_matthew), find all files
        and load the case automatically.

        Discovery order:
          1. Find the most recent .txt in Deepgram\\  → pass to _load_existing_case()
          2. Find any PDFs in source_docs\\           → trigger PDF intake

        Shows an error dialog if no Deepgram\\ subfolder or no .txt exists.
        """
        folder = os.path.normpath(folder)

        # ── 1. Find the most recent transcript .txt ───────────────────────────
        deepgram_dir = os.path.join(folder, "Deepgram")
        txt_path = None

        if os.path.isdir(deepgram_dir):
            txt_files = [
                f for f in os.listdir(deepgram_dir)
                if f.endswith(".txt")
                and not f.endswith("_corrected.txt")
                and not f.endswith("_renamed.txt")
            ]
            if txt_files:
                txt_files.sort(
                    key=lambda f: os.path.getmtime(os.path.join(deepgram_dir, f)),
                    reverse=True,
                )
                txt_path = os.path.join(deepgram_dir, txt_files[0])

        if not txt_path:
            from tkinter import messagebox
            messagebox.showerror(
                "No Transcript Found",
                f"No .txt transcript was found in:\n{deepgram_dir}\n\n"
                "Please select the case folder that contains a Deepgram\\ subfolder.",
            )
            return

        # ── 2. Load the transcript and case metadata ──────────────────────────
        self._load_existing_case(txt_path)

        # ── 3. Auto-load NOD / PDF from source_docs\ if present ──────────────
        source_docs_dir = os.path.join(folder, "source_docs")
        if os.path.isdir(source_docs_dir):
            pdf_files = [
                f for f in os.listdir(source_docs_dir)
                if f.lower().endswith(".pdf")
            ]
            if pdf_files:
                pdf_files.sort(
                    key=lambda f: os.path.getmtime(os.path.join(source_docs_dir, f)),
                    reverse=True,
                )
                pdf_path = os.path.join(source_docs_dir, pdf_files[0])
                logger.info("[LoadProject] Auto-loading PDF: %s", pdf_path)
                try:
                    self._handle_pdf_upload(pdf_path=pdf_path, auto_detected=True)
                except Exception as exc:
                    logger.warning("[LoadProject] PDF load failed: %s", exc)

        # ── 4. Restore saved transcription settings ───────────────────────────
        from core.job_config_manager import load_job_config
        job_config_data = load_job_config(folder)
        self._apply_saved_transcription_settings(job_config_data)
        self._auto_detect_source_docs()

        logger.info("[LoadProject] Project loaded from: %s", folder)

    def _load_existing_case(self, txt_path: str):
        """
        Restore a prior case from a saved transcript file.

        Expected structure:
            {base}\\{YYYY}\\{Mon}\\{CauseNumber}\\{last_first}\\Deepgram\\filename.txt
        """
        txt_path = os.path.normpath(txt_path)
        parts = txt_path.split(os.sep)

        cause_number = ""
        witness_last = ""
        witness_first = ""
        depo_date = ""

        try:
            if len(parts) >= 6:
                cause_number = parts[-4]
                witness_folder = parts[-3]
                name_parts = witness_folder.split("_")
                if len(name_parts) >= 2:
                    witness_last = name_parts[0].capitalize()
                    witness_first = name_parts[1].capitalize()
        except (IndexError, ValueError):
            pass

        deepgram_dir = os.path.dirname(txt_path)
        ufm_fields_data = {}

        # Load job_config.json from source_docs/ (sibling of Deepgram/)
        case_folder = str(Path(deepgram_dir).parent)
        job_config_data = {}
        try:
            from core.job_config_manager import load_job_config
            job_config_data = load_job_config(case_folder)
            ufm_fields_data = job_config_data.get("ufm_fields", {})
            if not ufm_fields_data:
                logger.info("[LoadCase] No ufm_fields in job_config.json — UFM review disabled")
        except Exception as exc:
            logger.warning("[LoadCase] Could not load job_config.json: %s", exc)

        if ufm_fields_data.get("depo_date"):
            depo_date = ufm_fields_data["depo_date"]

        self._reset_case_state()
        self._apply_saved_transcription_settings(job_config_data)

        if cause_number:
            self._cause_var.set(cause_number)
            self._cause_badge.configure(text="🟦 From path", text_color="#4488CC")
        if witness_last:
            self._lastname_var.set(witness_last)
            self._witness_badge.configure(text="🟦 From path", text_color="#4488CC")
        if witness_first:
            self._firstname_var.set(witness_first)
        if depo_date:
            self._date_var.set(depo_date)
            self._date_badge.configure(text="🟦 UFM", text_color="#4488CC")

        self._correction_mode = True
        self._loaded_transcript_path = txt_path
        self._loaded_case_folder = deepgram_dir
        self._ufm_fields = ufm_fields_data
        # _open_transcript reads _last_transcript_path; populate it so the
        # toolbar button works in the load-existing-case flow too.
        self._last_transcript_path = txt_path

        witness_display = f"{witness_first} {witness_last}".strip()
        if ufm_fields_data:
            self._extracted_case_data = self._build_case_data_from_ufm_fields(
                ufm_fields_data,
                witness_fallback=witness_display,
            )

        self._set_create_buttons(state="disabled", text="CORRECTION MODE — Transcription Disabled")

        self._current_case_path = os.path.dirname(deepgram_dir)
        self._update_path_preview()

        # Open Output Folder + Open Transcript are pickers now — already enabled.

        logger.info(
            "[CorrectionMode] Loaded transcript: %s | cause=%s witness=%s %s",
            txt_path, cause_number, witness_first, witness_last,
        )

    def _open_loaded_transcript(self):
        """Switch to Transcript tab and load the existing transcript file."""
        if not self._loaded_transcript_path:
            return
        self.load_transcript(self._loaded_transcript_path)
        self.set_status("Existing transcript loaded", "#7DD8E8")

    def load_case_folder(self, folder: str) -> None:
        """
        Public entry point for loading a case folder. Driven from the
        Transcript tab's 'Load Case' button.

        Populates this tab's case metadata (cause, witness, dates),
        auto-loads the NOD PDF from source_docs/, restores saved
        transcription settings from job_config.json, and pushes the
        most recent .txt transcript into the Transcript tab.
        """
        self._load_project_folder(folder)
        if self._loaded_transcript_path:
            self._open_loaded_transcript()

    def _clear_correction_mode(self):
        """Exit correction mode and reset case state for a fresh transcription."""
        self._correction_mode = False
        self._loaded_transcript_path = None
        self._loaded_case_folder = None

        self._set_create_buttons(state="normal", text=_PRIMARY_ACTION_TEXT)

        self._reset_case_state()

        logger.info("[CorrectionMode] Cleared — returned to new transcription mode.")

    def _on_create_transcript(self):
        """
        Trigger transcription from the Transcribe tab and switch to Transcript.
        """
        self.start_transcription()

    def _handle_pdf_upload(self, pdf_path: str | None = None, auto_detected: bool = False):
        """Load one PDF and extract case fields and confirmed spellings."""
        try:
            filepath = pdf_path or filedialog.askopenfilename(
                title="Select Deposition PDF",
                filetypes=[("PDF Files", "*.pdf"), ("All Files", "*.*")],
            )
            if not filepath:
                return

            saved_pdf_path = self._persist_source_doc(filepath)
            self._last_pdf_path = saved_pdf_path
            self._upload_pdf_btn.configure(state="disabled", text="Processing\u2026")
            self._extract_status_label.configure(text="Extracting case info\u2026", text_color="white")

            def _run():
                from core.pdf_extractor import extract_case_info_from_pdf

                results = extract_case_info_from_pdf(
                    saved_pdf_path,
                    progress_callback=lambda msg: self.after(0, self._append_transcript_log, msg),
                )
                self.after(0, self._apply_pdf_results, results, auto_detected)

            threading.Thread(target=_run, daemon=True).start()
        except Exception as exc:
            logger.exception("[TranscribeTab] PDF upload failed")
            messagebox.showerror("PDF Load Error", str(exc))

    def _apply_pdf_results(self, results: dict, auto_detected: bool = False):
        """Populate fields and confirmed spellings from one PDF extraction result."""

        self._upload_pdf_btn.configure(
            state="normal",
            text="PDF Auto-Detected" if auto_detected else "\U0001f4c4  Upload NOD / PDF",
            fg_color="#2A6F3A" if auto_detected else BTN_UTILITY_BLUE,
        )

        _BADGE = {
            "filename": ("\U0001f7e9 Filename", "#44AA66"),
            "regex":    ("\U0001f7e6 Regex", "#4488CC"),
            "ai":       ("\U0001f7e3 AI", "#9966CC"),
            "failed":   ("\u26a0\ufe0f Manual", "#888888"),
        }

        scanned = results.get("scanned", False)
        if scanned:
            self._extract_status_label.configure(
                text="This appears to be a scanned PDF. Please enter fields manually.",
                text_color="#CCAA44",
            )
            for badge in (self._cause_badge, self._witness_badge, self._date_badge):
                txt, col = _BADGE["failed"]
                badge.configure(text=txt, text_color=col)
            # Do NOT set _pdf_already_loaded — allow retry
            return

        # Mark loaded only on a successfully parsed PDF
        self._pdf_already_loaded = True

        # Cause Number
        cause_val, cause_src = results.get("cause_number", (None, "failed"))
        if cause_val and not self._cause_var.get().strip():
            self._cause_var.set(cause_val)
        txt, col = _BADGE.get(cause_src, _BADGE["failed"])
        self._cause_badge.configure(text=txt, text_color=col)

        # Witness Last Name
        witness_val, witness_src = results.get("witness_last", (None, "failed"))
        if witness_val and not self._lastname_var.get().strip():
            self._lastname_var.set(witness_val)
        txt, col = _BADGE.get(witness_src, _BADGE["failed"])
        self._witness_badge.configure(text=txt, text_color=col)

        # Witness First Name
        witness_first_val, _ = results.get("witness_first", (None, "failed"))
        if witness_first_val and not self._firstname_var.get().strip():
            self._firstname_var.set(witness_first_val)

        # Date
        date_val, date_src = results.get("date", (None, "failed"))
        if date_val and not self._date_var.get().strip():
            self._date_var.set(date_val)
            txt, col = _BADGE.get(date_src, _BADGE["failed"])
            self._date_badge.configure(text=txt, text_color=col)
        elif date_val and self._date_var.get().strip():
            txt, col = _BADGE.get(date_src, _BADGE["failed"])
            self._date_badge.configure(text=txt, text_color=col)

        intake_result = results.get("intake_result")
        self._pdf_keyterms = list(results.get("keyterms", []) or [])
        if intake_result:
            self._apply_intake_result_to_ui(intake_result)
            self._extract_status_label.configure(
                text="Case data extracted",
                text_color="#44FF44",
            )
            self._append_transcript_log("Case data extracted from PDF")
            if self._current_case_path:
                self._create_case_folders_now()
                if self._last_pdf_path:
                    try:
                        persisted_pdf = self._persist_source_doc(self._last_pdf_path)
                        if persisted_pdf:
                            self._last_pdf_path = persisted_pdf
                    except Exception as exc:
                        logger.warning("[SourceDocs] Could not persist PDF after extraction: %s", exc)
                self._save_intake_result_to_job_config(intake_result)
                logger.info("[UI] Saved %d deepgram_keyterms to job_config", len(self._pdf_keyterms))
                logger.info("[UI] job_config.json saved: %s", self._current_case_path)

        sources = [cause_src, witness_src, date_src]
        filled = sum(1 for s in sources if s != "failed")
        if not intake_result:
            self._extract_status_label.configure(
                text=f"Extracted {filled}/3 fields from PDF.",
                text_color="#44FF44" if filled == 3 else "#CCAA44",
            )

    def _open_review_dialog(self):
        self._auto_detect_source_docs()
        IntakeReviewDialog(self, self._extracted_case_data)

    def _set_source_doc_status(self, file_paths: list[str], keyterm_count: int | None = None, color: str = "#44FF44"):
        count = len(file_paths)
        if keyterm_count is None:
            text = self._summarize_loaded_filenames(file_paths)
        else:
            text = f"{count} document(s) loaded — {keyterm_count} keyterms extracted"
        self._extract_status_label.configure(text=text, text_color=color)

    def _apply_intake_result_to_ui(self, intake_result) -> None:
        deponent_name = ""
        if intake_result.deponents:
            deponent_name = str(intake_result.deponents[0].get("name", "")).strip()

        self._extracted_case_data = {
            "deposition_details": {
                "cause_number": intake_result.cause_number or "",
                "witness": deponent_name,
                "date": intake_result.deposition_date or "",
                "court": intake_result.court or "",
                "court_caption": intake_result.court or "",
                "case_style": intake_result.case_style or "",
                "method": intake_result.deposition_method or "",
                "ordered_by": intake_result.ordered_by or "",
                "amendment": intake_result.amendment or "",
            },
            "ordering_attorney": dict(intake_result.ordering_attorney or {}),
            "filing_attorney": dict(intake_result.filing_attorney or {}),
            "copy_attorneys": list(intake_result.copy_attorneys or []),
            "all_attorneys": [],
            "court_reporter": {
                "name": intake_result.reporter_name or "",
                "csr_number": intake_result.reporter_csr or "",
                "agency": intake_result.reporter_firm or "",
            },
            "discrepancies": [],
        }
        self._populate_case_fields(self._extracted_case_data)
        self._review_btn.configure(state="normal")
        self._confirmed_spellings = dict(intake_result.confirmed_spellings or {})
        self._speaker_map_suggestion = dict(intake_result.speaker_map_suggestion or {})
        self._pdf_keyterms = self._build_keyterms_from_intake(intake_result)

    def _save_intake_result_to_job_config(self, intake_result) -> None:
        if not self._current_case_path:
            return

        self._create_case_folders_now()
        self._sync_source_docs_to_case_folder()

        from core.job_config_manager import merge_and_save
        from core.ufm_field_mapper import map_intake_to_ufm

        ufm_fields = map_intake_to_ufm(self._extracted_case_data)
        merge_and_save(
            self._current_case_path,
            ufm_fields=ufm_fields,
            confirmed_spellings=dict(intake_result.confirmed_spellings or {}),
            deepgram_keyterms=self._pdf_keyterms or None,
            speaker_map_suggestion=dict(intake_result.speaker_map_suggestion or {}),
            intake_entity_counts=dict(intake_result.entity_counts or {}),
        )

    def _apply_source_docs_results(
        self,
        file_paths: list[str],
        combined_text: str,
        intake_result,
        keyterms: list[str],
        auto_detected: bool = False,
    ) -> None:
        self._source_doc_paths = list(file_paths)
        self._source_docs_text = combined_text
        self._source_docs_keyterms = list(keyterms or [])

        self._upload_reporter_notes_btn.configure(
            text="Docs Auto-Detected" if auto_detected else "Docs Loaded",
            fg_color="#2A6F3A",
            state="normal",
        )

        if intake_result:
            self._apply_intake_result_to_ui(intake_result)
            self._save_intake_result_to_job_config(intake_result)

        self._set_source_doc_status(file_paths, keyterm_count=len(self._source_docs_keyterms))
        self._append_transcript_log(
            f"{'Auto-detected' if auto_detected else 'Loaded'} source documents: "
            + ", ".join(Path(path).name for path in file_paths)
        )

    def _on_source_docs_load_failed(self, error_msg: str) -> None:
        self._source_docs_text = ""
        self._source_docs_keyterms = []
        self._source_doc_paths = []
        self._upload_reporter_notes_btn.configure(
            text="Load Failed",
            fg_color="#8B0000",
            state="normal",
        )
        self._extract_status_label.configure(
            text=f"Source document load failed: {error_msg}",
            text_color="#FF4444",
        )

    def _load_source_documents(self, file_paths: list[str], auto_detected: bool = False):
        supported_paths: list[str] = []
        for path in file_paths:
            if not self._is_supported_source_doc(path):
                logger.warning("[SourceDocs] Unsupported file skipped: %s", path)
                continue
            try:
                saved_path = str(path) if auto_detected else self._persist_source_doc(str(path))
            except Exception as exc:
                logger.warning("[SourceDocs] Could not persist %s: %s", path, exc)
                saved_path = str(path)
            supported_paths.append(saved_path)

        deduped_paths = list(dict.fromkeys(supported_paths))
        if not deduped_paths:
            return

        self._source_doc_paths = deduped_paths
        self._set_source_doc_status(deduped_paths)
        self._upload_reporter_notes_btn.configure(state="disabled", text="Processing…")

        def _run():
            from core.intake_parser import parse_intake_document
            from core.keyterm_extractor import extract_keyterms_from_text
            from core.source_docs_extractor import extract_text_from_files

            try:
                combined_text = extract_text_from_files([Path(path) for path in deduped_paths])
                intake_result = None
                if combined_text.strip():
                    intake_result = parse_intake_document(
                        "",
                        progress_callback=lambda msg: self.after(0, self._append_transcript_log, msg),
                        extracted_text=combined_text,
                    )
                keyterms = extract_keyterms_from_text(combined_text)
                self.after(
                    0,
                    self._apply_source_docs_results,
                    deduped_paths,
                    combined_text,
                    intake_result,
                    keyterms,
                    auto_detected,
                )
            except Exception as exc:
                logger.exception("[SourceDocs] Source document load failed")
                self.after(0, self._on_source_docs_load_failed, str(exc))

        def _start_worker():
            self._extract_status_label.configure(
                text=f"{len(deduped_paths)} document(s) loaded — extracting keyterms...",
                text_color="white",
            )
            threading.Thread(target=_run, daemon=True).start()

        self.after(250, _start_worker)

    def _upload_nod_and_notes(self):
        """Open a multi-select dialog for NOD and supporting notes documents."""
        filepaths = filedialog.askopenfilenames(
            title="Select NOD and Notes",
            filetypes=[
                ("Supported documents", "*.pdf *.docx *.txt"),
                ("PDF files", "*.pdf"),
                ("Word documents", "*.docx"),
                ("Text files", "*.txt"),
                ("All files", "*.*"),
            ],
        )
        if not filepaths:
            return

        try:
            self._load_source_documents(list(filepaths), auto_detected=False)
        except Exception as exc:
            self._source_docs_text = ""
            self._source_docs_keyterms = []
            self._source_doc_paths = []
            self._upload_reporter_notes_btn.configure(
                text="Load Failed",
                fg_color="#8B0000",
                state="normal",
            )
            messagebox.showerror("NOD and Notes Error", str(exc))

    # ── Transcription Flow ───────────────────────────────────────────────────

    def start_transcription(self):
        """Public entry point called by TranscriptTab."""
        try:
            if self._correction_mode:
                messagebox.showinfo(
                    "Correction Mode Active",
                    "A transcript is loaded for correction.\n"
                    "Click '✕ Clear' in the Load panel to start a new transcription.",
                )
                return

            # Validate file
            if not self._selected_file or not os.path.isfile(self._selected_file):
                messagebox.showerror("No file selected", "Please select an audio or video file first.")
                return

            if self._model_var.get() not in {"nova-3", "nova-3-medical"}:
                messagebox.showerror("Model Missing", "Please select a transcription model before continuing.")
                return

            self._auto_detect_source_docs()

            # Validate required filing fields
            cause = self._cause_var.get().strip()
            last = self._lastname_var.get().strip()
            first = self._firstname_var.get().strip()
            date = self._date_var.get().strip()

            missing = []
            if not cause:
                missing.append("Cause Number")
            if not last:
                missing.append("Last Name")
            if not first:
                missing.append("First Name")
            if not date:
                missing.append("Deposition Date")

            if missing:
                messagebox.showerror(
                    "Required Fields Missing",
                    "Please complete the following fields before transcribing:\n\n"
                    + "\n".join(f"  \u2022 {f}" for f in missing)
                    + "\n\nThese are required to save the transcript to the "
                    "correct folder."
                )
                return

            # Validate API key
            from config import DEEPGRAM_API_KEY

            if not DEEPGRAM_API_KEY or not DEEPGRAM_API_KEY.strip():
                messagebox.showerror(
                    "API Key Missing",
                    "DEEPGRAM_API_KEY is not set.\nAdd it to your .env file and restart.",
                )
                return

            self._append_transcript_log("Transcription started")
            self._append_transcript_log("Processing audio...")
            self._extract_status_label.configure(
                text="Transcription started.",
                text_color="#44FF44",
            )

            # Disable button
            self._running = True
            self._set_create_buttons(state="disabled", text="Transcribing...")
            self._open_folder_btn.configure(state="disabled")
            self._open_transcript_btn.configure(state="disabled")
            self.set_transcription_running()

            # Launch background thread
            thread = threading.Thread(target=self._run_job, daemon=True)
            thread.start()
        except Exception as exc:
            logger.exception("[TranscribeTab] start_transcription failed")
            self._running = False
            self._status_progress.set(0)
            self._set_create_buttons(state="normal", text=_PRIMARY_ACTION_TEXT)
            messagebox.showerror("Transcription Error", str(exc))

    def _run_job(self):
        from core.job_runner import run_transcription_job
        from core.keyterm_extractor import merge_keyterms

        self._create_case_folders_now()
        final_keyterms, _, _ = merge_keyterms(self._pdf_keyterms, self._source_docs_keyterms)
        if final_keyterms:
            self._append_transcript_log(f"Using {len(final_keyterms)} Deepgram keyterms")

        run_transcription_job(
            audio_path=self._selected_file,
            model=self._model_var.get(),
            quality=self._quality_var.get(),
            base_dir=self._base_dir_var.get(),
            cause_number=self._cause_var.get(),
            last_name=self._lastname_var.get(),
            first_name=self._firstname_var.get(),
            date_str=self._date_var.get(),
            keyterms=final_keyterms or None,
            confirmed_spellings=self._confirmed_spellings,
            ufm_fields=self._ufm_fields or None,
            progress_callback=self._on_progress,
            log_callback=self._on_log,
            done_callback=self._on_done,
        )

    # ── Callbacks (called from background thread, dispatch via after()) ──────

    def _on_progress(self, percent: float, message: str):
        self.after(0, self._update_progress, percent, message)

    def _on_log(self, message: str):
        self.after(0, self._append_transcript_log, message)

    def _on_done(self, result: dict):
        self.after(0, self._finish, result)

    def _update_progress(self, percent: float, message: str):
        normalized = percent / 100 if percent > 1 else percent
        self._status_progress.set(max(0, min(normalized, 0.85)))
        self._set_transcript_status(message, "white")

    def _finish(self, result: dict):
        self._running = False

        if result.get("success"):
            self._last_transcript_path = result.get("transcript_path")
            self._current_txt_path = result.get("transcript_path")
            self._transcript_text = result.get("transcript_text", "")
            self._last_output_dir = result.get("output_dir", "")
            # Open Output Folder + Open Transcript are pickers — already enabled.
            self._set_create_buttons(state="normal", text=_PRIMARY_ACTION_TEXT)
            self._open_folder_btn.configure(state="normal")
            self._open_transcript_btn.configure(state="normal")

            # Show speaker labels section
            self._show_speaker_section()

            if self._last_transcript_path and os.path.isfile(self._last_transcript_path):
                self.set_transcription_complete(
                    transcript_path=self._last_transcript_path,
                    folder_path=self._current_case_path or self._last_output_dir,
                )

            # Enable review button if case data has been extracted
            if self._extracted_case_data:
                self._review_btn.configure(state="normal")

            tier = result.get("audio_tier", "")
            if tier == "CLEAN":
                self._audio_tier_label.configure(text=" CLEAN audio", text_color="#44BB44")
            elif tier == "ENHANCED":
                self._audio_tier_label.configure(text=" ENHANCED processing", text_color="#DDAA00")
            elif tier == "RESCUE":
                self._audio_tier_label.configure(text=" RESCUE processing", text_color="#DD4444")
            else:
                self._audio_tier_label.configure(text="", text_color="gray")
            self._append_transcript_log("Transcription complete")
            self._start_clean_format(result)
        else:
            error_msg = result.get("error", "Unknown error")
            self._set_create_buttons(state="normal", text=_PRIMARY_ACTION_TEXT)
            self._audio_tier_label.configure(text="", text_color="gray")
            self.set_transcription_failed(error_msg)
            messagebox.showerror("Transcription Failed", error_msg)

    def _build_clean_format_case_meta(self) -> dict[str, Any]:
        from clean_format.formatter import build_case_meta_from_ufm
        from core.job_config_manager import load_job_config

        ufm_fields: dict[str, Any] = {}
        if self._current_case_path:
            config_data = load_job_config(self._current_case_path)
            ufm_fields = dict(config_data.get("ufm_fields", {}) or {})

        if not ufm_fields:
            witness_name = " ".join(
                part for part in (self._firstname_var.get().strip(), self._lastname_var.get().strip()) if part
            )
            ufm_fields = {
                "cause_number": self._cause_var.get().strip(),
                "court_caption": "",
                "county": "",
                "judicial_district": "",
                "depo_date": self._date_var.get().strip(),
                "depo_time_start": "",
                "depo_time_end": "",
                "witness_name": witness_name,
                "plaintiff_name": "",
                "defendant_name": "",
                "reporter_name": "",
                "csr_number": "",
                "plaintiff_counsel": [],
                "defense_counsel": [],
            }

        return build_case_meta_from_ufm(ufm_fields)

    def _start_clean_format(self, result: dict) -> None:
        self._status_progress.set(0.92)
        self._set_transcript_status("Formatting transcript...", "white")
        self._append_transcript_log("Formatting transcript...")
        self._set_create_buttons(state="disabled", text="Formatting...")

        thread = threading.Thread(
            target=self._run_clean_format_job,
            args=(result,),
            daemon=True,
        )
        thread.start()

    def _run_clean_format_job(self, result: dict) -> None:
        try:
            from clean_format import format_transcript, write_deposition_docx

            case_dir = Path(result.get("output_dir") or "")
            raw_path = Path(result.get("raw_txt_path") or result.get("transcript_path") or "")
            case_meta = self._build_clean_format_case_meta()
            raw_text = raw_path.read_text(encoding="utf-8")

            case_meta_path = case_dir / "case_meta.json"
            case_meta_path.write_text(
                json.dumps(case_meta, indent=2, ensure_ascii=False),
                encoding="utf-8",
            )

            formatted_text = format_transcript(raw_text, case_meta)
            witness_last = (case_meta.get("witness_name", "Witness").split() or ["Witness"])[-1]
            date_part = str(case_meta.get("deposition_date", "")).replace("/", "-").replace(",", "")
            docx_path = case_dir / f"{witness_last}_Deposition_{date_part}.docx"
            saved_path = write_deposition_docx(formatted_text, case_meta, docx_path)

            self.after(0, self._on_clean_format_done, {"success": True, "docx_path": saved_path})
        except Exception as exc:
            logger.exception("[TranscribeTab] clean_format failed")
            self.after(0, self._on_clean_format_done, {"success": False, "error": str(exc)})

    def _on_clean_format_done(self, result: dict) -> None:
        self._set_create_buttons(state="normal", text=_PRIMARY_ACTION_TEXT)

        if result.get("success"):
            self._formatted_docx_path = result.get("docx_path")
            self._last_transcript_path = self._formatted_docx_path

            # Verify the file actually landed on disk before claiming
            # success. NTFS can silently write into an alternate data
            # stream when a filename contains illegal chars (notably ':').
            # The writer sanitizes now, but defense in depth: never tell
            # the reporter their deposition is ready when it isn't there.
            reported_path = self._formatted_docx_path
            if not reported_path or not Path(reported_path).is_file():
                path_text = str(reported_path) if reported_path else "(no path returned)"
                logger.error(
                    "[TranscribeTab] clean_format reported success but file not found at: %s",
                    path_text,
                )
                self._append_transcript_log(
                    f"ERROR: document not found at reported path: {path_text}"
                )
                self._status_progress.set(0)
                self._set_transcript_status(
                    f"Document write failed — file not found at {path_text}",
                    "#FF4444",
                )
                messagebox.showerror(
                    "Document Write Failed",
                    "The transcription completed but no document was found at:\n\n"
                    f"{path_text}\n\n"
                    "Check the run log for details.",
                )
                return

            self._append_transcript_log(f"Deposition document written to: {self._formatted_docx_path}")
            self._status_progress.set(1)
            self._set_transcript_status(
                f"Deposition document written to: {self._formatted_docx_path}",
                "#44FF44",
            )
            if messagebox.askyesno(
                "Document Ready",
                f"Deposition document written to:\n{self._formatted_docx_path}\n\nOpen the document?",
            ):
                os.startfile(self._formatted_docx_path)
        else:
            error_msg = result.get("error", "Unknown error")
            self._status_progress.set(0)
            self._append_transcript_log(f"Formatting failed: {error_msg}")
            self._set_transcript_status(f"Formatting failed: {error_msg}", "#FF4444")

    # ── Speaker Label Methods ────────────────────────────────────────────────

    def _show_speaker_section(self):
        """Scan transcript for speaker IDs, rebuild rows, and show the card."""
        # Clear previous rows and entries
        for widget in self._speaker_rows_frame.winfo_children():
            widget.destroy()
        self._speaker_entries.clear()

        reference_text = _build_ui_speaker_reference_text(self._speaker_map_suggestion)
        self._speaker_hint_label.configure(
            text=(
                f"NOD suggestions loaded.  {reference_text}  |  Review before applying."
                if reference_text
                else ""
            )
        )

        # Find all unique speaker IDs in the transcript
        speakers = sorted(set(re.findall(r'Speaker (\d+):', self._transcript_text)))
        suggested_defaults = _build_ui_speaker_defaults(
            speakers,
            self._saved_speaker_map,
            self._speaker_map_suggestion,
        )
        quickfill_labels = _build_ui_quickfill_labels(self._speaker_map_suggestion)

        for sid in speakers:
            row = ctk.CTkFrame(self._speaker_rows_frame, fg_color="transparent")
            row.pack(fill="x", pady=2)

            main_row = ctk.CTkFrame(row, fg_color="transparent")
            main_row.pack(fill="x")

            ctk.CTkLabel(
                main_row, text=f"Speaker {sid}:", width=100, anchor="w",
                font=ctk.CTkFont(weight="bold"),
            ).pack(side="left", padx=(0, 8))

            entry = ctk.CTkEntry(main_row, placeholder_text="e.g. THE REPORTER")
            entry.pack(side="left", fill="x", expand=True)
            default_label = suggested_defaults.get(f"Speaker {sid}", "")
            if default_label:
                entry.insert(0, default_label)
            self._speaker_entries[f"Speaker {sid}"] = entry

            if quickfill_labels:
                actions_row = ctk.CTkFrame(row, fg_color="transparent")
                actions_row.pack(fill="x", padx=(108, 0), pady=(4, 0))

                ctk.CTkLabel(
                    actions_row,
                    text="Quick fill:",
                    font=ctk.CTkFont(size=11),
                    text_color="#7DAACC",
                ).pack(side="left", padx=(0, 6))

                for label in quickfill_labels:
                    ctk.CTkButton(
                        actions_row,
                        text=label,
                        height=24,
                        width=max(96, len(label) * 7),
                        fg_color=BTN_UTILITY_BLUE,
                        hover_color=BTN_UTILITY_BLUE_HOVER,
                        font=ctk.CTkFont(size=11),
                        command=lambda e=entry, value=label: self._set_speaker_entry_value(e, value),
                    ).pack(side="left", padx=(0, 6))

        self._speaker_card.grid(row=2, column=0, sticky="ew", padx=10, pady=(0, _SECTION_GAP_Y))

    @staticmethod
    def _set_speaker_entry_value(entry: ctk.CTkEntry, value: str):
        entry.delete(0, "end")
        entry.insert(0, value)

    def _apply_and_save_labels(self):
        """
        Apply speaker label assignments in two ways:

        1. Text replacement in the .txt transcript file
           (Speaker 0: → MR. GARCIA:, etc.)

        2. Persist the speaker_map to self._ufm_fields and write it to
           the ufm_fields.json file on disk so the corrections pipeline
           and AI corrector automatically use the correct speaker names.
        """
        if not self._current_txt_path:
            messagebox.showerror("No file", "No transcript file path available.")
            return

        # ── Build the speaker map from UI entries ────────────────────────────
        # speaker_map format expected by JobConfig: {int_id: "DISPLAY NAME"}
        # e.g. {0: "THE VIDEOGRAPHER", 1: "THE WITNESS", 2: "MR. GARCIA"}
        speaker_map: dict[int, str] = {}
        for original_label, entry in self._speaker_entries.items():
            replacement = " ".join(entry.get().split()).strip()
            if replacement:
                try:
                    sid = int(original_label.replace("Speaker ", "").strip())
                    speaker_map[sid] = replacement
                except ValueError:
                    pass

        # ── 1. Text replacement in the .txt file ─────────────────────────────
        text = _apply_speaker_labels_to_text(self._transcript_text, speaker_map)

        try:
            Path(self._current_txt_path).write_text(text, encoding="utf-8")
        except Exception as exc:
            messagebox.showerror("Save Failed", str(exc))
            return

        self._transcript_text = text

        # ── 2. Persist speaker_map to job_config.json → source_docs/ ─────────
        if speaker_map:
            self._ufm_fields["speaker_map"] = speaker_map
            self._ufm_fields["speaker_map_verified"] = True

            try:
                from core.job_config_manager import load_job_config, merge_and_save
                case_folder = str(Path(self._current_txt_path).parent.parent)
                job_config_data = load_job_config(case_folder)
                ufm = dict(job_config_data.get("ufm_fields", {}))
                ufm["speaker_map"] = {str(k): v for k, v in speaker_map.items()}
                ufm["speaker_map_verified"] = True
                merge_and_save(case_folder, ufm_fields=ufm)
                logger.info("[SpeakerLabels] Wrote speaker_map to job_config.json")
            except Exception as exc:
                logger.warning("[SpeakerLabels] Could not update job_config.json: %s", exc)

        # ── 3. Update UI ──────────────────────────────────────────────────────
        self.load_transcript(self._current_txt_path)
        self.set_status("Speaker labels applied", "#44FF44")
        self.append_log(
            f"Speaker labels applied: "
            + ", ".join(f"Speaker {k} → {v}" for k, v in speaker_map.items())
        )
        self._apply_save_btn.configure(
            text="\u2713  Labels Saved",
            fg_color="#1A5C1A",
        )
        self.after(2500, lambda: self._apply_save_btn.configure(
            text="\u2713  Apply Speaker Labels",
            fg_color=BTN_UTILITY_BLUE,
        ))

    # ── Extraction callback (called externally when AI extraction finishes) ──

    def set_extracted_case_data(self, data: dict):
        """Store AI-extracted case data and enable the review button."""
        self._extracted_case_data = data
        self._review_btn.configure(state="normal")
        self._populate_case_fields(data)
