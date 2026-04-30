"""
Extract text from supported source documents for intake and keyterm parsing.
"""

from __future__ import annotations

from pathlib import Path

from app_logging import get_logger
from core.pdf_extractor import extract_pdf_text

logger = get_logger(__name__)


def _extract_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def extract_text_from_files(file_paths: list[Path]) -> str:
    """
    Returns one concatenated string with file boundaries marked by:
      ===== {filename} =====
      {extracted text}
    """
    sections: list[str] = []

    for raw_path in file_paths:
        path = Path(raw_path)
        suffix = path.suffix.lower()

        try:
            if suffix == ".pdf":
                text = extract_pdf_text(str(path))
            elif suffix == ".docx":
                text = _extract_docx(path)
            elif suffix == ".txt":
                text = path.read_text(encoding="utf-8", errors="replace")
            else:
                logger.warning("[SourceDocsExtractor] Unsupported file type skipped: %s", path.name)
                continue
        except Exception as exc:
            logger.warning("[SourceDocsExtractor] Failed to extract %s: %s", path.name, exc)
            continue

        sections.append(f"===== {path.name} =====\n{text.strip()}")

    return "\n\n".join(section for section in sections if section.strip())
