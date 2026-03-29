"""
core/confidence_docx_exporter.py

Export a transcript review DOCX with confidence-highlighted words and a
summary page for low-confidence review.
"""

from __future__ import annotations

from pathlib import Path
from typing import Callable
import re

from docx import Document
from docx.shared import Pt, RGBColor

from core.word_data_loader import (
    CONFIDENCE_MEDIUM,
    CONFIDENCE_LOW,
    get_flagged_summary,
)


def _color_for(confidence: float):
    if confidence < CONFIDENCE_LOW:
        return RGBColor(0xFF, 0x8C, 0x00)
    if confidence < CONFIDENCE_MEDIUM:
        return RGBColor(0xCC, 0xCC, 0x00)
    return None


def export_confidence_docx(
    transcript_path: str,
    words: list[dict],
    output_path: str | None = None,
    progress_callback: Callable[[str], None] | None = None,
) -> str:
    def log(message: str) -> None:
        if progress_callback:
            progress_callback(message)

    source = Path(transcript_path)
    if not source.is_file():
        raise FileNotFoundError(f"Transcript not found: {transcript_path}")

    target = Path(output_path or source.with_name(f"{source.stem}_confidence_review.docx"))
    summary = get_flagged_summary(words)
    conf_index: dict[str, float] = {}
    for word in words or []:
        text = str(word.get("word") or word.get("text") or "").strip().lower()
        try:
            confidence = float(word.get("confidence") or 1.0)
        except (TypeError, ValueError):
            confidence = 1.0
        if text:
            conf_index[text] = min(conf_index.get(text, 1.0), confidence)

    lines = source.read_text(encoding="utf-8").splitlines()

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Courier New"
    normal.font.size = Pt(11)

    section = doc.sections[0]
    section.left_margin = int(1.25 * 1440)
    section.right_margin = int(1.0 * 1440)
    section.top_margin = int(1.0 * 1440)
    section.bottom_margin = int(1.0 * 1440)

    heading = doc.add_paragraph()
    heading.add_run("Confidence Review Summary").bold = True
    for line in [
        f"Transcript: {source.name}",
        f"Total words: {summary['total']}",
        f"Flagged words: {summary['flagged']}",
        f"Medium: {summary['medium']}",
        f"Low: {summary['low']}",
        f"Critical: {summary['critical']}",
    ]:
        doc.add_paragraph(line)

    doc.add_page_break()

    for raw_line in lines:
        line = raw_line.rstrip()
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)

        if not line:
            continue

        tokens = re.split(r"(\s+)", line)
        for token in tokens:
            run = para.add_run(token)
            run.font.name = "Courier New"
            run.font.size = Pt(11)
            key = token.strip().lower()
            if key:
                color = _color_for(conf_index.get(key, 1.0))
                if color:
                    run.font.color.rgb = color

    target.parent.mkdir(parents=True, exist_ok=True)
    log(f"Saving review DOCX: {target.name}")
    doc.save(target)
    return str(target)
