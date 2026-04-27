from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

from core.docx_formatter import (
    _iter_formatted_lines,
    build_docx_from_transcript_text,
    build_full_docx_from_text,
    build_full_output_docx_path,
    build_output_docx_path,
    format_full_transcript_to_docx,
    format_transcript_to_docx,
)
from spec_engine.models import JobConfig, LineType


def test_iter_formatted_lines_classifies_supported_line_types():
    text = (
        "\tQ.  Did you see that?\n"
        "\tA.  Yes.\n"
        "\t\t\tTHE REPORTER:  Please raise your right hand.\n"
        "(Whereupon, the witness was sworn.)\n"
        "[SCOPIST: FLAG 1: Verify amount]\n"
        "EXAMINATION\n"
        "BY MR. SMITH:\n"
        "Loose text"
    )

    result = list(_iter_formatted_lines(text))

    assert result == [
        (LineType.Q, "Did you see that?"),
        (LineType.A, "Yes."),
        (LineType.SP, "THE REPORTER:  Please raise your right hand."),
        (LineType.PN, "(Whereupon, the witness was sworn.)"),
        (LineType.FLAG, "[SCOPIST: FLAG 1: Verify amount]"),
        (LineType.HEADER, "EXAMINATION"),
        (LineType.BY, "BY MR. SMITH:"),
        (LineType.PLAIN, "Loose text"),
    ]


def test_build_docx_from_transcript_text_uses_spec_engine_document_defaults():
    text = (
        "\tQ.  Did you see that?\n"
        "\tA.  Yes.\n"
        "\t\t\tTHE REPORTER:  Please raise your right hand."
    )

    doc = build_docx_from_transcript_text(text)

    assert len(doc.paragraphs) == 3
    section = doc.sections[0]
    assert section.left_margin == Inches(1.25)
    assert section.right_margin == Inches(1.0)
    assert section.top_margin == Inches(1.0)
    assert section.bottom_margin == Inches(1.0)
    first_run = doc.paragraphs[0].runs[0]
    assert first_run.font.name == "Courier New"
    assert first_run.font.size == Pt(12)


def test_build_output_docx_path_strips_corrected_suffix():
    output = build_output_docx_path(r"C:\tmp\sample_corrected.txt")

    assert output.endswith("sample_formatted.docx")


def test_format_transcript_to_docx_creates_file_and_reports_progress(tmp_path):
    source = tmp_path / "sample_corrected.txt"
    source.write_text(
        "\tQ.  Did you see that?\n"
        "\tA.  Yes.\n"
        "\t\t\tTHE REPORTER:  Please raise your right hand.\n",
        encoding="utf-8",
    )
    output = tmp_path / "formatted.docx"
    progress = []

    saved_path = format_transcript_to_docx(
        str(source),
        output_path=str(output),
        progress_callback=progress.append,
    )

    assert saved_path == str(output)
    assert output.exists()
    assert progress == [
        "Reading transcript: sample_corrected.txt",
        "Saved DOCX: formatted.docx",
    ]

    doc = Document(saved_path)
    assert len(doc.paragraphs) == 3
    # Old-format input ("\tQ.  ") in the source file goes through
    # _iter_formatted_lines which parses it into LineType.Q and re-emits
    # via emit_q_line / _qa_visual_text — the Phase F'd emitter. Round-
    # trip behavior: pre-Phase-F transcripts on disk are normalized to
    # the new tab-tab format on DOCX export. Input strings above stay
    # in old format intentionally — they represent real artifacts users
    # have on disk; the test verifies the function handles them.
    assert "\tQ.\tDid you see that?" in doc.paragraphs[0].text
    assert "\tA.\tYes." in doc.paragraphs[1].text


def test_format_transcript_to_docx_raises_for_missing_source(tmp_path):
    missing = tmp_path / "missing.txt"

    try:
        format_transcript_to_docx(str(missing))
        assert False, "expected FileNotFoundError"
    except FileNotFoundError as exc:
        assert str(missing) in str(exc)


# ── Full DOCX path (title page + caption + body + certificate) ───────────────

def _minimal_job_config() -> JobConfig:
    """Build a JobConfig fleshed-out enough that the page writers don't
    raise. Real production configs come from job_config.json — this just
    exercises the function plumbing."""
    return JobConfig(
        cause_number="2025-CI-23267",
        plaintiff_name="Bryan Roque Reyes",
        defendant_names=["Peter Durai Singh"],
        case_style="Bryan Roque Reyes v. Peter Durai Singh",
        court_type="District Court",
        county="Bexar County",
        state="Texas",
        judicial_district="408TH",
        depo_date="April 23, 2026",
        witness_name="Peter Durai Singh",
        reporter_name="Miah Bardot",
        speaker_map={1: "MR. GONZALEZ", 2: "MR. PENA", 3: "MR. SINGH"},
        examining_attorney_id=1,
        witness_id=3,
        speaker_map_verified=True,
    )


def test_build_full_output_docx_path_strips_corrected_suffix():
    output = build_full_output_docx_path(r"C:\tmp\sample_corrected.txt")
    assert output.endswith("sample_full.docx")


def test_build_full_output_docx_path_distinct_from_shallow():
    # The full path must NOT collide with the shallow path's filename,
    # so a user can have both versions side-by-side for comparison.
    shallow = build_output_docx_path(r"C:\tmp\sample_corrected.txt")
    full = build_full_output_docx_path(r"C:\tmp\sample_corrected.txt")
    assert shallow != full
    assert shallow.endswith("_formatted.docx")
    assert full.endswith("_full.docx")


def _all_text(doc) -> str:
    """Walk paragraphs AND table cells. The page writers (title_page,
    caption, certificate) render content in 25-line bordered tables —
    a paragraph-only walk misses them entirely."""
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    parts.append(p.text)
    return "\n".join(parts)


def test_build_full_docx_from_text_includes_caption_and_certificate():
    job_config = _minimal_job_config()
    text = (
        "\tQ.  Did you see the collision?\n"
        "\tA.  Yes, sir.\n"
    )

    doc = build_full_docx_from_text(text, job_config)
    full_text = _all_text(doc)

    # Title page / caption emit the cause number (inside a table).
    assert "2025-CI-23267" in full_text
    # Witness intro emits the EXAMINATION header (in body paragraphs).
    assert "EXAMINATION" in full_text
    # BY MR. <examiner> line.
    assert "BY MR. GONZALEZ" in full_text
    # Body lines preserved.
    assert "Did you see the collision?" in full_text
    assert "Yes, sir." in full_text
    # Certificate page emits something with "Reporter" or CSR text.
    assert "Reporter" in full_text or "CERTIFICATE" in full_text.upper()


def test_build_full_docx_from_text_witness_name_in_intro():
    job_config = _minimal_job_config()
    doc = build_full_docx_from_text("\tQ.  X?\n\tA.  Y.\n", job_config)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    # Witness intro line: "PETER DURAI SINGH,"
    assert "PETER DURAI SINGH" in full_text


def test_format_full_transcript_to_docx_creates_file(tmp_path):
    source = tmp_path / "sample_corrected.txt"
    source.write_text(
        "\tQ.  Did you see that?\n"
        "\tA.  Yes.\n",
        encoding="utf-8",
    )
    output = tmp_path / "out_full.docx"
    progress = []

    saved_path = format_full_transcript_to_docx(
        str(source),
        _minimal_job_config(),
        output_path=str(output),
        progress_callback=progress.append,
    )

    assert saved_path == str(output)
    assert output.exists()
    # Progress reported the build steps.
    assert any("Building full DOCX" in msg for msg in progress)
    assert any("Saved full DOCX" in msg for msg in progress)


def test_format_full_transcript_to_docx_raises_for_missing_source(tmp_path):
    missing = tmp_path / "missing.txt"
    try:
        format_full_transcript_to_docx(str(missing), _minimal_job_config())
        assert False, "expected FileNotFoundError"
    except FileNotFoundError as exc:
        assert str(missing) in str(exc)


def test_full_docx_path_default_uses_full_suffix(tmp_path):
    # When no output_path is provided, the saved file lands at the
    # default <stem>_full.docx location next to the source.
    source = tmp_path / "deposition_corrected.txt"
    source.write_text("\tQ.  X?\n\tA.  Y.\n", encoding="utf-8")
    saved_path = format_full_transcript_to_docx(
        str(source),
        _minimal_job_config(),
    )
    assert saved_path.endswith("deposition_full.docx")
