from pathlib import Path
import sys

from docx import Document
from reportlab.pdfgen import canvas

sys.path.insert(0, str(Path(__file__).resolve().parents[2]))

from core.source_docs_extractor import extract_text_from_files


def _make_pdf(path: Path, lines: list[str]) -> Path:
    pdf = canvas.Canvas(str(path))
    y = 760
    for line in lines:
        pdf.drawString(72, y, line)
        y -= 18
    pdf.save()
    return path


def _make_docx(path: Path, paragraphs: list[str], table_rows: list[list[str]] | None = None) -> Path:
    doc = Document()
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    if table_rows:
        table = doc.add_table(rows=0, cols=len(table_rows[0]))
        for row_values in table_rows:
            row = table.add_row().cells
            for idx, value in enumerate(row_values):
                row[idx].text = value
    doc.save(path)
    return path


def test_source_docs_extractor_pdf(tmp_path):
    pdf_path = _make_pdf(tmp_path / "notice.pdf", ["Cause No. 2025-CI-19595", "Bianca Caram"])

    result = extract_text_from_files([pdf_path])

    assert "===== notice.pdf =====" in result
    assert "Cause No. 2025-CI-19595" in result
    assert "Bianca Caram" in result


def test_source_docs_extractor_docx(tmp_path):
    docx_path = _make_docx(
        tmp_path / "notes.docx",
        ["Doctor Bianca Caram", "Legacy Women's Health"],
        table_rows=[["Witness", "Bianca Caram"]],
    )

    result = extract_text_from_files([docx_path])

    assert "===== notes.docx =====" in result
    assert "Doctor Bianca Caram" in result
    assert "Witness" in result


def test_source_docs_extractor_txt(tmp_path):
    txt_path = tmp_path / "notes.txt"
    txt_path.write_text("Michelle Maloney\nBilly Dunnill", encoding="utf-8")

    result = extract_text_from_files([txt_path])

    assert "===== notes.txt =====" in result
    assert "Michelle Maloney" in result
    assert "Billy Dunnill" in result


def test_source_docs_extractor_mixed(tmp_path):
    pdf_path = _make_pdf(tmp_path / "notice.pdf", ["Cause No. DC-25-13430"])
    docx_path = _make_docx(tmp_path / "summary.docx", ["Bianca Caram", "Miah Bardot"])
    txt_path = tmp_path / "notes.txt"
    txt_path.write_text("Alex Video", encoding="utf-8")

    result = extract_text_from_files([pdf_path, docx_path, txt_path])

    assert "===== notice.pdf =====" in result
    assert "===== summary.docx =====" in result
    assert "===== notes.txt =====" in result
    assert "DC-25-13430" in result
    assert "Bianca Caram" in result
    assert "Alex Video" in result


def test_source_docs_extractor_unsupported_skips(tmp_path, caplog):
    txt_path = tmp_path / "notes.txt"
    txt_path.write_text("Valid text", encoding="utf-8")
    bad_path = tmp_path / "ignore.xyz"
    bad_path.write_text("skip me", encoding="utf-8")

    result = extract_text_from_files([bad_path, txt_path])

    assert "Valid text" in result
    assert "ignore.xyz" not in result
    assert "Unsupported file type skipped: ignore.xyz" in caplog.text


def test_source_docs_extractor_handles_corrupt_file(tmp_path, caplog):
    corrupt_pdf = tmp_path / "corrupt.pdf"
    corrupt_pdf.write_text("not a real pdf", encoding="utf-8")
    txt_path = tmp_path / "notes.txt"
    txt_path.write_text("Still works", encoding="utf-8")

    result = extract_text_from_files([corrupt_pdf, txt_path])

    assert "Still works" in result
    assert "===== notes.txt =====" in result
    assert "corrupt.pdf" in caplog.text
