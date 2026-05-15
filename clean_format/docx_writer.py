"""DOCX writer for clean-format deposition output.

This module handles block parsing and paragraph rendering only.
Document structure (caption, appearances, proceedings, certificate)
will be added via UFM templates in a later phase.

Line-type contract (matches formatter.py and prompt.py §11):

  Q/A     \tQ.\tbody        \tA.\tbody
  Label   \t\t\tLABEL:  body   (three tabs, colon, two spaces, body)
  Paren   \t\t\t\t(text)         (four tabs, parenthetical)
  Header  EXAMINATION  /  CROSS-EXAMINATION  (flush left, no tabs)
  BY-line BY MS.  MALONEY:             (flush left, ends with colon)

Tab-stop positions (360/900/1440/2160/2880 twips = 0.25"/0.625"/1.0"/1.5"/2.0"):
  \t     → 0.25"   (Q./A. letter)
  \t\t   → 0.625"  (Q./A. body start)
  \t\t\t → 1.0"    (speaker label position)
  \t\t\t\t → 1.5"  (parenthetical position)
"""

from __future__ import annotations

import re
import time
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX, WD_LINE_SPACING
from docx.shared import Inches, Pt, Twips

from clean_format.low_confidence_markers import split_into_runs

# ── Tab-stop and indent constants ─────────────────────────────────────────────

# All five UFM tab stops in twips.  Added to every transcript paragraph
# via _add_tab_stops() so the tab characters in formatter.py output land
# at the correct columns.
_TAB_STOPS_TWIPS = [360, 900, 1440, 2160, 2880]

# Q/A hanging-indent geometry:
#   First line at margin → \t → Q./A. at 0.25" → \t → body at 0.625".
#   Continuation wraps to 0.625" (body column).
_QA_LEFT_INDENT      = Inches(0.625)
_QA_FIRST_LINE_DELTA = Inches(-0.625)

# Speaker-label hanging-indent geometry:
#   \t\t\t → label at 1.0".  Continuation wraps to 1.0".
_SP_LEFT_INDENT      = Inches(1.0)
_SP_FIRST_LINE_DELTA = Inches(-1.0)

# Examination-header keywords (flush left, bold, centered).
_EXAM_HEADERS = frozenset({
    "EXAMINATION",
    "CROSS-EXAMINATION",
    "REDIRECT EXAMINATION",
    "RECROSS-EXAMINATION",
    "FURTHER EXAMINATION",
})

# Regex helpers
_SENTENCE_SPACE_RE      = re.compile(r"([.!?])\s+")
_ILLEGAL_FILENAME_CHARS = re.compile(r'[<>:"/\\|?*]|[\x00-\x1f]')
_WHITESPACE_RUN         = re.compile(r"\s+")
_TRAILING_AT_SUFFIX     = re.compile(r"\s+at\s+.*$", re.IGNORECASE)


# ── File helpers ───────────────────────────────────────────────────────────────

def sanitize_filename_component(value: str) -> str:
    """Strip Windows-illegal characters from a filename component."""
    cleaned = _ILLEGAL_FILENAME_CHARS.sub("", value or "")
    cleaned = _WHITESPACE_RUN.sub(" ", cleaned).strip().strip(". ")
    return cleaned or "document"


def safe_save(
    document: Document, path: Path, *, retries: int = 3, delay_seconds: float = 1.0
) -> None:
    """Retry transient Word lock failures before surfacing a clean error."""
    for attempt in range(retries):
        try:
            document.save(path)
            return
        except PermissionError:
            if attempt == retries - 1:
                break
            time.sleep(delay_seconds)
    raise PermissionError(f"File is open or locked:\n{path}")


def _format_date_for_filename(raw_date: str) -> str:
    candidate = _TRAILING_AT_SUFFIX.sub("", raw_date or "").strip()
    for fmt in ("%m/%d/%Y", "%Y-%m-%d", "%B %d, %Y"):
        try:
            return datetime.strptime(candidate, fmt).strftime("%Y-%m-%d")
        except ValueError:
            continue
    return (candidate or raw_date).replace("/", "-").replace(",", "")


# ── Document setup ─────────────────────────────────────────────────────────────

def _set_document_defaults(document: Document) -> None:
    """Apply Miah Bardot / UFM page layout and default font to a Document."""
    section = document.sections[0]
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)

    style = document.styles["Normal"]
    style.font.name = "Courier New"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    style.paragraph_format.space_after = Pt(0)


def _add_tab_stops(paragraph: Any) -> None:
    """Add all five UFM tab stops to a paragraph."""
    for twip_value in _TAB_STOPS_TWIPS:
        paragraph.paragraph_format.tab_stops.add_tab_stop(Twips(twip_value))


# ── Text helpers ───────────────────────────────────────────────────────────────

def _double_space_sentences(text: str) -> str:
    """Ensure two spaces after every sentence-ending . ! ? (Miah spec)."""
    return _SENTENCE_SPACE_RE.sub(r"\1  ", (text or "").strip())


def _add_marked_runs(paragraph: Any, text: str) -> None:
    """Append runs to paragraph, rendering ‹LC:word› tokens with yellow highlight.

    Marker characters are stripped; the wrapped token text is preserved.
    When text has no markers, produces a single unstyled run.
    """
    runs = split_into_runs(text)
    if not runs:
        return
    for chunk, is_low_confidence in runs:
        run = paragraph.add_run(chunk)
        if is_low_confidence:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


# ── Block parser ───────────────────────────────────────────────────────────────

def _parse_blocks(formatted_text: str) -> list[dict[str, str]]:
    """Parse formatter.py plain-text output into structured rendering blocks.

    Recognised kinds:

      "qa"        \tQ.\tbody  /  \tA.\tbody  (also legacy Q.\t / A.\t)
      "sp_label"  \t\t\tLABEL:  body         speaker label + body, same line
      "paren"     \t\t\t\t(text)              procedural parenthetical
      "header"    EXAMINATION / CROSS-EXAMINATION / ...  flush-left bold
      "by_line"   BY MS.  MALONEY:            flush-left, ends with colon
      "directive" everything else             three-tab indent, no label

    Accepts both the current formatter output (colon + two spaces after label)
    and the legacy format (colon + tab) for backward compatibility.
    """
    blocks: list[dict[str, str]] = []

    for raw_block in (formatted_text or "").split("\n\n"):
        lines = [ln.rstrip() for ln in raw_block.splitlines() if ln.strip()]
        if not lines:
            continue
        first = lines[0]

        # ── Q/A ──────────────────────────────────────────────────────────────
        # Strip any leading tab before checking so both \tQ.\t and Q.\t match.
        stripped = first.lstrip("\t")
        if stripped.startswith("Q.\t"):
            text = stripped[3:]
            for extra in lines[1:]:
                tail = extra.lstrip("\t").strip()
                if tail:
                    text += " " + tail
            blocks.append({"kind": "qa", "label": "Q.", "text": text})
            continue
        if stripped.startswith("A.\t"):
            text = stripped[3:]
            for extra in lines[1:]:
                tail = extra.lstrip("\t").strip()
                if tail:
                    text += " " + tail
            blocks.append({"kind": "qa", "label": "A.", "text": text})
            continue

        # ── Parenthetical ─────────────────────────────────────────────────────
        # Must be checked before the three-tab sp_label path.
        if first.startswith("\t\t\t\t"):
            blocks.append({"kind": "paren", "label": "", "text": first[4:].strip()})
            continue

        # ── Speaker label ─────────────────────────────────────────────────────
        # Three tabs + LABEL:  body  (new format, two spaces after colon)
        # or legacy LABEL:\tbody  (no leading tabs, tab after colon).
        if first.startswith("\t\t\t"):
            body = first[3:]
            colon_idx = body.find(":")
            if colon_idx > 0:
                label = body[:colon_idx + 1]           # includes colon
                text  = body[colon_idx + 1:].lstrip()  # strip leading spaces/tab
                blocks.append({"kind": "sp_label", "label": label, "text": text})
            else:
                blocks.append({"kind": "directive", "label": "", "text": body.strip()})
            continue

        if ":\t" in first:
            label, text = first.split(":\t", 1)
            blocks.append({"kind": "sp_label", "label": label + ":", "text": text})
            continue

        # ── Examination header ────────────────────────────────────────────────
        if first.strip() in _EXAM_HEADERS:
            blocks.append({"kind": "header", "label": first.strip(), "text": ""})
            continue

        # ── BY-line ───────────────────────────────────────────────────────────
        if first.startswith("BY ") and first.rstrip().endswith(":"):
            blocks.append({"kind": "by_line", "label": first.strip(), "text": ""})
            continue

        # ── Fallback directive ────────────────────────────────────────────────
        text = first.lstrip("\t").strip()
        for extra in lines[1:]:
            tail = extra.lstrip("\t").strip()
            if tail:
                text += "\n" + tail
        if text:
            blocks.append({"kind": "directive", "label": "", "text": text})

    return _merge_consecutive_sp_label_blocks(blocks)


def _merge_consecutive_sp_label_blocks(
    blocks: list[dict[str, str]],
) -> list[dict[str, str]]:
    """Merge consecutive same-speaker sp_label blocks.

    When the formatter splits a long speaker utterance across two paragraph
    separators, consecutive blocks with identical labels are joined so the
    label appears only once in the rendered document.
    """
    merged: list[dict[str, str]] = []
    for block in blocks:
        if (
            merged
            and block["kind"] == "sp_label"
            and merged[-1]["kind"] == "sp_label"
            and block["label"]
            and block["label"] == merged[-1]["label"]
        ):
            merged[-1]["text"] = merged[-1]["text"].rstrip() + "  " + block["text"].lstrip()
            continue
        merged.append(block)
    return merged


def _merge_consecutive_speaker_blocks(
    blocks: list[dict[str, str]],
) -> list[dict[str, str]]:
    """Backward-compatible alias for older callers/tests."""
    return _merge_consecutive_sp_label_blocks(blocks)


# ── Block renderer ─────────────────────────────────────────────────────────────

def _render_block(document: Document, block: dict[str, str]) -> None:
    """Render one parsed block as a Word paragraph with UFM formatting."""
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    paragraph.paragraph_format.space_after = Pt(0)
    _add_tab_stops(paragraph)

    kind  = block["kind"]
    label = block.get("label", "")
    text  = block.get("text", "")

    if kind == "qa":
        # Hanging indent: first line at margin, \t → Q./A. at 0.25",
        # \t → body at 0.625".  Continuation wraps to 0.625".
        paragraph.paragraph_format.left_indent      = _QA_LEFT_INDENT
        paragraph.paragraph_format.first_line_indent = _QA_FIRST_LINE_DELTA
        paragraph.add_run(f"\t{label}\t")
        _add_marked_runs(paragraph, text)

    elif kind == "sp_label":
        # Three tabs → label at 1.0" → body on the same line.
        # Label is already formatted by formatter.py (e.g. "MS.  MALONEY:").
        # Continuation wraps to 1.0".
        paragraph.paragraph_format.left_indent      = _SP_LEFT_INDENT
        paragraph.paragraph_format.first_line_indent = _SP_FIRST_LINE_DELTA
        run = paragraph.add_run(f"\t\t\t{label}  ")
        run.bold = True
        _add_marked_runs(paragraph, _double_space_sentences(text))

    elif kind == "paren":
        # Four tabs → parenthetical at 1.5".
        paragraph.paragraph_format.left_indent      = Inches(1.5)
        paragraph.paragraph_format.first_line_indent = Inches(-1.5)
        paragraph.add_run(f"\t\t\t\t{text}")

    elif kind == "header":
        # Examination header: centered, bold, no indent.
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.left_indent      = Inches(0)
        paragraph.paragraph_format.first_line_indent = Inches(0)
        paragraph.add_run(label).bold = True

    elif kind == "by_line":
        # BY-line: flush left, bold.
        paragraph.paragraph_format.left_indent      = Inches(0)
        paragraph.paragraph_format.first_line_indent = Inches(0)
        paragraph.add_run(label).bold = True

    else:
        # directive / fallback: three-tab indent, no label.
        paragraph.paragraph_format.left_indent      = _SP_LEFT_INDENT
        paragraph.paragraph_format.first_line_indent = _SP_FIRST_LINE_DELTA
        paragraph.add_run("\t\t\t")
        _add_marked_runs(paragraph, _double_space_sentences(text))


# ── Public rendering entry-point ───────────────────────────────────────────────

def render_transcript_body(document: Document, formatted_text: str) -> None:
    """Parse and render all blocks from formatter.py output into document.

    Called by the UFM template writer after it has set up page layout,
    caption, and appearances.  All tab stops and indentation geometry
    match prompt.py §11 and formatter.py output contract.
    """
    for block in _parse_blocks(formatted_text):
        _render_block(document, block)


# ── Bridge function — temporary until UFM templates are wired in ───────────────

def write_deposition_docx(
    formatted_text: str,
    case_meta: dict[str, Any],
    output_path: str | Path | None = None,
) -> str:
    """Create a DOCX file from clean-format text.

    TEMPORARY BRIDGE:  This function produces a body-only document —
    page layout, font, and tab stops are applied via _set_document_defaults,
    but the UFM cover page, caption table, appearances, and reporter's
    certificate are NOT included.  Those sections will be added by the
    UFM template writer in a later phase.  The function signature is
    intentionally identical to the original so existing callers in
    tab_transcribe.py and clean_format/__init__.py continue to work
    without modification.

    Replace this function with the full UFM template writer call when
    the templates are ready.
    """
    witness_last = (
        case_meta.get("witness_name", "Witness").split() or ["Witness"]
    )[-1]
    date_part = _format_date_for_filename(
        str(case_meta.get("deposition_date", ""))
    )
    if output_path:
        path = Path(output_path)
    else:
        path = Path(f"{witness_last}_Deposition_{date_part}.docx")

    suffix         = path.suffix or ".docx"
    sanitized_stem = sanitize_filename_component(path.stem)
    path           = path.with_name(f"{sanitized_stem}{suffix}")
    path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    _set_document_defaults(document)
    render_transcript_body(document, formatted_text)

    safe_save(document, path)
    return str(path)
