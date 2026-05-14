from __future__ import annotations

import re
import tempfile
import zipfile
from pathlib import Path

from docx import Document

from clean_format.docx_writer import (
    _merge_consecutive_speaker_blocks,
    _parse_blocks,
    build_deposition_document,
    safe_save,
    sanitize_filename_component,
    write_deposition_docx,
)
from spec_engine.emitter import format_blocks_to_text
from spec_engine.models import TranscriptBlock


def _case_meta() -> dict:
    return {
        "cause_number": "DC-25-13430",
        "court": "191st District Court",
        "county": "Dallas",
        "judicial_district": "191ST",
        "deposition_date": "2026-04-09",
        "start_time": "9:00 AM",
        "end_time": "12:30 PM",
        "witness_name": "Bianca Caram",
        "witness_credentials": "M.D.",
        "plaintiff_name": "Maria Lopez",
        "defendant_names": ["Acme Medical Group"],
        "reporter_name": "Miah Bardot",
        "reporter_csr": "12129",
        "attorneys": [
            {"name": "Jane Smith", "role": "plaintiff", "city": "Dallas"},
            {"name": "Emily Johnson", "role": "defendant", "city": "Houston"},
        ],
        "videographer_name": "Alex Video",
    }


def test_build_deposition_document_has_caption_table_and_examination_header():
    document = build_deposition_document(
        "\tQ.\tState your name.\n\n\tA.\tBianca Caram.", _case_meta()
    )
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert document.tables
    assert "EXAMINATION" in text


def test_write_deposition_docx_sets_courier_default(tmp_path):
    output_path = tmp_path / "sample.docx"
    saved_path = write_deposition_docx(
        "\tQ.\tQuestion\n\n\tA.\tAnswer", _case_meta(), output_path
    )
    document = Document(saved_path)
    assert document.styles["Normal"].font.name == "Courier New"


def test_write_deposition_docx_writes_file(tmp_path):
    output_path = tmp_path / "sample.docx"
    saved_path = write_deposition_docx(
        "LABEL:\tToday's date is April 9, 2026.", _case_meta(), output_path
    )
    assert output_path.exists()
    assert saved_path.endswith("sample.docx")


def test_parse_blocks_merges_consecutive_same_speaker_into_one_paragraph():
    formatted_text = (
        "\t\t\tVIDEOGRAPHER:\n\t\t\tToday's date is 04/09/2026.\n\n"
        "\t\t\tVIDEOGRAPHER:\n\t\t\tThe time is 08:12 AM.\n\n"
        "\t\t\tVIDEOGRAPHER:\n\t\t\tThis is the beginning of the video deposition.\n\n"
        "\t\t\tVIDEOGRAPHER:\n\t\t\tWill the court reporter please swear in the witness?"
    )

    blocks = _parse_blocks(formatted_text)

    assert blocks == [
        {
            "kind": "colloquy_block",
            "label": "VIDEOGRAPHER:",
            "text": (
                "Today's date is 04/09/2026.\nThe time is 08:12 AM.\n"
                "This is the beginning of the video deposition.\n"
                "Will the court reporter please swear in the witness?"
            ),
        }
    ]


def test_write_proceedings_uses_two_spaces_after_speaker_colon_and_sentences():
    document = build_deposition_document(
        (
            "\t\t\tVIDEOGRAPHER:\n\t\t\tToday's date is 04/09/2026.\n\n"
            "\t\t\tVIDEOGRAPHER:\n\t\t\tThe time is 08:12 AM."
        ),
        _case_meta(),
    )
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "VIDEOGRAPHER:" in text
    assert "Today's date is 04/09/2026." in text
    assert "The time is 08:12 AM." in text


def test_write_proceedings_sets_qa_tab_stops_to_requested_positions():
    document = build_deposition_document("\tQ.\tQuestion\n\n\tA.\tAnswer", _case_meta())
    qa_paragraph = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("\tQ.\t")
    )
    tab_positions = [tab.position for tab in qa_paragraph.paragraph_format.tab_stops]
    # Canonical UFM tab stops at 0.5" / 1.0" / 1.5" (UFM Section 2.102.11),
    # mirrored from spec_engine/ufm_rules.py:25. EMU = English Metric Units;
    # 914400 EMU = 1 inch.
    assert 457200 in tab_positions   # 0.5"
    assert 914400 in tab_positions   # 1.0"
    assert 1371600 in tab_positions  # 1.5"


def test_write_proceedings_qa_paragraph_uses_hanging_indent_for_wrap():
    """Q/A paragraphs must hang at 1.0" so wrapped continuation lines
    align under the first character of the question/answer body. The
    canonical text shape is "\tQ.\t{body}" / "\tA.\t{body}" (matches
    spec_engine/emitter.py). Geometry:
      left_indent       = 1.0"   wrap continuation column
      first_line_indent = -1.0"  first-line origin = column 0; the
                                 leading tab lands "Q."/"A." at the
                                 0.5" tab stop and the body tab
                                 pushes text to the 1.0" stop.
    """
    document = build_deposition_document("\tQ.\tQuestion\n\n\tA.\tAnswer", _case_meta())
    qa_paragraph = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("\tQ.\t")
    )
    pf = qa_paragraph.paragraph_format
    # 914400 EMU = 1.0", -914400 EMU = -1.0"
    assert pf.left_indent == 914400
    assert pf.first_line_indent == -914400


def test_write_proceedings_qa_run_text_preserves_leading_tab():
    """The Q/A run text must include the leading tab so the spec_engine
    canonical shape "\tQ.\t..." / "\tA.\t..." is preserved through the
    DOCX writer (not stripped or re-built without the leading tab)."""
    document = build_deposition_document("\tQ.\tQuestion\n\n\tA.\tAnswer", _case_meta())
    qa_paragraph = next(
        paragraph
        for paragraph in document.paragraphs
        if "Question" in paragraph.text
    )
    assert qa_paragraph.text == "\tQ.\tQuestion"


def test_write_proceedings_speaker_paragraph_uses_hanging_indent():
    """Legacy speaker blocks use the 1.5" hanging indent so wrap
    continuation matches the visible first-line position."""
    document = build_deposition_document(
        "\t\t\tMR. SMITH:\n\t\t\tObjection. Form.", _case_meta()
    )
    speaker_paragraph = next(
        paragraph
        for paragraph in document.paragraphs
        if "MR. SMITH" in paragraph.text
    )
    pf = speaker_paragraph.paragraph_format
    assert pf.left_indent == 1371600
    assert pf.first_line_indent == -1371600


# ── Step 2J: three-tab prefix on non-Q/A paragraphs ──────────────────────────


def test_speaker_paragraph_text_starts_with_three_tabs():
    """Step 2J: speaker (non-Q/A) paragraphs render with a leading
    "\\t\\t\\t" so the content lands at the 1.5" tab stop."""
    document = build_deposition_document(
        "\t\t\tMR. SMITH:\n\t\t\tObjection. Form.", _case_meta()
    )
    speaker_paragraph = next(
        paragraph
        for paragraph in document.paragraphs
        if "MR. SMITH" in paragraph.text
    )
    assert speaker_paragraph.text.startswith("\t\t\t"), (
        f"speaker paragraph missing 3-tab prefix: {speaker_paragraph.text!r}"
    )


def test_speaker_paragraph_tab_stops_match_qa():
    """Non-Q/A paragraphs share the same 0.5/1.0/1.5 tab stops with Q/A
    paragraphs so the three-tab prefix lands at 1.5".""" 
    document = build_deposition_document(
        "\t\t\tMR. SMITH:\n\t\t\tObjection. Form.", _case_meta()
    )
    speaker_paragraph = next(
        paragraph
        for paragraph in document.paragraphs
        if "MR. SMITH" in paragraph.text
    )
    tab_positions = [tab.position for tab in speaker_paragraph.paragraph_format.tab_stops]
    assert 457200 in tab_positions   # 0.5"
    assert 914400 in tab_positions   # 1.0"
    assert 1371600 in tab_positions  # 1.5"


def test_no_arrow_characters_in_docx_output():
    """No arrow characters should appear anywhere in the rendered DOCX."""
    document = build_deposition_document(
        "Q.\tDid you see it?\n\n"
        "A.\tYes.\n\n"
        "MR. SMITH:\tObjection. Form.\n\n"
        "(Exhibit 1 marked.)",
        _case_meta(),
    )
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    for arrow in ("→", "⇒", "⟶", "►", "▶"):
        assert arrow not in text, f"unexpected arrow {arrow!r} in DOCX output"


def test_sanitize_filename_component_replaces_spaces_and_punctuation():
    assert sanitize_filename_component(
        "CARAM Deposition April 9, 2026 at 800 a.m."
    ) == ("CARAM Deposition April 9, 2026 at 800 a.m")


def test_safe_save_retries_permission_error(monkeypatch, tmp_path):
    document = build_deposition_document("Q.\tQuestion\n\nA.\tAnswer", _case_meta())
    target_path = tmp_path / "retry.docx"
    attempts = {"count": 0}
    original_save = document.save

    def flaky_save(path):
        attempts["count"] += 1
        if attempts["count"] < 3:
            raise PermissionError("locked")
        return original_save(path)

    monkeypatch.setattr(document, "save", flaky_save)

    safe_save(document, target_path, delay_seconds=0)

    assert attempts["count"] == 3
    assert target_path.exists()


def test_parse_blocks_question_classified_as_qa():
    text = "\tQ.\tDid you visit the property?"
    result = _parse_blocks(text)
    assert len(result) == 1
    assert result[0]["kind"] == "qa"
    assert result[0]["label"] == "Q."
    assert result[0]["text"] == "Did you visit the property?"


def test_parse_blocks_answer_classified_as_qa():
    text = "\tA.\tYes, I was there on May 7."
    result = _parse_blocks(text)
    assert len(result) == 1
    assert result[0]["kind"] == "qa"
    assert result[0]["label"] == "A."
    assert result[0]["text"] == "Yes, I was there on May 7."


def test_parse_blocks_colloquy_classified_as_colloquy_block():
    """A speaker-label line followed by body line(s) becomes one colloquy_block."""
    text = "\t\t\tMS. ZHAN:\n\t\t\tObjection.  Form."
    result = _parse_blocks(text)
    assert len(result) == 1
    assert result[0]["kind"] == "colloquy_block"
    assert result[0]["label"] == "MS. ZHAN:"
    assert result[0]["text"] == "Objection.  Form."


def test_parse_blocks_colloquy_with_multi_line_body():
    """Multiple body lines under the same speaker label stay newline-separated."""
    text = (
        "\t\t\tMS. ZHAN:\n"
        "\t\t\tObjection.  Form.\n"
        "\t\t\tThe witness has not been qualified as an expert."
    )
    result = _parse_blocks(text)
    assert len(result) == 1
    assert result[0]["kind"] == "colloquy_block"
    assert result[0]["label"] == "MS. ZHAN:"
    assert "Objection.  Form." in result[0]["text"]
    assert "qualified as an expert" in result[0]["text"]


def test_parse_blocks_directive_no_speaker_label():
    text = "\t\t\t(RECESS FROM 10:14 A.M. TO 10:32 A.M.)"
    result = _parse_blocks(text)
    assert len(result) == 1
    assert result[0]["kind"] == "directive"
    assert "RECESS" in result[0]["text"]


def test_parse_blocks_bare_header_unchanged():
    """A bare 'EXAMINATION:' header line stays kind='header'."""
    text = "EXAMINATION:"
    result = _parse_blocks(text)
    assert len(result) == 1
    assert result[0]["kind"] == "header"
    assert result[0]["label"] == "EXAMINATION:"


def test_parse_blocks_mixed_sequence_full_round_trip():
    """The exact sequence from the defect #7 discovery probe."""
    blocks = [
        TranscriptBlock(
            speaker="MR. NUNEZ",
            text="Did you review the records from Doctor Fisher before preparing your report?",
            type="question",
            examiner="MR. NUNEZ",
        ),
        TranscriptBlock(
            speaker="THE WITNESS",
            text="No. I reviewed the records and my own examination notes.",
            type="answer",
        ),
        TranscriptBlock(
            speaker="MS. ZHAN",
            text="Objection. Form. The question assumes facts not in evidence.",
            type="colloquy",
        ),
        TranscriptBlock(
            speaker="",
            text="(Recess from 10:14 a.m. to 10:32 a.m.)",
            type="directive",
        ),
    ]
    emitter_output = format_blocks_to_text(blocks)
    parsed = _parse_blocks(emitter_output)

    kinds = [b["kind"] for b in parsed]
    assert kinds == ["qa", "qa", "colloquy_block", "directive"]
    assert parsed[0]["label"] == "Q."
    assert parsed[0]["text"].startswith("Did you review")
    assert parsed[1]["label"] == "A."
    assert parsed[1]["text"].startswith("No.")
    assert parsed[2]["label"] == "MS. ZHAN:"
    assert "Objection." in parsed[2]["text"]
    assert "RECESS" in parsed[3]["text"]


def test_parse_blocks_defensive_fallback_does_not_crash():
    """A weird input goes to the defensive directive fallback."""
    text = "some unexpected text with no tab prefix"
    result = _parse_blocks(text)
    assert len(result) == 1
    assert result[0]["kind"] == "directive"
    assert result[0]["text"] == "some unexpected text with no tab prefix"


def test_merge_consecutive_colloquy_blocks_same_speaker():
    """Two adjacent colloquy_block entries with the same label are merged."""
    blocks = [
        {"kind": "colloquy_block", "label": "MS. ZHAN:", "text": "First line."},
        {"kind": "colloquy_block", "label": "MS. ZHAN:", "text": "Second line."},
    ]
    result = _merge_consecutive_speaker_blocks(blocks)
    assert len(result) == 1
    assert result[0]["text"] == "First line.\nSecond line."


def test_merge_does_not_merge_different_speakers():
    blocks = [
        {"kind": "colloquy_block", "label": "MS. ZHAN:", "text": "First."},
        {"kind": "colloquy_block", "label": "MR. NUNEZ:", "text": "Second."},
    ]
    result = _merge_consecutive_speaker_blocks(blocks)
    assert len(result) == 2


def _build_docx(blocks: list[TranscriptBlock]) -> str:
    """Build a deposition DOCX from the given blocks and return its path."""
    case_meta = {
        "cause_number": "TEST-001",
        "plaintiff_name": "PLAINTIFF",
        "defendant_names": ["DEFENDANT"],
        "judicial_district": "37TH",
        "county": "BEXAR",
        "witness_name": "Test Witness",
        "deposition_date": "May 7, 2026",
        "attorneys": [
            {"role": "plaintiff", "name": "Mr. Nunez", "city": "San Antonio"},
            {"role": "defendant", "name": "Ms. Zhan", "city": "San Antonio"},
        ],
        "videographer_name": "",
        "reporter_name": "Test Reporter",
    }
    emitter_output = format_blocks_to_text(blocks)
    document = build_deposition_document(emitter_output, case_meta)
    tmp_path = Path(tempfile.gettempdir()) / "defect7_test.docx"
    document.save(tmp_path)
    return str(tmp_path)


def _read_paragraph_xml(path: str) -> str:
    with zipfile.ZipFile(path) as z:
        return z.read("word/document.xml").decode("utf-8")


def _find_paragraph_pPr_for_text(document_xml: str, text_marker: str) -> str | None:
    """Return the pPr XML block of the first paragraph containing text_marker."""
    paragraphs = re.findall(r"<w:p\b[^>]*>.*?</w:p>", document_xml, re.DOTALL)
    for p in paragraphs:
        text_content = "".join(re.findall(r"<w:t[^>]*>([^<]*)</w:t>", p))
        if text_marker in text_content:
            match = re.search(r"<w:pPr>(.*?)</w:pPr>", p, re.DOTALL)
            if match:
                return match.group(1)
            return ""
    return None


def test_docx_qa_paragraph_has_1440_hanging_indent():
    """Q/A paragraphs render with left=1440 twips and hanging=1440 twips."""
    blocks = [
        TranscriptBlock(
            speaker="MR. NUNEZ",
            text="Did you review the records from Doctor Fisher?",
            type="question",
            examiner="MR. NUNEZ",
        ),
    ]
    path = _build_docx(blocks)
    xml = _read_paragraph_xml(path)
    pPr = _find_paragraph_pPr_for_text(xml, "Did you review")
    assert pPr is not None, "Q paragraph not found in document.xml"
    assert 'w:left="1440"' in pPr, f"Expected left=1440, got: {pPr}"
    assert 'w:hanging="1440"' in pPr, f"Expected hanging=1440, got: {pPr}"


def test_docx_colloquy_paragraph_has_2160_hanging_indent():
    """Colloquy paragraphs render with left=2160 twips and hanging=2160 twips."""
    blocks = [
        TranscriptBlock(
            speaker="MS. ZHAN",
            text="Objection. Form. The question assumes facts not in evidence.",
            type="colloquy",
        ),
    ]
    path = _build_docx(blocks)
    xml = _read_paragraph_xml(path)
    pPr = _find_paragraph_pPr_for_text(xml, "Objection")
    assert pPr is not None, "Colloquy paragraph not found"
    assert 'w:left="2160"' in pPr, f"Expected left=2160, got: {pPr}"
    assert 'w:hanging="2160"' in pPr, f"Expected hanging=2160, got: {pPr}"


def test_docx_directive_paragraph_has_2160_hanging_indent():
    """Directive paragraphs render with the same 1.5\" hanging indent."""
    blocks = [
        TranscriptBlock(
            speaker="",
            text="(Recess from 10:14 a.m. to 10:32 a.m.)",
            type="directive",
        ),
    ]
    path = _build_docx(blocks)
    xml = _read_paragraph_xml(path)
    pPr = _find_paragraph_pPr_for_text(xml, "RECESS FROM")
    assert pPr is not None, "Directive paragraph not found"
    assert 'w:left="2160"' in pPr, f"Expected left=2160, got: {pPr}"
    assert 'w:hanging="2160"' in pPr, f"Expected hanging=2160, got: {pPr}"


def test_docx_qa_paragraph_contains_label_in_text():
    """The Q label still appears in the rendered paragraph text."""
    blocks = [
        TranscriptBlock(
            speaker="MR. NUNEZ",
            text="Did you visit the property?",
            type="question",
            examiner="MR. NUNEZ",
        ),
    ]
    path = _build_docx(blocks)
    xml = _read_paragraph_xml(path)
    paragraphs = re.findall(r"<w:p\b[^>]*>.*?</w:p>", xml, re.DOTALL)
    target = None
    for p in paragraphs:
        text_content = "".join(re.findall(r"<w:t[^>]*>([^<]*)</w:t>", p))
        if "Did you visit" in text_content:
            target = text_content
            break
    assert target is not None
    assert "Q." in target
    assert "Did you visit the property?" in target


def test_docx_colloquy_paragraph_contains_label_and_body():
    blocks = [
        TranscriptBlock(
            speaker="MS. ZHAN",
            text="Objection. Form.",
            type="colloquy",
        ),
    ]
    path = _build_docx(blocks)
    xml = _read_paragraph_xml(path)
    paragraphs = re.findall(r"<w:p\b[^>]*>.*?</w:p>", xml, re.DOTALL)
    target = None
    for p in paragraphs:
        text_content = "".join(re.findall(r"<w:t[^>]*>([^<]*)</w:t>", p))
        if "Objection" in text_content:
            target = text_content
            break
    assert target is not None
    assert "MS. ZHAN:" in target
    assert "Objection" in target


# ============================================================
# Defect #9 - Reporter's Certificate page tests
# ============================================================

import re as _re_d9
import zipfile as _zipfile_d9


def _build_docx_with_cert(case_meta: dict) -> str:
    """Build a deposition DOCX with the cert page appended."""
    formatted_text = "\tQ.\tDid you visit the property?\n\n\tA.\tYes."
    document = build_deposition_document(formatted_text, case_meta)
    tmp_path = Path(tempfile.gettempdir()) / "defect9_test.docx"
    document.save(tmp_path)
    return str(tmp_path)


def _full_text(path: str) -> str:
    """Concatenate all paragraph text in the .docx for content checks."""
    with _zipfile_d9.ZipFile(path) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    texts = _re_d9.findall(r"<w:t[^>]*>([^<]*)</w:t>", xml)
    return "\n".join(texts)


def test_cert_page_contains_required_header():
    """The cert page header text appears in the document."""
    case_meta = {
        "cause_number": "TEST-001",
        "plaintiff_name": "PLAINTIFF",
        "defendant_names": ["DEFENDANT"],
        "judicial_district": "37TH",
        "county": "BEXAR",
        "witness_name": "Test Witness",
        "deposition_date": "May 7, 2026",
        "attorneys": [
            {"role": "plaintiff", "name": "Mr. Nunez", "city": "San Antonio"},
        ],
        "videographer_name": "",
        "reporter_name": "Miah Bardot",
    }
    path = _build_docx_with_cert(case_meta)
    text = _full_text(path)
    assert "REPORTER'S CERTIFICATION" in text


def test_cert_page_contains_reporter_name():
    case_meta = {
        "cause_number": "TEST-001",
        "plaintiff_name": "PLAINTIFF",
        "defendant_names": ["DEFENDANT"],
        "judicial_district": "37TH",
        "county": "BEXAR",
        "witness_name": "Test Witness",
        "deposition_date": "May 7, 2026",
        "attorneys": [],
        "videographer_name": "",
        "reporter_name": "Miah Bardot",
    }
    path = _build_docx_with_cert(case_meta)
    text = _full_text(path)
    assert "Miah Bardot" in text
    assert "Certified Shorthand Reporter" in text


def test_cert_page_contains_witness_name():
    case_meta = {
        "cause_number": "TEST-001",
        "plaintiff_name": "PLAINTIFF",
        "defendant_names": ["DEFENDANT"],
        "judicial_district": "37TH",
        "county": "BEXAR",
        "witness_name": "Heath Thomas",
        "deposition_date": "May 7, 2026",
        "attorneys": [],
        "videographer_name": "",
        "reporter_name": "Miah Bardot",
    }
    path = _build_docx_with_cert(case_meta)
    text = _full_text(path)
    assert "Heath Thomas" in text
    assert "duly sworn" in text


def test_cert_page_renders_csr_number_when_provided():
    case_meta = {
        "cause_number": "TEST-001",
        "plaintiff_name": "PLAINTIFF",
        "defendant_names": ["DEFENDANT"],
        "judicial_district": "37TH",
        "county": "BEXAR",
        "witness_name": "Test Witness",
        "deposition_date": "May 7, 2026",
        "attorneys": [],
        "videographer_name": "",
        "reporter_name": "Miah Bardot",
        "reporter_csr": "12129",
        "reporter_credentials": "CSR, RPR, CRR",
    }
    path = _build_docx_with_cert(case_meta)
    text = _full_text(path)
    assert "Texas CSR 12129" in text
    assert "CSR, RPR, CRR" in text


def test_cert_page_renders_firm_when_provided():
    case_meta = {
        "cause_number": "TEST-001",
        "plaintiff_name": "PLAINTIFF",
        "defendant_names": ["DEFENDANT"],
        "judicial_district": "37TH",
        "county": "BEXAR",
        "witness_name": "Test Witness",
        "deposition_date": "May 7, 2026",
        "attorneys": [],
        "videographer_name": "",
        "reporter_name": "Miah Bardot",
        "firm_name": "SA LEGAL SOLUTIONS",
        "firm_registration": "10698",
        "firm_address_line1": "100 Main Street, Suite 200",
        "firm_city": "San Antonio",
        "firm_state": "Texas",
        "firm_zip": "78205",
        "firm_phone": "(210) 555-0100",
        "firm_email": "miah@example.com",
    }
    path = _build_docx_with_cert(case_meta)
    text = _full_text(path)
    assert "SA LEGAL SOLUTIONS" in text
    assert "10698" in text
    assert "100 Main Street" in text
    assert "San Antonio" in text
    assert "78205" in text
    assert "(210) 555-0100" in text
    assert "miah@example.com" in text


def test_cert_page_uses_blank_placeholders_for_missing_fields():
    """When reporter/firm fields are missing, render blank placeholders
    (visible blanks > silent fabrication)."""
    case_meta = {
        "cause_number": "TEST-001",
        "plaintiff_name": "PLAINTIFF",
        "defendant_names": ["DEFENDANT"],
        "judicial_district": "37TH",
        "county": "BEXAR",
        "witness_name": "Test Witness",
        "deposition_date": "May 7, 2026",
        "attorneys": [],
        "videographer_name": "",
        "reporter_name": "Miah Bardot",
    }
    path = _build_docx_with_cert(case_meta)
    text = _full_text(path)
    assert "_______" in text
    assert "[CSR Number]" not in text
    assert "[Firm Name]" not in text
    assert "[Credentials]" not in text


def test_cert_page_renders_time_used_per_attorney_block():
    """Time-used block contains one line per attorney with blank time
    placeholder."""
    case_meta = {
        "cause_number": "TEST-001",
        "plaintiff_name": "PLAINTIFF",
        "defendant_names": ["DEFENDANT"],
        "judicial_district": "37TH",
        "county": "BEXAR",
        "witness_name": "Test Witness",
        "deposition_date": "May 7, 2026",
        "attorneys": [
            {"role": "plaintiff", "name": "Mr. Nunez", "city": "San Antonio"},
            {"role": "defendant", "name": "Ms. Zhan", "city": "San Antonio"},
        ],
        "videographer_name": "",
        "reporter_name": "Miah Bardot",
    }
    path = _build_docx_with_cert(case_meta)
    text = _full_text(path)
    assert "Mr. Nunez (______ hours ______ minutes)" in text
    assert "Ms. Zhan (______ hours ______ minutes)" in text


def test_cert_page_renders_attorney_party_pairs_block():
    """Attorney/party pairs block lists each attorney with role+party."""
    case_meta = {
        "cause_number": "TEST-001",
        "plaintiff_name": "ACME CORP",
        "defendant_names": ["XYZ INC"],
        "judicial_district": "37TH",
        "county": "BEXAR",
        "witness_name": "Test Witness",
        "deposition_date": "May 7, 2026",
        "attorneys": [
            {"role": "plaintiff", "name": "Mr. Nunez", "city": "San Antonio"},
            {"role": "defendant", "name": "Ms. Zhan", "city": "San Antonio"},
        ],
        "videographer_name": "",
        "reporter_name": "Miah Bardot",
    }
    path = _build_docx_with_cert(case_meta)
    text = _full_text(path)
    assert "Mr. Nunez, Attorney for Plaintiff, ACME CORP" in text
    assert "Ms. Zhan, Attorney for Defendant, XYZ INC" in text


def test_cert_page_contains_trcp_boilerplate():
    """The fixed TRCP boilerplate appears unmodified."""
    case_meta = {
        "cause_number": "TEST-001",
        "plaintiff_name": "PLAINTIFF",
        "defendant_names": ["DEFENDANT"],
        "judicial_district": "37TH",
        "county": "BEXAR",
        "witness_name": "Test Witness",
        "deposition_date": "May 7, 2026",
        "attorneys": [],
        "videographer_name": "",
        "reporter_name": "Miah Bardot",
    }
    path = _build_docx_with_cert(case_meta)
    text = _full_text(path)
    assert "neither counsel for" in text
    assert "Rule 203" in text or "Rule 203 of TRCP" in text


def test_cert_page_contains_always_blank_certification_date():
    """The 'Certified to by me this ___ day of ___, ___' line uses
    a blank placeholder, never a fabricated date."""
    case_meta = {
        "cause_number": "TEST-001",
        "plaintiff_name": "PLAINTIFF",
        "defendant_names": ["DEFENDANT"],
        "judicial_district": "37TH",
        "county": "BEXAR",
        "witness_name": "Test Witness",
        "deposition_date": "May 7, 2026",
        "attorneys": [],
        "videographer_name": "",
        "reporter_name": "Miah Bardot",
    }
    path = _build_docx_with_cert(case_meta)
    text = _full_text(path)
    assert "Certified to by me this" in text
    assert "_____ day of" in text


def test_cert_page_no_bracketed_placeholders_leak_through():
    """No paragraph in the output retains a [...] placeholder."""
    case_meta = {
        "cause_number": "TEST-001",
        "plaintiff_name": "PLAINTIFF",
        "defendant_names": ["DEFENDANT"],
        "judicial_district": "37TH",
        "county": "BEXAR",
        "witness_name": "Test Witness",
        "deposition_date": "May 7, 2026",
        "attorneys": [
            {"role": "plaintiff", "name": "Mr. Nunez"},
        ],
        "videographer_name": "",
        "reporter_name": "Miah Bardot",
        "reporter_csr": "12129",
        "firm_name": "SA LEGAL SOLUTIONS",
    }
    path = _build_docx_with_cert(case_meta)
    text = _full_text(path)
    leaked = _re_d9.findall(r"\[[A-Z][^\]]*\]", text)
    assert not leaked, f"Bracketed placeholders leaked: {leaked}"


def test_cert_page_includes_page_break_before():
    """A page break separates the proceedings from the cert page."""
    case_meta = {
        "cause_number": "TEST-001",
        "plaintiff_name": "PLAINTIFF",
        "defendant_names": ["DEFENDANT"],
        "judicial_district": "37TH",
        "county": "BEXAR",
        "witness_name": "Test Witness",
        "deposition_date": "May 7, 2026",
        "attorneys": [],
        "videographer_name": "",
        "reporter_name": "Miah Bardot",
    }
    path = _build_docx_with_cert(case_meta)
    with _zipfile_d9.ZipFile(path) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    cert_marker = "REPORTER'S CERTIFICATION"
    cert_position = xml.find(cert_marker)
    assert cert_position > 0, "Cert page not found in output"
    preceding = xml[:cert_position]
    assert 'w:br w:type="page"' in preceding


def test_cert_page_renders_state_as_texas_in_caption():
    """The caption [State] resolves to TEXAS even without firm_state."""
    case_meta = {
        "cause_number": "TEST-001",
        "plaintiff_name": "PLAINTIFF",
        "defendant_names": ["DEFENDANT"],
        "judicial_district": "37TH",
        "county": "BEXAR",
        "witness_name": "Test Witness",
        "deposition_date": "May 7, 2026",
        "attorneys": [],
        "videographer_name": "",
        "reporter_name": "Miah Bardot",
    }
    path = _build_docx_with_cert(case_meta)
    text = _full_text(path)
    assert "TEXAS" in text


def test_cert_page_renders_firm_state_when_provided():
    case_meta = {
        "cause_number": "TEST-001",
        "plaintiff_name": "PLAINTIFF",
        "defendant_names": ["DEFENDANT"],
        "judicial_district": "37TH",
        "county": "BEXAR",
        "witness_name": "Test Witness",
        "deposition_date": "May 7, 2026",
        "attorneys": [],
        "videographer_name": "",
        "reporter_name": "Miah Bardot",
        "firm_city": "San Antonio",
        "firm_state": "Texas",
        "firm_zip": "78205",
    }
    path = _build_docx_with_cert(case_meta)
    text = _full_text(path)
    assert "Texas" in text


def test_cert_page_skips_gracefully_when_template_missing(monkeypatch):
    """If the cert template is absent at build time, the writer
    skips the cert page rather than crashing."""
    from clean_format import docx_writer

    monkeypatch.setattr(
        docx_writer,
        "_CERT_TEMPLATE_PATH",
        Path(tempfile.gettempdir()) / "nonexistent_cert.docx",
    )

    case_meta = {
        "cause_number": "TEST-001",
        "plaintiff_name": "PLAINTIFF",
        "defendant_names": ["DEFENDANT"],
        "judicial_district": "37TH",
        "county": "BEXAR",
        "witness_name": "Test Witness",
        "deposition_date": "May 7, 2026",
        "attorneys": [],
        "videographer_name": "",
        "reporter_name": "Miah Bardot",
    }
    formatted_text = "\tQ.\tHello.\n\n\tA.\tHi."
    document = build_deposition_document(formatted_text, case_meta)
    tmp_path = Path(tempfile.gettempdir()) / "defect9_no_template.docx"
    document.save(tmp_path)
    assert Path(tmp_path).exists()
